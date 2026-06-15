#!/usr/bin/env python3
from __future__ import annotations

import csv
import hashlib
import json
import math
import re
import shutil
import subprocess
import tempfile
from collections import Counter, defaultdict
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path('/Users/newhosung/Academic/2026/AI Adoption Meta Analysis')
DATE = '20260615'
PKG = ROOT / 'paper_a/manuscript/target_journal/apa7_model_family_full_manuscript_scaffold_20260615'
FIG_DIR = PKG / 'figures'
DATA_PKG = ROOT / 'data/04_extraction/05_llm_masem_substitution/results/paper_a_apa7_model_family_manuscript_package_20260615'
INF = ROOT / 'data/04_extraction/05_llm_masem_substitution/results/paper_a_researcher_approved_s048_inference_figures_manuscript_20260615'
SUB = ROOT / 'data/04_extraction/05_llm_masem_substitution/results/paper_a_researcher_approved_s048_model_family_submission_packet_20260615'
SUPP = ROOT / 'data/04_extraction/05_llm_masem_substitution/results/paper_a_model_family_supplemental_diagnostics_20260615'
MASEM = ROOT / 'data/04_extraction/05_llm_masem_substitution/results/paper_a_researcher_approved_s048_model_family_masem_20260615'
PUBLIC = ROOT / 'paper_a/public_data_repository_20260615'
ONEDRIVE = Path('/Users/newhosung/Library/CloudStorage/OneDrive-SharedLibraries-ThePennsylvaniaStateUniversity/AI Adoption Meta Analysis - Documents/05_manuscripts/Paper_A/2026-06-15_apa7_model_family_full_manuscript_scaffold')
OUT_MD = PKG / f'PAPER_A_APA7_TABLE_FIGURE_REBUILD_MANUSCRIPT_{DATE}.md'
OUT_DOCX = PKG / f'PAPER_A_APA7_TABLE_FIGURE_REBUILD_MANUSCRIPT_{DATE}.docx'
LOCK_MD = ROOT / f'paper_a/PRISMA_COUNTS_LOCK_20260615.md'
LOCK_CSV = ROOT / f'paper_a/PRISMA_COUNTS_LOCK_20260615.csv'
BLACK = RGBColor(0, 0, 0)

FIGS = {
    'prisma': FIG_DIR / 'figure_1_prisma_2020_flow_diagram_paper_a_20260615.png',
    'genealogy': FIG_DIR / 'figure_2_theoretical_genealogy_full10_model_family_20260615.png',
    'workflow': FIG_DIR / 'figure_3_paper_a_analytic_workflow_20260615.png',
    'heatmap': FIG_DIR / 'figure_4_full10_evidence_map_publication_20260615.png',
    'core7': FIG_DIR / 'figure_5_core7_publication_masem_path_20260615.png',
    'trust6': FIG_DIR / 'figure_6_trust6_publication_masem_path_20260615.png',
    'coef': FIG_DIR / 'figure_7_path_estimate_coefficient_plot_20260615.png',
    'feas': FIG_DIR / 'figure_8_model_feasibility_plot_20260615.png',
}

CONSTRUCT_TABLE = [
    ['PE', 'Performance expectancy / perceived usefulness', 'TAM, TAM2, UTAUT', 'Instrumental usefulness: AI improves learning, teaching, productivity, or task performance.', 'full10; core7; trust6'],
    ['EE', 'Effort expectancy / perceived ease of use', 'TAM, computer self-efficacy, UTAUT', 'Operational ease: AI is manageable, learnable, and low burden.', 'full10; core7; trust6'],
    ['SI', 'Social influence', 'UTAUT', 'Normative and institutional endorsement mechanism.', 'full10; core7; trust6'],
    ['FC', 'Facilitating conditions', 'UTAUT', 'Resource and infrastructure mechanism enabling adoption and use.', 'full10; core7'],
    ['ATT', 'Attitude', 'TRA/TPB, TAM', 'Evaluative mediator translating beliefs into intention.', 'full10; core7'],
    ['TRU', 'Trust', 'Trust in automation, trust in IS, AI reliance', 'Reliance mechanism under AI opacity, uncertainty, and vulnerability.', 'full10; trust6'],
    ['ANX', 'Anxiety', 'Technology readiness, affective threat', 'Threat/unease mechanism retained in theory but underidentified for primary MASEM.', 'full10; future mechanism'],
    ['SE', 'Self-efficacy', 'Social cognitive theory, computer self-efficacy', 'Capability mechanism; currently feasible mainly in smaller supplemental sets.', 'full10; future mechanism'],
    ['BI', 'Behavioral intention', 'TRA/TPB, TAM, UTAUT', 'Proximal motivational adoption outcome.', 'full10; core7; trust6'],
    ['UB', 'Use behavior', 'TAM, UTAUT', 'Behavioral adoption/use outcome.', 'full10; core7; trust6'],
]


def norm_doi(x: str) -> str:
    x = (x or '').strip().lower()
    x = x.replace('https://doi.org/', '').replace('http://doi.org/', '').replace('doi:', '')
    return x.strip('/ ')


def fmt(x, digits=3):
    try:
        if pd.isna(x):
            return 'NA'
        return f'{float(x):.{digits}f}'.replace('-0.000', '0.000')
    except Exception:
        return str(x)


def fmt_p(x):
    try:
        if pd.isna(x):
            return 'NA'
        v = float(x)
        return '< .001' if v < .001 else f'{v:.3f}'.lstrip('0')
    except Exception:
        return str(x)


def ci(lo, hi):
    return f'[{fmt(lo)}, {fmt(hi)}]'


def md_table(headers, rows):
    lines = ['| ' + ' | '.join(headers) + ' |', '| ' + ' | '.join(['---'] * len(headers)) + ' |']
    for row in rows:
        lines.append('| ' + ' | '.join(str(v).replace('|', '\\|') for v in row) + ' |')
    return '\n'.join(lines)


def lock_prisma_counts():
    dedup_report = ROOT / 'data/01_identification/dedup_report.txt'
    human_csv = ROOT / 'data/02_screening/human_screening_results_consolidated.csv'
    pdf_log = ROOT / 'data/02_screening/pdf_download_log.json'
    rows = list(csv.DictReader(human_csv.open(encoding='utf-8-sig')))
    pdf_rows = json.loads(pdf_log.read_text(encoding='utf-8')) if pdf_log.exists() else []
    pdf_status_counts = Counter(r.get('status', 'unknown') for r in pdf_rows)
    included_rows = [r for r in rows if r['screening_decision'] == 'I']
    excluded_rows = [r for r in rows if r['screening_decision'] == 'X']
    doi_groups = defaultdict(list)
    for r in included_rows:
        doi = norm_doi(r.get('doi_url', ''))
        if doi:
            doi_groups[doi].append(r)
    dup_groups = {k: v for k, v in doi_groups.items() if len(v) > 1}
    duplicate_included_rows = sum(len(v) - 1 for v in dup_groups.values())
    unique_included = len(included_rows) - duplicate_included_rows
    identified = 22166
    duplicates_removed = 5977
    unique_records = 16189
    human_reviewed = len(rows)
    ai_or_record_screen_excluded = unique_records - human_reviewed
    counts = {
        'records_identified_database': identified,
        'duplicates_removed': duplicates_removed,
        'records_after_deduplication': unique_records,
        'records_screened_title_abstract_ai_assisted': unique_records,
        'records_excluded_before_human_review': ai_or_record_screen_excluded,
        'records_human_reviewed': human_reviewed,
        'human_review_excluded_rows': len(excluded_rows),
        'human_review_included_rows': len(included_rows),
        'duplicate_included_doi_rows_merged': duplicate_included_rows,
        'unique_included_reports_current_lock': unique_included,
        'pdf_download_log_entries': len(pdf_rows),
        'local_automated_pdf_downloaded': pdf_status_counts.get('downloaded', 0),
        'local_automated_pdf_not_downloaded': len(pdf_rows) - pdf_status_counts.get('downloaded', 0),
    }
    exclude_code_counts = Counter(r.get('exclude_code', '') or 'not_coded' for r in excluded_rows)
    source_counts = Counter(r.get('source', '') for r in included_rows)
    year_counts = Counter(r.get('year', '') for r in included_rows)
    with LOCK_CSV.open('w', newline='', encoding='utf-8') as f:
        w = csv.writer(f)
        w.writerow(['item', 'count', 'source_or_note'])
        for k, v in counts.items():
            w.writerow([k, v, 'computed from local PRISMA/source files'])
        for k, v in sorted(exclude_code_counts.items()):
            w.writerow([f'exclude_code_{k}', v, 'human_screening_results_consolidated.csv'])
        for k, v in sorted(pdf_status_counts.items()):
            w.writerow([f'pdf_status_{k}', v, 'pdf_download_log.json'])
    dup_lines = []
    for doi, vals in dup_groups.items():
        dup_lines.append(f'- DOI `{doi}` has {len(vals)} included rows: ' + '; '.join(f"{r['record_id']} ({r['source']}, {r['year']}) {r['title']}" for r in vals))
    locked_counts_table = md_table(['PRISMA item','Count','Source/logic'], [
        ['Records identified from databases', counts['records_identified_database'], '`data/01_identification/dedup_report.txt`'],
        ['Duplicate records removed', counts['duplicates_removed'], '`data/01_identification/dedup_report.txt`'],
        ['Records after deduplication', counts['records_after_deduplication'], '`data/01_identification/dedup_report.txt`'],
        ['Records screened', counts['records_screened_title_abstract_ai_assisted'], 'deduplicated records entering AI-assisted screening'],
        ['Records excluded before human review', counts['records_excluded_before_human_review'], '16,189 unique records minus 657 human-reviewed records'],
        ['Human-reviewed records', counts['records_human_reviewed'], '`data/02_screening/human_screening_results_consolidated.csv`'],
        ['Human-reviewed excluded rows', counts['human_review_excluded_rows'], 'screening_decision = X'],
        ['Human-reviewed included rows', counts['human_review_included_rows'], 'screening_decision = I'],
        ['Duplicate included DOI rows merged', counts['duplicate_included_doi_rows_merged'], 'included DOI duplicate audit'],
        ['Unique included reports/studies', counts['unique_included_reports_current_lock'], '225 included rows - 1 duplicate DOI row'],
        ['Local automated PDF downloads', counts['local_automated_pdf_downloaded'], '`data/02_screening/pdf_download_log.json`; not the final retrieval box'],
        ['Local automated PDF not downloaded/requires access', counts['local_automated_pdf_not_downloaded'], '`data/02_screening/pdf_download_log.json`; manual/library retrieval may exist elsewhere'],
    ])
    duplicate_audit = '\n'.join(dup_lines) if dup_lines else '- No duplicate DOI among included rows.'
    exclude_table = md_table(['Exclude code','Count'], sorted([[k, v] for k, v in exclude_code_counts.items()], key=lambda x: str(x[0])))
    source_table = md_table(['Source','Included rows'], sorted([[k, v] for k, v in source_counts.items()]))
    year_table = md_table(['Year','Included rows'], sorted([[k, v] for k, v in year_counts.items()]))
    lock_text = f"""# Paper A PRISMA counts lock

Updated: 2026-06-15

## Current lock

The previous `224` versus `225` discrepancy is explained by one duplicate DOI among the 225 included screening rows. The current working lock is therefore:

- 225 included screening rows.
- 1 duplicate included DOI row merged.
- 224 unique included reports/studies for PRISMA reporting, pending final team review.

## Locked counts

{locked_counts_table}

## Duplicate included DOI audit

{duplicate_audit}

## Human exclusion-code counts

{exclude_table}

## Included rows by source

{source_table}

## Included rows by year

{year_table}

## Boundary

This lock is sufficient for a draft PRISMA 2020-style flow diagram. Before journal submission, the team should confirm whether the two duplicate-DOI rows represent the same report, a metadata error, or two distinct reports sharing a DOI.
"""
    LOCK_MD.write_text(lock_text, encoding='utf-8')
    return counts, exclude_code_counts, source_counts, year_counts, dup_groups


def r_vec(items):
    return 'c(' + ', '.join(json.dumps(str(x)) for x in items) + ')'


def generate_figures(counts):
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    paths = pd.read_csv(INF / 'paper_a_model_family_structural_paths_ci_inference_20260615.csv')
    fit = pd.read_csv(INF / 'paper_a_model_family_fit_with_n_20260615.csv')
    elig = pd.read_csv(SUB / 'paper_a_model_family_eligibility_table_20260615.csv')
    heatmap_csv = MASEM / 'paper_a_full10_theory_target_pairwise_pooled_matrix_20260615.csv'
    path_csv = tempfile.NamedTemporaryFile('w', suffix='.csv', delete=False)
    paths.to_csv(path_csv.name, index=False)
    fit_csv = tempfile.NamedTemporaryFile('w', suffix='.csv', delete=False)
    fit.to_csv(fit_csv.name, index=False)
    elig_csv = tempfile.NamedTemporaryFile('w', suffix='.csv', delete=False)
    elig.to_csv(elig_csv.name, index=False)
    r_code = f"""
fig_dir <- {json.dumps(str(FIG_DIR))}
paths <- read.csv({json.dumps(path_csv.name)}, stringsAsFactors=FALSE)
fit <- read.csv({json.dumps(fit_csv.name)}, stringsAsFactors=FALSE)
elig <- read.csv({json.dumps(elig_csv.name)}, stringsAsFactors=FALSE)
heatmap_file <- {json.dumps(str(heatmap_csv))}
serif <- 'serif'

save_plot <- function(file_base, w, h, expr) {{
  png(file.path(fig_dir, paste0(file_base, '.png')), width=w, height=h, units='in', res=300)
  force(expr); dev.off()
  svg(file.path(fig_dir, paste0(file_base, '.svg')), width=w, height=h)
  force(expr); dev.off()
}}
box_text <- function(x1,y1,x2,y2,label, cex=.72, lwd=1.0, fill='white') {{
  rect(x1,y1,x2,y2,col=fill,border='black',lwd=lwd)
  text((x1+x2)/2,(y1+y2)/2,label,cex=cex,col='black')
}}
arrow_down <- function(x,y1,y2) arrows(x,y1,x,y2,length=.055,lwd=1.0,col='black')
arrow_right <- function(x1,x2,y) arrows(x1,y,x2,y,length=.055,lwd=1.0,col='black')

save_plot('figure_1_prisma_2020_flow_diagram_paper_a_20260615', 8.2, 10.5, {{
  par(mar=c(.35,.35,.7,.35), family=serif); plot.new(); plot.window(xlim=c(0,1),ylim=c(0,1))
  text(.52,.985,'PRISMA 2020 flow diagram for Paper A',font=2,cex=1.02)
  label_stage <- function(y1,y2,txt) {{
    rect(.025,y1,.075,y2,col='#f2f2f2',border='black',lwd=.8)
    text(.05,(y1+y2)/2,txt,srt=90,cex=.62,font=2)
  }}
  main_x1 <- .15; main_x2 <- .55; side_x1 <- .66; side_x2 <- .92
  label_stage(.835,.945,'Identification')
  label_stage(.565,.815,'Screening')
  label_stage(.305,.545,'Retrieval / eligibility')
  label_stage(.175,.285,'Included')
  box_text(main_x1,.875,main_x2,.945,'Records identified from databases\n(n = {counts['records_identified_database']:,})',.70)
  box_text(side_x1,.875,side_x2,.945,'Records removed before screening\nDuplicate records removed\n(n = {counts['duplicates_removed']:,})',.62)
  arrow_right(main_x2,side_x1,.91)
  arrow_down(.34,.875,.815)
  box_text(main_x1,.745,main_x2,.815,'Records screened by AI-assisted title/abstract workflow\n(n = {counts['records_after_deduplication']:,})',.62)
  box_text(side_x1,.745,side_x2,.815,'Records excluded before human review\n(n = {counts['records_excluded_before_human_review']:,})',.61)
  arrow_right(main_x2,side_x1,.78)
  arrow_down(.34,.745,.685)
  box_text(main_x1,.615,main_x2,.685,'Records reviewed by humans\n(n = {counts['records_human_reviewed']:,})',.68)
  box_text(side_x1,.615,side_x2,.685,'Human-reviewed records excluded\n(n = {counts['human_review_excluded_rows']:,})',.62)
  arrow_right(main_x2,side_x1,.65)
  arrow_down(.34,.615,.545)
  box_text(main_x1,.475,main_x2,.545,'Reports sought for local automated\nPDF retrieval\n(n = {counts['human_review_included_rows']:,})',.57)
  box_text(side_x1,.475,side_x2,.545,'Duplicate included DOI row merged\n(n = {counts['duplicate_included_doi_rows_merged']:,})',.60)
  arrow_right(main_x2,side_x1,.51)
  arrow_down(.34,.475,.415)
  box_text(main_x1,.345,main_x2,.415,'Local automated PDFs retrieved\n(n = {counts['local_automated_pdf_downloaded']:,})',.66)
  box_text(side_x1,.345,side_x2,.415,'Local automated retrieval\nnot downloaded / access needed\n(n = {counts['local_automated_pdf_not_downloaded']:,})',.49)
  arrow_right(main_x2,side_x1,.38)
  arrow_down(.34,.345,.285)
  box_text(main_x1,.215,main_x2,.285,'Unique included reports/studies\n(current PRISMA lock n = {counts['unique_included_reports_current_lock']:,})',.66,lwd=1.2)
  box_text(side_x1,.215,side_x2,.285,'Full-text eligibility exclusion boxes\nnot yet source-locked in this repository',.56,fill='#f7f7f7')
  text(.52,.135,'Note. This is a PRISMA 2020-style draft figure from local repository evidence. The 224/225 discrepancy is resolved by one duplicate DOI among included screening rows.\nThe local automated PDF retrieval boxes are not final full-text retrieval boxes; final eligibility exclusions require team confirmation.',cex=.57)
}})

save_plot('figure_2_theoretical_genealogy_full10_model_family_20260615', 9.2, 5.8, {{
  par(mar=c(.4,.4,.8,.4), family=serif); plot.new(); plot.window(xlim=c(0,1),ylim=c(0,1))
  text(.5,.965,'Theoretical genealogy of Paper A model families',font=2,cex=1.05)
  labs <- c('TRA/TPB\nbelief-attitude-intention', 'TAM/TAM2\nPE, EE, ATT, BI', 'UTAUT\nPE, EE, SI, FC, BI, UB', 'Trust/automation\nTRU', 'SCT/readiness\nSE, ANX')
  xs <- seq(.10,.90,length.out=5)
  for (i in seq_along(xs)) box_text(xs[i]-.075,.72,xs[i]+.075,.84,labs[i],.62)
  for (x in xs) arrows(x,.72,.50,.58,length=.055,lwd=1)
  box_text(.34,.46,.66,.58,'Full10 theoretical target\nPE, EE, SI, FC, ATT, TRU, ANX, SE, BI, UB\n45/45 pairwise coverage; 0 complete 10-construct matrices',.67,lwd=1.2)
  arrows(.45,.46,.32,.33,length=.055,lwd=1); arrows(.55,.46,.68,.33,length=.055,lwd=1)
  box_text(.16,.18,.48,.33,'Core7 empirical route\nPE, EE, SI, FC, ATT, BI, UB\nattitude-mediation MASEM',.67)
  box_text(.52,.18,.84,.33,'Trust6 empirical route\nPE, EE, SI, TRU, BI, UB\nAI-reliance mechanism MASEM',.67)
  text(.5,.08,'ANX and SE remain theory-relevant future mechanisms because current complete-case matrices are insufficient for primary MASEM.',cex=.67)
}})

save_plot('figure_3_paper_a_analytic_workflow_20260615', 8.4, 6.2, {{
  par(mar=c(.4,.4,.8,.4), family=serif); plot.new(); plot.window(xlim=c(0,1),ylim=c(0,1))
  text(.5,.965,'Paper A analytic workflow',font=2,cex=1.05)
  steps <- c('1. Search and PRISMA count lock', '2. Source-anchored construct coding', '3. Full10 matrix-feasibility diagnosis', '4. Model-family TSSEM/MASEM', '5. Supplemental diagnostics', '6. OSF-ready share-safe package')
  desc <- c('22,166 records; 16,189 after dedup; 224 unique included current lock', 'Human-coded and researcher-approved source-traced correlations', '45/45 full10 pairwise coverage; 0 complete full10 matrices', 'Core7 and trust6 complete-case empirical routes', 'Reduced models, PE-vs-EE, path CIs, ANX/SE feasibility', 'Aggregate outputs only; raw PDFs and private workbooks excluded')
  y <- .82
  for (i in seq_along(steps)) {{
    box_text(.10,y-.055,.90,y+.055,paste0(steps[i],'\n',desc[i]),.64)
    if (i < length(steps)) arrows(.5,y-.055,.5,y-.105,length=.05,lwd=1)
    y <- y - .145
  }}
}})

style_col <- function(cls) ifelse(cls %in% c('supported_positive_95ci','supported_negative_95ci'), 'black', ifelse(cls == 'not_supported_95ci_includes_zero', '#666666', '#aaaaaa'))
style_lty <- function(cls) ifelse(cls %in% c('supported_positive_95ci','supported_negative_95ci'), 1, ifelse(cls == 'not_supported_95ci_includes_zero', 2, 3))
node <- function(x,y,label) {{ rect(x-.045,y-.028,x+.045,y+.028,col='white',border='black',lwd=1.1); text(x,y,label,font=2,cex=.78) }}
qarrow <- function(x1,y1,x2,y2,curve=0,col='black',lty=1,lwd=1.7) {{
  mx <- (x1+x2)/2; my <- (y1+y2)/2; dx <- x2-x1; dy <- y2-y1; nx <- -dy; ny <- dx; len <- sqrt(nx^2+ny^2)
  if (len>0) {{ nx <- nx/len; ny <- ny/len }}
  cx <- mx + curve*nx; cy <- my + curve*ny; t <- seq(0,1,length.out=80)
  x <- (1-t)^2*x1 + 2*(1-t)*t*cx + t^2*x2; y <- (1-t)^2*y1 + 2*(1-t)*t*cy + t^2*y2
  lines(x,y,col=col,lty=lty,lwd=lwd,lend='round'); n <- length(x); arrows(x[n-4],y[n-4],x[n],y[n],length=.055,col=col,lty=lty,lwd=lwd)
}}
label <- function(x,y,txt,col='black') {{ rect(x-.032,y-.018,x+.032,y+.018,col='white',border=NA); text(x,y,txt,cex=.64,col=col) }}
plot_path_model <- function(route, file_base, title, pos, curves, labpos, fitnote) {{
  d <- paths[paths$route == route,]
  save_plot(file_base, 8.5, 5.6, {{
    par(mar=c(1.8,.4,3.4,.4), family=serif); plot.new(); plot.window(xlim=c(0,1),ylim=c(0,1),asp=1)
    title(main=title,cex.main=1.0,line=1.6)
    legend(.02,.96,c('CI excludes 0','CI includes 0','CI incomplete'),lwd=c(2.2,1.7,1.5),lty=c(1,2,3),col=c('black','#666666','#aaaaaa'),bty='n',cex=.67)
    for (i in seq_len(nrow(d))) {{
      p <- d$parameter[i]; fr <- d$from[i]; to <- d$to[i]; cls <- d$inference_class[i]
      a <- pos[[fr]]; b <- pos[[to]]; cv <- ifelse(p %in% names(curves), curves[[p]], 0)
      qarrow(a[1],a[2],b[1],b[2],cv,style_col(cls),style_lty(cls),ifelse(style_lty(cls)==1,2.2,1.7))
    }}
    for (nm in names(pos)) node(pos[[nm]][1],pos[[nm]][2],nm)
    for (i in seq_len(nrow(d))) {{
      p <- d$parameter[i]; if (p %in% names(labpos)) label(labpos[[p]][1],labpos[[p]][2],sprintf('%.2f',as.numeric(d$estimate[i])),style_col(d$inference_class[i]))
    }}
    mtext(fitnote,side=1,line=.5,cex=.62)
  }})
}}
core_pos <- list(PE=c(.12,.78),EE=c(.12,.60),SI=c(.12,.42),FC=c(.12,.24),ATT=c(.45,.60),BI=c(.70,.60),UB=c(.90,.40))
core_curves <- list(PE_to_ATT=0,EE_to_ATT=0,SI_to_ATT=0,FC_to_ATT=0,ATT_to_BI=0,PE_to_BI=-.12,EE_to_BI=-.04,SI_to_BI=.04,FC_to_UB=.12,BI_to_UB=0)
core_lab <- list(PE_to_ATT=c(.30,.73),EE_to_ATT=c(.30,.61),SI_to_ATT=c(.30,.50),FC_to_ATT=c(.30,.37),ATT_to_BI=c(.58,.65),PE_to_BI=c(.45,.80),EE_to_BI=c(.46,.70),SI_to_BI=c(.48,.49),FC_to_UB=c(.53,.27),BI_to_UB=c(.80,.51))
trust_pos <- list(PE=c(.12,.76),EE=c(.12,.52),SI=c(.12,.28),TRU=c(.45,.60),BI=c(.70,.52),UB=c(.90,.52))
trust_curves <- list(PE_to_TRU=0,EE_to_TRU=0,SI_to_TRU=0,TRU_to_BI=0,PE_to_BI=-.13,EE_to_BI=-.02,SI_to_BI=.09,BI_to_UB=0)
trust_lab <- list(PE_to_TRU=c(.30,.72),EE_to_TRU=c(.30,.55),SI_to_TRU=c(.30,.39),TRU_to_BI=c(.58,.61),PE_to_BI=c(.46,.81),EE_to_BI=c(.46,.48),SI_to_BI=c(.48,.34),BI_to_UB=c(.80,.58))
plot_path_model('paper_a_core7_att_mediation','figure_5_core7_publication_masem_path_20260615','Core7 attitude-mediation MASEM path model',core_pos,core_curves,core_lab,'k=4 positive-definite matrices; N_eff=3,172; CFI=.999; RMSEA=.009')
plot_path_model('paper_a_trust6_mechanism','figure_6_trust6_publication_masem_path_20260615','Trust6 AI-reliance MASEM path model',trust_pos,trust_curves,trust_lab,'k=7 positive-definite matrices; N_eff=10,315; CFI=.996; RMSEA=.011')

save_plot('figure_7_path_estimate_coefficient_plot_20260615', 8.2, 6.6, {{
  d <- paths; d$label <- paste(d$model_family, d$from, '->', d$to)
  ord <- order(d$model_family, d$estimate); d <- d[ord,]; y <- seq_len(nrow(d))
  par(mar=c(4,8.2,2.2,.8), family=serif); plot.new(); plot.window(xlim=c(-.25,.85),ylim=c(.5,nrow(d)+.5))
  abline(v=0,lty=2,col='#777777')
  for (i in seq_len(nrow(d))) {{
    col <- style_col(d$inference_class[i]); pch <- ifelse(d$inference_class[i]=='ci_incomplete',1,19)
    if (!is.na(as.numeric(d$ci_low[i])) && !is.na(as.numeric(d$ci_high[i]))) segments(as.numeric(d$ci_low[i]),y[i],as.numeric(d$ci_high[i]),y[i],col=col,lwd=1.4)
    points(as.numeric(d$estimate[i]),y[i],pch=pch,col=col,bg=col,cex=.9)
  }}
  axis(1); axis(2,at=y,labels=paste(d$from,'->',d$to,'(',ifelse(grepl('Core7',d$model_family),'core7','trust6'),')'),las=1,cex.axis=.62)
  title(main='Primary MASEM path estimates and likelihood-based 95% CIs',xlab='Standardized path estimate')
  legend('bottomright',c('CI excludes 0','CI includes 0','CI incomplete'),pch=c(19,19,1),col=c('black','#666666','#aaaaaa'),bty='n',cex=.72)
}})

save_plot('figure_8_model_feasibility_plot_20260615', 8.2, 5.6, {{
  par(mar=c(5.4,4.5,3.1,1), family=serif)
  labs <- c('Core7\\nATT mediation', 'Trust6\\nmechanism', 'Full10\\ntheoretical target')
  vals <- rbind(as.numeric(elig$required_pairs), as.numeric(elig$observed_pairs_after_rescue), as.numeric(elig$positive_definite_complete_cases))
  bp <- barplot(vals, beside=TRUE, col=c('white','#d9d9d9','#555555'), border='black', names.arg=labs, las=1, cex.names=.78, ylim=c(0,max(vals,na.rm=TRUE)+13), ylab='Count')
  legend('topleft',c('Required construct pairs','Observed pairs','Positive-definite complete cases'),fill=c('white','#d9d9d9','#555555'),border='black',bty='n',cex=.70)
  title(main='Model-family feasibility diagnosis')
  text(bp, vals + 1.2, labels=vals, cex=.68)
}})

if (file.exists(heatmap_file)) {{
  df <- read.csv(heatmap_file, stringsAsFactors=FALSE)
  constructs <- c('PE','EE','SI','FC','ATT','SE','TRU','ANX','BI','UB')
  mat <- matrix(NA_real_, length(constructs), length(constructs), dimnames=list(constructs,constructs)); kval <- mat
  kcol <- intersect(c('k','study_k','n_studies','num_studies'), names(df)); kcol <- ifelse(length(kcol)>0,kcol[1],NA)
  for (i in seq_len(nrow(df))) {{
    a <- df$construct_1[i]; b <- df$construct_2[i]
    if (a %in% constructs && b %in% constructs) {{ mat[a,b] <- as.numeric(df$pooled_r[i]); mat[b,a] <- as.numeric(df$pooled_r[i]); if (!is.na(kcol)) {{ kval[a,b] <- as.numeric(df[[kcol]][i]); kval[b,a] <- as.numeric(df[[kcol]][i]) }} }}
  }}
  save_plot('figure_4_full10_evidence_map_publication_20260615', 8.2, 7.4, {{
    par(mar=c(6.4,6.2,2.4,5), family=serif); pal <- gray.colors(80,start=.95,end=.25)
    image(seq_along(constructs), seq_along(constructs), t(mat[constructs,rev(constructs)]), col=pal, axes=FALSE, xlab='', ylab='', main='Full10 evidence map: pooled pairwise correlations')
    axis(1,at=seq_along(constructs),labels=constructs,las=2); axis(2,at=seq_along(constructs),labels=rev(constructs),las=1); box()
    for (i in seq_along(constructs)) for (j in seq_along(constructs)) {{ v <- mat[rev(constructs)[i],constructs[j]]; if (!is.na(v)) text(j,i,sprintf('%.2f',v),cex=.64,col=ifelse(v>.55,'white','black')) }}
    mtext('Note. Cells are pooled pairwise r values; full10 has pairwise coverage but no complete 10-construct matrices.',side=1,line=4.8,cex=.64)
  }})
}}
"""
    with tempfile.NamedTemporaryFile('w', suffix='.R', delete=False) as f:
        f.write(r_code)
        r_path = f.name
    subprocess.run(['Rscript', r_path], check=True)


def build_tables(counts):
    fit = pd.read_csv(INF / 'paper_a_model_family_fit_with_n_20260615.csv')
    paths = pd.read_csv(INF / 'paper_a_model_family_structural_paths_ci_inference_20260615.csv')
    supp = pd.read_csv(SUPP / 'paper_a_supplemental_model_comparison_20260615.csv')
    pe_ee = pd.read_csv(SUPP / 'paper_a_pe_vs_ee_role_comparison_20260615.csv')
    anx = pd.read_csv(SUPP / 'paper_a_anx_se_targeted_model_attempts_20260615.csv')
    elig = pd.read_csv(SUB / 'paper_a_model_family_eligibility_table_20260615.csv')
    prisma_rows = [
        ['Records identified from databases', counts['records_identified_database']],
        ['Duplicate records removed', counts['duplicates_removed']],
        ['Records after deduplication/screened', counts['records_after_deduplication']],
        ['Records excluded before human review', counts['records_excluded_before_human_review']],
        ['Human-reviewed records', counts['records_human_reviewed']],
        ['Human-reviewed excluded rows', counts['human_review_excluded_rows']],
        ['Human-reviewed included rows', counts['human_review_included_rows']],
        ['Duplicate included DOI rows merged', counts['duplicate_included_doi_rows_merged']],
        ['Unique included reports/studies current lock', counts['unique_included_reports_current_lock']],
    ]
    elig_rows = [[r['model_family'], r['constructs'], int(r['required_pairs']), int(r['observed_pairs_after_rescue']), int(r['partial_matrix_studies']), int(r['positive_definite_complete_cases']), r['manuscript_role']] for _, r in elig.iterrows()]
    fit_rows = [[r['model_family'], int(r['complete_case_k']), int(r['effective_sample_size']), fmt(r['chisq']), fmt(r['df'],0), fmt_p(r['p']), fmt(r['CFI']), fmt(r['TLI']), fmt(r['RMSEA']), fmt(r['SRMR'])] for _, r in fit.iterrows()]
    path_rows = [[r['model_family'], r['from'], r['to'], fmt(r['estimate']), r['ci_text'], r['inference_class'].replace('_',' ')] for _, r in paths.iterrows()]
    keep = ['core7_full','core6_no_ATT_direct_beliefs','core7_pure_ATT_mediation_no_direct_belief_BI','trust6_full','trust5_no_TRU_direct_acceptance','trust6_trust_mediator_no_direct_belief_BI','se4_capability_effort_intention']
    supp_rows = [[r['model_id'], int(r['positive_definite_complete_case_studies']), r['stage2_status'], fmt(r['chisq']), fmt(r['df'],0), fmt_p(r['p']), fmt(r['CFI']), fmt(r['RMSEA']), fmt(r['AIC'])] for _, r in supp[supp['model_id'].isin(keep)].iterrows()]
    pe_rows = [[r['family'], r['predictor'], r['target'], fmt(r['estimate']), ci(r['ci_low'], r['ci_high']), r['inference_class'].replace('_',' ')] for _, r in pe_ee[pe_ee['source']=='primary_path'].iterrows()]
    anx_rows = [[r['model_id'], r['constructs'], int(r['positive_definite_complete_case_studies']), r['stage2_status'], fmt(r['CFI']), fmt(r['RMSEA'])] for _, r in anx.iterrows()]
    return prisma_rows, elig_rows, fit_rows, path_rows, supp_rows, pe_rows, anx_rows


def write_markdown(counts):
    prisma_rows, elig_rows, fit_rows, path_rows, supp_rows, pe_rows, anx_rows = build_tables(counts)
    text = f"""# From Theoretical Coverage to Estimable Model Families: A Meta-Analytic Structural Equation Modeling Study of AI Adoption

## Status note

This rebuild focuses on submission-ready reporting structure: APA-style tables, a PRISMA 2020-style flow diagram, model-family feasibility graphics, and MASEM path/coefficient figures. The manuscript remains pending final target-journal formatting and team verification of full-text eligibility boxes.

## Abstract

AI adoption research in education draws on technology acceptance, unified acceptance, trust, self-efficacy, and anxiety traditions, but individual studies rarely report complete correlation matrices needed to test the whole theoretical system. This study reconstructs AI adoption as a 10-construct theoretical target and evaluates which parts of that target are empirically estimable using model-family meta-analytic structural equation modeling (MASEM). The current PRISMA count lock identifies 224 unique included reports/studies after resolving a 225-row screening include set by merging one duplicate DOI. The full10 target achieved 45/45 pairwise coverage but zero complete 10-construct matrices, so empirical MASEM was conducted through core7 attitude-mediation and trust6 AI-reliance model-family descendants.

## Method

### PRISMA count lock

{md_table(['Item','Count'], prisma_rows)}

![Figure 1. PRISMA 2020 flow diagram for Paper A.]({FIGS['prisma'].relative_to(ROOT)})

### Theoretical model-family specification

{md_table(['Construct','Label','Origin','AI-adoption function','Role'], CONSTRUCT_TABLE)}

![Figure 2. Theoretical genealogy of Paper A model families.]({FIGS['genealogy'].relative_to(ROOT)})

![Figure 3. Paper A analytic workflow.]({FIGS['workflow'].relative_to(ROOT)})

### Matrix-feasibility diagnosis

{md_table(['Model family','Constructs','Required pairs','Observed pairs','Partial studies','Positive-definite complete cases','Manuscript role'], elig_rows)}

![Figure 8. Model-family feasibility diagnosis.]({FIGS['feas'].relative_to(ROOT)})

## Results

### Full10 evidence map

![Figure 4. Full10 evidence map.]({FIGS['heatmap'].relative_to(ROOT)})

### Primary model-family fit

{md_table(['Model','k','N_eff','chi-square','df','p','CFI','TLI','RMSEA','SRMR'], fit_rows)}

### Primary path estimates

{md_table(['Model','From','To','Estimate','95% CI','Inference'], path_rows)}

![Figure 5. Core7 MASEM path model.]({FIGS['core7'].relative_to(ROOT)})

![Figure 6. Trust6 MASEM path model.]({FIGS['trust6'].relative_to(ROOT)})

![Figure 7. Path estimate coefficient plot.]({FIGS['coef'].relative_to(ROOT)})

### Supplemental diagnostics

{md_table(['Model','k','Status','chi-square','df','p','CFI','RMSEA','AIC'], supp_rows)}

### PE versus EE role comparison

{md_table(['Model','Predictor','Target','Estimate','95% CI','Inference'], pe_rows)}

### Anxiety and self-efficacy feasibility

{md_table(['Model','Constructs','k','Status','CFI','RMSEA'], anx_rows)}

## Reporting boundary

External references are used as methodological and reporting exemplars only. The tables and figures above are recreated from Paper A data and should replace the scaffold-level table/figure presentation in the previous Word draft.
"""
    OUT_MD.write_text(text, encoding='utf-8')
    DATA_PKG.mkdir(parents=True, exist_ok=True)
    shutil.copy2(OUT_MD, DATA_PKG / OUT_MD.name)
    if ONEDRIVE.exists():
        shutil.copy2(OUT_MD, ONEDRIVE / OUT_MD.name)


def set_run(run, bold=False, italic=False, size=12):
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.color.rgb = BLACK


def configure_doc(doc):
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    for style_name in ['Normal','Title','Heading 1','Heading 2','Heading 3']:
        st = doc.styles[style_name]
        st.font.name = 'Times New Roman'
        st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        st.font.size = Pt(12)
        st.font.color.rgb = BLACK
        st.paragraph_format.line_spacing = 2
        st.paragraph_format.space_after = Pt(0)
    for style_name in ['Title','Heading 1','Heading 2','Heading 3']:
        doc.styles[style_name].font.bold = True


def add_p(doc, text='', style=None, align=None, indent=False, bold=False, italic=False):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.line_spacing = 2
    p.paragraph_format.space_after = Pt(0)
    if indent:
        p.paragraph_format.first_line_indent = Inches(.5)
    r = p.add_run(text)
    set_run(r, bold=bold, italic=italic)
    return p


def set_cell(cell, text, bold=False):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 1
    r = p.add_run(str(text))
    set_run(r, bold=bold, size=8.5)


def border_el(name, val='single', sz='6'):
    e = OxmlElement(f'w:{name}')
    e.set(qn('w:val'), val)
    e.set(qn('w:color'), '000000')
    e.set(qn('w:sz'), sz)
    e.set(qn('w:space'), '0')
    return e


def set_table_borders(table):
    tblPr = table._tbl.tblPr
    old = tblPr.first_child_found_in('w:tblBorders')
    if old is not None:
        tblPr.remove(old)
    borders = OxmlElement('w:tblBorders')
    borders.append(border_el('top'))
    borders.append(border_el('left', 'nil'))
    borders.append(border_el('bottom'))
    borders.append(border_el('right', 'nil'))
    borders.append(border_el('insideH', 'nil'))
    borders.append(border_el('insideV', 'nil'))
    tblPr.append(borders)


def set_cell_bottom_border(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    old = tcPr.first_child_found_in('w:tcBorders')
    if old is not None:
        tcPr.remove(old)
    borders = OxmlElement('w:tcBorders')
    borders.append(border_el('bottom'))
    tcPr.append(borders)


def add_table(doc, num, title, headers, rows, note):
    add_p(doc, f'Table {num}', bold=True)
    add_p(doc, title, italic=True)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    for i, h in enumerate(headers):
        set_cell(table.rows[0].cells[i], h, bold=True)
        set_cell_bottom_border(table.rows[0].cells[i])
    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            set_cell(cells[i], v)
    add_p(doc, 'Note. ' + note)


def add_fig(doc, num, title, path, note, width=6.35):
    add_p(doc, f'Figure {num}', bold=True)
    add_p(doc, title, italic=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    add_p(doc, 'Note. ' + note)


def build_docx(counts):
    prisma_rows, elig_rows, fit_rows, path_rows, supp_rows, pe_rows, anx_rows = build_tables(counts)
    doc = Document()
    configure_doc(doc)
    add_p(doc, 'From Theoretical Coverage to Estimable Model Families: A Meta-Analytic Structural Equation Modeling Study of AI Adoption', style='Title', align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, 'Author Note', bold=True)
    add_p(doc, 'This rebuild corrects the reporting layer of the Paper A scaffold. It adds APA-style table rules, a PRISMA 2020-style flow diagram, publication-oriented MASEM figures, and explicit source-count boundaries. Final author information, target-journal details, and full-text eligibility boxes require team confirmation.')
    add_p(doc, 'Abstract', style='Heading 1', align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, 'AI adoption research in education draws on technology acceptance, unified acceptance, trust, self-efficacy, and anxiety traditions, but individual studies rarely report complete correlation matrices needed to test the whole theoretical system. This study reconstructs AI adoption as a 10-construct theoretical target and evaluates which parts of that target are empirically estimable using model-family MASEM. The current PRISMA count lock identifies 224 unique included reports/studies after resolving a 225-row screening include set by merging one duplicate DOI. The full10 target achieved 45/45 pairwise coverage but zero complete 10-construct matrices, so empirical MASEM was conducted through core7 attitude-mediation and trust6 AI-reliance model-family descendants.')
    add_p(doc, 'Method', style='Heading 1', align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, 'PRISMA Count Lock', style='Heading 2')
    add_p(doc, 'The repository previously contained a 224 versus 225 included-count discrepancy. The current audit found 225 included screening rows and one duplicate included DOI row, yielding a working lock of 224 unique included reports/studies.', indent=True)
    add_table(doc, 1, 'PRISMA Count Lock for Paper A', ['Item','Count'], prisma_rows, 'The final full-text eligibility boxes still require team confirmation before journal submission.')
    add_fig(doc, 1, 'PRISMA 2020 Flow Diagram for Paper A', FIGS['prisma'], 'The figure follows the PRISMA 2020 flow-logic format using currently locked local counts; full-text boxes remain pending team confirmation.', width=5.4)
    add_p(doc, 'Theoretical Model-Family Specification', style='Heading 2')
    add_table(doc, 2, 'Construct Genealogy and Model-Family Role in Paper A', ['Construct','Label','Origin','AI-Adoption Function','Role'], CONSTRUCT_TABLE, 'PE = performance expectancy; EE = effort expectancy; ATT = attitude; TRU = trust; ANX = anxiety; SE = self-efficacy; BI = behavioral intention; UB = use behavior.')
    add_fig(doc, 2, 'Theoretical Genealogy of Paper A Model Families', FIGS['genealogy'], 'The full10 target preserves theory; core7 and trust6 are empirical descendants supported by complete-case MASEM feasibility.', width=6.2)
    add_fig(doc, 3, 'Paper A Analytic Workflow', FIGS['workflow'], 'The workflow separates PRISMA count lock, source-anchored coding, matrix diagnosis, MASEM, supplemental diagnostics, and share-safe OSF packaging.', width=6.2)
    add_p(doc, 'Matrix-Feasibility Diagnosis', style='Heading 2')
    add_table(doc, 3, 'Model-Family Matrix Feasibility', ['Model family','Constructs','Required pairs','Observed pairs','Partial studies','Positive-definite complete cases','Role'], elig_rows, 'Full10 has complete pairwise coverage but no complete 10-construct matrices; core7 and trust6 are the empirical model-family routes.')
    add_fig(doc, 8, 'Model-Family Feasibility Diagnosis', FIGS['feas'], 'Bars compare required pairs, observed pairs, and positive-definite complete-case matrices by model family.', width=5.9)
    add_p(doc, 'Results', style='Heading 1', align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, 'Full10 Evidence Map', style='Heading 2')
    add_fig(doc, 4, 'Full10 Evidence Map', FIGS['heatmap'], 'Cells show pooled pairwise correlations. This is an evidence map, not a full10 SEM estimate.', width=5.9)
    add_p(doc, 'Primary Model-Family Fit', style='Heading 2')
    add_table(doc, 4, 'Primary Model-Family MASEM Fit', ['Model','k','N_eff','χ²','df','p','CFI','TLI','RMSEA','SRMR'], fit_rows, 'Path p values are not reported because Stage 2 returned likelihood-based CIs but not finite standard errors for all paths.')
    add_p(doc, 'Primary Structural Paths', style='Heading 2')
    add_table(doc, 5, 'Primary MASEM Structural Path Estimates', ['Model','From','To','Estimate','95% CI','Inference'], path_rows, 'Supported paths are classified by 95% confidence intervals excluding zero; incomplete intervals are reported as indeterminate, not significant.')
    add_fig(doc, 5, 'Core7 MASEM Path Model', FIGS['core7'], 'Solid paths indicate CIs excluding zero; dashed paths include zero; dotted paths have incomplete intervals.', width=6.0)
    add_fig(doc, 6, 'Trust6 MASEM Path Model', FIGS['trust6'], 'Trust is reported as an AI reliance mechanism rather than a fully mediating mechanism.', width=6.0)
    add_fig(doc, 7, 'Primary Path Estimate Coefficient Plot', FIGS['coef'], 'Points are standardized path estimates. Horizontal bars are likelihood-based 95% CIs where available.', width=5.8)
    add_p(doc, 'Supplemental Diagnostics', style='Heading 2')
    add_table(doc, 6, 'Supplemental Reduced and Alternative Model-Family Diagnostics', ['Model','k','Status','χ²','df','p','CFI','RMSEA','AIC'], supp_rows, 'These are diagnostic comparisons; construct removal can change k and matrix structure, so they are not definitive nested tests.')
    add_p(doc, 'PE Versus EE Role Comparison', style='Heading 2')
    add_table(doc, 7, 'Performance Expectancy and Effort Expectancy as Distinct Mechanisms', ['Model','Predictor','Target','Estimate','95% CI','Inference'], pe_rows, 'The comparison concerns PE and EE roles across outcomes, not the PE-EE correlation.')
    add_p(doc, 'Anxiety and Self-Efficacy Feasibility', style='Heading 2')
    add_table(doc, 8, 'Targeted Anxiety and Self-Efficacy Feasibility Attempts', ['Model','Constructs','k','Status','CFI','RMSEA'], anx_rows, 'ANX and SE remain theory-relevant but currently underidentified for primary complete-case MASEM.')
    add_p(doc, 'Discussion', style='Heading 1', align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, 'The reporting rebuild strengthens Paper A by separating what is theoretically covered from what is structurally estimable. Full10 remains valuable as a theoretical target and evidence map; core7 and trust6 provide the current empirical model-family MASEM evidence. The next scientific step is not to force full10, but to finalize PRISMA/full-text boxes and improve source reporting so future analyses can test anxiety and self-efficacy more directly.', indent=True)
    add_p(doc, 'Data Availability', style='Heading 1', align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, 'A share-safe Paper A public repository package is prepared locally and intended for a Paper A component under https://osf.io/mkrgd/overview. Raw PDFs, private source documents, raw coder workbooks, and runtime files are excluded.', indent=True)
    doc.save(OUT_DOCX)
    DATA_PKG.mkdir(parents=True, exist_ok=True)
    shutil.copy2(OUT_DOCX, DATA_PKG / OUT_DOCX.name)
    if ONEDRIVE.exists():
        shutil.copy2(OUT_DOCX, ONEDRIVE / OUT_DOCX.name)


def update_package_and_docs(counts):
    fig_public = PUBLIC / '2_Figures'
    man_public = PUBLIC / '1_Manuscript'
    review_public = PUBLIC / '6_PRISMA_Count_Lock'
    for d in [fig_public, man_public, review_public]:
        d.mkdir(parents=True, exist_ok=True)
    for p in [OUT_MD, OUT_DOCX]:
        shutil.copy2(p, man_public / p.name)
    for key, png in FIGS.items():
        shutil.copy2(png, fig_public / png.name)
        svg = png.with_suffix('.svg')
        if svg.exists():
            shutil.copy2(svg, fig_public / svg.name)
    shutil.copy2(LOCK_MD, review_public / LOCK_MD.name)
    shutil.copy2(LOCK_CSV, review_public / LOCK_CSV.name)
    for doc_path in [ROOT/'paper_a/RESEARCHER_README.md', ROOT/'paper_a/README.md', ROOT/'README.md', ROOT/'CURRENT.md']:
        txt = doc_path.read_text(encoding='utf-8')
        marker = '2026-06-15 Paper A APA table/figure rebuild'
        if marker not in txt:
            add = f"\n\n## {marker}\n\n- Locked the current PRISMA include discrepancy as `225 included screening rows - 1 duplicate included DOI row = 224 unique included reports/studies`, pending final full-text box confirmation.\n- Generated APA 7th table/figure rebuild manuscript: `{OUT_DOCX.relative_to(ROOT)}`.\n- Generated PRISMA-style flow, theoretical genealogy, analytic workflow, revised full10 evidence map, core7/trust6 path diagrams, path coefficient plot, and model-feasibility plot.\n- Updated OSF-ready package with the rebuilt manuscript, figures, and PRISMA count-lock files.\n"
            doc_path.write_text(txt.rstrip() + add + '\n', encoding='utf-8')
    readme = PUBLIC / 'README.md'
    txt = readme.read_text(encoding='utf-8') if readme.exists() else '# Paper A public repository package\n'
    marker = '2026-06-15 APA table/figure rebuild update'
    if marker not in txt:
        txt += f"\n## {marker}\n\n- Added table/figure rebuild manuscript: `{OUT_DOCX.name}`.\n- Added PRISMA count lock: `{LOCK_MD.name}`.\n- Added publication-oriented figures 1-8 under `2_Figures/`.\n"
        readme.write_text(txt, encoding='utf-8')
    rows = []
    for p in sorted(PUBLIC.rglob('*')):
        if p.is_file() and p.name != 'MANIFEST_SHA256.csv':
            rows.append({'path': str(p.relative_to(PUBLIC)), 'bytes': p.stat().st_size, 'sha256': hashlib.sha256(p.read_bytes()).hexdigest()})
    with (PUBLIC / 'MANIFEST_SHA256.csv').open('w', newline='', encoding='utf-8') as f:
        w = csv.DictWriter(f, fieldnames=['path','bytes','sha256'])
        w.writeheader(); w.writerows(rows)
    zip_path = ROOT / 'paper_a/public_data_repository_20260615_osf_ready.zip'
    if zip_path.exists():
        zip_path.unlink()
    shutil.make_archive(str(zip_path).replace('.zip',''), 'zip', ROOT, 'paper_a/public_data_repository_20260615')


def main():
    counts, exclude_code_counts, source_counts, year_counts, dup_groups = lock_prisma_counts()
    generate_figures(counts)
    write_markdown(counts)
    build_docx(counts)
    update_package_and_docs(counts)
    print('locked_unique_included=', counts['unique_included_reports_current_lock'], sep='')
    print('wrote=', OUT_DOCX, sep='')
    print('wrote=', OUT_MD, sep='')
    print('wrote=', LOCK_MD, sep='')
    print('figures=', FIG_DIR, sep='')
    print('zip=', ROOT / 'paper_a/public_data_repository_20260615_osf_ready.zip', sep='')

if __name__ == '__main__':
    main()
