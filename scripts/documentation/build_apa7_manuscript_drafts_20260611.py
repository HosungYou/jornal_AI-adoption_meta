#!/usr/bin/env python3
"""Build APA 7 professional-style manuscript draft shells for Paper A and Paper B.

The drafts intentionally reserve Literature Review and Discussion for team
authors while preserving Introduction, Method/Analysis, Results, tables, and
claim-boundary notes from the current project evidence.
"""

from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT_A = ROOT / "paper_a" / "manuscript"
OUT_B = ROOT / "paper_b" / "manuscript"
DATE = "2026-06-11"


def set_run_font(run, *, size: float = 12, bold: bool = False, italic: bool = False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, size=12)


def set_cell_text(cell, text: str, *, bold: bool = False):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.text = ""
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, bold=bold)


def configure_document(doc: Document, running_title: str):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 2.0
    normal.paragraph_format.space_after = Pt(0)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.line_spacing = 2.0
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
    styles["Heading 1"].font.size = Pt(12)
    styles["Heading 1"].font.bold = True
    styles["Heading 2"].font.size = Pt(12)
    styles["Heading 2"].font.bold = True
    styles["Heading 3"].font.size = Pt(12)
    styles["Heading 3"].font.bold = True
    styles["Heading 3"].font.italic = True

    header = section.header
    for paragraph in list(header.paragraphs):
        paragraph._element.getparent().remove(paragraph._element)
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    table.autofit = False
    table.allow_autofit = False
    table.columns[0].width = Inches(5.75)
    table.columns[1].width = Inches(0.75)
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "nil")
        borders.append(tag)
    tbl_pr.append(borders)
    left = table.cell(0, 0).paragraphs[0]
    left.paragraph_format.space_after = Pt(0)
    left.paragraph_format.line_spacing = 1.0
    left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = left.add_run(running_title[:50].upper())
    set_run_font(run, size=12)
    right = table.cell(0, 1).paragraphs[0]
    right.paragraph_format.space_after = Pt(0)
    right.paragraph_format.line_spacing = 1.0
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(right)


def add_para(
    doc: Document,
    text: str = "",
    *,
    align=None,
    bold: bool = False,
    italic: bool = False,
    indent: bool = True,
    size: float = 12,
):
    paragraph = doc.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.line_spacing = 2.0
    paragraph.paragraph_format.space_after = Pt(0)
    if indent and text:
        paragraph.paragraph_format.first_line_indent = Inches(0.5)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return paragraph


def add_heading(doc: Document, text: str, level: int):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.line_spacing = 2.0
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.first_line_indent = None
    if level == 1:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.text = ""
    run = paragraph.add_run(text)
    set_run_font(run, size=12, bold=True, italic=(level == 3))


def add_title_page(doc: Document, title: str, subtitle: str, author: str, affiliation: str, note: str):
    for _ in range(3):
        doc.add_paragraph()
    p = add_para(doc, title, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, indent=False)
    p.paragraph_format.space_after = Pt(12)
    add_para(doc, subtitle, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    add_para(doc, author, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    add_para(doc, affiliation, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    add_para(doc, f"Draft date: {DATE}", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    doc.add_page_break()
    add_heading(doc, "Author Note", 1)
    add_para(doc, note, indent=True)
    doc.add_page_break()


def add_abstract(doc: Document, abstract: str, keywords: str):
    add_heading(doc, "Abstract", 1)
    add_para(doc, abstract, indent=False)
    p = add_para(doc, "", indent=False)
    r = p.add_run("Keywords: ")
    set_run_font(r, italic=True)
    r2 = p.add_run(keywords)
    set_run_font(r2)
    doc.add_page_break()


def add_bullets(doc: Document, items: Iterable[str]):
    for item in items:
        paragraph = doc.add_paragraph(style=None)
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.paragraph_format.line_spacing = 2.0
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(f"- {item}")
        set_run_font(run)


def add_numbered(doc: Document, items: Iterable[str]):
    for i, item in enumerate(items, start=1):
        paragraph = doc.add_paragraph(style=None)
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.paragraph_format.line_spacing = 2.0
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(f"{i}. {item}")
        set_run_font(run)


def add_table(doc: Document, number: str, title: str, headers: list[str], rows: list[list[str]]):
    p_num = add_para(doc, number, bold=True, indent=False)
    p_num.paragraph_format.keep_with_next = True
    p_title = add_para(doc, title, italic=True, indent=False)
    p_title.paragraph_format.keep_with_next = True
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = OxmlElement("w:tcW")
            tc_w.set(qn("w:w"), str(int(9360 / len(headers))))
            tc_w.set(qn("w:type"), "dxa")
            tc_pr.append(tc_w)
    add_para(doc, "Note. Draft table; update after lead analysis lock where indicated.", indent=False, size=10, italic=True)


def add_references(doc: Document, refs: list[str]):
    add_heading(doc, "References", 1)
    for ref in refs:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.first_line_indent = Inches(-0.5)
        paragraph.paragraph_format.line_spacing = 2.0
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(ref)
        set_run_font(run)


def write_markdown(path: Path, title: str, sections: list[dict], refs: list[str]):
    lines: list[str] = [f"# {title}", "", f"Draft date: {DATE}", ""]
    for block in sections:
        kind = block["type"]
        if kind == "h1":
            lines += [f"## {block['text']}", ""]
        elif kind == "h2":
            lines += [f"### {block['text']}", ""]
        elif kind == "h3":
            lines += [f"#### {block['text']}", ""]
        elif kind == "p":
            lines += [block["text"], ""]
        elif kind == "bullets":
            lines += [f"- {item}" for item in block["items"]] + [""]
        elif kind == "numbered":
            lines += [f"{i}. {item}" for i, item in enumerate(block["items"], start=1)] + [""]
        elif kind == "table":
            headers = block["headers"]
            lines += [f"**{block['number']}**", "", f"*{block['title']}*", ""]
            lines.append("| " + " | ".join(headers) + " |")
            lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
            for row in block["rows"]:
                lines.append("| " + " | ".join(row) + " |")
            lines += ["", "Note. Draft table; update after lead analysis lock where indicated.", ""]
    if refs:
        lines += ["## References", ""]
        lines += [f"- {ref}" for ref in refs]
        lines.append("")
    path.write_text("\n".join(lines), encoding="utf-8")


def build_docx(
    path: Path,
    *,
    running_title: str,
    title: str,
    subtitle: str,
    author: str,
    affiliation: str,
    note: str,
    abstract: str,
    keywords: str,
    sections: list[dict],
    refs: list[str],
):
    doc = Document()
    configure_document(doc, running_title)
    add_title_page(doc, title, subtitle, author, affiliation, note)
    add_abstract(doc, abstract, keywords)
    add_para(doc, title, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, indent=False)
    for block in sections:
        kind = block["type"]
        if kind == "h1":
            add_heading(doc, block["text"], 1)
        elif kind == "h2":
            add_heading(doc, block["text"], 2)
        elif kind == "h3":
            add_heading(doc, block["text"], 3)
        elif kind == "p":
            add_para(doc, block["text"], indent=block.get("indent", True))
        elif kind == "bullets":
            add_bullets(doc, block["items"])
        elif kind == "numbered":
            add_numbered(doc, block["items"])
        elif kind == "table":
            add_table(doc, block["number"], block["title"], block["headers"], block["rows"])
    if refs:
        add_references(doc, refs)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


PAPER_A_REFS = [
    "Ajzen, I. (1991). The theory of planned behavior. Organizational Behavior and Human Decision Processes, 50(2), 179-211.",
    "Blut, M., Chong, A., Tsiga, Z., & Venkatesh, V. (2022). Meta-analysis of the unified theory of acceptance and use of technology (UTAUT). Journal of the Association for Information Systems, 23(1), 13-95.",
    "Cheung, M. W.-L. (2015). Meta-analysis: A structural equation modeling approach. Wiley.",
    "Davis, F. D. (1989). Perceived usefulness, perceived ease of use, and user acceptance of information technology. MIS Quarterly, 13(3), 319-340.",
    "Fishbein, M., & Ajzen, I. (1975). Belief, attitude, intention and behavior: An introduction to theory and research. Addison-Wesley.",
    "Jak, S., & Cheung, M. W.-L. (2020). Meta-analytic structural equation modeling with moderating effects on SEM parameters. Psychological Methods, 25(4), 430-449.",
    "Scherer, R., Siddiq, F., & Tondeur, J. (2019). The technology acceptance model (TAM): A meta-analytic structural equation modeling approach. Computers & Education, 128, 13-35.",
    "Venkatesh, V., Morris, M. G., Davis, G. B., & Davis, F. D. (2003). User acceptance of information technology: Toward a unified view. MIS Quarterly, 27(3), 425-478.",
]

PAPER_B_REFS = [
    "Cheung, M. W.-L. (2015). Meta-analysis: A structural equation modeling approach. Wiley.",
    "Jak, S., & Cheung, M. W.-L. (2020). Meta-analytic structural equation modeling with moderating effects on SEM parameters. Psychological Methods, 25(4), 430-449.",
    "Page, M. J., McKenzie, J. E., Bossuyt, P. M., Boutron, I., Hoffmann, T. C., Mulrow, C. D., Shamseer, L., Tetzlaff, J. M., Akl, E. A., Brennan, S. E., Chou, R., Glanville, J., Grimshaw, J. M., Hrobjartsson, A., Lalu, M. M., Li, T., Loder, E. W., Mayo-Wilson, E., McDonald, S., ... Moher, D. (2021). The PRISMA 2020 statement: An updated guideline for reporting systematic reviews. BMJ, 372, n71.",
]


paper_a_sections = [
    {"type": "h1", "text": "Draft Scope and Team Boundary"},
    {
        "type": "p",
        "text": "This file drafts the portions of Paper A that can be assembled from the current project evidence without taking over the team-authored Literature Review or Discussion. The Literature Review and Discussion headings are retained only as insertion points for team authors. Final structural path estimates, indirect effects, moderator estimates, and fit statistics remain lead-analysis insertion points.",
    },
    {"type": "h1", "text": "Introduction"},
    {
        "type": "p",
        "text": "Artificial intelligence tools are now embedded in higher education through large language models, intelligent tutoring systems, automated assessment systems, writing assistants, recommendation tools, and analytics platforms. Their spread has produced a rapidly expanding empirical literature on adoption, acceptance, and use. Yet this literature remains difficult to interpret cumulatively because studies draw from overlapping but nonidentical acceptance frameworks, measure different subsets of constructs, and report evidence in formats that do not directly support a single structural synthesis.",
    },
    {
        "type": "p",
        "text": "Meta-analytic structural equation modeling (MASEM) is well suited to this problem because it can synthesize study-level correlation matrices and test a theory-guided network of relationships. In AI adoption research, this is especially important. Performance expectancy, effort expectancy, social influence, facilitating conditions, attitude, self-efficacy, behavioral intention, and use behavior provide continuity with TAM, TPB, and UTAUT. At the same time, trust in AI and AI anxiety capture psychological features of AI systems that are not fully reducible to general usefulness or ease-of-use beliefs.",
    },
    {
        "type": "p",
        "text": "The present study therefore develops a MASEM of AI adoption in higher education that integrates traditional technology acceptance constructs with AI-specific psychological constructs. The working model treats attitude as a theoretically meaningful mediator rather than assuming that the parsimonious UTAUT exclusion of attitude applies unchanged to AI adoption contexts. It also tests whether trust and anxiety contribute to behavioral intention beyond standard acceptance predictors.",
    },
    {
        "type": "p",
        "text": "The study is designed to answer four research questions: (a) whether core TAM/UTAUT structural paths hold in the pooled AI adoption evidence base, (b) whether AI-specific trust and anxiety add explanatory value beyond traditional predictors, (c) whether attitude mediates effects from performance and effort expectancy to behavioral intention, and (d) whether year, cultural context, education level, and AI tool type moderate focal relationships.",
    },
    {"type": "h1", "text": "Literature Review [Reserved for Team Contribution]"},
    {
        "type": "p",
        "text": "[Reserved. Team authors should draft the theory synthesis, construct definitions, prior TAM/UTAUT MASEM evidence, AI trust/anxiety evidence, and hypothesis-development prose here. This draft does not supply Literature Review prose.]",
        "indent": False,
    },
    {"type": "h1", "text": "Method"},
    {"type": "h2", "text": "Design and Reporting Standards"},
    {
        "type": "p",
        "text": "Paper A is the parent meta-analysis for the AI adoption evidence-synthesis project. It uses systematic-review procedures to identify eligible studies and applies MASEM to synthesize construct-level relationships. Reporting is organized for PRISMA 2020 and APA Meta-Analysis Reporting Standards alignment, with a separate public or supplementary package expected for search records, coding rules, analytic scripts, and model outputs.",
    },
    {"type": "h2", "text": "Search, Screening, and Eligibility"},
    {
        "type": "p",
        "text": "The documented search workflow yielded 22,166 records. After deduplication, 16,189 records remained for screening. The screening workflow used tiered AI-assisted and human-verification procedures. The current Paper A proposal brief reports 224 included empirical studies at the proposal checkpoint; the repository README separately marks the final full-text MASEM-eligible count as a lead-analysis value to be locked before submission. The manuscript should harmonize this count before journal submission.",
    },
    {
        "type": "p",
        "text": "Studies were eligible when they examined AI technology adoption, acceptance, or use in higher education; reported quantitative relationships among at least two focal constructs; included undergraduate students, graduate students, instructors, faculty, or comparable higher-education users; and provided direct correlations, matrix tables, Fornell-Larcker or related evidence, HTMT or sufficient statistics that could be evaluated for MASEM readiness. Studies without usable target construct-pair evidence were excluded from the primary MASEM input while preserving audit trails for sensitivity and source-risk review.",
    },
    {"type": "h2", "text": "Constructs and Model Architecture"},
    {
        "type": "p",
        "text": "The primary structural model uses 10 constructs: performance expectancy, effort expectancy, social influence, facilitating conditions, attitude, self-efficacy, trust in AI, AI anxiety, behavioral intention, and use behavior. The planned structural architecture places performance expectancy and effort expectancy upstream of attitude, attitude upstream of behavioral intention, and behavioral intention upstream of use behavior. Social influence, self-efficacy, trust, and anxiety are modeled as additional antecedents of behavioral intention, and facilitating conditions are modeled as an antecedent of use behavior.",
    },
    {"type": "h2", "text": "Coding, Source Adjudication, and Analysis-Ready Inputs"},
    {
        "type": "p",
        "text": "The extraction workflow separates raw coder workbooks, pre-adjudication disagreement records, source-document decisions, analysis-ready direct-r inputs, expanded direct-r-form inputs, and converted/source-statistic sensitivity inputs. Raw coder workbooks are preserved as returns and are not overwritten by downstream consensus or analysis scripts.",
    },
    {
        "type": "p",
        "text": "The 2026-06-05 Paper1 analysis-ready package defines three analysis sets. The primary set keeps source-reported or source-adjudicated direct-r rows for the main MASEM input. The expanded set retains all human-consensus direct-r-form rows while explicitly marking model-derived or converted rows. The sensitivity set keeps beta/path/source-statistic converted rows for sensitivity analysis rather than treating them as direct-r equivalents.",
    },
    {"type": "h2", "text": "Meta-Analytic Structural Equation Modeling Plan"},
    {
        "type": "p",
        "text": "The primary analysis will use two-stage MASEM. Stage 1 will pool study-level correlation matrices using a random-effects model and full-information handling of incomplete construct-pair coverage when supported by the software and data structure. Stage 2 will fit the prespecified structural model to the pooled matrix using weighted least squares. Moderator analyses will be evaluated with one-stage MASEM or equivalent meta-regression methods when the moderator data are sufficiently complete.",
    },
    {
        "type": "p",
        "text": "Sensitivity analyses will evaluate whether conclusions change when expanded direct-r-form rows, converted beta/path/source-statistic rows, source-risk exclusions, and influence diagnostics are included or excluded. The S072 ANX-EE r = 1.0 row is excluded from the primary model input and retained only as a trace/influence diagnostic.",
    },
    {"type": "h1", "text": "Results"},
    {"type": "h2", "text": "Search and Screening Status"},
    {
        "type": "p",
        "text": "The current project records document 22,166 identified records and 16,189 records after deduplication. The screening pipeline further reduced records through keyword filtering, AI-assisted classification, and human verification. The lead-analysis manuscript should insert the final PRISMA counts for full-text assessment, exclusions with reasons, and final MASEM-eligible studies after the final inclusion file is locked.",
    },
    {
        "type": "table",
        "number": "Table 1",
        "title": "Current Paper A Search and Screening Counts",
        "headers": ["Stage", "Count", "Status for manuscript"],
        "rows": [
            ["Records identified", "22,166", "Ready to report"],
            ["After deduplication", "16,189", "Ready to report"],
            ["After keyword filter", "3,274", "Repository README checkpoint"],
            ["AI-assisted human review queue", "1,457", "Repository README checkpoint"],
            ["Included studies", "224 proposal checkpoint; final count to lock", "Lead confirmation required before submission"],
        ],
    },
    {"type": "h2", "text": "Analysis-Ready Evidence Base"},
    {
        "type": "p",
        "text": "At the 2026-06-05 analysis-ready checkpoint, the primary direct-r input contained 822 rows, the expanded direct-r-form input contained 1,303 rows, the converted/source-statistic sensitivity input contained 481 rows, and the long stacked file contained 2,606 rows. Within the primary file, 805 rows were direct-r-like and 17 rows remained non-direct or source-statistic review candidates.",
    },
    {
        "type": "table",
        "number": "Table 2",
        "title": "Paper A Analysis-Ready Input Sets",
        "headers": ["Input set", "Rows", "Manuscript role"],
        "rows": [
            ["Primary direct-r", "822", "Main MASEM input after final source and sample-size policy checks"],
            ["Expanded direct-r-form", "1,303", "Sensitivity or expanded coverage layer"],
            ["Converted beta/path/source-statistic", "481", "Sensitivity layer, not direct-r equivalence"],
            ["Long stacked file", "2,606", "Audit and workflow trace layer"],
        ],
    },
    {"type": "h2", "text": "Pre-Model Quality-Control Findings"},
    {
        "type": "p",
        "text": "The pre-model QC report identified 55 construct-pair coverage rows and no duplicate study-pair keys. It also flagged moderator missingness and numeric QC issues. Sample-size and moderator missingness remain the main analysis-readiness risk for moderator and N-weighted SEM claims. The S072 ANX-EE r = 1.0 row is treated as an influence diagnostic and excluded from the primary model input.",
    },
    {
        "type": "table",
        "number": "Table 3",
        "title": "Paper A Pre-Model QC Items",
        "headers": ["QC item", "Current count or rate", "Draft interpretation"],
        "rows": [
            ["Construct-pair coverage rows", "55", "Supports pair-level coverage reporting"],
            ["Duplicate study-pair keys", "0", "No duplicate key blocker in QC report"],
            ["Moderator missingness fields", "7", "Moderator claims require missingness caveats"],
            ["Numeric QC flagged rows", "2,482", "Driven primarily by missing sample-size numeric fields"],
            ["r = +/-1 influence check", "2 rows: S072 ANX-EE in primary and expanded", "Exclude from primary; retain as diagnostic"],
        ],
    },
    {"type": "h2", "text": "Primary MASEM Results [Lead Analysis Insertion Point]"},
    {
        "type": "p",
        "text": "[Insert Stage 1 pooled correlation matrix, heterogeneity estimates, Stage 2 path coefficients, indirect effects, model fit, and sensitivity results after the lead-analysis run is locked. Do not convert this placeholder into substantive conclusions until the final TSSEM/OSMASEM specification, sample-size policy, and source-risk rules are finalized.]",
        "indent": False,
    },
    {"type": "h1", "text": "Discussion [Reserved for Team Contribution]"},
    {
        "type": "p",
        "text": "[Reserved. Team authors should draft interpretation, theoretical implications, practical implications, limitations, and future research after the lead-analysis Results are inserted.]",
        "indent": False,
    },
]


paper_b_sections = [
    {"type": "h1", "text": "Draft Scope and Team Boundary"},
    {
        "type": "p",
        "text": "This file drafts the Introduction, Method/Analysis, Results, claim boundary, and data-availability components for Paper B. It does not draft the team-authored Literature Review or Discussion. The manuscript is framed as an LLM augmentation and validation study for evidence synthesis, not as an LLM replacement claim or vendor-ranking benchmark.",
    },
    {"type": "h1", "text": "Introduction"},
    {
        "type": "p",
        "text": "Evidence synthesis for meta-analytic structural equation modeling requires extraction decisions that are more demanding than simple article summarization. A usable evidence record must distinguish source-reported direct correlations from converted statistics, map constructs consistently, preserve source provenance, identify source-absence cases, handle human-coder disagreement, and maintain enough sample-size and matrix information to support downstream SEM weighting and model fitting.",
    },
    {
        "type": "p",
        "text": "Large language models may help with this work, but their value depends on the unit of evaluation. Treating thousands of heterogeneous task units as one accuracy denominator would obscure the difference between low-consequence metadata, high-consequence direct-r extraction, source-risk triage, and downstream substitution risk. A defensible validation study therefore requires a source-anchored adjudicated human reference standard, locked model outputs, prespecified task-family scoring rules, and claim boundaries that separate workflow augmentation from autonomous replacement.",
    },
    {
        "type": "p",
        "text": "Paper B evaluates whether a prespecified LLM workflow can augment MASEM-ready extraction in an AI adoption evidence-synthesis project. The primary workflow is Codex GPT-5.5. Claude Sonnet and Gemini 3 Flash are used only as supplementary cross-model sensitivity and triage signals. The study asks whether the workflow can recover extraction targets, characterize errors by source condition, prioritize expert review, and preserve downstream MASEM or TSSEM conclusions under bounded diagnostic substitution checks.",
    },
    {"type": "h1", "text": "Literature Review [Reserved for Team Contribution]"},
    {
        "type": "p",
        "text": "[Reserved. Team authors should draft prior work on LLM-assisted systematic reviews, extraction automation, human-in-the-loop evidence synthesis, and benchmark limitations. This draft does not supply Literature Review prose.]",
        "indent": False,
    },
    {"type": "h1", "text": "Method"},
    {"type": "h2", "text": "Corpus, Data States, and Reference Standard"},
    {
        "type": "p",
        "text": "The parent corpus is an AI adoption in higher education MASEM project. Paper B uses the validation and extraction subset derived from that corpus. The workflow separates raw independent human coder workbooks, pre-adjudication human-human disagreement queues, source-document adjudication decisions, a frozen source-anchored adjudicated human reference layer, locked LLM outputs, and downstream diagnostic analysis files.",
    },
    {
        "type": "p",
        "text": "The canonical human-consensus package is the OneDrive folder Paper2_Human_Final_Consensus_20260605_v2, with Git reference-freeze and scoring artifacts derived downstream. The post-freeze full-corpus reference contains 213 studies and preserves caveats rather than silently rewriting raw coder workbooks. The legacy model-explicit denominator-family package contains 8,783 task units and remains useful as pre-full-corpus reproducibility evidence, but final full-corpus result claims are governed by the 2026-06-09 post-freeze Step 5 gate.",
    },
    {"type": "h2", "text": "Task Families and Scoring Rules"},
    {
        "type": "p",
        "text": "Task units are not interpreted as one accuracy denominator. Each row is assigned a denominator family and scoring eligibility rule. Direct-r extraction rows, converted or source-statistic numeric rows, metadata rows, human-review decision rows, source-absence rows, duplicate-source exclusions, blank/absence consensus rows, and trace rows are scored or interpreted separately.",
    },
    {
        "type": "p",
        "text": "Direct-r numeric extraction is scored with an absolute error tolerance of 0.005. Metadata extraction is reported using strict exact match and relaxed normalized match. Abstentions on scorable rows count as incorrect and are reported as workflow behavior. Blank/absence consensus and human-disagreement trace rows are interpreted as triage evidence rather than final content-accuracy rows.",
    },
    {"type": "h2", "text": "Model Scope and Locked Outputs"},
    {
        "type": "p",
        "text": "Codex GPT-5.5 is the primary prespecified workflow. Claude Sonnet and Gemini 3 Flash are retained as supplementary sensitivity and triage evidence only. Clean model-explicit locked outputs are available for Codex GPT-5.5, Claude Sonnet, and Gemini 3 Flash across 7,859 task units in the legacy model-explicit package. Earlier Claude default-unspecified rows are retained only as audit provenance after the Sonnet backfill.",
    },
    {"type": "h2", "text": "Analysis Plan"},
    {
        "type": "p",
        "text": "RQ1 evaluates extraction validity by denominator family and task stratum. RQ2 classifies errors by source condition, source-type status, denominator family, and downstream consequence. RQ3 evaluates whether model behavior, cross-model disagreement, source-risk flags, and human-disagreement traces prioritize expert review. Downstream substitution analyses are reported only as bounded diagnostics: they test whether expert-reviewed LLM-assisted inputs change human-reference pooled correlations or TSSEM paths under the current eligible subset.",
    },
    {
        "type": "p",
        "text": "The approved missing-N rule excludes rows without source-supported numeric sample size from N-weighted TSSEM/MASEM weighting unless later source checks supply numeric N. The deterministic reconciliation layer fills numeric N for 741 of 804 legacy primary rerun rows and excludes the remaining 63 rows from N-weighted SEM weighting. Therefore, all-row SEM wording is prohibited unless numeric N is completed for every SEM input row.",
    },
    {"type": "h1", "text": "Results"},
    {"type": "h2", "text": "Reference and Locked-Output Coverage"},
    {
        "type": "p",
        "text": "The legacy model-explicit package contains 8,783 task units. Codex GPT-5.5, Claude Sonnet, and Gemini 3 Flash each have clean model-explicit locked outputs for 7,859 eligible task units. These rows support denominator-family scoring and supplementary cross-model triage, not a single pooled vendor-ranking denominator.",
    },
    {
        "type": "table",
        "number": "Table 1",
        "title": "Paper B Data States and Claim Roles",
        "headers": ["Data state", "Current evidence", "Claim role"],
        "rows": [
            ["Frozen full-corpus reference", "213 studies frozen on 2026-06-09", "Current governing reference layer"],
            ["Legacy task-unit package", "8,783 task units", "Pre-full-corpus reproducibility and denominator-family evidence"],
            ["Clean model-explicit outputs", "7,859 rows per model", "RQ1-RQ3 task-family scoring and sensitivity"],
            ["Bounded source-rendered M1-R shard", "90 rows", "Staged diagnostic only, not full-corpus accuracy"],
            ["Core-6 TSSEM diagnostic", "15 complete-case studies", "Subset substitution-stability diagnostic"],
        ],
    },
    {"type": "h2", "text": "RQ1: Extraction Validity by Task Family"},
    {
        "type": "p",
        "text": "For the primary Codex GPT-5.5 workflow in the legacy model-explicit package, source-reported direct-r extraction contained 323 scored rows. Codex matched 3 rows within the 0.005 tolerance and abstained on 320 rows. The 43 source_blank_direct_r rows were retained in the direct-r extraction family but all received abstentions. Converted numeric strata included 30 beta rows, 53 beta/path-converted human-consensus rows, and 5 numeric source-statistic rows; all were abstained under the current locked prompt/input condition.",
    },
    {
        "type": "p",
        "text": "Metadata performance varied by field. Codex achieved exact and relaxed match for all scored source-type rows and statistic-count rows, and 16 of 19 study-design rows. Other metadata families, including AI type, user role, country, and first author, showed high abstention rates. These results describe task-family workflow behavior under locked inputs rather than general model capability.",
    },
    {
        "type": "table",
        "number": "Table 2",
        "title": "Primary Codex GPT-5.5 RQ1 Numeric and Metadata Results",
        "headers": ["Task stratum", "Scored rows", "Correct", "Abstentions", "Interpretation"],
        "rows": [
            ["Source-reported direct-r", "323", "3", "320", "Not safe for autonomous numeric substitution"],
            ["Source-blank direct-r", "43", "0", "43", "Direct-r family with weaker source evidence"],
            ["Converted beta", "30", "0", "30", "High-consequence numeric stratum"],
            ["Beta/path converted by human consensus", "53", "0", "53", "High-consequence numeric stratum"],
            ["Numeric source-statistic converted by human consensus", "5", "0", "5", "High-consequence numeric stratum"],
            ["Metadata source type", "18", "18", "0", "Strong field-specific performance"],
            ["Metadata study design", "19", "16", "0", "Relatively strong field-specific performance"],
        ],
    },
    {"type": "h2", "text": "RQ2: Error Taxonomy and Source Conditions"},
    {
        "type": "p",
        "text": "The dominant RQ2 pattern was abstention on scorable rows. Codex produced 320 abstentions for source-reported direct-r rows, 43 abstentions for source-blank direct-r rows, and 88 abstentions across converted/source-type numeric strata. Metadata extraction also contained 381 abstentions on scorable source-evidence rows and 29 metadata mismatches. These patterns show why high-consequence MASEM numeric rows cannot be replaced automatically under the current workflow.",
    },
    {
        "type": "p",
        "text": "Blank/absence consensus and human-disagreement trace rows behaved differently from final content-accuracy rows. They are useful because they show over-answering, abstention, and cross-model inconsistency in cases where source evidence is absent, ambiguous, or disputed. They should not be collapsed with direct-r or metadata accuracy.",
    },
    {"type": "h2", "text": "RQ3: Human-Review Triage and Cross-Model Sensitivity"},
    {
        "type": "p",
        "text": "The RQ3 triage analysis used the full 8,783 task-unit reference universe and left-joined locked model rows where available. It classified 1,196 rows as P0 expert-review numeric or MASEM tasks, 649 rows as P1 source or human-disagreement review tasks, 483 rows as P1 general review-signal tasks, 6,412 rows as P2 blank-behavior audit rows, 1 row as a P2 scoring-completeness check, and 42 rows as low priority after the primary workflow check.",
    },
    {
        "type": "p",
        "text": "Cross-model behavior disagreement appeared in 6,592 task units and is interpreted only as a review-prioritization signal. Human-disagreement traces appeared in 467 task units, source or trace risk appeared in 1,525 task units, and 924 task units were reference-only rows without locked model output. These signals support targeted review triage, not model ranking.",
    },
    {
        "type": "table",
        "number": "Table 3",
        "title": "RQ3 Review Priority Counts",
        "headers": ["Review priority", "Task units", "Manuscript interpretation"],
        "rows": [
            ["P0 expert-review numeric or MASEM", "1,196", "High-consequence numeric review"],
            ["P1 source or human-disagreement review", "649", "Source-risk or human-disagreement review"],
            ["P1 general review signal", "483", "Other review-prioritization signal"],
            ["P2 blank-behavior audit", "6,412", "Workflow behavior, not final accuracy"],
            ["P2 scoring-completeness check", "1", "Completeness audit"],
            ["P3 low priority after primary check", "42", "Low priority under current scoring"],
        ],
    },
    {"type": "h2", "text": "Post-Freeze Source-Rendered Diagnostic Evidence"},
    {
        "type": "p",
        "text": "After the 2026-06-09 full-corpus reference freeze, a source-rendered Step 5 workflow reached full target source-rendering coverage and executed a bounded 90-row M1-R shard with Codex GPT-5.5. The shard produced 90 of 90 locked rows, model_cli_error = 0, source quote policy violations = 0, 65 nonblank answers, and 25 abstentions. Exception-aware generic numeric scoring was 15 of 30 for direct/source-r rows, 27 of 30 for latent or construct correlations, and 13 of 30 for secondary beta/path rows. This is staged diagnostic evidence only and must not be interpreted as full-corpus accuracy.",
    },
    {"type": "h2", "text": "Downstream MASEM and TSSEM Diagnostic"},
    {
        "type": "p",
        "text": "The human-reference MASEM baseline is the Paper1 tiered primary model-ready file with 804 rows. The pre-tiered primary file contains 822 rows and is retained for audit. Expanded and converted inputs contain 1,303 and 481 rows, respectively, and are treated as sensitivity layers.",
    },
    {
        "type": "p",
        "text": "The P0/P1 expert-review layer covered 1,845 task units, including 1,196 P0 numeric/MASEM rows and 649 P1 source or human-disagreement rows. The expert-reviewed LLM-assisted primary input contains 804 rows. It applies 3 exact numeric replacements, all of which match the frozen human-reference values, yielding 0 nonzero value deltas relative to the human-reference baseline.",
    },
    {
        "type": "p",
        "text": "At the pooled-correlation level, the primary expert-reviewed LLM-assisted input has maximum absolute mean-r delta = 0.000000 and no structural edges with nonzero change. Source-risk exclusion and converted-input augmentation are sensitivity diagnostics, with maximum absolute mean-r deltas of 0.407000 and 0.116229, respectively, and nonzero changes on 9 structural edges in each sensitivity layer.",
    },
    {
        "type": "p",
        "text": "The bounded R/metaSEM TSSEM diagnostic used N-weighted eligible rows and the six-construct complete-case subset PE, EE, SI, FC, BI, and UB. Fifteen studies reported all 15 pairwise correlations for this subset, yielding 225 aggregated pair rows. Stage 1 random-effects TSSEM and Stage 2 path models converged for both the human-reference baseline and the expert-reviewed LLM-assisted input. The maximum absolute pooled-correlation delta was 0.00000000, and the structural paths were identical across scenarios: PE to BI = 0.376578, EE to BI = 0.271255, SI to BI = 0.242604, FC to UB = 0.222908, and BI to UB = 0.566349. Model fit was also identical: chi-square = 3.554181, df = 4, p = 0.469688, CFI = 1.000000, RMSEA = 0.000000, and SRMR = 0.025199.",
    },
    {
        "type": "table",
        "number": "Table 4",
        "title": "Downstream Substitution and TSSEM Diagnostic",
        "headers": ["Diagnostic layer", "Current result", "Permitted claim"],
        "rows": [
            ["Expert-reviewed primary input", "804 rows; 3 exact replacements; 0 nonzero value deltas", "No primary input change relative to human reference"],
            ["Pooled-correlation rerun", "Max abs mean-r delta = 0.000000", "Primary pooled-r unchanged"],
            ["Source-risk exclusion sensitivity", "Max abs mean-r delta = 0.407000; 9 edges changed", "Sensitivity only"],
            ["Converted-input augmentation sensitivity", "Max abs mean-r delta = 0.116229; 9 edges changed", "Sensitivity only"],
            ["Core-6 TSSEM diagnostic", "15 studies; paths and fit identical", "Subset diagnostic stability only"],
        ],
    },
    {"type": "h2", "text": "Claim Boundary"},
    {
        "type": "p",
        "text": "The current results support a bounded augmentation claim: the workflow is useful for structured locked-output evaluation and review triage, but current locked outputs are not sufficient for unsupervised MASEM substitution. The deterministic expert-reviewed rerun supports a narrow claim that the primary LLM-assisted input made no nonzero pooled-correlation changes relative to the human-reference baseline. The bounded core-6 TSSEM diagnostic supports subset stability for PE, EE, SI, FC, BI, and UB only. Broader statements about all constructs, all rows, indirect effects, or substantive SEM conclusions require the final approved TSSEM/MASEM specification and source-supported numeric N for every SEM input row.",
    },
    {"type": "h1", "text": "Discussion [Reserved for Team Contribution]"},
    {
        "type": "p",
        "text": "[Reserved. Team authors should draft interpretation of augmentation value, limitations of current locked outputs, implications for human-in-the-loop extraction, and future workflow improvements after the lead approves the final claim boundary.]",
        "indent": False,
    },
    {"type": "h1", "text": "Data and Code Availability"},
    {
        "type": "p",
        "text": "The share-safe Paper 2 public repository is available at https://osf.io/mkrgd/overview. It contains prompts, schemas, scoring rules, manifest-registered locked model outputs, derived analysis outputs, scripts, reporting checklists, and decision records. Raw article PDFs, raw human coder workbooks, and private OneDrive-only working materials are excluded. The public package should be interpreted with the same claim boundary used in this manuscript.",
    },
]


def main():
    OUT_A.mkdir(parents=True, exist_ok=True)
    OUT_B.mkdir(parents=True, exist_ok=True)

    paper_a_docx = OUT_A / "PAPER_A_APA7_INTRO_METHOD_RESULTS_DRAFT_20260611.docx"
    paper_a_md = OUT_A / "PAPER_A_APA7_INTRO_METHOD_RESULTS_DRAFT_20260611.md"
    paper_b_docx = OUT_B / "PAPER_B_APA7_INTRO_METHOD_RESULTS_DRAFT_20260611.docx"
    paper_b_md = OUT_B / "PAPER_B_APA7_INTRO_METHOD_RESULTS_DRAFT_20260611.md"

    paper_a_title = "AI Adoption in Higher Education: A Meta-Analytic Structural Equation Modeling Study"
    paper_b_title = "Can a Prespecified LLM Workflow Augment MASEM-Ready Evidence Extraction?"

    build_docx(
        paper_a_docx,
        running_title="AI Adoption MASEM",
        title=paper_a_title,
        subtitle="APA 7 Professional Draft Excluding Literature Review and Discussion",
        author="Hosung You",
        affiliation="College of Education, The Pennsylvania State University",
        note=(
            "This draft preserves the lead-analysis boundary. Literature Review and Discussion "
            "are reserved for team authors. Final model estimates and submission-ready claims "
            "must be inserted only after the lead-analysis run is locked."
        ),
        abstract=(
            "This draft prepares the Introduction, Method/Analysis, and Results scaffold for a "
            "meta-analytic structural equation modeling study of AI adoption in higher education. "
            "The project search workflow identified 22,166 records and retained 16,189 after "
            "deduplication. The current analysis-ready package contains 822 primary direct-r rows, "
            "1,303 expanded direct-r-form rows, and 481 converted beta/path/source-statistic "
            "sensitivity rows. The manuscript positions the study as a theory-guided MASEM "
            "integrating traditional technology acceptance constructs with trust in AI and AI "
            "anxiety. Final pooled correlations, path coefficients, indirect effects, moderator "
            "tests, and fit statistics remain lead-analysis insertion points."
        ),
        keywords=(
            "artificial intelligence adoption, higher education, meta-analytic structural "
            "equation modeling, technology acceptance, trust in AI, AI anxiety"
        ),
        sections=paper_a_sections,
        refs=PAPER_A_REFS,
    )
    write_markdown(paper_a_md, paper_a_title, paper_a_sections, PAPER_A_REFS)

    build_docx(
        paper_b_docx,
        running_title="LLM Extraction Validation",
        title=paper_b_title,
        subtitle="APA 7 Professional Draft Excluding Literature Review and Discussion",
        author="Hosung You",
        affiliation="College of Education, The Pennsylvania State University",
        note=(
            "This draft preserves the approved claim boundary: Paper B is an LLM augmentation "
            "and validation study for evidence synthesis, not an LLM replacement study or "
            "vendor-ranking benchmark. Literature Review and Discussion are reserved for team authors."
        ),
        abstract=(
            "This draft prepares the Introduction, Method/Analysis, and Results scaffold for a "
            "source-anchored validation study of LLM-assisted MASEM-ready extraction. The workflow "
            "uses a source-anchored adjudicated human reference layer, locked model outputs, and "
            "task-family scoring rather than a single pooled accuracy denominator. In the legacy "
            "model-explicit package, Codex GPT-5.5, Claude Sonnet, and Gemini 3 Flash each have "
            "7,859 clean locked task-unit rows; the primary Codex workflow matched 3 of 323 "
            "source-reported direct-r rows within tolerance and abstained on 320. Downstream "
            "expert-reviewed substitution produced 0 nonzero primary value deltas, and a bounded "
            "core-6 TSSEM diagnostic showed identical paths and fit in a 15-study complete-case "
            "subset. Full-corpus and all-construct SEM claims remain gated."
        ),
        keywords=(
            "large language models, evidence synthesis, MASEM, extraction validation, "
            "human adjudication, task-family scoring"
        ),
        sections=paper_b_sections,
        refs=PAPER_B_REFS,
    )
    write_markdown(paper_b_md, paper_b_title, paper_b_sections, PAPER_B_REFS)

    print(paper_a_docx)
    print(paper_a_md)
    print(paper_b_docx)
    print(paper_b_md)


if __name__ == "__main__":
    main()
