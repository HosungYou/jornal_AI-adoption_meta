#!/usr/bin/env python3
from __future__ import annotations

import csv
import shutil
from pathlib import Path
from textwrap import dedent

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path('/Users/newhosung/Academic/2026/AI Adoption Meta Analysis')
DATE = '20260615'
PKG = ROOT / 'paper_a/manuscript/target_journal/apa7_model_family_full_manuscript_scaffold_20260615'
DATA_PKG = ROOT / 'data/04_extraction/05_llm_masem_substitution/results/paper_a_apa7_model_family_manuscript_package_20260615'
SUPP = ROOT / 'data/04_extraction/05_llm_masem_substitution/results/paper_a_model_family_supplemental_diagnostics_20260615'
ONEDRIVE = Path('/Users/newhosung/Library/CloudStorage/OneDrive-SharedLibraries-ThePennsylvaniaStateUniversity/AI Adoption Meta Analysis - Documents/05_manuscripts/Paper_A/2026-06-15_apa7_model_family_full_manuscript_scaffold')
REFS = PKG / f'PAPER_A_EXPANDED_REFERENCE_BANK_{DATE}.csv'
FIG_DIR = PKG / 'figures'
CURRENT = ROOT / 'CURRENT.md'
README = PKG / f'README_PAPER_A_APA7_MODEL_FAMILY_MANUSCRIPT_PACKAGE_{DATE}.md'

OUT_MD = PKG / f'PAPER_A_APA7_REVISED_MODEL_FAMILY_MANUSCRIPT_WITH_SUPPLEMENTAL_DIAGNOSTICS_{DATE}.md'
OUT_DOCX = PKG / f'PAPER_A_APA7_REVISED_MODEL_FAMILY_MANUSCRIPT_WITH_SUPPLEMENTAL_DIAGNOSTICS_{DATE}.docx'
DATA_MD = DATA_PKG / OUT_MD.name
DATA_DOCX = DATA_PKG / OUT_DOCX.name
ONEDRIVE_MD = ONEDRIVE / OUT_MD.name
ONEDRIVE_DOCX = ONEDRIVE / OUT_DOCX.name

CONSTRUCT_TABLE = [
    ['PE', 'Performance expectancy / perceived usefulness', 'TAM, TAM2, UTAUT', 'Instrumental outcome belief: AI is valuable because it improves learning, teaching, productivity, or task performance.', 'full10; core7; trust6'],
    ['EE', 'Effort expectancy / perceived ease of use', 'TAM, computer self-efficacy, UTAUT', 'Operational-friction belief: AI is acceptable when use is manageable, learnable, and not cognitively burdensome.', 'full10; core7; trust6'],
    ['SI', 'Social influence', 'UTAUT', 'Normative/institutional pressure and endorsement mechanism.', 'full10; core7; trust6'],
    ['FC', 'Facilitating conditions', 'UTAUT', 'Resource and infrastructure mechanism that enables evaluation and enacted use.', 'full10; core7'],
    ['ATT', 'Attitude', 'TRA/TPB, TAM', 'Evaluative mediator translating beliefs into intention.', 'full10; core7'],
    ['TRU', 'Trust', 'Trust in automation, trust in IS, AI reliance', 'AI-specific reliance mechanism under opacity, autonomy, uncertainty, and vulnerability.', 'full10; trust6'],
    ['ANX', 'Anxiety', 'Technology readiness, affective threat', 'Threat/unease mechanism; retained but underidentified for complete-case primary MASEM.', 'full10; future mechanism'],
    ['SE', 'Self-efficacy', 'Social cognitive theory, computer self-efficacy', 'Capability mechanism; retained but mostly feasible only in smaller supplemental sets.', 'full10; future mechanism'],
    ['BI', 'Behavioral intention', 'TRA/TPB, TAM, UTAUT', 'Proximal motivational outcome.', 'full10; core7; trust6'],
    ['UB', 'Use behavior', 'TAM, UTAUT', 'Behavioral adoption outcome.', 'full10; core7; trust6'],
]


def fmt(x, digits=3):
    try:
        if pd.isna(x):
            return 'NA'
        return f'{float(x):.{digits}f}'.replace('-0.000', '0.000')
    except Exception:
        return str(x)


def fmt_p(x):
    try:
        if pd.isna(x):
            return 'NA'
        v = float(x)
        if v < .001:
            return '< .001'
        return f'{v:.3f}'.lstrip('0')
    except Exception:
        return str(x)


def ci(lo, hi):
    return f'[{fmt(lo)}, {fmt(hi)}]'


def read_refs():
    rows = []
    with REFS.open(newline='', encoding='utf-8') as f:
        for row in csv.DictReader(f):
            rows.append(row)
    return rows


def refs_text(refs):
    lines = []
    for row in sorted(refs, key=lambda r: r.get('citation', '')):
        citation = row.get('citation', '').strip()
        url = row.get('doi_or_url', '').strip()
        if citation:
            lines.append(f'{citation} {url}'.strip())
    return '\n\n'.join(lines)


def md_table(headers, rows):
    lines = ['| ' + ' | '.join(headers) + ' |', '| ' + ' | '.join(['---'] * len(headers)) + ' |']
    for row in rows:
        lines.append('| ' + ' | '.join(str(v).replace('|', '\\|').replace('\n', ' ') for v in row) + ' |')
    return '\n'.join(lines)


def table_model_fit(df):
    keep = ['core7_full','core6_no_ATT_direct_beliefs','core7_pure_ATT_mediation_no_direct_belief_BI','trust6_full','trust5_no_TRU_direct_acceptance','trust6_trust_mediator_no_direct_belief_BI','se4_capability_effort_intention']
    sub = df[df['model_id'].isin(keep)].copy()
    rows = []
    for _, r in sub.iterrows():
        rows.append([r['model_id'], int(r['positive_definite_complete_case_studies']), r['stage2_status'], fmt(r['chisq']), fmt(r['df'],0), fmt_p(r['p']), fmt(r['CFI']), fmt(r['TLI']), fmt(r['RMSEA']), fmt(r['SRMR']), fmt(r['AIC'])])
    return rows


def build_markdown():
    refs = read_refs()
    primary_fit = pd.read_csv(PKG / 'tables/paper_a_model_family_fit_with_n_20260615.csv')
    primary_paths = pd.read_csv(PKG / 'tables/paper_a_model_family_structural_paths_ci_inference_20260615.csv')
    model_comp = pd.read_csv(SUPP / 'paper_a_supplemental_model_comparison_20260615.csv')
    pe_ee = pd.read_csv(SUPP / 'paper_a_pe_vs_ee_role_comparison_20260615.csv')
    anx_scan = pd.read_csv(SUPP / 'paper_a_anx_se_complete_case_feasibility_scan_20260615.csv')
    anx_attempts = pd.read_csv(SUPP / 'paper_a_anx_se_targeted_model_attempts_20260615.csv')
    omitted = pd.read_csv(SUPP / 'paper_a_full10_omitted_pair_diagnostic_20260615.csv')

    core = primary_fit[primary_fit['route'] == 'paper_a_core7_att_mediation'].iloc[0]
    trust = primary_fit[primary_fit['route'] == 'paper_a_trust6_mechanism'].iloc[0]
    se_effort = anx_attempts[anx_attempts['model_id'] == 'se4_capability_effort_intention'].iloc[0]
    feasible_counts = anx_scan['feasibility'].value_counts().to_dict()
    future_pairs = int((omitted['diagnostic_priority'] == 'future_mechanism_feasibility_or_pairwise_only').sum())
    omitted_primary = int((~omitted['in_primary_structural_path']).sum())

    construct_md = md_table(['Construct', 'Construct Label', 'Theoretical Origin', 'AI-Adoption Function', 'Model Role'], CONSTRUCT_TABLE)
    primary_fit_md = md_table(['Model family', 'k', 'N_eff', 'chi-square', 'df', 'p', 'CFI', 'TLI', 'RMSEA', 'SRMR'], [
        ['Core7 attitude mediation', int(core['complete_case_k']), int(core['effective_sample_size']), fmt(core['chisq']), fmt(core['df'],0), fmt_p(core['p']), fmt(core['CFI']), fmt(core['TLI']), fmt(core['RMSEA']), fmt(core['SRMR'])],
        ['Trust6 AI trust mechanism', int(trust['complete_case_k']), int(trust['effective_sample_size']), fmt(trust['chisq']), fmt(trust['df'],0), fmt_p(trust['p']), fmt(trust['CFI']), fmt(trust['TLI']), fmt(trust['RMSEA']), fmt(trust['SRMR'])],
    ])
    supp_md = md_table(['Diagnostic model', 'k', 'Status', 'chi-square', 'df', 'p', 'CFI', 'TLI', 'RMSEA', 'SRMR', 'AIC'], table_model_fit(model_comp))
    pe_ee_primary_rows = []
    for _, r in pe_ee[pe_ee['source'] == 'primary_path'].iterrows():
        pe_ee_primary_rows.append([r['family'], r['predictor'], r['target'], fmt(r['estimate']), ci(r['ci_low'], r['ci_high']), r['inference_class']])
    pe_ee_md = md_table(['Model', 'Predictor', 'Target', 'Estimate', '95% CI', 'Inference'], pe_ee_primary_rows)
    anx_rows = []
    for _, r in anx_attempts.iterrows():
        anx_rows.append([r['model_id'], r['constructs'], int(r['positive_definite_complete_case_studies']), r['stage2_status'], fmt(r['CFI']), fmt(r['RMSEA']), r['rationale']])
    anx_md = md_table(['Model', 'Constructs', 'k', 'Status', 'CFI', 'RMSEA', 'Rationale'], anx_rows)

    text = f"""\
# From Theoretical Coverage to Estimable Model Families: A Meta-Analytic Structural Equation Modeling Study of AI Adoption

**Author Note**  
Author affiliations, acknowledgments, funding, conflicts of interest, data availability, and repository policy will be completed after the team confirms the final target journal and sharing boundary.

## Abstract

Artificial intelligence (AI) adoption research draws on technology acceptance, unified acceptance, trust, self-efficacy, and anxiety traditions, but primary studies rarely report the complete matrix needed to test the full theoretical system. This study reconstructs AI adoption as a full 10-construct theoretical target and evaluates which parts of that target are empirically estimable using model-family meta-analytic structural equation modeling (MASEM). The full10 target was generated by combining deductive theory reconstruction, inductive construct mapping from the extracted literature, and matrix-feasibility diagnosis. It achieved complete pairwise coverage across all 45 construct pairs but had no complete same-study 10-construct matrices. Therefore, empirical MASEM was conducted through estimable model-family descendants: a core7 attitude-mediation model and a trust6 AI-reliance model. The core7 model fit well, chi-square(5) = {fmt(core['chisq'])}, p = {fmt_p(core['p'])}, CFI = {fmt(core['CFI'])}, RMSEA = {fmt(core['RMSEA'])}. The trust6 model also fit well, chi-square(4) = {fmt(trust['chisq'])}, p = {fmt_p(trust['p'])}, CFI = {fmt(trust['CFI'])}, RMSEA = {fmt(trust['RMSEA'])}. Supplemental diagnostics showed that performance expectancy and effort expectancy operate as distinct usefulness and effort mechanisms rather than interchangeable belief indicators, that trust is best interpreted as a direct AI reliance mechanism rather than a fully mediating mechanism, and that anxiety/self-efficacy remain theoretically important but underidentified for primary complete-case MASEM. The study contributes a transparent framework for preserving broad AI adoption theory while reporting only empirically defensible structural models.

**Keywords:** artificial intelligence adoption, technology acceptance, trust in AI systems, attitude mediation, MASEM, TSSEM

## Introduction

AI adoption has become a central issue in education because AI-supported systems increasingly mediate writing, feedback, tutoring, learning analytics, assessment support, and administrative decision making. The literature has expanded quickly, but its theoretical organization remains uneven. Many studies cite technology acceptance or unified acceptance models, but each primary study typically measures only part of the broader adoption system. This creates a field-level problem: there is enough evidence to map many associations, but not necessarily enough complete matrix evidence to estimate a single omnibus structural model.

The present study treats that problem as theoretically meaningful rather than merely technical. AI adoption theory has inherited constructs from reasoned action and planned behavior, the technology acceptance model (TAM), TAM2, the unified theory of acceptance and use of technology (UTAUT), trust in automation, computer self-efficacy, technology readiness, and AI anxiety traditions (Ajzen, 1991; Bandura, 1977, 1997; Davis, 1989; Davis et al., 1989; Lee & See, 2004; Venkatesh & Davis, 2000; Venkatesh et al., 2003, 2012). These traditions imply a broad system of beliefs, evaluations, reliance judgments, capability judgments, threat responses, intentions, and behavior. Yet the system cannot be tested by simply drawing every theoretically plausible path and forcing it into sparse data.

### Theoretical Evolution of Adoption Models

The earliest acceptance logic begins with the proposition that beliefs shape attitudes, attitudes shape intentions, and intentions shape behavior. The theory of reasoned action and theory of planned behavior gave this sequence its general social-psychological form (Fishbein & Ajzen, 1975; Ajzen, 1991). TAM translated that logic into technology use by positioning perceived usefulness and perceived ease of use as central beliefs that shape attitude and behavioral intention (Davis, 1989; Davis et al., 1989). TAM2 and UTAUT expanded this structure by adding social influence, facilitating conditions, performance expectancy, effort expectancy, and moderators of technology acceptance (Venkatesh & Davis, 2000; Venkatesh et al., 2003, 2012).

AI adoption requires a further extension. Unlike many earlier educational technologies, AI systems are often opaque, probabilistic, autonomous, and embedded in consequential educational judgments. Therefore, acceptance is not only a question of whether a system is useful or easy to use. Users must also decide whether the system is trustworthy, whether they are capable of using it appropriately, and whether it evokes anxiety or threat. Trust in automation and information systems research provides the theoretical basis for treating trust as a reliance mechanism, whereas social cognitive theory and technology-readiness research support self-efficacy and anxiety as capability and threat mechanisms (Bandura, 1977, 1997; Compeau & Higgins, 1995; Gefen et al., 2003; Hancock et al., 2011; Lee & See, 2004; McKnight et al., 2002; Parasuraman, 2000).

### Present Study

This study asks how much of the broad AI adoption system can be tested with current source-anchored meta-analytic data. The answer requires separating three levels of evidence: theoretical coverage, pairwise meta-analytic evidence, and estimable structural evidence. The full10 target preserves theoretical coverage. The core7 and trust6 models test estimable structural mechanisms. Supplemental diagnostics evaluate whether reduced or alternative models change the interpretation of the primary model family.

## Theoretical Framework

### Constructing the Full10 Theoretical Target

The full10 target was generated through a combined deductive, inductive, and meta-analytic process. Deductively, the model begins with constructs repeatedly specified in TAM, TAM2, UTAUT, trust in automation, self-efficacy, and anxiety traditions. Inductively, the extracted AI adoption literature was mapped into harmonized construct families. Meta-analytically, the resulting construct families were evaluated for pairwise coverage and complete-case matrix feasibility. The full10 target therefore did not emerge as an arbitrary list. It is the intersection of theory history, observed construct use, and the matrix structure required for MASEM.

The target includes performance expectancy/perceived usefulness, effort expectancy/perceived ease of use, social influence, facilitating conditions, attitude, trust, anxiety, self-efficacy, behavioral intention, and use behavior. All 45 pairwise relations among these constructs were observed at least once in the source-anchored dataset. This justifies full10 as a theoretical evidence map. However, no study reported all 10 constructs in a complete same-study matrix. Thus, full10 is not currently an estimable complete-case SEM result.

### Performance Expectancy and Effort Expectancy as Distinct Belief Mechanisms

Performance expectancy and effort expectancy should not be collapsed into a single generic acceptance belief. Performance expectancy is an instrumental value judgment: users ask whether AI improves learning, teaching, productivity, quality, or task performance. Effort expectancy is an operational-friction judgment: users ask whether the system is understandable, manageable, learnable, and low burden. In education, this distinction is theoretically important because a tool can be powerful but difficult to use, or easy to use but not pedagogically valuable.

Accordingly, the supplemental diagnostics compare PE and EE as distinct mechanisms. This is not a test of the PE-EE correlation. It is a comparison of how usefulness-based and effort-based beliefs behave across attitude, intention, trust, and reduced model-family specifications. The primary results show positive PE and EE paths in the expected direction, but the strength of inference differs across targets. EE -> BI is supported in the trust6 model, whereas several PE paths are directionally positive but interval-incomplete. Therefore, Paper A should frame PE and EE as conceptually distinct mechanisms whose relative empirical roles remain partly model-dependent.

### Core7 as the Estimable Acceptance Backbone

Core7 is the empirical descendant of full10 that preserves the central TAM/UTAUT acceptance backbone while satisfying complete-case MASEM requirements. It includes PE, EE, social influence, facilitating conditions, attitude, behavioral intention, and use behavior. It is not a convenience model produced by dropping difficult constructs. It is the largest defensible complete-case model family that retains belief, evaluative, intention, and behavior components.

The core7 model allows attitude to mediate belief and condition effects on intention, while also allowing direct belief-to-intention paths. This is theoretically appropriate because AI users may form intentions both through evaluative attitudes and through direct judgments of utility, effort, or social expectation.

### Trust6 as an AI-Specific Reliance Mechanism

Trust6 is the empirical descendant of full10 that tests whether trust adds an AI-specific reliance mechanism. The model includes PE, EE, social influence, trust, behavioral intention, and use behavior. Trust is expected to matter because AI systems create reliance under uncertainty. Users may judge an AI system as useful or easy to use but still hesitate to rely on it if its outputs are opaque, unstable, biased, or misaligned with educational values.

The supplemental diagnostics show why trust should not be written merely as a full mediator. A trust-only mediation model that removed direct PE/EE/SI -> BI paths fit poorly. Thus, trust is better framed as a central AI reliance predictor that operates alongside usefulness, effort, and social influence, rather than as the sole channel through which acceptance beliefs affect intention.

### Anxiety and Self-Efficacy as Future Mechanisms

Anxiety and self-efficacy remain part of the full10 target because they represent threat and capability mechanisms. The present evidence does not justify deleting them from AI adoption theory. However, the complete-case matrix structure does not support their inclusion in the primary full10 MASEM. Supplemental feasibility scans found {feasible_counts.get('potential_primary_or_supplemental_candidate', 0)} ANX/SE-inclusive construct sets with at least four positive-definite complete-case matrices, but these were concentrated in smaller self-efficacy sets. Targeted anxiety models with attitude, trust, intention, and use behavior had zero positive-definite complete-case matrices. Therefore, anxiety and self-efficacy should be described as theoretically specified but empirically underidentified mechanisms.

### Table 1

*Construct Genealogy and Model-Family Role in Paper A*

{construct_md}

## Method

### Design

The study used two-stage meta-analytic structural equation modeling (TSSEM/MASEM) with a model-family strategy. The design separated the full theoretical target from estimable empirical model families. The full10 model was used to evaluate construct coverage and pairwise evidence. Complete-case MASEM was then conducted for model-family members that had positive-definite same-study correlation matrices.

### Source-Anchored Data and Construct Harmonization

The analysis used the source-anchored adjudicated human reference standard with researcher-approved S048 additions. Construct labels from primary studies were harmonized into the 10 target families. Correlations were retained when they were source-traceable and had usable sample-size information. The current analysis-ready input contained 836 rows. Paper B reference-standard claims remain separate and unchanged.

### Primary Model-Family MASEM

The primary empirical model family consisted of core7 and trust6. Core7 tested the attitude-mediation acceptance backbone. Trust6 tested the AI-specific trust mechanism. Stage 1 synthesized complete-case correlation matrices using random-effects TSSEM. Stage 2 fit the prespecified structural model to the pooled matrix. Model fit was evaluated using chi-square, CFI, TLI, RMSEA, SRMR, AIC, and BIC. Path-level interpretation relied on likelihood-based 95% confidence intervals because finite standard errors and z-based p values were not available for all paths.

### Supplemental Diagnostics

Supplemental analyses were designed to test whether the selected model family was defensible. First, reduced and alternative models removed PE, EE, SI, FC, ATT, or TRU, or changed mediation assumptions. Second, PE and EE were compared as distinct usefulness and effort mechanisms. Third, all full10 pairs omitted from the primary structural paths were catalogued. Fourth, ANX/SE-inclusive construct sets were scanned for complete-case feasibility and targeted ANX/SE model attempts were fit when possible. These diagnostics are sensitivity analyses, not a mechanical model-selection tournament, because construct removal can alter the complete-case study set.

## Results

### Full10 Evidence Map

The full10 target achieved complete pairwise coverage across all 45 construct pairs. This supports its role as the theoretical evidence map. However, zero studies contained a complete 10-construct matrix, so the full10 target was not estimable as a single complete-case SEM. Among the 45 full10 pairs, {omitted_primary} were not represented as primary structural paths in core7 or trust6. Of those, {future_pairs} involved anxiety or self-efficacy and are therefore central to the future-mechanism agenda.

### Primary Model-Family Fit

Table 2 reports the primary empirical MASEM results. The core7 model fit well, and the trust6 model also fit well. These results support the model-family route: the study can report estimable structural evidence without overstating the full10 target as an omnibus SEM.

### Table 2

*Primary Model-Family MASEM Fit*

{primary_fit_md}

### Primary Structural Paths

In core7, supported paths included FC -> ATT, SI -> BI, ATT -> BI, FC -> UB, and BI -> UB. PE -> ATT and SI -> ATT had intervals that included zero, and EE -> ATT, PE -> BI, and EE -> BI had incomplete likelihood-based intervals. Thus, the model supports the attitude-intention-behavior chain while leaving several belief-specific paths indeterminate.

In trust6, supported paths included EE -> BI, TRU -> BI, and BI -> UB. EE -> TRU, SI -> TRU, and SI -> BI had intervals that included zero, whereas PE -> TRU and PE -> BI were directionally positive but interval-incomplete. These results support trust as an AI reliance mechanism and suggest that effort expectancy may have a more stable direct link to intention than performance expectancy in this trust-focused complete-case family.

### Supplemental Model Comparison

Table 3 reports selected supplemental models. Removing FC weakened core7 fit, which supports retaining facilitating conditions in the acceptance backbone. A direct-belief model without ATT fit well but used a much larger complete-case set, so it should be treated as evidence that direct belief-intention pathways are also present, not as a replacement for the attitude-mediation model. A pure attitude-mediation model that removed direct PE/EE/SI -> BI paths fit worse, indicating that direct belief-intention paths should remain in the primary core7 specification.

For trust6, removing trust produced an acceptable classic direct-acceptance model, but the trust6 model remains theoretically important because TRU -> BI was supported in the primary model. The trust-only mediation model fit poorly, indicating that trust should not be framed as the sole mediator of PE, EE, and SI effects. Instead, trust functions as a direct AI reliance mechanism alongside standard acceptance beliefs.

### Table 3

*Supplemental Reduced and Alternative Model-Family Diagnostics*

{supp_md}

### PE Versus EE Role Comparison

Table 4 reports the primary PE and EE paths. The comparison supports a conceptual distinction rather than a simple winner-takes-all ranking. PE captures instrumental value, whereas EE captures operational friction. In the trust6 model, EE -> BI was supported, whereas PE -> BI and PE -> TRU were positive but interval-incomplete. In the core7 model, both PE and EE paths to ATT/BI were positive but not confirmatory under the current CI rule. Reduced model fits were similar when either PE or EE was removed, so the manuscript should not claim that one belief universally dominates the other. The defensible claim is that PE and EE represent different adoption mechanisms whose support depends on the target mechanism and complete-case model family.

### Table 4

*Performance Expectancy and Effort Expectancy as Distinct Mechanisms*

{pe_ee_md}

### Anxiety and Self-Efficacy Feasibility

Table 5 reports targeted ANX/SE models. Self-efficacy has more near-term empirical potential than anxiety in the current dataset. Several SE-inclusive subsets have adequate complete-case counts in the feasibility scan, and the SE -> EE -> BI/UB targeted model was minimally estimable with two positive-definite complete-case matrices. However, this model had status 6 and should be reported only as exploratory. Anxiety-inclusive models involving ATT, TRU, BI, and UB had zero positive-definite complete-case matrices and should not be forced into primary MASEM.

### Table 5

*Targeted Anxiety and Self-Efficacy Feasibility Attempts*

{anx_md}

## Discussion

The revised Paper A argument is stronger if it avoids presenting reduced models as compromises. The full10 target is the theoretical reconstruction of AI adoption; core7 and trust6 are empirically estimable descendants of that target. This distinction is the central contribution. It allows the manuscript to retain a broad theory while respecting the matrix requirements of TSSEM/MASEM.

The results support attitude and trust as different mechanisms. Attitude captures evaluative mediation within the traditional acceptance backbone. Trust captures AI-specific reliance under opacity and uncertainty. Supplemental diagnostics refine this claim: direct belief-intention pathways should remain in core7, and trust should be treated as a direct reliance predictor rather than the sole mediator of usefulness, effort, and social influence.

The PE-versus-EE comparison is also important. PE and EE are not interchangeable. PE concerns expected value; EE concerns the burden of use. Current evidence does not support a universal dominance claim, but it shows that their roles differ across model families. This distinction should be emphasized in the theoretical framework and Discussion.

Finally, anxiety and self-efficacy are not rejected. They are retained as future mechanisms because the full10 map shows pairwise evidence, but complete-case structural evidence is not yet strong enough for primary MASEM. This becomes a methodological contribution: the paper shows how meta-analysis can identify not only what is supported, but also what cannot yet be structurally tested.

## References

{refs_text(refs)}
"""
    OUT_MD.write_text(text, encoding='utf-8')
    DATA_MD.parent.mkdir(parents=True, exist_ok=True)
    ONEDRIVE_MD.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(OUT_MD, DATA_MD)
    shutil.copy2(OUT_MD, ONEDRIVE_MD)


def set_cell_text(cell, text, bold=False, italic=False):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 1
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(9)


def set_no_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = tblPr.first_child_found_in('w:tblBorders')
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tblPr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = 'w:' + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn('w:val'), 'nil')


def add_page_number(section):
    p = section.header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    fld1 = OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'PAGE'
    fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'end')
    run._r.append(fld1); run._r.append(instr); run._r.append(fld2)


def configure_doc(doc):
    section = doc.sections[0]
    for attr in ['top_margin','bottom_margin','left_margin','right_margin']:
        setattr(section, attr, Inches(1))
    add_page_number(section)
    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 2
    normal.paragraph_format.space_after = Pt(0)
    for name in ['Title','Heading 1','Heading 2','Heading 3']:
        st = doc.styles[name]
        st.font.name = 'Times New Roman'
        st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        st.font.size = Pt(12)
        st.font.bold = True
        st.paragraph_format.line_spacing = 2
        st.paragraph_format.space_after = Pt(0)


def p(doc, text='', style=None, align=None, indent=False, bold=False, italic=False):
    para = doc.add_paragraph(style=style)
    if align is not None:
        para.alignment = align
    para.paragraph_format.line_spacing = 2
    para.paragraph_format.space_after = Pt(0)
    if indent:
        para.paragraph_format.first_line_indent = Inches(.5)
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return para


def add_doc_table(doc, number, title, headers, rows, note=None):
    p(doc, f'Table {number}', align=WD_ALIGN_PARAGRAPH.LEFT, bold=True)
    p(doc, title, align=WD_ALIGN_PARAGRAPH.LEFT, italic=True)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_no_table_borders(table)
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    if note:
        p(doc, f'Note. {note}', indent=False)


def markdown_to_docx():
    md = OUT_MD.read_text(encoding='utf-8')
    doc = Document()
    configure_doc(doc)
    in_table = False
    table_lines = []
    table_counter = 0
    skip_table = False
    pending_table_title = None
    lines = md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if stripped.startswith('|'):
            i += 1
            continue
        if not stripped:
            i += 1
            continue
        if stripped.startswith('# '):
            p(doc, stripped[2:], style='Title', align=WD_ALIGN_PARAGRAPH.CENTER)
        elif stripped.startswith('## '):
            heading = stripped[3:]
            p(doc, heading, style='Heading 1', align=WD_ALIGN_PARAGRAPH.CENTER)
        elif stripped.startswith('### Table'):
            p(doc, stripped[4:], align=WD_ALIGN_PARAGRAPH.LEFT, bold=True)
        elif stripped.startswith('### '):
            p(doc, stripped[4:], style='Heading 2', align=WD_ALIGN_PARAGRAPH.LEFT)
        elif stripped.startswith('*') and stripped.endswith('*') and '|' not in stripped:
            p(doc, stripped.strip('*'), align=WD_ALIGN_PARAGRAPH.LEFT, italic=True)
        elif stripped.startswith('**Keywords:**') or stripped.startswith('**Author Note**'):
            p(doc, stripped.replace('**',''), indent=False)
        else:
            is_ref = False
            p(doc, stripped.replace('**',''), indent=not stripped.startswith('Author affiliations'))
        i += 1

    doc.add_page_break()
    # APA-styled tables from structured data
    add_doc_table(doc, 1, 'Construct Genealogy and Model-Family Role in Paper A', ['Construct','Label','Origin','AI-Adoption Function','Role'], CONSTRUCT_TABLE, 'PE = performance expectancy; EE = effort expectancy; ATT = attitude; TRU = trust; ANX = anxiety; SE = self-efficacy; BI = behavioral intention; UB = use behavior.')
    doc.add_page_break()
    primary_fit = pd.read_csv(PKG / 'tables/paper_a_model_family_fit_with_n_20260615.csv')
    primary_rows = []
    for _, r in primary_fit.iterrows():
        primary_rows.append([r['model_family'], int(r['complete_case_k']), int(r['effective_sample_size']), fmt(r['chisq']), fmt(r['df'],0), fmt_p(r['p']), fmt(r['CFI']), fmt(r['TLI']), fmt(r['RMSEA']), fmt(r['SRMR'])])
    add_doc_table(doc, 2, 'Primary Model-Family MASEM Fit', ['Model','k','N_eff','χ²','df','p','CFI','TLI','RMSEA','SRMR'], primary_rows, 'Full10 is reported as a theoretical evidence map rather than as an estimable complete-case SEM.')
    doc.add_page_break()
    model_comp = pd.read_csv(SUPP / 'paper_a_supplemental_model_comparison_20260615.csv')
    add_doc_table(doc, 3, 'Supplemental Reduced and Alternative Model-Family Diagnostics', ['Model','k','Status','χ²','df','p','CFI','TLI','RMSEA','SRMR','AIC'], table_model_fit(model_comp), 'Reduced models are diagnostic and not definitive nested tests because construct removal can change complete-case k and matrix structure.')
    doc.add_page_break()
    pe_ee = pd.read_csv(SUPP / 'paper_a_pe_vs_ee_role_comparison_20260615.csv')
    pe_rows = []
    for _, r in pe_ee[pe_ee['source'] == 'primary_path'].iterrows():
        pe_rows.append([r['family'], r['predictor'], r['target'], fmt(r['estimate']), ci(r['ci_low'], r['ci_high']), r['inference_class']])
    add_doc_table(doc, 4, 'Performance Expectancy and Effort Expectancy as Distinct Mechanisms', ['Model','Predictor','Target','Estimate','95% CI','Inference'], pe_rows, 'The comparison concerns PE and EE roles across targets, not the PE-EE correlation.')
    doc.add_page_break()
    anx_attempts = pd.read_csv(SUPP / 'paper_a_anx_se_targeted_model_attempts_20260615.csv')
    anx_rows = []
    for _, r in anx_attempts.iterrows():
        anx_rows.append([r['model_id'], r['constructs'], int(r['positive_definite_complete_case_studies']), r['stage2_status'], fmt(r['CFI']), fmt(r['RMSEA'])])
    add_doc_table(doc, 5, 'Targeted Anxiety and Self-Efficacy Feasibility Attempts', ['Model','Constructs','k','Status','CFI','RMSEA'], anx_rows, 'ANX/SE models are feasibility diagnostics and should not replace the primary model-family MASEM.')
    doc.add_page_break()
    p(doc, 'Figures', style='Heading 1', align=WD_ALIGN_PARAGRAPH.CENTER)
    figs = [
        ('Figure 1', 'Full10 theoretical evidence map.', FIG_DIR / 'figure_1_full10_theoretical_evidence_map_heatmap_ci_20260615.png'),
        ('Figure 2', 'Core7 attitude-mediation MASEM path diagram.', FIG_DIR / 'figure_2_core7_att_mediation_masem_path_ci_20260615.png'),
        ('Figure 3', 'Trust6 mechanism MASEM path diagram.', FIG_DIR / 'figure_3_trust6_mechanism_masem_path_ci_20260615.png'),
    ]
    for idx, (num, title, path) in enumerate(figs):
        if idx:
            doc.add_page_break()
        p(doc, num, align=WD_ALIGN_PARAGRAPH.LEFT, bold=True)
        p(doc, title, align=WD_ALIGN_PARAGRAPH.LEFT, italic=True)
        if path.exists():
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.add_run().add_picture(str(path), width=Inches(6.4))
        p(doc, 'Note. Path diagrams classify paths using likelihood-based 95% confidence intervals; incomplete intervals are descriptive rather than confirmatory.', indent=False)
    doc.save(OUT_DOCX)
    shutil.copy2(OUT_DOCX, DATA_DOCX)
    shutil.copy2(OUT_DOCX, ONEDRIVE_DOCX)


def update_project_docs():
    note = f"""\n## 2026-06-15 Paper A revised APA7 manuscript with supplemental diagnostics\n\n- Generated revised manuscript and Word file with stronger model-history/theory narrative, full10 genealogy, PE-vs-EE role comparison, core7/trust6 model-family justification, reduced-model diagnostics, omitted-pair diagnostic framing, and ANX/SE feasibility results.\n- Revised Markdown: `paper_a/manuscript/target_journal/apa7_model_family_full_manuscript_scaffold_20260615/{OUT_MD.name}`.\n- Revised Word: `paper_a/manuscript/target_journal/apa7_model_family_full_manuscript_scaffold_20260615/{OUT_DOCX.name}`.\n- Supplemental diagnostics: `data/04_extraction/05_llm_masem_substitution/results/paper_a_model_family_supplemental_diagnostics_20260615/`.\n- Interpretation boundary: reduced models are diagnostic/sensitivity analyses, not definitive nested chi-square model-selection tests, because construct removal can change complete-case k and matrix structure.\n"""
    text = CURRENT.read_text(encoding='utf-8') if CURRENT.exists() else '# CURRENT\n'
    if '2026-06-15 Paper A revised APA7 manuscript with supplemental diagnostics' not in text:
        CURRENT.write_text(text.rstrip() + '\n' + note, encoding='utf-8')
    if README.exists():
        readme = README.read_text(encoding='utf-8')
        add = f"\n- `{OUT_MD.name}`\n- `{OUT_DOCX.name}`\n- `supplemental_diagnostics/`\n"
        if OUT_DOCX.name not in readme:
            README.write_text(readme.rstrip() + '\n\n## Revised manuscript with supplemental diagnostics\n' + add, encoding='utf-8')


def main():
    build_markdown()
    markdown_to_docx()
    update_project_docs()
    print(f'Wrote revised Markdown: {OUT_MD}')
    print(f'Wrote revised Word: {OUT_DOCX}')
    print(f'Copied revised outputs to data package and OneDrive mirror')

if __name__ == '__main__':
    main()
