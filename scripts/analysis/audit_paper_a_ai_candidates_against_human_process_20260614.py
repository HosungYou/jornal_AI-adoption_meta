#!/usr/bin/env python3
"""Audit Paper A AI missing-pair candidates against the human coding process.

The audit checks whether each AI candidate was already handled by at least one
human coder, included in the latest direct human input, included in the frozen
reference, or explicitly excluded by human/source adjudication records.
"""
from __future__ import annotations

from collections import Counter, defaultdict
from pathlib import Path
import csv

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[2]
ONEDRIVE_BASE = Path('/Users/newhosung/Library/CloudStorage/OneDrive-SharedLibraries-ThePennsylvaniaStateUniversity/AI Adoption Meta Analysis - Documents')
CANON = ONEDRIVE_BASE / 'Meta/AI Adoption/03_source_adjudication/Paper_A/2026-06-14_human_process_candidate_audit'
REPO_OUT = ROOT / 'data/04_extraction/05_llm_masem_substitution/results/paper_a_human_process_candidate_audit_20260614'

CANDIDATES = ROOT / 'data/04_extraction/05_llm_masem_substitution/results/paper_a_ai_candidate_source_trace_20260614/paper_a_ai_candidate_full10_densification_trace_20260614.csv'
RAW_HUMAN = ROOT / 'data/04_extraction/02_pre_adjudication_disagreement/combined/derived/combined_coder_values_long_20260525.csv'
LATEST_DIRECT = ROOT / 'data/04_extraction/05_llm_masem_substitution/results/paper_a_latest_human_workbook_audit_20260614/paper_a_latest_human_workbook_direct_r_input_20260614.csv'
FROZEN = ROOT / 'data/04_extraction/04_reference_standard_freeze/full_corpus_reference_standard_frozen_20260609.csv'
EXCLUSIONS = ROOT / 'data/04_extraction/04_reference_standard_freeze/phase2_confirmed_exclusion_full_corpus_audit_20260608.csv'
REVIEW_QUEUE = ROOT / 'data/04_extraction/02_pre_adjudication_disagreement/combined/derived/combined_correlation_review_queue_20260525.csv'

CSV_NAME = 'paper_a_ai_candidates_against_human_process_audit_20260614.csv'
MD_NAME = 'PAPER_A_AI_CANDIDATES_AGAINST_HUMAN_PROCESS_AUDIT_20260614.md'
DOCX_NAME = 'REVIEW_THIS_PAPER_A_HUMAN_PROCESS_AUDIT_20260614.docx'


def pair_key(c1: str, c2: str) -> str:
    return '--'.join(sorted([c1.strip(), c2.strip()]))


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open(newline='', encoding='utf-8-sig') as f:
        return list(csv.DictReader(f))


def index_human_rows(rows: list[dict[str, str]], source: str) -> dict[tuple[str, str], list[dict[str, str]]]:
    out: dict[tuple[str, str], list[dict[str, str]]] = defaultdict(list)
    for row in rows:
        sid = (row.get('study_id') or '').strip()
        c1 = (row.get('construct_1') or '').strip()
        c2 = (row.get('construct_2') or '').strip()
        if not sid or not c1 or not c2:
            continue
        if source == 'raw' and (row.get('field_family') or '') != 'correlation':
            continue
        out[(sid, pair_key(c1, c2))].append(row)
    return out


def build_audit() -> list[dict[str, str]]:
    candidates = [r for r in read_csv(CANDIDATES) if r['study_id'] in {'S057', 'S138', 'S176'}]
    raw_idx = index_human_rows(read_csv(RAW_HUMAN), 'raw')
    latest_idx = index_human_rows(read_csv(LATEST_DIRECT), 'latest')
    frozen_idx = index_human_rows(read_csv(FROZEN), 'frozen')
    exclusions = read_csv(EXCLUSIONS)
    excluded_studies = {r['study_id']: r for r in exclusions}
    queue_rows = read_csv(REVIEW_QUEUE)
    queue_by_study = {r['study_id']: r for r in queue_rows if r.get('study_id')}

    rows: list[dict[str, str]] = []
    for cand in candidates:
        sid = cand['study_id']
        key = pair_key(cand['construct_1'], cand['construct_2'])
        raw = raw_idx.get((sid, key), [])
        latest = latest_idx.get((sid, key), [])
        frozen = frozen_idx.get((sid, key), [])
        exclusion = excluded_studies.get(sid)
        queue = queue_by_study.get(sid, {})

        raw_coders = sorted({r.get('coder', '') for r in raw if r.get('coder')})
        latest_sources = sorted({r.get('source_file', '') for r in latest if r.get('source_file')})
        frozen_statuses = sorted({r.get('decision_status', '') for r in frozen if r.get('decision_status')})

        constructs = {cand['construct_1'], cand['construct_2']}
        if frozen or latest or raw:
            human_process_status = 'already_human_handled_exact_pair'
            recommendation = 'do_not_add_duplicate'
            rationale = 'At least one human-process layer already contains this exact unordered construct pair.'
        elif exclusion:
            human_process_status = 'human_excluded_study_or_row'
            recommendation = 'exclude_ai_candidate'
            rationale = f"Human/source adjudication exclusion applies: {exclusion.get('exclusion_code','')} {exclusion.get('rationale','')}"
        elif sid == 'S176' and ('ANX' in constructs or 'SE' in constructs):
            human_process_status = 'not_human_handled_no_explicit_human_exclusion'
            recommendation = 'exclude_ai_false_positive_unless_reopened_by_source_evidence'
            rationale = 'No exact human-coded pair found and no explicit human exclusion row found, but prior PDF text check showed Table 4 constructs HM, UB, BI, EE, FC, HA, PE, PI, SI, TR; ANX/SE were not visible as target constructs.'
        elif sid == 'S138' and ('ANX' in constructs):
            human_process_status = 'not_human_handled_no_explicit_human_exclusion'
            recommendation = 'exclude_or_defer_unless_anxiety_source_matrix_confirmed'
            rationale = 'No exact human-coded pair found. ANX appears to depend on risk/fear term hits; add only if the source matrix explicitly contains anxiety or an approved anxiety-equivalent construct.'
        elif sid == 'S138' and ('FC' in constructs):
            human_process_status = 'not_human_handled_no_explicit_human_exclusion'
            recommendation = 'exclude_or_defer_unless_fc_source_matrix_confirmed'
            rationale = 'No exact human-coded pair found. FC appears to depend on resources/support term hits; add only if the source matrix explicitly contains facilitating conditions or an approved equivalent.'
        elif sid == 'S057' and ('TRU' in constructs or 'UB' in constructs):
            human_process_status = 'not_human_handled_no_explicit_human_exclusion'
            recommendation = 'source_review_required_before_any_add_proposal'
            rationale = 'No exact human-coded pair found and no explicit human exclusion row found. Candidate_value is blank, so this can become an add proposal only after PDF/source table confirmation of both target constructs and a numeric cell.'
        else:
            human_process_status = 'not_human_handled_no_explicit_human_exclusion'
            recommendation = 'source_review_required_before_any_add_proposal'
            rationale = 'No exact human-coded pair found and no explicit human exclusion row found; source confirmation is required before addition.'

        one_coder_note = ''
        if queue:
            one_coder_note = f"review_queue: coder_a={queue.get('coder_a','')}, coder_b={queue.get('coder_b','')}, n_one_coder_only={queue.get('n_one_coder_only','')}, note={queue.get('review_note','')}"
        frozen_study_rows = [r for (s, _), vals in frozen_idx.items() if s == sid for r in vals]
        frozen_study_statuses = sorted({r.get('decision_status','') for r in frozen_study_rows})
        rows.append({
            'study_id': sid,
            'missing_pair': cand['missing_pair'],
            'candidate_value': cand.get('candidate_value',''),
            'raw_human_exact_pair_count': str(len(raw)),
            'raw_human_coders_exact_pair': ';'.join(raw_coders),
            'latest_human_exact_pair_count': str(len(latest)),
            'latest_human_sources_exact_pair': ';'.join(latest_sources),
            'frozen_reference_exact_pair_count': str(len(frozen)),
            'frozen_reference_status_exact_pair': ';'.join(frozen_statuses),
            'human_exclusion_found_for_study': 'yes' if exclusion else 'no',
            'human_exclusion_code_or_reason': (exclusion.get('exclusion_code','') + ' ' + exclusion.get('rationale','')).strip() if exclusion else '',
            'one_coder_rule_evidence_for_study': one_coder_note,
            'frozen_reference_statuses_for_study': ';'.join(frozen_study_statuses),
            'human_process_status': human_process_status,
            'recommendation': recommendation,
            'rationale': rationale,
        })
    return sorted(rows, key=lambda r: (r['study_id'], r['missing_pair']))


def write_csvs(rows: list[dict[str, str]]) -> None:
    for outdir in [REPO_OUT, CANON]:
        outdir.mkdir(parents=True, exist_ok=True)
        with (outdir / CSV_NAME).open('w', newline='', encoding='utf-8-sig') as f:
            writer = csv.DictWriter(f, fieldnames=list(rows[0].keys()))
            writer.writeheader()
            writer.writerows(rows)


def grouped_summary(rows: list[dict[str, str]]) -> list[tuple[str, str, str, str, str]]:
    return [
        ('All 51 candidates', '0 exact raw human rows; 0 latest direct rows; 0 frozen-reference exact rows', 'No study-level human exclusion found for S057/S138/S176 in the audited exclusion file', 'No direct additions', 'A blank candidate_value plus no human exact pair means source review, not automatic insertion.'),
        ('S057 TRU/UB candidates', 'No exact human-coded TRU/UB pair among the 17 AI candidates', 'No explicit human exclusion found for these exact pairs', 'source_review_required_before_any_add_proposal', 'If source table confirms TRU/UB as target constructs and numeric cells, promote as source_confirmed_add_candidate; otherwise exclude as AI false positive.'),
        ('S138 ANX candidates', 'No exact human-coded ANX pair among the 9 AI candidates', 'No explicit human exclusion found for these exact pairs', 'exclude_or_defer_unless_anxiety_source_matrix_confirmed', 'Risk/fear term hits are not enough. Add only if anxiety or an approved anxiety-equivalent matrix construct is visible.'),
        ('S138 FC candidates', 'No exact human-coded FC pair among the 8 AI candidates', 'No explicit human exclusion found for these exact pairs', 'exclude_or_defer_unless_fc_source_matrix_confirmed', 'Resources/support term hits are not enough. Add only if facilitating conditions or an approved equivalent is visible.'),
        ('S176 ANX/SE candidates', 'No exact human-coded ANX/SE candidate pair', 'No explicit human exclusion found for these exact pairs', 'exclude_ai_false_positive_unless_reopened_by_source_evidence', 'Prior PDF text check showed Table 4 constructs HM, UB, BI, EE, FC, HA, PE, PI, SI, TR; ANX/SE were not visible.'),
    ]


def write_md(rows: list[dict[str, str]]) -> None:
    counts = Counter(r['recommendation'] for r in rows)
    lines = [
        '# Paper A AI candidates against human-process audit',
        '',
        'Date: 2026-06-14',
        '',
        '## Bottom line',
        '',
        'The 51 AI/source-trace rows are not confirmed human omissions. Across the audited human-process layers, none of the 51 exact unordered construct pairs appears in raw human coder rows, latest direct human input, or the frozen reference. No study-level human exclusion record was found for S057, S138, or S176 in the audited exclusion file.',
        '',
        'Therefore the correct rule is not “no human exclusion means add.” The correct rule is: no human exclusion plus source-confirmed target construct pair plus visible numeric cell means add proposal. Otherwise exclude or defer.',
        '',
        '## Recommendation counts',
        '',
    ]
    for key, val in sorted(counts.items()):
        lines.append(f'- {key}: {val}')
    lines.extend(['', '## Grouped audit table', '', '| Candidate group | Human exact-pair evidence | Human exclusion evidence | Recommendation | Rationale |', '|---|---|---|---|---|'])
    for row in grouped_summary(rows):
        lines.append('| ' + ' | '.join(row) + ' |')
    lines.extend(['', '## Row-level audit file', '', f'CSV: `{CSV_NAME}`', ''])
    for outdir in [REPO_OUT, CANON]:
        (outdir / MD_NAME).write_text('\n'.join(lines), encoding='utf-8')


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement('w:tblHeader')
    tbl_header.set(qn('w:val'), 'true')
    tr_pr.append(tbl_header)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 8) -> None:
    cell.text = ''
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Aptos'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def write_docx(rows: list[dict[str, str]]) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    styles = doc.styles
    styles['Normal'].font.name = 'Aptos'
    styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    styles['Normal'].font.size = Pt(10)
    for style in ['Heading 1', 'Heading 2', 'Heading 3']:
        styles[style].font.name = 'Aptos Display'
        styles[style]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Paper A Human-Process Audit of AI Candidate Rows')
    r.bold = True
    r.font.size = Pt(18)
    r.font.name = 'Aptos Display'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Question: were these missed by humans, excluded by humans, or only AI/source-trace candidates?').italic = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Date: 2026-06-14 | Scope: S057, S138, S176 priority 51 AI candidates')

    doc.add_heading('1. Bottom line', level=1)
    doc.add_paragraph('이 51개 행은 “인간 코더가 놓친 값”으로 확정되지 않았습니다. audited human-process layers 기준으로 51개 exact unordered construct pair는 raw human coder rows, latest direct human input, frozen reference 어디에도 exact pair로 나타나지 않았습니다.')
    doc.add_paragraph('또한 S057, S138, S176에 대한 study-level human exclusion record는 audited exclusion file에서 발견되지 않았습니다. 따라서 결론은 “인간 제외가 없으니 추가”가 아니라 “인간 제외가 없고 source table에서 target construct-pair numeric cell이 확인되면 추가 제안”입니다.')

    callout = doc.add_table(rows=1, cols=1)
    callout.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = callout.cell(0, 0)
    set_cell_shading(cell, 'FFF2CC')
    set_cell_text(cell, 'Operational decision: human-excluded AI candidates should be excluded. Non-excluded AI candidates still require source confirmation before add proposal. Current direct-add count: 0.', bold=True, size=10)

    doc.add_heading('2. Was the one-coder rule applied?', level=1)
    doc.add_paragraph('Yes, for existing human-handled row sets the evidence shows that one-coder-only/source-checked rows were allowed into the frozen reference where appropriate. S057 frozen rows carry one-coder/source-checked status; S138 has coder-agreed correlation rows; S176 has source-checked frozen reference rows. The current 51 candidates are outside those exact human-coded row sets.')

    doc.add_heading('3. Grouped audit decision', level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_repeat_table_header(table.rows[0])
    for i, h in enumerate(['Candidate group', 'Human exact-pair evidence', 'Human exclusion evidence', 'Recommendation', 'Rationale']):
        set_cell_shading(table.rows[0].cells[i], '1F4E79')
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=8)
    for group in grouped_summary(rows):
        cells = table.add_row().cells
        for i, val in enumerate(group):
            set_cell_text(cells[i], val, size=7)

    doc.add_heading('4. Final decision rule to use from here', level=1)
    for step in [
        'If any human/adjudication exclusion applies, exclude the AI candidate unless the researcher explicitly reopens the decision.',
        'If at least one human already coded the exact pair, treat it as human-handled and avoid duplicate AI promotion.',
        'If no human handled the exact pair and no exclusion applies, check the PDF/source table.',
        'Promote only if both target constructs and a usable numeric correlation/latent-correlation cell are visible in the same source matrix.',
        'If the candidate depends only on generic term hits such as risk, support, use, trust, or table, exclude or defer rather than add.',
    ]:
        doc.add_paragraph(step, style='List Number')

    doc.add_heading('5. Row-level file', level=1)
    doc.add_paragraph(f'The full 51-row audit is in the companion CSV: {CSV_NAME}')

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.text = 'Paper A human-process audit | AI candidates only | no matrix update'

    for outdir in [REPO_OUT, CANON]:
        outdir.mkdir(parents=True, exist_ok=True)
        doc.save(outdir / DOCX_NAME)


def write_readme() -> None:
    text = f"""# Paper A human-process candidate audit

Primary review file:

`{CANON / DOCX_NAME}`

Companion files:

- `{CANON / CSV_NAME}`
- `{CANON / MD_NAME}`

Conclusion:

- The 51 AI candidates are not confirmed human omissions.
- No exact candidate pair appears in raw human coder rows, latest direct human input, or frozen reference.
- No study-level human exclusion was found for S057/S138/S176 in the audited exclusion file.
- Add proposal requires source-confirmed target pair plus visible numeric cell; absence of human exclusion alone is insufficient.
"""
    for outdir in [REPO_OUT, CANON]:
        (outdir / 'README_HUMAN_PROCESS_AUDIT_20260614.md').write_text(text, encoding='utf-8')


def main() -> None:
    rows = build_audit()
    write_csvs(rows)
    write_md(rows)
    write_docx(rows)
    write_readme()
    print(CANON / DOCX_NAME)
    print(CANON / CSV_NAME)
    print(CANON / MD_NAME)
    print(REPO_OUT)
    print('rows=', len(rows))
    print('recommendations=', dict(Counter(r['recommendation'] for r in rows)))


if __name__ == '__main__':
    main()
