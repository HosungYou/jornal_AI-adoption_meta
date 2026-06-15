#!/usr/bin/env python3
from __future__ import annotations

import csv
import re
import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path('/Users/newhosung/Academic/2026/AI Adoption Meta Analysis')
DATE = '20260615'
PKG = ROOT / 'paper_a/manuscript/target_journal/apa7_model_family_full_manuscript_scaffold_20260615'
DATA_PKG = ROOT / 'data/04_extraction/05_llm_masem_substitution/results/paper_a_apa7_model_family_manuscript_package_20260615'
ONEDRIVE = Path('/Users/newhosung/Library/CloudStorage/OneDrive-SharedLibraries-ThePennsylvaniaStateUniversity/AI Adoption Meta Analysis - Documents/05_manuscripts/Paper_A/2026-06-15_apa7_model_family_full_manuscript_scaffold')
SUPP = ROOT / 'data/04_extraction/05_llm_masem_substitution/results/paper_a_model_family_supplemental_diagnostics_20260615'
REFS = PKG / f'PAPER_A_EXPANDED_REFERENCE_BANK_{DATE}.csv'
SOURCE_MD = PKG / f'PAPER_A_APA7_REVISED_MODEL_FAMILY_MANUSCRIPT_WITH_SUPPLEMENTAL_DIAGNOSTICS_{DATE}.md'
OUT_DOCX = PKG / f'PAPER_A_APA7_INLINE_TABLES_FIGURES_BLACK_FONT_{DATE}.docx'
DATA_DOCX = DATA_PKG / OUT_DOCX.name
ONEDRIVE_DOCX = ONEDRIVE / OUT_DOCX.name
CURRENT = ROOT / 'CURRENT.md'
README = PKG / f'README_PAPER_A_APA7_MODEL_FAMILY_MANUSCRIPT_PACKAGE_{DATE}.md'
FIG_DIR = PKG / 'figures'

BLACK = RGBColor(0, 0, 0)

CONSTRUCT_TABLE = [
    ['PE', 'Performance expectancy / perceived usefulness', 'TAM, TAM2, UTAUT', 'Instrumental outcome belief: AI improves learning, teaching, productivity, or performance.', 'full10; core7; trust6'],
    ['EE', 'Effort expectancy / perceived ease of use', 'TAM, computer self-efficacy, UTAUT', 'Operational-friction belief: AI is manageable, learnable, and low burden.', 'full10; core7; trust6'],
    ['SI', 'Social influence', 'UTAUT', 'Normative/institutional pressure and endorsement mechanism.', 'full10; core7; trust6'],
    ['FC', 'Facilitating conditions', 'UTAUT', 'Resource and infrastructure mechanism enabling evaluation and enacted use.', 'full10; core7'],
    ['ATT', 'Attitude', 'TRA/TPB, TAM', 'Evaluative mediator translating beliefs into intention.', 'full10; core7'],
    ['TRU', 'Trust', 'Trust in automation, trust in IS, AI reliance', 'AI-specific reliance mechanism under opacity, autonomy, uncertainty, and vulnerability.', 'full10; trust6'],
    ['ANX', 'Anxiety', 'Technology readiness, affective threat', 'Threat/unease mechanism retained but underidentified for primary complete-case MASEM.', 'full10; future mechanism'],
    ['SE', 'Self-efficacy', 'Social cognitive theory, computer self-efficacy', 'Capability mechanism; feasible mainly in smaller supplemental sets.', 'full10; future mechanism'],
    ['BI', 'Behavioral intention', 'TRA/TPB, TAM, UTAUT', 'Proximal motivational outcome.', 'full10; core7; trust6'],
    ['UB', 'Use behavior', 'TAM, UTAUT', 'Behavioral adoption outcome.', 'full10; core7; trust6'],
]

KNOWN_ITALIC_SUBSTRINGS = [
    'Frontiers in Psychology', 'Information Systems Frontiers', 'Computers & Education', 'Computers in Human Behavior',
    'Human Factors', 'MIS Quarterly', 'Management Science', 'Information Systems Research', 'SAGE Open',
    'Online Learning', 'International Journal of STEM Education', 'Psychological Methods', 'Behavior Research Methods',
    'Research Synthesis Methods', 'Structural Equation Modeling', 'Prevention Science', 'Journal of the Academy of Marketing Science',
    'Academy of Management Annals', 'Organizational Behavior and Human Decision Processes', 'Telematics and Informatics',
    'Interactive Learning Environments', 'Journal of Organizational and End User Computing', 'Routledge', 'Wiley', 'Free Press',
    'Diffusion of innovations', 'Belief, attitude, intention, and behavior', 'Self-efficacy: The exercise of control',
]


def fmt(x, digits=3):
    try:
        if pd.isna(x):
            return 'NA'
        return f'{float(x):.{digits}f}'.replace('-0.000', '0.000')
    except Exception:
        return str(x)


def fmt_p(x):
    try:
        if pd.isna(x):
            return 'NA'
        v = float(x)
        if v < .001:
            return '< .001'
        return f'{v:.3f}'.lstrip('0')
    except Exception:
        return str(x)


def ci(lo, hi):
    return f'[{fmt(lo)}, {fmt(hi)}]'


def set_run(run, bold=False, italic=False, size=12, font='Times New Roman'):
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = BLACK


def configure(doc):
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    header_p = sec.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header_p.add_run()
    fld1 = OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'PAGE'
    fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'end')
    r._r.append(fld1); r._r.append(instr); r._r.append(fld2)
    for style_name in ['Normal','Title','Heading 1','Heading 2','Heading 3','List Bullet']:
        style = doc.styles[style_name]
        style.font.name = 'Times New Roman'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        style.font.size = Pt(12)
        style.font.color.rgb = BLACK
        style.paragraph_format.line_spacing = 2
        style.paragraph_format.space_after = Pt(0)
    for style_name in ['Title','Heading 1','Heading 2','Heading 3']:
        doc.styles[style_name].font.bold = True


def add_p(doc, text='', style=None, align=None, indent=False, bold=False, italic=False):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.line_spacing = 2
    p.paragraph_format.space_after = Pt(0)
    if indent:
        p.paragraph_format.first_line_indent = Inches(.5)
    r = p.add_run(text)
    set_run(r, bold=bold, italic=italic)
    return p


def no_vertical_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = tblPr.first_child_found_in('w:tblBorders')
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tblPr.append(borders)
    settings = {
        'top': ('single', '000000', '8'),
        'bottom': ('single', '000000', '8'),
        'insideH': ('single', '000000', '4'),
        'left': ('nil', '000000', '0'),
        'right': ('nil', '000000', '0'),
        'insideV': ('nil', '000000', '0'),
    }
    for edge, (val, color, size) in settings.items():
        element = borders.find(qn('w:' + edge))
        if element is None:
            element = OxmlElement('w:' + edge)
            borders.append(element)
        element.set(qn('w:val'), val)
        element.set(qn('w:color'), color)
        element.set(qn('w:sz'), size)


def add_cell(cell, text, bold=False):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 1
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    set_run(r, bold=bold, size=9)


def add_apa_table(doc, number, title, headers, rows, note=None):
    add_p(doc, f'Table {number}', bold=True)
    add_p(doc, title, italic=True)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_vertical_borders(table)
    for i, h in enumerate(headers):
        add_cell(table.rows[0].cells[i], h, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            add_cell(cells[i], value)
    if note:
        add_p(doc, 'Note. ' + note)


def add_apa_figure(doc, number, title, path, note=None):
    add_p(doc, f'Figure {number}', bold=True)
    add_p(doc, title, italic=True)
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(6.3))
    if note:
        add_p(doc, 'Note. ' + note)


def build_tables():
    primary_fit = pd.read_csv(PKG / 'tables/paper_a_model_family_fit_with_n_20260615.csv')
    primary_rows = []
    for _, r in primary_fit.iterrows():
        primary_rows.append([r['model_family'], int(r['complete_case_k']), int(r['effective_sample_size']), fmt(r['chisq']), fmt(r['df'],0), fmt_p(r['p']), fmt(r['CFI']), fmt(r['TLI']), fmt(r['RMSEA']), fmt(r['SRMR'])])
    model_comp = pd.read_csv(SUPP / 'paper_a_supplemental_model_comparison_20260615.csv')
    keep = ['core7_full','core6_no_ATT_direct_beliefs','core7_pure_ATT_mediation_no_direct_belief_BI','trust6_full','trust5_no_TRU_direct_acceptance','trust6_trust_mediator_no_direct_belief_BI','se4_capability_effort_intention']
    supp_rows = []
    for _, r in model_comp[model_comp['model_id'].isin(keep)].iterrows():
        supp_rows.append([r['model_id'], int(r['positive_definite_complete_case_studies']), r['stage2_status'], fmt(r['chisq']), fmt(r['df'],0), fmt_p(r['p']), fmt(r['CFI']), fmt(r['TLI']), fmt(r['RMSEA']), fmt(r['SRMR']), fmt(r['AIC'])])
    pe_ee = pd.read_csv(SUPP / 'paper_a_pe_vs_ee_role_comparison_20260615.csv')
    pe_rows = []
    for _, r in pe_ee[pe_ee['source'] == 'primary_path'].iterrows():
        pe_rows.append([r['family'], r['predictor'], r['target'], fmt(r['estimate']), ci(r['ci_low'], r['ci_high']), r['inference_class']])
    anx = pd.read_csv(SUPP / 'paper_a_anx_se_targeted_model_attempts_20260615.csv')
    anx_rows = []
    for _, r in anx.iterrows():
        anx_rows.append([r['model_id'], r['constructs'], int(r['positive_definite_complete_case_studies']), r['stage2_status'], fmt(r['CFI']), fmt(r['RMSEA'])])
    return {
        '1': (1, 'Construct Genealogy and Model-Family Role in Paper A', ['Construct','Label','Origin','AI-Adoption Function','Role'], CONSTRUCT_TABLE, 'PE = performance expectancy; EE = effort expectancy; ATT = attitude; TRU = trust; ANX = anxiety; SE = self-efficacy; BI = behavioral intention; UB = use behavior.'),
        '2': (2, 'Primary Model-Family MASEM Fit', ['Model','k','N_eff','χ²','df','p','CFI','TLI','RMSEA','SRMR'], primary_rows, 'Full10 is reported as a theoretical evidence map rather than an estimable complete-case SEM.'),
        '3': (3, 'Supplemental Reduced and Alternative Model-Family Diagnostics', ['Model','k','Status','χ²','df','p','CFI','TLI','RMSEA','SRMR','AIC'], supp_rows, 'Reduced models are diagnostic and not definitive nested tests because construct removal can change complete-case k and matrix structure.'),
        '4': (4, 'Performance Expectancy and Effort Expectancy as Distinct Mechanisms', ['Model','Predictor','Target','Estimate','95% CI','Inference'], pe_rows, 'The comparison concerns PE and EE roles across targets, not the PE-EE correlation.'),
        '5': (5, 'Targeted Anxiety and Self-Efficacy Feasibility Attempts', ['Model','Constructs','k','Status','CFI','RMSEA'], anx_rows, 'ANX/SE models are feasibility diagnostics and should not replace the primary model-family MASEM.'),
    }


def add_reference(doc, text):
    url = ''
    if ' http' in text:
        text, url = text.rsplit(' http', 1)
        url = 'http' + url
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    p.paragraph_format.left_indent = Inches(0.5)
    remaining = text
    while remaining:
        matches = [(remaining.find(s), s) for s in KNOWN_ITALIC_SUBSTRINGS if remaining.find(s) >= 0]
        if not matches:
            r = p.add_run(remaining)
            set_run(r)
            break
        idx, substr = min(matches, key=lambda x: x[0])
        if idx > 0:
            r = p.add_run(remaining[:idx])
            set_run(r)
        r = p.add_run(substr)
        set_run(r, italic=True)
        remaining = remaining[idx + len(substr):]
    if url:
        r = p.add_run(' ' + url)
        set_run(r)


def add_references(doc):
    refs = []
    with REFS.open(newline='', encoding='utf-8') as f:
        for row in csv.DictReader(f):
            citation = row.get('citation','').strip()
            url = row.get('doi_or_url','').strip()
            if citation:
                refs.append((citation + (' ' + url if url else '')).strip())
    for ref in sorted(set(refs)):
        add_reference(doc, ref)


def build():
    doc = Document()
    configure(doc)
    tables = build_tables()
    inserted_fig1 = inserted_fig23 = False
    lines = SOURCE_MD.read_text(encoding='utf-8').splitlines()
    in_refs = False
    skip_table_block = False
    i = 0
    while i < len(lines):
        s = lines[i].strip()
        if not s:
            i += 1
            continue
        if in_refs:
            i += 1
            continue
        m = re.match(r'^### Table (\d+)$', s)
        if m:
            key = m.group(1)
            if key in tables:
                add_apa_table(doc, *tables[key])
            i += 1
            while i < len(lines):
                nxt = lines[i].strip()
                if nxt.startswith('## ') or nxt.startswith('### '):
                    break
                i += 1
            continue
        if s.startswith('|') or s.startswith('*') and s.endswith('*') and 'Table' in s:
            i += 1
            continue
        if s.startswith('# '):
            add_p(doc, s[2:], style='Title', align=WD_ALIGN_PARAGRAPH.CENTER)
        elif s.startswith('## References'):
            add_p(doc, 'References', style='Heading 1', align=WD_ALIGN_PARAGRAPH.CENTER)
            add_references(doc)
            in_refs = True
        elif s.startswith('## '):
            add_p(doc, s[3:], style='Heading 1', align=WD_ALIGN_PARAGRAPH.CENTER)
        elif s.startswith('### '):
            heading = s[4:]
            add_p(doc, heading, style='Heading 2')
            if heading == 'Full10 Evidence Map' and not inserted_fig1:
                add_apa_figure(doc, 1, 'Full10 theoretical evidence map.', FIG_DIR / 'figure_1_full10_theoretical_evidence_map_heatmap_ci_20260615.png', 'Cells show pairwise random-effects pooled correlations and k. This is an evidence map, not a full10 SEM estimate.')
                inserted_fig1 = True
            if heading == 'Primary Structural Paths' and not inserted_fig23:
                add_apa_figure(doc, 2, 'Core7 attitude-mediation MASEM path diagram.', FIG_DIR / 'figure_2_core7_att_mediation_masem_path_ci_20260615.png', 'Solid paths indicate likelihood-based 95% CIs excluding zero; dashed paths include zero; dotted paths have incomplete intervals.')
                add_apa_figure(doc, 3, 'Trust6 mechanism MASEM path diagram.', FIG_DIR / 'figure_3_trust6_mechanism_masem_path_ci_20260615.png', 'Solid paths indicate likelihood-based 95% CIs excluding zero; dashed paths include zero; dotted paths have incomplete intervals.')
                inserted_fig23 = True
        elif s.startswith('**Keywords:**'):
            add_p(doc, s.replace('**',''))
        elif s.startswith('**Author Note**'):
            add_p(doc, 'Author Note', bold=True)
        elif s.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.line_spacing = 2
            r = p.add_run(s[2:])
            set_run(r)
        elif s.startswith('*') and s.endswith('*'):
            add_p(doc, s.strip('*'), italic=True)
        else:
            add_p(doc, s.replace('**',''), indent=True)
        i += 1
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    shutil.copy2(OUT_DOCX, DATA_DOCX)
    shutil.copy2(OUT_DOCX, ONEDRIVE_DOCX)
    note = f"""\n## 2026-06-15 Paper A inline APA7 Word manuscript\n\n- Generated inline APA 7 Word manuscript with tables and figures inserted in the body near first mention: `paper_a/manuscript/target_journal/apa7_model_family_full_manuscript_scaffold_20260615/{OUT_DOCX.name}`.\n- Formatting rule applied: black manuscript font, APA-style table/figure numbering and titles, body-embedded tables/figures, and reference italics where reference-bank metadata supports run-level styling.\n"""
    text = CURRENT.read_text(encoding='utf-8') if CURRENT.exists() else '# CURRENT\n'
    if '2026-06-15 Paper A inline APA7 Word manuscript' not in text:
        CURRENT.write_text(text.rstrip() + '\n' + note, encoding='utf-8')
    if README.exists():
        readme = README.read_text(encoding='utf-8')
        if OUT_DOCX.name not in readme:
            README.write_text(readme.rstrip() + f'\n- `{OUT_DOCX.name}`\n', encoding='utf-8')
    print(f'Wrote inline APA7 Word manuscript: {OUT_DOCX}')
    print(f'Copied to data package: {DATA_DOCX}')
    print(f'Copied to OneDrive: {ONEDRIVE_DOCX}')

if __name__ == '__main__':
    build()
