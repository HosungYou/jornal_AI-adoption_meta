#!/usr/bin/env python3
"""Build and structure the Paper B near-submission collaboration package."""

from __future__ import annotations

import csv
import hashlib
import math
import os
import shutil
import textwrap
import time
import zipfile
from collections import Counter
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.datavalidation import DataValidation


DATE = "20260619"
ROOT = Path(
    "/Users/newhosung/Library/CloudStorage/OneDrive-SharedLibraries-ThePennsylvaniaStateUniversity/"
    "AI Adoption Meta Analysis - Documents"
)
PAPER_B = ROOT / "02_Paper_B"
REPO = ROOT / "90_repository_mirror" / "journal_AI-adoption_meta"
RESULTS = REPO / "data" / "04_extraction" / "05_llm_masem_substitution" / "results"
LOCKED = REPO / "data" / "04_extraction" / "05_llm_masem_substitution" / "locked_outputs"
OUT = Path("/Users/newhosung/Documents/Codex/2026-06-18/google-drive-onedrive-library-cloudstorage-googledrive/outputs")


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def move_if_exists(src: Path, dst: Path) -> None:
    if not src.exists() or src.resolve() == dst.resolve():
        return
    ensure_dir(dst.parent)
    if dst.exists():
        merge_tree(src, dst)
        return
    shutil.move(str(src), str(dst))


def merge_tree(src: Path, dst: Path) -> None:
    """Merge an already-created legacy folder into its intended active folder."""

    ensure_dir(dst)
    for item in list(src.iterdir()):
        target = dst / item.name
        if not target.exists():
            shutil.move(str(item), str(target))
            continue

        if item.is_dir() and target.is_dir():
            merge_tree(item, target)
            continue

        if item.is_file() and target.is_file() and sha256(item) == sha256(target):
            item.unlink()
            continue

        suffix = 1
        while True:
            if item.is_file():
                candidate = target.with_name(f"{target.stem}__legacy_{suffix}{target.suffix}")
            else:
                candidate = target.with_name(f"{target.name}__legacy_{suffix}")
            if not candidate.exists():
                shutil.move(str(item), str(candidate))
                break
            suffix += 1

    try:
        src.rmdir()
    except OSError:
        pass


def read_csv(path: Path) -> list[dict[str, str]]:
    if not path.exists():
        return []
    with path.open(newline="", encoding="utf-8-sig") as handle:
        return list(csv.DictReader(handle))


def write_text(path: Path, text: str) -> None:
    ensure_dir(path.parent)
    path.write_text(text.strip() + "\n", encoding="utf-8")


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def replace_file_with_retry(src: Path, dst: Path, attempts: int = 5) -> None:
    """Copy a completed local artifact into cloud storage with short retries."""

    ensure_dir(dst.parent)
    last_error: BaseException | None = None
    for attempt in range(attempts):
        staging = dst.with_name(f".{dst.name}.staging")
        try:
            shutil.copy2(src, staging)
            staging.replace(dst)
            return
        except (OSError, TimeoutError) as exc:
            last_error = exc
            try:
                if staging.exists():
                    staging.unlink()
            except OSError:
                pass
            time.sleep(2 * (attempt + 1))
    if last_error:
        raise last_error


def locate_preferred(paths: list[Path]) -> Path:
    for path in paths:
        if path.exists():
            return path
    return paths[0]


def restructure_paper_b() -> dict[str, Path]:
    """Move the legacy Paper B directories into the active 20260619 structure."""

    paths = {
        "start": PAPER_B / "00_START_HERE",
        "working": PAPER_B / "01_working_manuscript",
        "tracking": PAPER_B / "02_B_requirements_and_tracking",
        "reference": PAPER_B / "03_B_reference_standard",
        "analysis": PAPER_B / "04_B_analysis_inputs_outputs",
        "manuscript": PAPER_B / "05_B_manuscript_and_repository",
        "source_adj": PAPER_B / "06_B_source_adjudication",
        "archive": PAPER_B / "99_archive",
    }
    for p in paths.values():
        ensure_dir(p)

    move_if_exists(PAPER_B / "01_reference_standard_candidates", paths["reference"] / "reference_standard_candidates")
    move_if_exists(PAPER_B / "02_analysis_inputs_outputs", paths["analysis"])
    move_if_exists(
        PAPER_B / "03_B_requirements_and_status",
        paths["tracking"] / "legacy_status_docs_20260618",
    )
    move_if_exists(PAPER_B / "04_B_manuscript_and_repository", paths["manuscript"])
    move_if_exists(PAPER_B / "05_B_source_adjudication", paths["source_adj"])

    return paths


def archive_old_manuscript_candidates(paths: dict[str, Path]) -> None:
    manuscript_dir = paths["manuscript"] / "repository_paper_b" / "manuscript"
    archive_dir = paths["archive"] / "legacy_paper_b_manuscript_candidates_20260619"
    ensure_dir(archive_dir)
    if not manuscript_dir.exists():
        return

    keep_names = {"target_journal", "README.md", "figures"}
    for item in list(manuscript_dir.iterdir()):
        if item.name in keep_names:
            continue
        if item.name.startswith("render_paper_b_rsm_target_"):
            continue
        if item.suffix.lower() in {".docx", ".md"} or item.name.startswith("render_"):
            dst = archive_dir / item.name
            if not dst.exists():
                shutil.move(str(item), str(dst))


def full_m1r_summary() -> pd.DataFrame:
    scored = pd.read_csv(RESULTS / "paper_b_full_corpus_m1_raw_full_scored_20260612.csv")
    rows = []
    labels = {
        "primary_latent_or_construct_correlation_with_source_type_flag": "Latent/source-flagged r",
        "primary_direct_r_or_source_reported_correlation": "Direct/source-reported r",
        "secondary_beta_or_path_converted_effect_size": "Converted beta/path",
    }
    for family, group in scored.groupby("denominator_family", sort=False):
        total = len(group)
        scored_rows = int((group["score_status"] == "scored").sum())
        abstentions = int((group["score_status"] == "scored_abstention").sum())
        correct = int(group.loc[group["score_status"] == "scored", "is_correct"].sum())
        incorrect = scored_rows - correct
        rows.append(
            {
                "family_key": family,
                "family": labels.get(family, family),
                "rows_total": total,
                "scored_rows": scored_rows,
                "correct_rows": correct,
                "incorrect_scored_rows": incorrect,
                "abstention_rows": abstentions,
                "accuracy_scored_only": correct / scored_rows if scored_rows else math.nan,
                "accuracy_all_scorable": correct / total if total else math.nan,
            }
        )
    return pd.DataFrame(rows)


def rq3_summary() -> tuple[pd.DataFrame, pd.DataFrame]:
    signal = pd.read_csv(RESULTS / "initial_execution_validation_20260612" / "paper_b_rq3_signal_validation_20260612.csv")
    priority = pd.read_csv(RESULTS / "initial_execution_validation_20260612" / "paper_b_rq3_review_priority_summary_20260612.csv")
    dedup = signal.drop_duplicates(subset=["signal"]).copy()
    return dedup, priority


def masem_gate_summary() -> pd.DataFrame:
    return pd.read_csv(RESULTS / "paper2_masem_matrix_construct_set_completeness_20260612.csv")


def sparse_probe_summary() -> pd.DataFrame:
    path = RESULTS / "paper_b_sparse_tssem_probe_20260612" / "paper_b_sparse_tssem_probe_summary_20260612.csv"
    if path.exists():
        return pd.read_csv(path)
    return pd.DataFrame()


def create_figures(paths: dict[str, Path]) -> dict[str, Path]:
    fig_dir = paths["working"] / "figures"
    ensure_dir(fig_dir)

    outputs: dict[str, Path] = {}

    from PIL import Image, ImageDraw, ImageFont

    def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
        candidates = [
            "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
            "/Library/Fonts/Arial Bold.ttf" if bold else "/Library/Fonts/Arial.ttf",
            "/System/Library/Fonts/Helvetica.ttc",
        ]
        for candidate in candidates:
            try:
                return ImageFont.truetype(candidate, size=size)
            except Exception:
                pass
        return ImageFont.load_default()

    def save_png(img: Image.Image, path: Path) -> None:
        img.save(path, "PNG", dpi=(300, 300))

    # Figure 1: workflow.
    w, h = 2400, 900
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    title_font = font(48, bold=True)
    label_font = font(34, bold=True)
    note_font = font(28)
    draw.text((w // 2, 70), "Source-anchored Paper B validation workflow", anchor="mm", fill="#1A202C", font=title_font)
    steps = [
        ("1", "Raw human coder\nfreeze"),
        ("2", "Human-human\ndisagreement"),
        ("3", "Source-document\nadjudication"),
        ("4", "Source-anchored\nreference freeze"),
        ("5", "Locked LLM outputs\nand diagnostics"),
    ]
    box_w, box_h = 380, 260
    gap = 48
    start_x = 120
    y0 = 285
    for idx, (num, label) in enumerate(steps):
        x0 = start_x + idx * (box_w + gap)
        x1 = x0 + box_w
        y1 = y0 + box_h
        draw.rounded_rectangle((x0, y0, x1, y1), radius=28, fill="#EDF2F7", outline="#2D3748", width=5)
        draw.text((x0 + box_w / 2, y0 + 70), num, anchor="mm", fill="#2B6CB0", font=font(44, bold=True))
        for line_i, line in enumerate(label.split("\n")):
            draw.text((x0 + box_w / 2, y0 + 140 + line_i * 42), line, anchor="mm", fill="#1A202C", font=label_font)
        if idx < len(steps) - 1:
            ax0 = x1 + 8
            ax1 = x1 + gap - 8
            ay = y0 + box_h / 2
            draw.line((ax0, ay, ax1, ay), fill="#2D3748", width=5)
            draw.polygon([(ax1, ay), (ax1 - 22, ay - 14), (ax1 - 22, ay + 14)], fill="#2D3748")
    note = "Claims move forward only after reference freeze and exception-aware gates are explicit."
    draw.text((w // 2, 735), note, anchor="mm", fill="#4A5568", font=note_font)
    out = fig_dir / f"figure_1_paper_b_source_anchored_workflow_{DATE}.png"
    save_png(img, out)
    outputs["workflow"] = out

    # Figure 2: denominator-family score profile.
    m1 = full_m1r_summary()
    labels = m1["family"].tolist()
    w, h = 2200, 1100
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    draw.text((w // 2, 70), "Full-corpus M1-R outcomes by denominator family", anchor="mm", fill="#1A202C", font=title_font)
    left, top, bar_w, bar_h = 650, 230, 1250, 120
    max_total = int(m1["rows_total"].max())
    colors = [("#2F855A", "Correct"), ("#C05621", "Incorrect among scored"), ("#718096", "Abstention")]
    for i, (_, row) in enumerate(m1.iterrows()):
        y = top + i * 210
        draw.text((90, y + bar_h / 2), labels[i], anchor="lm", fill="#1A202C", font=font(32, bold=True))
        x = left
        segments = [
            (int(row["correct_rows"]), "#2F855A"),
            (int(row["incorrect_scored_rows"]), "#C05621"),
            (int(row["abstention_rows"]), "#718096"),
        ]
        for value, color in segments:
            seg_w = int(bar_w * value / max_total)
            if seg_w > 0:
                draw.rectangle((x, y, x + seg_w, y + bar_h), fill=color)
                x += seg_w
        draw.rectangle((left, y, left + bar_w, y + bar_h), outline="#2D3748", width=3)
        draw.text((left + bar_w + 40, y + bar_h / 2), f'n={int(row["rows_total"])}', anchor="lm", fill="#1A202C", font=font(30))
    legend_x, legend_y = 720, 910
    for idx, (color, label) in enumerate(colors):
        x = legend_x + idx * 420
        draw.rectangle((x, legend_y, x + 48, legend_y + 48), fill=color)
        draw.text((x + 65, legend_y + 24), label, anchor="lm", fill="#1A202C", font=font(28))
    out = fig_dir / f"figure_2_m1r_denominator_family_outcomes_{DATE}.png"
    save_png(img, out)
    outputs["m1r"] = out

    # Figure 3: triage signals.
    signal, _priority = rq3_summary()
    selected = signal[
        signal["signal"].isin(
            [
                "cross_model_behavior_disagreement",
                "blank_behavior_family",
                "primary_not_scored",
                "primary_incorrect",
                "primary_abstention",
                "source_or_trace_risk",
                "human_disagreement_trace",
                "high_consequence_numeric",
            ]
        )
    ].copy()
    selected["signal_label"] = selected["signal"].str.replace("_", " ")
    selected = selected.sort_values("flagged_n")
    w, h = 2200, 1400
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    draw.text((w // 2, 70), "Review-triage signals in the model-overlap task universe", anchor="mm", fill="#1A202C", font=title_font)
    left, top, bar_w, bar_h = 780, 180, 1100, 82
    max_flagged = int(selected["flagged_n"].max())
    for i, (_, row) in enumerate(selected.iterrows()):
        y = top + i * 140
        label = str(row["signal_label"])
        draw.text((80, y + bar_h / 2), label, anchor="lm", fill="#1A202C", font=font(27, bold=True))
        width = int(bar_w * int(row["flagged_n"]) / max_flagged)
        draw.rectangle((left, y, left + width, y + bar_h), fill="#2B6CB0")
        draw.rectangle((left, y, left + bar_w, y + bar_h), outline="#CBD5E0", width=2)
        draw.text((left + width + 35, y + bar_h / 2), f'{int(row["flagged_n"])}', anchor="lm", fill="#1A202C", font=font(28))
    out = fig_dir / f"figure_3_rq3_review_triage_signals_{DATE}.png"
    save_png(img, out)
    outputs["triage"] = out

    # Figure 4: downstream matrix gates.
    gate = masem_gate_summary()
    w, h = 2200, 1100
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    draw.text((w // 2, 70), "MASEM/TSSEM construct-set identification gate", anchor="mm", fill="#1A202C", font=title_font)
    left, bottom, bar_w, max_h = 250, 840, 260, 560
    max_cc = max(1, int(gate["complete_case_studies"].max()))
    for i, (_, row) in enumerate(gate.iterrows()):
        x = left + i * 370
        value = int(row["complete_case_studies"])
        color = "#2F855A" if value > 1 else "#C05621" if value == 1 else "#718096"
        height = int(max_h * value / max_cc)
        draw.rectangle((x, bottom - height, x + bar_w, bottom), fill=color)
        draw.rectangle((x, bottom - max_h, x + bar_w, bottom), outline="#CBD5E0", width=2)
        draw.text((x + bar_w / 2, bottom - height - 28), str(value), anchor="mm", fill="#1A202C", font=font(30, bold=True))
        label = construct_set_label(row["construct_set"], multiline=True)
        for line_i, line in enumerate(label.split("\n")):
            draw.text((x + bar_w / 2, bottom + 42 + line_i * 30), line, anchor="mm", fill="#1A202C", font=font(24))
    draw.text((45, bottom - max_h / 2 - 22), "Complete-case", anchor="lm", fill="#1A202C", font=font(24))
    draw.text((45, bottom - max_h / 2 + 16), "studies", anchor="lm", fill="#1A202C", font=font(24))
    out = fig_dir / f"figure_4_masem_identification_gate_{DATE}.png"
    save_png(img, out)
    outputs["masem_gate"] = out
    return outputs


def fmt_pct(value: float) -> str:
    if pd.isna(value):
        return ""
    return f"{value * 100:.1f}%"


CONSTRUCT_SET_LABELS = {
    "core6_legacy_tssem_diagnostic": "Core 6 legacy TSSEM",
    "core7_add_att": "Core 7 plus ATT",
    "core8_add_tru": "Core 8 plus TRU",
    "core9_add_anx": "Core 9 plus ANX",
    "theory_target_10": "Theory target 10",
}

GATE_LABELS = {
    "eligible_for_bounded_tssem_diagnostic": "Bounded",
    "not_identified_as_complete_case_model": "Not identified",
}


def construct_set_label(value: str, multiline: bool = False) -> str:
    label = CONSTRUCT_SET_LABELS.get(str(value), str(value).replace("_", " "))
    if multiline:
        return label.replace(" plus ", "\n+ ").replace(" legacy ", "\nlegacy ")
    return label


def gate_label(value: str) -> str:
    return GATE_LABELS.get(str(value), str(value).replace("_", " "))


def gate_display_table(gate: pd.DataFrame) -> pd.DataFrame:
    display = gate[
        [
            "construct_set",
            "construct_count",
            "required_pairs",
            "covered_pairs",
            "missing_pairs",
            "complete_case_studies",
            "identification_gate",
        ]
    ].copy()
    display["construct_set"] = display["construct_set"].map(construct_set_label)
    display["missing_pairs"] = display["missing_pairs"].map(lambda x: "" if pd.isna(x) else x)
    display["identification_gate"] = display["identification_gate"].map(gate_label)
    return display


def table_lines(df: pd.DataFrame, columns: list[str]) -> str:
    header = "| " + " | ".join(columns) + " |"
    sep = "| " + " | ".join(["---"] * len(columns)) + " |"
    lines = [header, sep]
    for _, row in df.iterrows():
        lines.append("| " + " | ".join(str(row.get(col, "")) for col in columns) + " |")
    return "\n".join(lines)


def build_docs(paths: dict[str, Path], figures: dict[str, Path]) -> dict[str, Path]:
    m1 = full_m1r_summary()
    signal, priority = rq3_summary()
    gate = masem_gate_summary()

    m1_md = m1.assign(
        accuracy_scored_only=m1["accuracy_scored_only"].map(fmt_pct),
        accuracy_all_scorable=m1["accuracy_all_scorable"].map(fmt_pct),
    )[
        [
            "family",
            "rows_total",
            "scored_rows",
            "correct_rows",
            "incorrect_scored_rows",
            "abstention_rows",
            "accuracy_scored_only",
            "accuracy_all_scorable",
        ]
    ]
    signal_short = signal[
        [
            "signal",
            "scope_n",
            "flagged_n",
            "review_needed_n",
            "precision_review_needed",
            "recall_review_needed",
            "review_burden_share",
            "baseline_review_needed_rate",
            "precision_lift_vs_baseline",
        ]
    ].copy()
    for col in ["precision_review_needed", "recall_review_needed", "review_burden_share", "baseline_review_needed_rate", "precision_lift_vs_baseline"]:
        signal_short[col] = signal_short[col].map(lambda x: "" if pd.isna(x) else f"{x:.3f}")

    debrief = f"""
# Paper B Analysis Structure Debrief

Date: 2026-06-19

## One-sentence Answer

The 2,043-row M1-R result is not a replacement claim. It is the completed source-packet-required Step 5 evidence showing how the prespecified Codex GPT-5.5 workflow performs across denominator families, abstentions, and exception gates after the 213-study source-anchored reference was frozen.

## Analysis Structure

1. Raw human coder workbooks are preserved as read-only trace evidence.
2. Human-human disagreements are summarized before any model comparison.
3. Source-document adjudication creates the source-anchored human reference layer.
4. The reference layer is frozen and documented before Step 5.
5. Locked LLM outputs are evaluated against the reference by task family, not as one pooled corpus-wide denominator.
6. Cross-model disagreement, source-risk flags, human-disagreement traces, abstentions, and missing model rows are interpreted as review-triage evidence.
7. Downstream MASEM/TSSEM is reported only as a bounded diagnostic. Current defensible lane is core-6; broader core7/core8/full-10 probes do not support stronger all-construct or all-row SEM substitution claims.

## What the 2,043-row Full-corpus M1-R Result Brings

- It resolves the previous Step 5 coverage blocker: the full source-rendered corpus was run across nine locked shards.
- It provides denominator-family evidence for three task strata: direct/source-reported correlations, latent/source-flagged correlations, and converted beta/path effects.
- It exposes abstention behavior rather than hiding it inside a pooled score.
- It supplies a defensible RSM-style validation result when paired with explicit exception handling, model provenance, and public-code/public-data boundaries.
- It does not authorize model-vendor ranking, autonomous replacement, or all-row SEM stability claims.

## Full-corpus M1-R Results by Denominator Family

{table_lines(m1_md, list(m1_md.columns))}

Interpretation: scored-only accuracy describes numeric agreement among rows where the model emitted a scorable value. All-scorable accuracy treats abstentions as incorrect workflow outcomes, which is the more conservative validation view for high-consequence extraction.

## RQ3 Triage Signal Summary

{table_lines(signal_short, list(signal_short.columns))}

Cross-model disagreement has high precision because the review-needed base rate is already extremely high. Its value is therefore not vendor ranking, but surfacing where the workflow should route attention together with source-risk and human-disagreement traces.

## MASEM/TSSEM Gate Summary

{table_lines(gate_display_table(gate), ["construct_set", "construct_count", "required_pairs", "covered_pairs", "missing_pairs", "complete_case_studies", "identification_gate"])}

## Figure Package

- Figure 1: `{figures["workflow"].name}`
- Figure 2: `{figures["m1r"].name}`
- Figure 3: `{figures["triage"].name}`
- Figure 4: `{figures["masem_gate"].name}`
"""

    status = f"""
# Paper B Near-submission Readiness Status

Date: 2026-06-19

## Current Decision State

Target journal: Research Synthesis Methods.

Current manuscript objective: near-submission draft with Results and Discussion treated as the strongest sections; Introduction, Literature Review, and Methods remain editable drafts that need final author review.

## Ready Now

- The source-anchored reference route is documented and the v2 consensus packet is the current official reference candidate.
- Full-corpus M1-R Step 5 has completed: 2,043 rows, nine locked shards, 0 duplicate task IDs, and 0 post-repair model CLI failures.
- Denominator-family scoring is available for direct/source-reported r, latent/source-flagged r, and converted beta/path rows.
- RQ3 triage files are available for cross-model disagreement, source risk, human-disagreement traces, abstentions, and missing model coverage.
- Core-6 TSSEM diagnostic evidence exists and broader sparse probes have been documented.

## Not Yet Submission-safe Without Final PI Review

- Do not report one pooled full-corpus accuracy denominator.
- Do not frame Codex, Claude, and Gemini as a vendor-ranking result.
- Do not claim autonomous replacement of human coders.
- Do not claim all-construct/all-row SEM substitution stability.
- Do not upload raw PDFs, raw human coder workbooks, or private source packets into public repository materials.

## Active Files

- Working manuscript: `../01_working_manuscript/Paper_B_작업원고_TRACK_CHANGES_{DATE}.docx`
- Untouched prior target draft: `../01_working_manuscript/Paper_B_원본_DO_NOT_EDIT_{DATE}.docx`
- Task board: `Paper_B_Researcher_Task_Board_{DATE}.xlsx`
- Table/Figure workbook: `Paper_B_Tables_Figures_Package_{DATE}.xlsx`
- Analysis debrief: `PAPER_B_ANALYSIS_STRUCTURE_DEBRIEF_{DATE}.md`
"""

    crosswalk = """
# Research Synthesis Methods Crosswalk for Paper B

Date: 2026-06-19

## Current RSM-facing Fit

Paper B directly fits the Research Synthesis Methods GenAI evaluation guidance because it reports research design, source data conditions, prompts/locked outputs, validation, quality assurance, reproducibility, and claim limitations.

## Crosswalk

| RSM/GenAI expectation | Paper B current evidence | Remaining action |
| --- | --- | --- |
| Research design and dataset characteristics | 213-study source-anchored reference; denominator-family task units; full-corpus Step 5 M1-R run | Keep counts and subset definitions in Methods and Results |
| Variables/task families | Direct/source-reported r, source-flagged latent r, converted beta/path rows, metadata/trace families | Do not collapse families into one pooled score |
| Prompt/model documentation | Locked output shards and model provenance are preserved | Add model access date/API/workflow parameter note before final submission |
| Validation | Source-anchored human reference standard and exception-aware scoring | Keep human adjudication and source-risk gate visible |
| Evaluation metrics | Scored-only and all-scorable accuracy, abstentions, exception gates, triage precision/recall | Report conservative interpretation first |
| Reproducibility | Scripts, locked outputs, manifests, OSF/public repository boundary | Refresh public manifest/checksums if 2026-06-12 full-corpus files are redistributed |
| Transparency and disclosure | Claim-boundary notes and no-replacement framing | Add final AI-use disclosure and conflicts/funding text |
"""

    start_here = f"""
# Paper B - Start Here

Updated: 2026-06-19

This folder is now the active Paper B workspace. Paper A requirements, Paper A Track Changes files, and Paper A APA/JARS tracking are intentionally kept outside this directory.

## Use These First

1. `01_working_manuscript/Paper_B_작업원고_TRACK_CHANGES_{DATE}.docx`
2. `01_working_manuscript/Paper_B_원본_DO_NOT_EDIT_{DATE}.docx`
3. `02_B_requirements_and_tracking/Paper_B_Researcher_Task_Board_{DATE}.xlsx`
4. `02_B_requirements_and_tracking/PAPER_B_ANALYSIS_STRUCTURE_DEBRIEF_{DATE}.md`
5. `02_B_requirements_and_tracking/PAPER_B_SUBMISSION_READINESS_STATUS_{DATE}.md`

## Current Paper B Boundary

Paper B is a Research Synthesis Methods manuscript about source-anchored validation of a prespecified LLM workflow for MASEM-ready evidence extraction. It supports workflow augmentation, review triage, denominator-family validation, and bounded downstream diagnostics. It does not support a pooled accuracy score, vendor ranking, autonomous replacement, or all-construct/all-row SEM stability.

## Current Folder Map

- `00_START_HERE/`: entry point and short instructions.
- `01_working_manuscript/`: editable Paper B manuscript and figures.
- `02_B_requirements_and_tracking/`: task board, status, debrief, RSM crosswalk, and legacy status docs.
- `03_B_reference_standard/`: Paper2 human final consensus and reference-standard candidates.
- `04_B_analysis_inputs_outputs/`: analysis inputs, outputs, and working packages.
- `05_B_manuscript_and_repository/`: manuscript repository and public-data repository materials.
- `06_B_source_adjudication/`: source-adjudication evidence.
- `99_archive/`: prior non-canonical manuscript drafts and legacy candidates.
"""

    outputs = {
        paths["tracking"] / f"PAPER_B_ANALYSIS_STRUCTURE_DEBRIEF_{DATE}.md": debrief,
        paths["tracking"] / f"PAPER_B_SUBMISSION_READINESS_STATUS_{DATE}.md": status,
        paths["tracking"] / f"PAPER_B_RSM_REQUIREMENTS_CROSSWALK_{DATE}.md": crosswalk,
        paths["start"] / f"PAPER_B_START_HERE_{DATE}.md": start_here,
        paths["start"] / "PAPER_B_START_HERE_20260618.md": start_here,
        OUT / f"PAPER_B_ANALYSIS_STRUCTURE_DEBRIEF_{DATE}.md": debrief,
        OUT / f"PAPER_B_SUBMISSION_READINESS_STATUS_{DATE}.md": status,
    }
    for path, text in outputs.items():
        write_text(path, text)
    return {path.name: path for path in outputs}


def add_para(doc: Document, text: str = "", style: str | None = None, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.bold = bold


def set_table_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "BFBFBF")


def add_table(doc: Document, title: str, headers: list[str], rows: list[list[object]]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    run = p.add_run(title)
    run.bold = True
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = True
    set_table_borders(table)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True)
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()


def add_picture(doc: Document, path: Path, caption: str) -> None:
    doc.add_picture(str(path), width=Inches(6.2))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(caption)
    run.italic = True


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color in [
        ("Heading 1", 16, "000000"),
        ("Heading 2", 14, "000000"),
        ("Heading 3", 12, "434343"),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True


def manuscript_sections(m1: pd.DataFrame, signal: pd.DataFrame, priority: pd.DataFrame, gate: pd.DataFrame) -> dict[str, list[str]]:
    return {
        "Introduction": [
            "Data extraction for meta-analytic structural equation modeling (MASEM) requires more than article summarization. A valid record must identify source-reported correlations, distinguish converted statistics from direct correlations, map constructs, preserve sample-size eligibility, and keep enough provenance to support downstream model fitting.",
            "Large language models (LLMs) may reduce review burden, but evidence-synthesis methodology journals increasingly require validation designs that are transparent about model versioning, prompts, source conditions, human oversight, and reproducibility. Research Synthesis Methods now explicitly welcomes ML and GenAI evaluations in systematic review and meta-analysis contexts when the research design, validation, and reporting are sufficiently transparent.",
            "This study evaluates a prespecified locked-output LLM workflow for MASEM-ready extraction in an AI adoption evidence-synthesis project. The contribution is not a model leaderboard. Instead, we ask whether a source-anchored LLM workflow can augment human review by structuring numeric extraction, uncertainty, source risk, and downstream substitution diagnostics without collapsing heterogeneous task families into one accuracy score.",
        ],
        "Literature Review": [
            "Recent evidence-synthesis studies show that LLMs can support parts of data extraction, but performance depends strongly on the task, source material, and human verification design. Bibliographic or low-consequence metadata extraction is not equivalent to recovering a correlation matrix that will later feed TSSEM or OSMASEM.",
            "The relevant methodological literature therefore points toward four requirements for this paper: a defensible human reference standard, locked and auditable model outputs, task-family denominators, and downstream consequence checks. Paper B follows that logic by separating direct correlations, source-type-flagged latent correlations, converted beta/path evidence, metadata and trace rows, and source-risk cases.",
            "For MASEM specifically, numeric agreement is only one layer. A row that appears numerically close may still be unusable if it is a beta coefficient, an HTMT value, a duplicate-source trace, an ineligible source statistic, or a value without a source-supported sample size. The literature review should be expanded further before submission, but the Results and Discussion below already use this task-contingent framing.",
        ],
        "Method": [
            "The validation corpus was derived from the AI adoption in higher education MASEM project. The workflow preserved raw independent human coder workbooks, generated human-human disagreement queues, adjudicated source-document disagreements, and then froze a source-anchored human reference layer before locked model evaluation.",
            "The official current reference candidate is the Paper2 Human Final Consensus v2 packet prepared on 2026-06-05. The post-freeze full-corpus gate used a 213-study source-anchored reference, source-rendered task packets, and a full-corpus Step 5 shell containing 2,043 M1-R task rows.",
            "The primary prespecified model for the full-corpus Step 5 run was Codex GPT-5.5. The full run was executed across nine locked shards. Claude Sonnet and Gemini 3 Flash outputs were retained for cross-model disagreement and triage sensitivity, not for vendor ranking.",
            "Evaluation was conducted by denominator family. Direct/source-reported correlations, latent or construct correlations with source-type flags, and secondary beta/path converted effects were interpreted separately. Abstentions on scorable rows were counted as incorrect workflow outcomes. Fifteen exception-layer rows were interpreted by gate status rather than included in a generic full-corpus accuracy numerator.",
        ],
        "Results": [
            "The full-corpus M1-R blocker was resolved on 2026-06-12. The source-packet-required run covered 2,043 eligible task units across nine locked shards, with no duplicate task identifiers and no post-repair model CLI failures. The run therefore provides full-corpus Step 5 evidence, but only under denominator-family and exception-aware interpretation.",
            "Among latent/source-flagged correlation rows, 715 of 931 rows were scorable, 672 were correct, and 216 were abstentions. Among direct/source-reported correlation rows, 572 of 697 rows were scorable, 517 were correct, and 125 were abstentions. Among secondary beta/path converted rows, 338 of 415 rows were scorable, 153 were correct, and 77 were abstentions.",
            "Scored-only accuracy was highest for the two primary correlation strata and weakest for converted beta/path rows. This pattern is substantively important: the model can reproduce many source-rendered numeric correlations, but converted or model-derived evidence remains a sensitivity stratum rather than a safe replacement layer.",
            "RQ3 triage evidence showed that cross-model disagreement flagged 6,592 of 7,859 multi-model task units. Precision for identifying review-needed units was 0.999, recall was 0.843, and the review burden share was 0.839. Because the baseline review-needed rate was 0.994, cross-model disagreement should not be interpreted as a high-lift standalone detector. Its value is in organizing review attention together with source-risk, human-disagreement, abstention, and model-coverage signals.",
            "The downstream diagnostic remained bounded. The core-6 construct set retained the defensible TSSEM/MASEM diagnostic lane. Core7 and core8 probes were attempted but were too sparse for stronger main-text extension: core7_add_att had only three complete-case studies and failed the Stage 2 path-model probe because the asymptotic covariance matrix was not positive definite; core8_add_tru had only one complete-case study.",
        ],
        "Discussion": [
            "The results support a workflow-augmentation claim. A locked LLM workflow can structure MASEM-ready extraction review, preserve model provenance, expose abstention behavior, and route expert attention to high-consequence numeric and source-risk rows. The evidence does not support an autonomous replacement claim.",
            "The denominator-family results show why a single pooled accuracy score would be misleading. Direct/source-reported correlations and latent/source-flagged correlations carried different scoring implications, and converted beta/path rows were much less stable. A pooled score would obscure exactly the distinctions that matter for evidence synthesis.",
            "The high cross-model-disagreement precision should be read cautiously. Because nearly all multi-model task units were review-needed under the current rules, precision is inflated by the base rate. The signal is still useful operationally because it clusters review work, but it does not prove that disagreement alone efficiently identifies rare errors.",
            "The bounded core-6 diagnostic is equally important. In the expert-reviewed primary input, high-risk rows were retained under the human reference rather than autonomously replaced, and the diagnostic subset therefore did not change relative to the human-reference primary input. That is a conservative but methodologically useful result: the workflow can document where substitution is not authorized, not merely where a model appears correct.",
            "For Research Synthesis Methods readers, the main contribution is a transparent evaluation design for LLM use in complex evidence synthesis. The study shows how to combine source-anchored adjudication, locked outputs, exception-aware scoring, denominator-family reporting, and downstream diagnostic boundaries. The public repository should expose share-safe data, code, prompts, manifests, and checksums while excluding raw PDFs, raw human workbooks, and private source packets.",
            "The remaining limitation is that the Introduction, Literature Review, and Methods need final author polishing and reference completion. The Results and Discussion are now aligned with the current evidence boundaries: no vendor ranking, no pooled accuracy denominator, no autonomous replacement claim, and no all-construct/all-row SEM stability claim.",
        ],
    }


def build_manuscript_docx(paths: dict[str, Path], figures: dict[str, Path]) -> tuple[Path, Path, Path]:
    target_source = locate_preferred(
        [
            paths["manuscript"]
            / "repository_paper_b"
            / "manuscript"
            / "target_journal"
            / "PAPER_B_RESEARCH_SYNTHESIS_METHODS_TARGET_DRAFT_20260612.docx",
            REPO / "paper_b" / "manuscript" / "target_journal" / "PAPER_B_RESEARCH_SYNTHESIS_METHODS_TARGET_DRAFT_20260612.docx",
        ]
    )

    original = paths["working"] / f"Paper_B_원본_DO_NOT_EDIT_{DATE}.docx"
    if target_source.exists() and not original.exists():
        shutil.copy2(target_source, original)

    m1 = full_m1r_summary()
    signal, priority = rq3_summary()
    gate = masem_gate_summary()

    doc = Document()
    style_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Can a Prespecified LLM Workflow Augment MASEM-Ready Evidence Extraction?")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Target journal: Research Synthesis Methods").italic = True

    doc.add_heading("Abstract", level=1)
    add_para(
        doc,
        "Data extraction for meta-analytic structural equation modeling (MASEM) requires reviewers to recover numeric source evidence, map constructs, distinguish direct correlations from converted statistics, and preserve provenance for downstream model fitting. We evaluated a prespecified locked-output LLM workflow against a source-anchored adjudicated human reference standard in an AI adoption evidence-synthesis project. Task units were analyzed by denominator family rather than as one pooled accuracy denominator. The full-corpus M1-R run covered 2,043 source-rendered task rows across nine locked shards, with no duplicate task IDs or post-repair model CLI failures. Primary correlation strata showed higher scored-only accuracy than converted beta/path rows, while abstentions and exception-layer rows identified where automated substitution was not authorized. Cross-model disagreement was retained as a review-triage signal but not as a vendor-ranking metric. Downstream evidence supported only a bounded core-6 TSSEM diagnostic; broader core7/core8 probes were too sparse for stronger SEM claims. The results support workflow augmentation, denominator-family validation, and human-supervised review triage, not autonomous replacement.",
    )
    add_para(doc, "Keywords: evidence synthesis; data extraction; large language models; MASEM; validation; human-in-the-loop")

    doc.add_heading("Highlights", level=1)
    for heading, text in [
        ("What is already known", "LLMs can support parts of systematic review workflows, but data extraction accuracy is task-dependent and requires human verification."),
        ("What is new", "This study evaluates a locked LLM workflow for MASEM-ready extraction using source-anchored adjudication, denominator-family scoring, and downstream diagnostics."),
        ("Potential impact for Research Synthesis Methods readers", "The workflow shows how to evaluate LLM extraction without pooled accuracy overreach, vendor ranking, or replacement claims."),
    ]:
        doc.add_heading(heading, level=2)
        add_para(doc, text)

    sections = manuscript_sections(m1, signal, priority, gate)
    for heading in ["Introduction", "Literature Review", "Method"]:
        doc.add_heading(heading, level=1)
        for para in sections[heading]:
            add_para(doc, para)

    doc.add_heading("Results", level=1)
    for para in sections["Results"][:1]:
        add_para(doc, para)

    add_picture(
        doc,
        figures["workflow"],
        "Figure 1. Source-anchored Paper B validation workflow. Claims move forward only after reference freeze and exception-aware gates.",
    )

    data_state_rows = [
        ["Frozen full-corpus reference", "213 studies; frozen 2026-06-09", "Governing reference layer"],
        ["Full-corpus M1-R", "2,043 rows; nine locked shards", "Denominator-family Step 5 evidence"],
        ["Exception layer", "15 gate-status rows", "Prevents generic pooled accuracy overreach"],
        ["Core-6 TSSEM diagnostic", "16 complete-case studies in 2026-06-12 matrix audit; prior TSSEM diagnostic reported 15 studies", "Bounded downstream diagnostic"],
        ["Core7/core8 probes", "3 and 1 complete-case studies", "Sparse-probe evidence, not main SEM extension"],
    ]
    add_table(doc, "Table 1. Data states and claim roles", ["Data state", "Current evidence", "Claim role"], data_state_rows)

    m1_rows = []
    for _, r in m1.iterrows():
        m1_rows.append(
            [
                r["family"],
                int(r["rows_total"]),
                int(r["scored_rows"]),
                int(r["correct_rows"]),
                int(r["incorrect_scored_rows"]),
                int(r["abstention_rows"]),
                fmt_pct(r["accuracy_scored_only"]),
                fmt_pct(r["accuracy_all_scorable"]),
            ]
        )
    add_table(
        doc,
        "Table 2. Full-corpus M1-R outcomes by denominator family",
        ["Family", "Rows", "Scored", "Correct", "Incorrect", "Abstained", "Scored-only accuracy", "All-scorable accuracy"],
        m1_rows,
    )
    add_picture(
        doc,
        figures["m1r"],
        "Figure 2. Denominator-family profile for the 2,043-row full-corpus M1-R run. Converted beta/path rows remain a sensitivity stratum.",
    )

    for para in sections["Results"][1:4]:
        add_para(doc, para)

    triage_rows = []
    selected_signals = [
        "cross_model_behavior_disagreement",
        "blank_behavior_family",
        "primary_not_scored",
        "primary_incorrect",
        "primary_abstention",
        "source_or_trace_risk",
        "human_disagreement_trace",
        "high_consequence_numeric",
    ]
    for _, r in signal[signal["signal"].isin(selected_signals)].iterrows():
        triage_rows.append(
            [
                r["signal"].replace("_", " "),
                int(r["flagged_n"]),
                int(r["true_positive_n"]),
                int(r["false_positive_n"]),
                f'{r["precision_review_needed"]:.3f}' if not pd.isna(r["precision_review_needed"]) else "",
                f'{r["recall_review_needed"]:.3f}' if not pd.isna(r["recall_review_needed"]) else "",
            ]
        )
    add_table(
        doc,
        "Table 3. Review-triage signal validation",
        ["Signal", "Flagged", "True positives", "False positives", "Precision", "Recall"],
        triage_rows,
    )
    add_picture(doc, figures["triage"], "Figure 3. Main review-triage signal counts in the multi-model task universe.")

    for para in sections["Results"][4:]:
        add_para(doc, para)

    gate_rows = []
    for _, r in gate.iterrows():
        gate_rows.append(
            [
                construct_set_label(r["construct_set"]),
                int(r["construct_count"]),
                int(r["required_pairs"]),
                int(r["covered_pairs"]),
                "" if pd.isna(r["missing_pairs"]) else r["missing_pairs"],
                int(r["complete_case_studies"]),
                gate_label(r["identification_gate"]),
            ]
        )
    add_table(
        doc,
        "Table 4. MASEM/TSSEM construct-set identification gates",
        ["Construct set", "k", "Required pairs", "Covered pairs", "Missing", "Complete cases", "Gate"],
        gate_rows,
    )
    add_picture(doc, figures["masem_gate"], "Figure 4. Construct-set identification gate for downstream TSSEM/MASEM reporting.")

    doc.add_heading("Discussion", level=1)
    for para in sections["Discussion"]:
        add_para(doc, para)

    doc.add_heading("Data Availability", level=1)
    add_para(
        doc,
        "The share-safe Paper B public repository is available at https://osf.io/mkrgd/overview. If the 2026-06-12 full-corpus artifacts are redistributed as part of the public archive, the public manifest, checksums, and copied file set should be refreshed. Raw PDFs, raw human coder workbooks, private source packets, and unapproved raw model outputs should remain outside the public package.",
    )

    doc.add_heading("References", level=1)
    references = [
        "Farotimi, O., Dunn, A., Van Lissa, C. J., Polanin, J. R., Mavridis, D., & Pigott, T. D. (2026). Guidance for manuscript submissions testing the use of generative AI for systematic review and meta-analysis. Research Synthesis Methods, 17, 237-239. https://doi.org/10.1017/rsm.2025.10058",
        "Flemyng, E., & Noel-Storr, A. (2025). Responsible AI in Evidence Synthesis (RAISE): Guidance and recommendations (version 2). OSF. https://osf.io/cqa82",
        "Gartlehner, G., Kahwati, L., Hilscher, R., Thomas, I., Kugley, S., Crotty, K., Viswanathan, M., Nussbaumer-Streit, B., Booth, G., Erskine, N., Konet, A., & Chew, R. (2024). Data extraction for evidence synthesis using a large language model: A proof-of-concept study. Research Synthesis Methods, 15, 576-589. https://doi.org/10.1002/jrsm.1710",
        "Cheung, M. W.-L. (2015). Meta-analytic structural equation modeling. Wiley.",
        "Jak, S., & Cheung, M. W.-L. (2020). Meta-analytic structural equation modeling with moderating effects on SEM parameters. Psychological Methods.",
        "Page, M. J., McKenzie, J. E., Bossuyt, P. M., Boutron, I., Hoffmann, T. C., Mulrow, C. D., et al. (2021). The PRISMA 2020 statement. BMJ, 372, n71.",
    ]
    for ref in references:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)

    working = paths["working"] / f"Paper_B_작업원고_TRACK_CHANGES_{DATE}.docx"
    local_working = OUT / "tmp" / working.name
    ensure_dir(local_working.parent)
    doc.save(local_working)
    replace_file_with_retry(local_working, working)
    enable_track_changes(working)

    md = paths["working"] / f"Paper_B_작업원고_TRACK_CHANGES_{DATE}.md"
    write_text(
        md,
        "# Can a Prespecified LLM Workflow Augment MASEM-Ready Evidence Extraction?\n\n"
        "This markdown companion is generated alongside the DOCX working manuscript. Use the DOCX as the authoritative editable manuscript.\n\n"
        "Results and Discussion are written to current evidence boundaries; Introduction, Literature Review, and Method remain author-review drafts.",
    )

    return original, working, md


def enable_track_changes(docx_path: Path) -> None:
    settings_name = "word/settings.xml"
    tmp = OUT / "tmp" / f"{docx_path.stem}.track.tmp.docx"
    ensure_dir(tmp.parent)
    with zipfile.ZipFile(docx_path, "r") as src, zipfile.ZipFile(tmp, "w", compression=zipfile.ZIP_DEFLATED) as dst:
        seen: set[str] = set()
        for item in src.infolist():
            if item.filename in seen:
                continue
            seen.add(item.filename)
            data = src.read(item.filename)
            if item.filename == settings_name:
                xml = data.decode("utf-8")
                if "w:trackRevisions" not in xml:
                    xml = xml.replace("</w:settings>", '<w:trackRevisions w:val="true"/></w:settings>')
                data = xml.encode("utf-8")
            dst.writestr(item, data)
    replace_file_with_retry(tmp, docx_path)


def build_workbooks(paths: dict[str, Path]) -> tuple[Path, Path]:
    m1 = full_m1r_summary()
    signal, priority = rq3_summary()
    gate = masem_gate_summary()

    task_rows = [
        ["B1", "R1", "High", "Workflow order and reference-standard claim boundary", "Check that raw freeze, disagreement, source adjudication, reference freeze, and model comparison are not conflated.", "Not started", "R1/B1_PaperB_workflow_order_claim_boundary_20260619.md"],
        ["B2", "R2", "High", "Human-human disagreement synthesis", "Summarize coder differences by source condition, construct mapping, sample choice, statistic type, and downstream consequence.", "Not started", "R2/B2_human_disagreement_synthesis_20260619.xlsx"],
        ["B3", "R3", "High", "Source-adjudication gap queue", "Identify rows where source-document adjudication or page/table locator evidence remains insufficient.", "Not started", "R3/B3_source_adjudication_gap_queue_20260619.xlsx"],
        ["B4", "R4", "High", "M1-R results table and appendix package", "Convert 2,043-row M1-R denominator-family results into manuscript and appendix-ready tables.", "Not started", "R4/B4_m1r_results_package_20260619.xlsx"],
        ["B5", "R1", "High", "SEM claim-boundary review", "Confirm core-6 diagnostic language and block all-construct/all-row substitution claims.", "Not started", "R1/B5_sem_claim_boundary_review_20260619.docx"],
        ["B6", "R4", "Medium", "Tables/Figures reader check", "Check table order, captions, figure fit, and reader-facing interpretation.", "Not started", "R4/B6_tables_figures_reader_check_20260619.xlsx"],
        ["B7", "PI", "High", "Results final pass", "Verify Results text against 2026-06-12 full-corpus files and 2026-06-11/12 TSSEM diagnostics.", "Ready for PI review", "PI/B7_results_final_pass_20260619.docx"],
        ["B8", "PI", "High", "Discussion final pass", "Verify no pooled accuracy, vendor-ranking, replacement, or all-row SEM claim remains.", "Ready for PI review", "PI/B8_discussion_final_pass_20260619.docx"],
        ["B9", "PI/R4", "High", "RSM/GenAI requirements crosswalk", "Complete AI-use disclosure, model access date, prompt/parameter note, code/data availability, and public repository boundary.", "In progress", "PI/B9_rsm_genai_crosswalk_20260619.xlsx"],
    ]

    def style_ws(ws):
        header_fill = PatternFill("solid", fgColor="D9EAF7")
        thin = Side(style="thin", color="BFBFBF")
        border = Border(top=thin, left=thin, right=thin, bottom=thin)
        for row in ws.iter_rows():
            for cell in row:
                cell.alignment = Alignment(wrap_text=True, vertical="top")
                cell.border = border
        for cell in ws[1]:
            cell.fill = header_fill
            cell.font = Font(bold=True)
        ws.freeze_panes = "A2"
        for col in range(1, ws.max_column + 1):
            max_len = max(len(str(ws.cell(row=r, column=col).value or "")) for r in range(1, ws.max_row + 1))
            ws.column_dimensions[get_column_letter(col)].width = min(max(max_len + 2, 12), 42)

    board = Workbook()
    ws = board.active
    ws.title = "Paper B Tasks"
    ws.append(["Task ID", "Owner", "Priority", "Task", "Instruction", "Status", "Output path"])
    for row in task_rows:
        ws.append(row)
    dv = DataValidation(type="list", formula1='"Not started,In progress,Review requested,Ready for PI review,Complete,Blocked"', allow_blank=False)
    ws.add_data_validation(dv)
    dv.add(f"F2:F{ws.max_row}")
    style_ws(ws)

    ws2 = board.create_sheet("Summary")
    ws2.append(["Item", "Value"])
    ws2.append(["Target journal", "Research Synthesis Methods"])
    ws2.append(["Current goal", "Near-submission draft; Results/Discussion strongest sections"])
    ws2.append(["Official reference candidate", "Paper2 Human Final Consensus v2"])
    ws2.append(["Full-corpus M1-R rows", 2043])
    ws2.append(["Allowed main claim", "LLM workflow augmentation and review triage"])
    ws2.append(["Blocked claims", "Pooled accuracy, vendor ranking, autonomous replacement, all-row SEM stability"])
    style_ws(ws2)

    ws3 = board.create_sheet("RSM Crosswalk")
    ws3.append(["Requirement", "Current evidence", "Next action"])
    for row in [
        ["Research design", "213-study source-anchored reference; 2,043-row full-corpus M1-R", "Keep dataset/subset counts explicit"],
        ["Evaluation", "Denominator-family accuracy, abstention, exception gates", "Report conservative all-scorable interpretation"],
        ["Validation", "Human reference standard and source-document adjudication", "Add final PI-approved reference freeze wording"],
        ["Reproducibility", "Locked shards, scripts, manifests, checksums", "Refresh public repository manifest if 2026-06-12 files are public"],
        ["Transparency", "Claim-boundary docs", "Add final AI disclosure and model access dates"],
    ]:
        ws3.append(row)
    style_ws(ws3)

    task_board = paths["tracking"] / f"Paper_B_Researcher_Task_Board_{DATE}.xlsx"
    local_board = OUT / "tmp" / task_board.name
    ensure_dir(local_board.parent)
    board.save(local_board)
    replace_file_with_retry(local_board, task_board)

    package = Workbook()
    ws = package.active
    ws.title = "M1R Results"
    ws.append(list(m1.columns))
    for _, row in m1.iterrows():
        ws.append([row[c] for c in m1.columns])
    style_ws(ws)

    ws = package.create_sheet("RQ3 Signals")
    ws.append(list(signal.columns))
    for _, row in signal.iterrows():
        ws.append([row[c] for c in signal.columns])
    style_ws(ws)

    ws = package.create_sheet("Review Priority")
    ws.append(list(priority.columns))
    for _, row in priority.iterrows():
        ws.append([row[c] for c in priority.columns])
    style_ws(ws)

    ws = package.create_sheet("MASEM Gates")
    gate_display = gate_display_table(gate)
    ws.append(list(gate_display.columns))
    for _, row in gate_display.iterrows():
        ws.append([row[c] for c in gate_display.columns])
    style_ws(ws)

    table_package = paths["tracking"] / f"Paper_B_Tables_Figures_Package_{DATE}.xlsx"
    local_table_package = OUT / "tmp" / table_package.name
    ensure_dir(local_table_package.parent)
    package.save(local_table_package)
    replace_file_with_retry(local_table_package, table_package)
    return task_board, table_package


def validate_docx(path: Path) -> bool:
    with zipfile.ZipFile(path) as zf:
        bad = zf.testzip()
    return bad is None


def create_manifest(paths: dict[str, Path], artifacts: list[Path]) -> Path:
    manifest = OUT / f"PAPER_B_PACKAGE_MANIFEST_{DATE}.csv"
    ensure_dir(manifest.parent)
    with manifest.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerow(["artifact", "path", "bytes", "sha256"])
        for artifact in artifacts:
            if artifact.exists() and artifact.is_file():
                writer.writerow([artifact.name, str(artifact), artifact.stat().st_size, sha256(artifact)])
    return manifest


def main() -> None:
    paths = restructure_paper_b()
    archive_old_manuscript_candidates(paths)
    figures = create_figures(paths)
    docs = build_docs(paths, figures)
    original, working, md = build_manuscript_docx(paths, figures)
    board, table_package = build_workbooks(paths)

    artifacts = [
        *figures.values(),
        *docs.values(),
        original,
        working,
        md,
        board,
        table_package,
    ]
    manifest = create_manifest(paths, artifacts)

    print(f"Paper B package built under: {PAPER_B}")
    print(f"Working manuscript: {working}")
    print(f"Original manuscript: {original}")
    print(f"Task board: {board}")
    print(f"Table package: {table_package}")
    print(f"Manifest: {manifest}")
    print(f"Working DOCX valid zip: {validate_docx(working)}")
    if original.exists():
        print(f"Original DOCX valid zip: {validate_docx(original)}")


if __name__ == "__main__":
    main()
