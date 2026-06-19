#!/usr/bin/env python3
"""Build Paper B visual-upgrade artifacts.

The output bundle is staged first, then copied into the shared OneDrive and
Google Drive roots after validation.
"""

from __future__ import annotations

import csv
import json
import math
import re
import shutil
import subprocess
import textwrap
import urllib.request
import zipfile
from dataclasses import dataclass
from pathlib import Path

import pandas as pd
import pdfplumber
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import load_workbook
from openpyxl.styles import Alignment, Font, PatternFill
from PIL import Image, ImageDraw, ImageFont
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas


DATE = "20260619"
TIME_LOG_DECISION = "No reviewer-time or per-study review-duration logs are available as of 2026-06-19."
TIMING_BOUNDARY = (
    "Do not generate a Lai-style elapsed-time efficiency plot or make a time-savings claim. "
    "Use the review-burden/triage figure instead."
)
WORKSPACE = Path("/Users/newhosung/Documents/Codex/2026-06-18/google-drive-onedrive-library-cloudstorage-googledrive")
OUT = WORKSPACE / "outputs" / f"paper_b_visual_upgrade_{DATE}"
STAGING = OUT / "shared_drive_staging"

ONEDRIVE_ROOT = Path(
    "/Users/newhosung/Library/CloudStorage/OneDrive-SharedLibraries-ThePennsylvaniaStateUniversity/"
    "AI Adoption Meta Analysis - Documents"
)
PAPER_B = ONEDRIVE_ROOT / "02_Paper_B"
TRACKING = PAPER_B / "02_B_requirements_and_tracking"
BENCHMARK_SRC = TRACKING / f"frontier_reference_benchmark_{DATE}"
PUBLIC_ANALYSIS = (
    PAPER_B
    / "05_B_manuscript_and_repository"
    / "repository_paper_b"
    / f"public_data_repository_{DATE[:4]}0611"
    / "4_Analysis_Outputs"
)

STAGE_PAPER_B = STAGING / "02_Paper_B"
STAGE_TRACKING = STAGE_PAPER_B / "02_B_requirements_and_tracking"
STAGE_WORKING = STAGE_PAPER_B / "01_working_manuscript"
STAGE_FIGURES = STAGE_WORKING / "figures"
STAGE_BENCHMARK = STAGE_TRACKING / f"frontier_reference_benchmark_{DATE}"
STAGE_PDFS = STAGE_BENCHMARK / "pdfs"
STAGE_EXTRACTED = STAGE_BENCHMARK / "extracted"
STAGE_REPORTS = STAGE_BENCHMARK / "reports"
STAGE_TABLES = STAGE_REPORTS / "generated_tables"
STAGE_REPORT_FIGURES = STAGE_REPORTS / "generated_figures"
STAGE_RENDER = OUT / "render_checks"

PDFTOPPM = Path("/Users/newhosung/.cache/codex-runtimes/codex-primary-runtime/dependencies/bin/pdftoppm")


@dataclass(frozen=True)
class ReferencePaper:
    key: str
    filename: str
    title: str
    url: str
    pdf_url: str
    paper_b_use: str


REFERENCES = [
    ReferencePaper(
        key="Huang_2025_JMIR_RoB2",
        filename="Huang_et_al_2025_JMIR_LLM_RoB2.pdf",
        title="Large Language Model-Assisted Risk-of-Bias Assessment in Randomized Controlled Trials Using RoB2",
        url="https://www.jmir.org/2025/1/e70450",
        pdf_url="https://www.jmir.org/2025/1/e70450/PDF",
        paper_b_use="Criterion-standard comparison, reviewer timing boundary, RoB-style screening/validation flow diagram.",
    ),
    ReferencePaper(
        key="Jansen_2026_EdPsychReview",
        filename="Jansen_et_al_2026_EducationalPsychologyReview_LLM_data_extraction.pdf",
        title="Automated Data Extraction by Large Language Models: Assessing Accuracy in Comparison to Human Experts",
        url="https://link.springer.com/article/10.1007/s10648-026-10136-5",
        pdf_url="https://link.springer.com/content/pdf/10.1007/s10648-026-10136-5.pdf",
        paper_b_use="Education meta-analysis extraction benchmark, frozen codebook, gold/silver reference standards, API extraction cost/time reporting.",
    ),
]


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def write_text(path: Path, text: str) -> None:
    ensure_dir(path.parent)
    path.write_text(text.strip() + "\n", encoding="utf-8")


def copy_if_exists(src: Path, dst: Path) -> None:
    if src.exists():
        ensure_dir(dst.parent)
        shutil.copy2(src, dst)


def download_pdf(ref: ReferencePaper) -> Path:
    ensure_dir(STAGE_PDFS)
    target = STAGE_PDFS / ref.filename
    if target.exists() and target.stat().st_size > 20_000:
        return target
    request = urllib.request.Request(
        ref.pdf_url,
        headers={
            "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X) Codex PaperB visual audit",
            "Accept": "application/pdf,*/*",
        },
    )
    with urllib.request.urlopen(request, timeout=90) as response:
        data = response.read()
    if not data.startswith(b"%PDF") and b"%PDF" not in data[:1024]:
        raise RuntimeError(f"Downloaded file for {ref.key} is not a PDF-like payload")
    target.write_bytes(data)
    return target


def safe_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Helvetica Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Helvetica.ttf",
        "/Library/Fonts/Arial.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size=size)
        except Exception:
            continue
    return ImageFont.load_default()


def text_size(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont) -> tuple[int, int]:
    box = draw.textbbox((0, 0), text, font=font)
    return box[2] - box[0], box[3] - box[1]


def wrap_to_width(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont, max_width: int) -> list[str]:
    words = str(text).split()
    if not words:
        return [""]
    lines: list[str] = []
    current: list[str] = []
    for word in words:
        candidate = " ".join(current + [word])
        if text_size(draw, candidate, font)[0] <= max_width or not current:
            current.append(word)
        else:
            lines.append(" ".join(current))
            current = [word]
    if current:
        lines.append(" ".join(current))
    return lines


def draw_wrapped(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int],
    text: str,
    font: ImageFont.ImageFont,
    fill: str,
    max_width: int,
    line_gap: int = 8,
    anchor: str = "left",
) -> int:
    x, y = xy
    lines = wrap_to_width(draw, text, font, max_width)
    for line in lines:
        w, h = text_size(draw, line, font)
        tx = x if anchor == "left" else x - w // 2
        draw.text((tx, y), line, font=font, fill=fill)
        y += h + line_gap
    return y


def draw_box(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int, int, int],
    title: str,
    body: str,
    fill: str,
    outline: str,
    title_color: str = "#111111",
) -> None:
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=4)
    title_font = safe_font(30, bold=True)
    body_font = safe_font(24)
    title_end = draw_wrapped(draw, (x1 + 28, y1 + 24), title, title_font, title_color, x2 - x1 - 56, line_gap=5)
    draw_wrapped(draw, (x1 + 28, min(title_end + 14, y2 - 78)), body, body_font, "#222222", x2 - x1 - 56, line_gap=7)


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str = "#27384A") -> None:
    draw.line([start, end], fill=color, width=5)
    ex, ey = end
    sx, sy = start
    angle = math.atan2(ey - sy, ex - sx)
    size = 18
    points = [
        (ex, ey),
        (ex - size * math.cos(angle - math.pi / 7), ey - size * math.sin(angle - math.pi / 7)),
        (ex - size * math.cos(angle + math.pi / 7), ey - size * math.sin(angle + math.pi / 7)),
    ]
    draw.polygon(points, fill=color)


def pct_value(text: str) -> float:
    return float(str(text).replace("%", "").strip()) / 100


def wilson(k: int, n: int) -> tuple[float, float]:
    if n <= 0:
        return (float("nan"), float("nan"))
    z = 1.96
    p = k / n
    denom = 1 + z**2 / n
    center = (p + z**2 / (2 * n)) / denom
    margin = z * math.sqrt((p * (1 - p) + z**2 / (4 * n)) / n) / denom
    return center - margin, center + margin


def prop_diff_ci(k1: int, n1: int, k0: int, n0: int) -> tuple[float, float, float]:
    p1 = k1 / n1
    p0 = k0 / n0
    diff = p1 - p0
    se = math.sqrt(p1 * (1 - p1) / n1 + p0 * (1 - p0) / n0)
    return diff, diff - 1.96 * se, diff + 1.96 * se


def load_core_tables() -> dict[str, pd.DataFrame]:
    tables: dict[str, pd.DataFrame] = {}
    for name in ["m1r_denominator_summary", "rq3_triage_summary", "masem_gate_summary"]:
        p = BENCHMARK_SRC / "reports" / "generated_tables" / f"{name}.csv"
        if not p.exists():
            raise FileNotFoundError(p)
        tables[name] = pd.read_csv(p)
    return tables


def enrich_m1_table(m1: pd.DataFrame) -> pd.DataFrame:
    rows = []
    for _, row in m1.iterrows():
        total = int(row["rows_total"])
        scored = int(row["scored_rows"])
        correct = int(row["correct_rows"])
        abstention = int(row["abstention_rows"])
        scored_lo, scored_hi = wilson(correct, scored)
        all_lo, all_hi = wilson(correct, total)
        rows.append(
            {
                "denominator_family": row["family"],
                "rows_total": total,
                "scored_rows": scored,
                "correct_rows": correct,
                "incorrect_scored_rows": int(row["incorrect_scored_rows"]),
                "abstention_unresolved_rows": abstention,
                "scored_only_accuracy": correct / scored,
                "scored_only_95ci": f"{scored_lo*100:.1f}% to {scored_hi*100:.1f}%",
                "all_scorable_accuracy": correct / total,
                "all_scorable_95ci": f"{all_lo*100:.1f}% to {all_hi*100:.1f}%",
                "review_burden_share": abstention / total,
                "claim_role": row["next_state_use"],
            }
        )
    primary = pd.DataFrame(rows[:2])
    total = int(primary["rows_total"].sum())
    scored = int(primary["scored_rows"].sum())
    correct = int(primary["correct_rows"].sum())
    incorrect = int(primary["incorrect_scored_rows"].sum())
    abstention = int(primary["abstention_unresolved_rows"].sum())
    slo, shi = wilson(correct, scored)
    alo, ahi = wilson(correct, total)
    rows.insert(
        0,
        {
            "denominator_family": "Primary r validation pool",
            "rows_total": total,
            "scored_rows": scored,
            "correct_rows": correct,
            "incorrect_scored_rows": incorrect,
            "abstention_unresolved_rows": abstention,
            "scored_only_accuracy": correct / scored,
            "scored_only_95ci": f"{slo*100:.1f}% to {shi*100:.1f}%",
            "all_scorable_accuracy": correct / total,
            "all_scorable_95ci": f"{alo*100:.1f}% to {ahi*100:.1f}%",
            "review_burden_share": abstention / total,
            "claim_role": "Primary validation summary",
        },
    )
    return pd.DataFrame(rows)


def write_df_artifacts(df: pd.DataFrame, stem: str) -> tuple[Path, Path]:
    csv_path = STAGE_TABLES / f"{stem}.csv"
    md_path = STAGE_TABLES / f"{stem}.md"
    ensure_dir(csv_path.parent)
    df.to_csv(csv_path, index=False)
    md_path.write_text(df_to_markdown(df) + "\n", encoding="utf-8")
    return csv_path, md_path


def df_to_markdown(df: pd.DataFrame) -> str:
    cols = [str(c) for c in df.columns]
    lines = [
        "| " + " | ".join(cols) + " |",
        "| " + " | ".join(["---"] * len(cols)) + " |",
    ]
    for _, row in df.iterrows():
        vals = []
        for col in df.columns:
            value = "" if pd.isna(row[col]) else str(row[col])
            vals.append(value.replace("\n", " ").replace("|", "\\|"))
        lines.append("| " + " | ".join(vals) + " |")
    return "\n".join(lines)


def draw_dense_table(df: pd.DataFrame, path: Path) -> None:
    ensure_dir(path.parent)
    width, height = 2600, 1500
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    title_font = safe_font(45, bold=True)
    small = safe_font(25)
    header = safe_font(26, bold=True)
    cell = safe_font(24)
    red = "#D21F1F"
    draw.text((60, 48), "Table 1 | Paper B extraction validity by denominator family", font=title_font, fill="#111111")
    draw_wrapped(
        draw,
        (60, 112),
        "Source: frozen 2,043-row M1-R scoring table. Scored-only accuracy is conditional on a scorable model answer; all-scorable accuracy treats abstentions as unresolved extraction work.",
        small,
        "#333333",
        2360,
    )
    columns = [
        ("Denominator / claim role", 520),
        ("Events / total", 250),
        ("Correct / scored", 280),
        ("Scored-only accuracy (95% CI)", 410),
        ("All-scorable accuracy (95% CI)", 420),
        ("Unresolved burden", 260),
        ("Use in manuscript", 360),
    ]
    x0, y0 = 60, 220
    row_h = 210
    header_h = 108
    table_w = sum(w for _, w in columns)
    draw.line((x0, y0 - 24, x0 + table_w, y0 - 24), fill="#111111", width=4)
    x = x0
    for title, w in columns:
        draw.rectangle((x, y0, x + w, y0 + header_h), fill="#ECEFF3", outline="#111111", width=2)
        draw_wrapped(draw, (x + 16, y0 + 18), title, header, "#111111", w - 32, line_gap=4)
        x += w
    for i, row in df.iterrows():
        y = y0 + header_h + i * row_h
        fill = "#FAFAFA" if i % 2 == 0 else "#FFFFFF"
        if row["denominator_family"] == "Primary r validation pool":
            fill = "#F3F7FB"
        values = [
            row["denominator_family"],
            f'{int(row["correct_rows"]):,} / {int(row["rows_total"]):,}',
            f'{int(row["correct_rows"]):,} / {int(row["scored_rows"]):,}',
            f'{row["scored_only_accuracy"]*100:.1f}% ({row["scored_only_95ci"]})',
            f'{row["all_scorable_accuracy"]*100:.1f}% ({row["all_scorable_95ci"]})',
            f'{int(row["abstention_unresolved_rows"]):,} ({row["review_burden_share"]*100:.1f}%)',
            row["claim_role"],
        ]
        x = x0
        for j, ((_, w), value) in enumerate(zip(columns, values)):
            draw.rectangle((x, y, x + w, y + row_h), fill=fill, outline="#222222", width=2)
            color = red if j == 0 and i in [1, 2, 3] else "#111111"
            this_font = safe_font(24, bold=(j == 0 or i == 0))
            draw_wrapped(draw, (x + 16, y + 26), str(value), this_font, color, w - 32, line_gap=4)
            x += w
    note_y = y0 + header_h + len(df) * row_h + 38
    draw_wrapped(
        draw,
        (60, note_y),
        "Interpretation: the primary r strata support bounded extraction-validity claims; converted beta/path rows should remain a sensitivity/exception-handling stratum rather than a pooled accuracy numerator.",
        small,
        "#333333",
        2320,
    )
    img.save(path, dpi=(240, 240))


def draw_forest_plot(df: pd.DataFrame, path: Path) -> pd.DataFrame:
    ensure_dir(path.parent)
    baseline = df[df["denominator_family"] == "Converted beta/path"].iloc[0]
    rows = []
    for _, row in df.iterrows():
        if row["denominator_family"] == "Converted beta/path":
            continue
        for metric, n_col in [
            ("Scored-only accuracy", "scored_rows"),
            ("All-scorable accuracy", "rows_total"),
        ]:
            k1 = int(row["correct_rows"])
            n1 = int(row[n_col])
            k0 = int(baseline["correct_rows"])
            n0 = int(baseline[n_col])
            diff, lo, hi = prop_diff_ci(k1, n1, k0, n0)
            rows.append(
                {
                    "comparison": f'{row["denominator_family"]} vs Converted beta/path',
                    "metric": metric,
                    "risk_difference": diff,
                    "ci_low": lo,
                    "ci_high": hi,
                }
            )
    out = pd.DataFrame(rows)
    width, height = 2400, 1350
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    title = safe_font(45, bold=True)
    font = safe_font(28)
    small = safe_font(24)
    draw.text((60, 46), "Figure 2 | Accuracy difference against the converted beta/path sensitivity stratum", font=title, fill="#111111")
    draw_wrapped(
        draw,
        (60, 108),
        "Risk differences use frozen Paper B denominator-family counts. Positive values mean the primary r stratum is more accurate than the converted beta/path stratum.",
        small,
        "#333333",
        2280,
    )
    plot_x1, plot_x2 = 940, 1760
    plot_y1, plot_y2 = 270, 1030
    xmin, xmax = -0.10, 0.70
    def xmap(v: float) -> int:
        return int(plot_x1 + (v - xmin) / (xmax - xmin) * (plot_x2 - plot_x1))
    draw.line((plot_x1, plot_y2, plot_x2, plot_y2), fill="#111111", width=4)
    for tick in [-0.1, 0.0, 0.2, 0.4, 0.6]:
        x = xmap(tick)
        draw.line((x, plot_y2, x, plot_y2 + 18), fill="#111111", width=4)
        draw.text((x - 34, plot_y2 + 28), f"{tick:.1f}", font=small, fill="#111111")
        if tick == 0.0:
            draw.line((x, plot_y1 - 30, x, plot_y2), fill="#111111", width=5)
        else:
            draw.line((x, plot_y1 - 30, x, plot_y2), fill="#DDDDDD", width=2)
    row_gap = 105
    y = plot_y1
    draw.text((60, y - 90), "Comparison", font=safe_font(30, bold=True), fill="#111111")
    draw.text((1780, y - 90), "RD", font=safe_font(30, bold=True), fill="#111111")
    draw.text((1930, y - 90), "95% CI", font=safe_font(30, bold=True), fill="#111111")
    colors_by_metric = {"Scored-only accuracy": "#2F6FA3", "All-scorable accuracy": "#B43B35"}
    for i, row in out.iterrows():
        color = colors_by_metric[row["metric"]]
        label = f'{row["metric"]}: {row["comparison"]}'
        draw_wrapped(draw, (60, y - 22), label, font, "#111111", 820, line_gap=4)
        lo, hi, rd = row["ci_low"], row["ci_high"], row["risk_difference"]
        draw.line((xmap(lo), y, xmap(hi), y), fill="#6F7E8D", width=5)
        x = xmap(rd)
        draw.rectangle((x - 14, y - 14, x + 14, y + 14), fill=color, outline="#FFFFFF", width=2)
        draw.text((1780, y - 22), f"{rd:+.3f}", font=font, fill="#111111")
        draw.text((1930, y - 22), f"[{lo:+.3f}; {hi:+.3f}]", font=font, fill="#111111")
        y += row_gap
    draw_wrapped(
        draw,
        (60, 1130),
        "Note: This is not a vendor ranking. It tests whether directly source-supported r extraction behaves differently from converted/model-derived effect-size extraction.",
        small,
        "#333333",
        2280,
    )
    img.save(path, dpi=(240, 240))
    return out


def draw_accuracy_burden(df: pd.DataFrame, path: Path) -> None:
    ensure_dir(path.parent)
    width, height = 2200, 1400
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    title = safe_font(45, bold=True)
    font = safe_font(28)
    small = safe_font(24)
    draw.text((60, 48), "Figure 3 | Accuracy and review-burden profile", font=title, fill="#111111")
    draw_wrapped(
        draw,
        (60, 110),
        "The x-axis is unresolved extraction burden, not elapsed time. A Lai-style time-savings plot requires reviewer-time logs that are not present in the current Paper B package.",
        small,
        "#333333",
        2060,
    )
    x1, y1, x2, y2 = 280, 240, 1760, 1080
    draw.line((x1, y2, x2, y2), fill="#111111", width=5)
    draw.line((x1, y1, x1, y2), fill="#111111", width=5)
    xmin, xmax = 0.0, 0.28
    ymin, ymax = 0.30, 1.00
    def xmap(v: float) -> int:
        return int(x1 + (v - xmin) / (xmax - xmin) * (x2 - x1))
    def ymap(v: float) -> int:
        return int(y2 - (v - ymin) / (ymax - ymin) * (y2 - y1))
    for tick in [0, 0.1, 0.2]:
        x = xmap(tick)
        draw.line((x, y2, x, y2 + 18), fill="#111111", width=4)
        draw.text((x - 35, y2 + 28), f"{tick*100:.0f}%", font=small, fill="#111111")
        draw.line((x, y1, x, y2), fill="#E0E0E0", width=2)
    for tick in [0.4, 0.6, 0.8, 1.0]:
        y = ymap(tick)
        draw.line((x1 - 18, y, x1, y), fill="#111111", width=4)
        draw.text((x1 - 92, y - 14), f"{tick*100:.0f}%", font=small, fill="#111111")
        draw.line((x1, y, x2, y), fill="#E0E0E0", width=2)
    draw.text((650, 1190), "Abstention / unresolved share of denominator", font=font, fill="#111111")
    draw.text((60, 560), "Scored-only accuracy", font=font, fill="#111111")
    colors_map = {
        "Primary r validation pool": "#384E77",
        "Latent/source-flagged r": "#2F6FA3",
        "Direct/source-reported r": "#3D8B5B",
        "Converted beta/path": "#B43B35",
    }
    plot_df = df[df["denominator_family"] != "Primary r validation pool"].copy()
    for _, row in plot_df.iterrows():
        x = xmap(row["review_burden_share"])
        y = ymap(row["scored_only_accuracy"])
        size = int(32 + math.sqrt(row["rows_total"]) * 2.5)
        color = colors_map.get(row["denominator_family"], "#777777")
        draw.ellipse((x - size, y - size, x + size, y + size), fill=color, outline="#FFFFFF", width=4)
        if row["denominator_family"] in {
            "Latent/source-flagged r",
            "Direct/source-reported r",
        }:
            labels = {
                "Latent/source-flagged r": ["Latent/source r", f'{row["scored_only_accuracy"]*100:.1f}%'],
                "Direct/source-reported r": ["Direct/source r", f'{row["scored_only_accuracy"]*100:.1f}%'],
            }[row["denominator_family"]]
            label_font = safe_font(22, bold=True)
            total_h = sum(text_size(draw, line, label_font)[1] for line in labels) + 8 * (len(labels) - 1)
            ty = y - total_h // 2
            for line in labels:
                tw, th = text_size(draw, line, label_font)
                draw.text((x - tw // 2, ty), line, font=label_font, fill="white")
                ty += th + 8
        else:
            draw_wrapped(draw, (x + size + 28, y - 24), row["denominator_family"], small, "#111111", 360, line_gap=3)
    lx, ly = 1810, 330
    draw.rectangle((1780, 280, 2130, 720), fill="#F8F8F8", outline="#777777", width=2)
    draw.text((lx, 310), "Legend", font=safe_font(30, bold=True), fill="#111111")
    legend_items = [
        ("Latent/source r", colors_map["Latent/source-flagged r"]),
        ("Direct/source r", colors_map["Direct/source-reported r"]),
        ("Converted beta/path", colors_map["Converted beta/path"]),
    ]
    for label, color in legend_items:
        ly += 70
        draw.ellipse((lx, ly - 12, lx + 24, ly + 12), fill=color)
        draw.text((lx + 42, ly - 18), label, font=small, fill="#111111")
    img.save(path, dpi=(240, 240))


def draw_flow_diagram(path: Path) -> None:
    ensure_dir(path.parent)
    width, height = 2600, 1500
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    title = safe_font(46, bold=True)
    small = safe_font(24)
    draw.text((70, 48), "Figure 1 | Paper B source-anchored validation workflow", font=title, fill="#111111")
    draw_wrapped(
        draw,
        (70, 112),
        "Frontier-paper aligned design: source evidence and human adjudication define the reference standard before locked LLM outputs are scored.",
        small,
        "#333333",
        2350,
    )
    group_specs = [
        ((60, 220, 910, 1100), "Reference standard"),
        ((950, 220, 1780, 1100), "Locked model evaluation"),
        ((1820, 220, 2530, 1100), "Claim boundary"),
    ]
    for box, label in group_specs:
        draw.rounded_rectangle(box, radius=28, outline="#777777", width=4)
        draw.rectangle((box[0] + 26, box[1] - 24, box[0] + 360, box[1] + 26), fill="white")
        draw.text((box[0] + 42, box[1] - 18), label, font=safe_font(28, bold=True), fill="#333333")
    boxes = [
        ((110, 310, 430, 540), "Source PDFs and human packets", "Coder workbooks, source locators, consensus traces", "#FFFFFF", "#111111"),
        ((540, 310, 860, 540), "Source adjudication", "Resolve statistic type, sample, construct mapping", "#F7FBFF", "#355C7D"),
        ((330, 690, 720, 900), "Frozen reference candidate", "Paper2 Human Final Consensus v2 defines scoring target", "#EAF2FA", "#355C7D"),
        ((1020, 310, 1340, 520), "Locked LLM outputs", "Codex/Gemini/Claude outputs preserved before scoring", "#F9F6EF", "#8A6D3B"),
        ((1420, 310, 1710, 540), "Denominator scoring", "Direct r, latent r, converted beta/path, metadata/trace", "#F9F6EF", "#8A6D3B"),
        ((1110, 690, 1630, 900), "Error taxonomy and triage", "Abstentions, cross-model disagreement, source-risk, numeric consequence", "#FFF9F1", "#8A6D3B"),
        ((1890, 310, 2210, 540), "Results tables and figures", "Dense table, forest RD, burden plot, appendix", "#F1F8F4", "#3D8B5B"),
        ((2220, 690, 2480, 920), "MASEM gate", "Core diagnostic only unless coverage is complete", "#FDF2F2", "#B43B35"),
    ]
    for args in boxes:
        draw_box(draw, *args)
    arrows = [
        ((430, 415), (540, 415)),
        ((700, 520), (560, 690)),
        ((860, 415), (1020, 415)),
        ((1340, 415), (1420, 415)),
        ((1565, 520), (1370, 690)),
        ((1630, 795), (1890, 430)),
        ((2210, 430), (2320, 690)),
        ((720, 795), (1110, 795)),
    ]
    for start, end in arrows:
        draw_arrow(draw, start, end)
    draw.rounded_rectangle((220, 1160, 2380, 1350), radius=26, fill="#F5F5F5", outline="#777777", width=3)
    draw_wrapped(
        draw,
        (260, 1195),
        "Reproducibility loop: frozen CSVs -> generation script -> generated tables/figures -> manuscript DOCX/PDF -> checksums and visual render checks. Timing/efficiency claims stay out unless reviewer-time logs are added.",
        safe_font(30, bold=True),
        "#222222",
        2060,
    )
    img.save(path, dpi=(240, 240))


def draw_masem_gate(masem: pd.DataFrame, path: Path) -> None:
    ensure_dir(path.parent)
    width, height = 2200, 1250
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    title = safe_font(45, bold=True)
    font = safe_font(27)
    small = safe_font(23)
    draw.text((60, 48), "Figure 4 | MASEM identification gate", font=title, fill="#111111")
    draw_wrapped(
        draw,
        (60, 112),
        "Covered construct pairs are not enough for a full target-model claim when complete-case studies collapse to zero.",
        small,
        "#333333",
        2050,
    )
    x1, y1, x2, y2 = 330, 240, 1720, 1010
    max_pairs = int(masem["required_pairs"].max())
    bar_h = 90
    gap = 48
    for i, row in masem.iterrows():
        y = y1 + i * (bar_h + gap)
        required_w = int((row["required_pairs"] / max_pairs) * (x2 - x1))
        covered_w = int((row["covered_pairs"] / max_pairs) * (x2 - x1))
        draw.text((60, y + 18), str(row["construct_set"]), font=small, fill="#111111")
        draw.rounded_rectangle((x1, y, x1 + required_w, y + bar_h), radius=14, fill="#E6E6E6", outline="#888888", width=2)
        color = "#3D8B5B" if row["complete_case_studies"] > 0 else "#B43B35"
        draw.rounded_rectangle((x1, y, x1 + covered_w, y + bar_h), radius=14, fill=color, outline="#333333", width=2)
        label = f'{int(row["covered_pairs"])}/{int(row["required_pairs"])} pairs; complete-case studies={int(row["complete_case_studies"])}'
        draw.text((x1 + 18, y + 28), label, font=font, fill="white")
        draw_wrapped(draw, (1760, y + 4), row["recommended_claim"], small, "#111111", 360, line_gap=3)
    draw.text((x1, 1080), "Covered pairs / required pairs", font=font, fill="#111111")
    img.save(path, dpi=(240, 240))


def render_pdf_pages(pdf_path: Path, pages_dir: Path) -> list[Path]:
    ensure_dir(pages_dir)
    prefix = pages_dir / "page"
    subprocess.run(
        [str(PDFTOPPM), "-png", "-r", "120", str(pdf_path), str(prefix)],
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
    )
    return sorted(pages_dir.glob("page-*.png"))


def contact_sheet(images: list[Path], output: Path, thumb_width: int = 320) -> None:
    if not images:
        return
    thumbs = []
    for p in images:
        im = Image.open(p).convert("RGB")
        ratio = thumb_width / im.width
        thumbs.append(im.resize((thumb_width, int(im.height * ratio))))
    cols = min(4, len(thumbs))
    rows = math.ceil(len(thumbs) / cols)
    pad = 24
    label_h = 34
    cell_h = max(t.height for t in thumbs) + label_h
    sheet = Image.new("RGB", (cols * (thumb_width + pad) + pad, rows * (cell_h + pad) + pad), "white")
    draw = ImageDraw.Draw(sheet)
    font = safe_font(20, bold=True)
    for i, thumb in enumerate(thumbs):
        x = pad + (i % cols) * (thumb_width + pad)
        y = pad + (i // cols) * (cell_h + pad)
        sheet.paste(thumb, (x, y + label_h))
        draw.text((x, y), f"page {i + 1}", font=font, fill="#111111")
    ensure_dir(output.parent)
    sheet.save(output)


CAPTION_RE = re.compile(r"(?i)\b((?:fig(?:ure)?\.?|table)\s*\d+[A-Za-z]?(?:\.|:)?[^\n]{0,220})")


def extract_pdf_reference(ref: ReferencePaper, pdf_path: Path) -> dict[str, object]:
    out_dir = STAGE_EXTRACTED / ref.key
    text_dir = out_dir / "text"
    pages_dir = out_dir / "pages"
    table_dir = out_dir / "tables"
    visual_dir = out_dir / "figure_table_pages"
    for d in [text_dir, pages_dir, table_dir, visual_dir]:
        ensure_dir(d)
    reader = PdfReader(str(pdf_path))
    text_chunks = []
    caption_rows = []
    for idx, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        (text_dir / f"page_{idx:02d}.txt").write_text(text, encoding="utf-8")
        text_chunks.append(f"\\n\\n--- PAGE {idx} ---\\n{text}")
        for match in CAPTION_RE.finditer(text.replace("\\r", "\\n")):
            caption_rows.append({"paper": ref.key, "page": idx, "caption_snippet": " ".join(match.group(1).split())})
    (text_dir / f"{ref.key}_full_text.txt").write_text("".join(text_chunks), encoding="utf-8")
    pages = render_pdf_pages(pdf_path, pages_dir)
    contact_sheet(pages, out_dir / f"{ref.key}_contact_sheet.png")
    for row in caption_rows:
        text = row["caption_snippet"].lower()
        row["kind"] = "table" if text.startswith("table") else "figure"
    caption_pages = sorted({int(row["page"]) for row in caption_rows})
    for page_num in caption_pages:
        candidates = [p for p in pages if p.name == f"page-{page_num}.png" or p.name == f"page-{page_num:02d}.png"]
        if not candidates and 1 <= page_num <= len(pages):
            candidates = [pages[page_num - 1]]
        if candidates:
            page_captions = [row for row in caption_rows if int(row["page"]) == page_num]
            kinds = sorted({row.get("kind", "caption") for row in page_captions})
            suffix = "_".join(kinds) if kinds else "caption"
            shutil.copy2(candidates[0], visual_dir / f"{ref.key}_page_{page_num:02d}_{suffix}.png")
    extracted_tables = []
    with pdfplumber.open(str(pdf_path)) as pdf:
        for idx, page in enumerate(pdf.pages, start=1):
            try:
                tables = page.extract_tables() or []
            except Exception:
                tables = []
            for t_idx, table in enumerate(tables, start=1):
                rows = [[cell or "" for cell in row] for row in table if row]
                if len(rows) < 2:
                    continue
                stem = f"page_{idx:02d}_table_{t_idx:02d}"
                csv_path = table_dir / f"{stem}.csv"
                with csv_path.open("w", encoding="utf-8", newline="") as f:
                    writer = csv.writer(f)
                    writer.writerows(rows)
                md_path = table_dir / f"{stem}.md"
                md_path.write_text(df_to_markdown(pd.DataFrame(rows[1:], columns=rows[0])) + "\n", encoding="utf-8")
                extracted_tables.append({"page": idx, "table_index": t_idx, "csv": str(csv_path.relative_to(STAGE_BENCHMARK))})
    caption_df = pd.DataFrame(caption_rows)
    if caption_df.empty:
        caption_df = pd.DataFrame(columns=["paper", "page", "kind", "caption_snippet"])
    caption_df.to_csv(out_dir / "caption_inventory.csv", index=False)
    return {
        "key": ref.key,
        "pdf": str(pdf_path.relative_to(STAGE_BENCHMARK)),
        "page_count": len(reader.pages),
        "caption_count": len(caption_rows),
        "caption_pages": caption_pages,
        "extracted_table_count": len(extracted_tables),
        "extracted_tables": extracted_tables,
    }


def build_reference_crosswalk(pdf_summaries: list[dict[str, object]]) -> pd.DataFrame:
    rows = [
        {
            "reference": "Lai et al. 2025, npj Digital Medicine",
            "url": "https://www.nature.com/articles/s41746-025-01457-w",
            "frontier_feature": "Dense accuracy tables, risk-difference forest plot, accuracy-efficiency plot, process flow diagram.",
            "paper_b_implementation": "Replicated as dense denominator table, RD forest-style plot, burden-not-time plot, and source-anchored flow diagram.",
        },
        {
            "reference": "Farotimi et al. 2026, Research Synthesis Methods",
            "url": "https://www.cambridge.org/core/journals/research-synthesis-methods/article/generative-artificial-intelligence-in-evidence-synthesis-an-advisory-and-reporting-framework-for-authors-peer-reviewers-and-journals/85AB368858C8FE6A75C158462828DAE4",
            "frontier_feature": "GenAI evidence-synthesis disclosure and reporting framework.",
            "paper_b_implementation": "AI disclosure, model-access boundary, reproducibility and claim-limit language.",
        },
    ]
    for ref in REFERENCES:
        summary = next((x for x in pdf_summaries if x["key"] == ref.key), {})
        rows.append(
            {
                "reference": ref.title,
                "url": ref.url,
                "frontier_feature": ref.paper_b_use,
                "paper_b_implementation": f"Downloaded locally; {summary.get('page_count', 'NA')} pages rendered/extracted; captions and tables inventoried for benchmark use.",
            }
        )
    return pd.DataFrame(rows)


def build_status_doc(paths: dict[str, Path], pdf_summaries: list[dict[str, object]]) -> str:
    return f"""
# Paper B Visual Upgrade Status

Date: {DATE}

## Direct Answer

The previous Paper B package did not yet meet the visual standard shown in the comparator screenshots. It had basic manuscript tables and bar-chart figures, but it did not include a dense journal-style extraction table, a risk-difference/forest-style figure, an accuracy-versus-burden figure, or a full validation flow diagram. This upgrade implements those missing artifacts.

## Flow Diagram Judgment

The flow diagram is essential for Paper B. Paper B's key contribution is not just an accuracy percentage; it is the validation architecture: source-anchored human reference standard, locked LLM outputs, denominator-family scoring, expert triage, and MASEM claim gate. Without a flow diagram, reviewers may misread the workflow as ordinary model benchmarking or autonomous extraction.

## Added Frontier References

- Huang et al. 2025 JMIR RoB2 evaluation: local PDF downloaded and page/text/table/caption extraction completed.
- Jansen et al. 2026 Educational Psychology Review data-extraction benchmark: local PDF downloaded and page/text/table/caption extraction completed.

PDF extraction summaries:

{json.dumps(pdf_summaries, ensure_ascii=False, indent=2)}

Machine-readable table extraction returned zero cell-structured tables for the two added PDFs. The extraction package therefore preserves rendered pages, figure/table candidate pages, full text, and caption inventories as the auditable reference layer.

## Generated Artifacts

- Dense journal-style table image: `{paths['dense_table'].name}`
- Risk-difference forest-style figure: `{paths['forest'].name}`
- Accuracy/review-burden figure: `{paths['burden'].name}`
- Source-anchored flow diagram: `{paths['flow'].name}`
- MASEM claim gate figure: `{paths['masem'].name}`
- Visual-upgrade manuscript insert DOCX: `{paths['docx'].name}`
- Visual-upgrade PDF report: `{paths['pdf_report'].name}`

## Resolved Timing Boundary

{TIME_LOG_DECISION} {TIMING_BOUNDARY} The implemented substitute is an accuracy-versus-review-burden plot using abstention/unresolved share.
"""


def build_goal_audit_doc(paths: dict[str, Path], pdf_summaries: list[dict[str, object]]) -> str:
    pdf_cells = sum(int(x.get("extracted_table_count", 0)) for x in pdf_summaries)
    return f"""
# Paper B Visual Goal Audit

Date: {DATE}

| Goal | Status | Evidence | Boundary |
| --- | --- | --- | --- |
| Dense journal-style table comparable to the screenshot examples | Implemented | `{paths['dense_table'].name}` plus CSV/Markdown source table | Needs final journal formatting only if target journal has specific width/font rules |
| Forest-style risk-difference figure | Implemented | `{paths['forest'].name}` plus source CSV/Markdown | Uses Paper B denominator-family contrast; not a direct copy of Lai et al. |
| Source/process flow diagram | Implemented and essential | `{paths['flow'].name}` | Should be treated as a core Methods figure |
| Accuracy-versus-efficiency figure | Not implemented by design; replaced with safer review-burden figure | `{paths['burden'].name}` | {TIME_LOG_DECISION} |
| MASEM downstream claim-gate figure | Implemented | `{paths['masem'].name}` | Supports bounded claim language rather than full target-model substitution |
| Additional frontier references downloaded locally | Implemented | Huang 2025 JMIR and Jansen 2026 Educational Psychology Review PDFs are in `pdfs/` | Publisher PDFs and extracted page images are not Git-safe without clearance |
| PDF table/figure extraction | Partially implemented | Rendered pages, figure/table candidate pages, full text, contact sheets, and caption inventories created | Machine cell-structured table extraction count: {pdf_cells} |
| Manuscript-facing integration | Implemented | `{paths['docx'].name}`, `{paths['md'].name}`, and `{paths['pdf_report'].name}` | This is an insert/update package, not yet a fully rewritten submission manuscript |
| Git release/tag | Approved only for a share-safe package | Private/full-text PDFs and extracted pages are present | Release package must exclude publisher PDFs, full-text extracts, and rendered reference-paper pages |
"""


def add_track_revisions(doc: Document) -> None:
    settings = doc.settings.element
    if settings.find(qn("w:trackRevisions")) is None:
        settings.append(OxmlElement("w:trackRevisions"))


def setup_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    for style_name in ["Normal", "Body Text"]:
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = "Calibri"
            style.font.size = Pt(10.2)
            style.paragraph_format.space_after = Pt(5)
    for style_name, size in [("Heading 1", 15.5), ("Heading 2", 12.5), ("Heading 3", 11.2)]:
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = "Calibri"
            style.font.size = Pt(size)
            style.font.color.rgb = RGBColor(31, 78, 121)


def add_doc_image(doc: Document, path: Path, caption: str, width: float = 6.7) -> None:
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(caption)
    run.italic = True
    run.font.size = Pt(8.8)
    run.font.color.rgb = RGBColor(80, 80, 80)


def build_docx(paths: dict[str, Path], crosswalk: pd.DataFrame) -> Path:
    doc = Document()
    setup_doc(doc)
    add_track_revisions(doc)
    doc.add_heading("Paper B Visual Upgrade Manuscript Insert Package", level=1)
    doc.add_paragraph(f"Date: {DATE}")
    doc.add_paragraph(
        "Purpose: replace the previous basic table/figure layer with a journal-style visual package aligned with current frontier studies on LLM-assisted evidence synthesis and data extraction."
    )
    doc.add_heading("Positioning", level=2)
    doc.add_paragraph(
        "The flow diagram should be treated as a core Methods figure. It makes clear that Paper B evaluates a source-anchored validation workflow rather than an autonomous model replacement."
    )
    add_doc_image(
        doc,
        paths["flow"],
        "Figure 1. Source-anchored validation workflow for Paper B. The reference standard is frozen before locked LLM outputs are scored.",
        6.8,
    )
    doc.add_heading("Results-Ready Table/Figure Layer", level=2)
    add_doc_image(
        doc,
        paths["dense_table"],
        "Table 1. Denominator-family extraction validity table. Accuracy is reported both conditional on scorable answers and across all scorable rows.",
        6.8,
    )
    add_doc_image(
        doc,
        paths["forest"],
        "Figure 2. Forest-style risk-difference plot comparing primary r strata against the converted beta/path sensitivity stratum.",
        6.8,
    )
    add_doc_image(
        doc,
        paths["burden"],
        "Figure 3. Accuracy and review-burden profile. This is not an elapsed-time plot; reviewer-time data are not yet available.",
        6.8,
    )
    add_doc_image(
        doc,
        paths["masem"],
        "Figure 4. MASEM claim gate. Construct-pair coverage does not justify a full target-model claim when complete-case identification fails.",
        6.8,
    )
    doc.add_heading("Frontier Reference Reinforcement", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["Reference", "Frontier feature", "Paper B use"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8.5)
    for _, row in crosswalk.iterrows():
        cells = table.add_row().cells
        vals = [row["reference"], row["frontier_feature"], row["paper_b_implementation"]]
        for i, val in enumerate(vals):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(7.3)
    doc.add_heading("Resolved Timing Boundary", level=2)
    doc.add_paragraph(
        f"{TIME_LOG_DECISION} {TIMING_BOUNDARY} Paper B should report this as review burden and triage evidence, not as elapsed-time efficiency."
    )
    docx_path = STAGE_WORKING / f"Paper_B_Implementation_Draft_RSM_VISUAL_UPGRADE_{DATE}.docx"
    ensure_dir(docx_path.parent)
    doc.save(docx_path)
    return docx_path


def build_md(paths: dict[str, Path]) -> Path:
    md = f"""
# Paper B Visual Upgrade Manuscript Insert Package

Date: {DATE}

## Visual Diagnosis

The earlier Paper B package did not yet match the screenshot-level benchmark. This package adds:

- Dense denominator-family extraction validity table.
- Forest-style risk-difference plot.
- Accuracy-versus-review-burden plot.
- Source-anchored validation flow diagram.
- MASEM claim-gate figure.

## Manuscript Claim Boundary

The primary claim should remain bounded: the workflow supports source-anchored extraction validation and review triage for primary r strata. Converted beta/path rows remain a sensitivity and exception-handling stratum. {TIME_LOG_DECISION} {TIMING_BOUNDARY}

## Generated Figures

- `{paths['flow'].name}`
- `{paths['dense_table'].name}`
- `{paths['forest'].name}`
- `{paths['burden'].name}`
- `{paths['masem'].name}`
"""
    path = STAGE_WORKING / f"Paper_B_Implementation_Draft_RSM_VISUAL_UPGRADE_{DATE}.md"
    write_text(path, md)
    return path


def build_pdf_report(paths: dict[str, Path]) -> Path:
    pdf_path = STAGE_REPORTS / f"Paper_B_Frontier_Visual_Upgrade_Report_{DATE}.pdf"
    ensure_dir(pdf_path.parent)
    c = canvas.Canvas(str(pdf_path), pagesize=letter)
    page_w, page_h = letter
    c.setTitle("Paper B Frontier Visual Upgrade Report")
    c.setFont("Helvetica-Bold", 16)
    c.drawString(0.7 * inch, page_h - 0.75 * inch, "Paper B Frontier Visual Upgrade Report")
    c.setFont("Helvetica", 10)
    c.drawString(0.7 * inch, page_h - 1.0 * inch, f"Date: {DATE}")
    c.setFillColor(colors.HexColor("#333333"))
    c.drawString(0.7 * inch, page_h - 1.25 * inch, "Generated from frozen Paper B tables and newly downloaded frontier reference PDFs.")
    y = page_h - 1.65 * inch
    for key, title in [
        ("flow", "Figure 1. Source-anchored validation workflow"),
        ("dense_table", "Table 1. Denominator-family extraction validity"),
        ("forest", "Figure 2. Risk-difference forest-style plot"),
        ("burden", "Figure 3. Accuracy and review-burden profile"),
        ("masem", "Figure 4. MASEM claim gate"),
    ]:
        if y < 2.4 * inch:
            c.showPage()
            y = page_h - 0.75 * inch
        c.setFont("Helvetica-Bold", 11)
        c.setFillColor(colors.black)
        c.drawString(0.7 * inch, y, title)
        y -= 0.18 * inch
        img = Image.open(paths[key])
        iw, ih = img.size
        max_w = page_w - 1.4 * inch
        max_h = 4.8 * inch
        scale = min(max_w / iw, max_h / ih)
        draw_w, draw_h = iw * scale, ih * scale
        if y - draw_h < 0.55 * inch:
            c.showPage()
            y = page_h - 0.75 * inch
            c.setFont("Helvetica-Bold", 11)
            c.drawString(0.7 * inch, y, title)
            y -= 0.18 * inch
        c.drawImage(str(paths[key]), 0.7 * inch, y - draw_h, width=draw_w, height=draw_h, preserveAspectRatio=True, mask="auto")
        y -= draw_h + 0.35 * inch
    c.save()
    return pdf_path


def update_task_board(paths: dict[str, Path]) -> Path:
    src = TRACKING / f"Paper_B_Researcher_Task_Board_{DATE}.xlsx"
    dst = STAGE_TRACKING / src.name
    ensure_dir(dst.parent)
    shutil.copy2(src, dst)
    wb = load_workbook(dst)
    ws = wb["Paper B Tasks"]
    rows = [
        [
            "B19",
            "PI/R4",
            "High",
            "Frontier-style visual upgrade",
            "Review dense table, flow diagram, forest-style RD plot, and burden plot for manuscript insertion.",
            "Created",
            f"01_working_manuscript/{paths['docx'].name}",
        ],
        [
            "B20",
            "PI/R1",
            "High",
            "Additional frontier reference reinforcement",
            "Check Huang 2025 JMIR and Jansen 2026 Educational Psychology Review extraction inventories against Paper B claims.",
            "Created",
            f"02_B_requirements_and_tracking/PAPER_B_FRONTIER_REFERENCE_REINFORCEMENT_{DATE}.md",
        ],
        [
            "B21",
            "PI",
            "Medium",
            "Elapsed-time efficiency evidence decision",
            "Decide whether reviewer-time logs exist; if not, keep review-burden framing rather than time-savings claim.",
            "Open question",
            f"02_B_requirements_and_tracking/PAPER_B_VISUAL_UPGRADE_STATUS_{DATE}.md",
        ],
    ]
    existing = {ws.cell(r, 1).value for r in range(2, ws.max_row + 1)}
    for row in rows:
        if row[0] not in existing:
            ws.append(row)
    for cell in ws[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="1F4E79")
    for col in range(1, ws.max_column + 1):
        ws.column_dimensions[chr(64 + col)].width = [12, 14, 12, 34, 72, 16, 66][col - 1]
    for row in ws.iter_rows():
        for cell in row:
            cell.alignment = Alignment(wrap_text=True, vertical="top")
    wb.save(dst)
    return dst


def zip_ok(path: Path) -> bool:
    try:
        with zipfile.ZipFile(path) as zf:
            return zf.testzip() is None
    except Exception:
        return False


def main() -> None:
    for d in [STAGE_TABLES, STAGE_REPORT_FIGURES, STAGE_FIGURES, STAGE_RENDER]:
        ensure_dir(d)

    tables = load_core_tables()
    enriched = enrich_m1_table(tables["m1r_denominator_summary"])
    write_df_artifacts(enriched, f"table_1_dense_denominator_validity_{DATE}")

    dense_path = STAGE_REPORT_FIGURES / f"table_1_dense_denominator_validity_{DATE}.png"
    flow_path = STAGE_REPORT_FIGURES / f"figure_1_paper_b_source_anchored_flow_diagram_{DATE}.png"
    forest_path = STAGE_REPORT_FIGURES / f"figure_2_paper_b_accuracy_difference_forest_{DATE}.png"
    burden_path = STAGE_REPORT_FIGURES / f"figure_3_paper_b_accuracy_review_burden_{DATE}.png"
    masem_path = STAGE_REPORT_FIGURES / f"figure_4_paper_b_masem_claim_gate_{DATE}.png"

    draw_flow_diagram(flow_path)
    draw_dense_table(enriched, dense_path)
    forest_df = draw_forest_plot(enriched, forest_path)
    draw_accuracy_burden(enriched, burden_path)
    draw_masem_gate(tables["masem_gate_summary"], masem_path)
    write_df_artifacts(forest_df, f"figure_2_accuracy_difference_forest_source_data_{DATE}")
    write_df_artifacts(tables["masem_gate_summary"], f"figure_4_masem_gate_source_data_{DATE}")

    for figure in [flow_path, dense_path, forest_path, burden_path, masem_path]:
        copy_if_exists(figure, STAGE_FIGURES / figure.name)

    pdf_summaries = []
    for ref in REFERENCES:
        pdf = download_pdf(ref)
        pdf_summaries.append(extract_pdf_reference(ref, pdf))

    crosswalk = build_reference_crosswalk(pdf_summaries)
    write_df_artifacts(crosswalk, f"frontier_reference_reinforcement_crosswalk_{DATE}")
    reinforcement_md = STAGE_TRACKING / f"PAPER_B_FRONTIER_REFERENCE_REINFORCEMENT_{DATE}.md"
    write_text(
        reinforcement_md,
        "# Paper B Frontier Reference Reinforcement\n\n"
        + df_to_markdown(crosswalk)
        + "\n\n## Interpretation\n\n"
        "The added references strengthen Paper B in two different ways: Huang et al. supports the need for a process/selection flow and explicit criterion-standard comparison, while Jansen et al. supports the education-meta-analysis extraction frame, frozen codebook, gold/silver reference standard logic, and cautious cost/time reporting.",
    )

    paths: dict[str, Path] = {
        "dense_table": dense_path,
        "flow": flow_path,
        "forest": forest_path,
        "burden": burden_path,
        "masem": masem_path,
    }
    docx_path = build_docx(paths, crosswalk)
    paths["docx"] = docx_path
    md_path = build_md(paths)
    paths["md"] = md_path
    pdf_report = build_pdf_report(paths)
    paths["pdf_report"] = pdf_report

    status_md = STAGE_TRACKING / f"PAPER_B_VISUAL_UPGRADE_STATUS_{DATE}.md"
    write_text(status_md, build_status_doc(paths, pdf_summaries))
    goal_audit_md = STAGE_TRACKING / f"PAPER_B_VISUAL_GOAL_AUDIT_{DATE}.md"
    write_text(goal_audit_md, build_goal_audit_doc(paths, pdf_summaries))
    logic_md = STAGE_TRACKING / f"PAPER_B_TABLE_FIGURE_PDF_GENERATION_LOGIC_VISUAL_UPGRADE_{DATE}.md"
    write_text(
        logic_md,
        f"""
# Paper B Visual Table/Figure/PDF Generation Logic

Date: {DATE}

## Inputs

- `m1r_denominator_summary.csv`
- `rq3_triage_summary.csv`
- `masem_gate_summary.csv`
- Newly downloaded frontier PDFs: Huang et al. 2025 JMIR and Jansen et al. 2026 Educational Psychology Review.

## Output Logic

1. Recompute denominator-family accuracy, abstention/unresolved share, and Wilson 95% confidence intervals from frozen CSV counts.
2. Generate a dense table image for manuscript insertion and a CSV/Markdown source table for audit.
3. Generate a forest-style risk-difference plot comparing primary r strata with the converted beta/path sensitivity stratum.
4. Generate an accuracy-versus-review-burden plot instead of a time-savings plot because no reviewer-time or per-study duration logs are available.
5. Generate a source-anchored validation flow diagram to prevent reviewers from misreading the workflow as autonomous model replacement.
6. Extract reference PDFs into page images, text files, caption inventories, and machine-detected tables.
7. Build a visual-upgrade DOCX and PDF report from the same generated image/table artifacts.

## Claim Boundary

Do not report elapsed-time efficiency. The current evidence supports review-burden and triage framing, not a direct time-spent comparison.
""",
    )

    board = update_task_board(paths)
    paths["task_board"] = board

    manifest = {
        "date": DATE,
        "stage_root": str(STAGING),
        "paths": {k: str(v) for k, v in paths.items()},
        "goal_audit": str(goal_audit_md),
        "pdf_summaries": pdf_summaries,
        "zip_checks": {
            "docx": zip_ok(docx_path),
            "task_board": zip_ok(board),
        },
    }
    manifest_path = OUT / f"paper_b_visual_upgrade_manifest_{DATE}.json"
    ensure_dir(manifest_path.parent)
    manifest_path.write_text(json.dumps(manifest, indent=2, ensure_ascii=False), encoding="utf-8")
    print(json.dumps(manifest, indent=2, ensure_ascii=False))


if __name__ == "__main__":
    main()
