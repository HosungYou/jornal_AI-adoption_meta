#!/usr/bin/env python3
"""Build target-journal manuscript planning packets for Paper A and Paper B.

The generated artifacts are claim-bounded. They incorporate current journal
instructions, exemplar-PDF decomposition, local workflow status, and manuscript
handoff structure without inventing unavailable SEM/full-corpus results.
"""

from __future__ import annotations

import csv
import math
import os
import re
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPO = Path(__file__).resolve().parents[2]
OUT = REPO / "docs" / "07_manuscript_exemplars" / "20260612"
PAPER_A_OUT = REPO / "paper_a" / "manuscript" / "target_journal"
PAPER_B_OUT = REPO / "paper_b" / "manuscript" / "target_journal"
PAPER_B_RESULTS = REPO / "data" / "04_extraction" / "05_llm_masem_substitution" / "results"
PAPER_B_PACKET_DIR = (
    REPO
    / "data"
    / "04_extraction"
    / "07_paper_c_harness_benchmark"
    / "private"
    / "source_renderings_20260609_full_coverage"
    / "source_packets"
)
PAPER_B_SHELL = (
    REPO
    / "data"
    / "04_extraction"
    / "05_llm_masem_substitution"
    / "full_corpus_step5_task_unit_shell_20260609.csv"
)
PRIVATE_DOC_ROOT = os.environ.get("PRIVATE_AI_ADOPTION_DOCUMENTS_ROOT", "")
PAPER_A_PRIMARY = (
    Path(PRIVATE_DOC_ROOT)
    / "Meta/AI Adoption/Paper1_MASEM_Working_20260605/09_model_ready_tiered_freeze/"
    / "paper1_direct_r_primary_model_ready_tiered_freeze_20260605.csv"
    if PRIVATE_DOC_ROOT
    else Path("<PRIVATE_AI_ADOPTION_DOCUMENTS_ROOT>")
    / "Meta/AI Adoption/Paper1_MASEM_Working_20260605/09_model_ready_tiered_freeze/"
    / "paper1_direct_r_primary_model_ready_tiered_freeze_20260605.csv"
)


CONSTRUCTS = ["PE", "EE", "SI", "FC", "ATT", "SE", "TRU", "ANX", "BI", "UB"]


def ensure_dirs() -> None:
    for path in [OUT, PAPER_A_OUT, PAPER_B_OUT, PAPER_B_RESULTS]:
        path.mkdir(parents=True, exist_ok=True)


def read_csv(path: Path) -> list[dict[str, str]]:
    if not path.exists():
        return []
    with path.open(newline="", encoding="utf-8-sig") as handle:
        return list(csv.DictReader(handle))


def paper_a_readiness() -> dict[str, object]:
    rows = read_csv(PAPER_A_PRIMARY)
    pair_set = {
        tuple(sorted((CONSTRUCTS[i], CONSTRUCTS[j])))
        for i in range(len(CONSTRUCTS))
        for j in range(i + 1, len(CONSTRUCTS))
    }
    usable = []
    missing_n = 0
    bad_r = 0
    pair_counts: Counter[tuple[str, str]] = Counter()
    study_pairs: dict[str, set[tuple[str, str]]] = defaultdict(set)
    n_studies_with_numeric_n = set()
    for row in rows:
        c1 = row.get("construct_1", "").strip()
        c2 = row.get("construct_2", "").strip()
        if c1 not in CONSTRUCTS or c2 not in CONSTRUCTS or c1 == c2:
            continue
        pair = tuple(sorted((c1, c2)))
        if pair not in pair_set:
            continue
        try:
            rval = float(row.get("r_numeric", ""))
        except ValueError:
            continue
        if abs(rval) >= 1:
            bad_r += 1
            continue
        try:
            nval = float(row.get("sample_size_numeric", ""))
        except ValueError:
            nval = math.nan
        if not math.isfinite(nval) or nval <= 0:
            missing_n += 1
        else:
            n_studies_with_numeric_n.add(row.get("study_id", ""))
        usable.append(row)
        sid = row.get("study_id", "")
        study_pairs[sid].add(pair)
        pair_counts[pair] += 1

    least_pairs = sorted(pair_counts.items(), key=lambda item: (item[1], item[0]))[:12]
    return {
        "source_path": str(PAPER_A_PRIMARY),
        "rows": len(rows),
        "usable_rows": len(usable),
        "bad_abs_ge_1": bad_r,
        "missing_n_rows": missing_n,
        "studies": len(study_pairs),
        "numeric_n_studies": len(n_studies_with_numeric_n),
        "covered_pairs": len(pair_counts),
        "total_pairs": len(pair_set),
        "complete_10construct_studies": sum(1 for pairs in study_pairs.values() if len(pairs) == len(pair_set)),
        "studies_ge_15_pairs": sum(1 for pairs in study_pairs.values() if len(pairs) >= 15),
        "least_pairs": [("-".join(pair), count) for pair, count in least_pairs],
    }


def paper_b_readiness() -> dict[str, object]:
    rows = read_csv(PAPER_B_SHELL)
    family_counts = Counter(row.get("denominator_family", "") for row in rows)
    packet_count = len(list(PAPER_B_PACKET_DIR.glob("*.txt"))) if PAPER_B_PACKET_DIR.exists() else 0
    return {
        "shell_exists": PAPER_B_SHELL.exists(),
        "shell_rows": len(rows),
        "family_counts": family_counts,
        "packet_dir": str(PAPER_B_PACKET_DIR),
        "packet_count": packet_count,
    }


def md_table(headers: list[str], rows: list[list[object]]) -> str:
    lines = [
        "| " + " | ".join(headers) + " |",
        "| " + " | ".join(["---"] * len(headers)) + " |",
    ]
    for row in rows:
        lines.append("| " + " | ".join(str(cell) for cell in row) + " |")
    return "\n".join(lines)


def build_contents(a: dict[str, object], b: dict[str, object]) -> dict[Path, str]:
    shell_family_rows = [[k or "(blank)", v] for k, v in b["family_counts"].most_common()]
    least_pair_rows = [[pair, count] for pair, count in a["least_pairs"]]

    source_manifest = f"""# Exemplar and Guideline Manifest

Date: 2026-06-12

Scope: target-journal and exemplar-paper evidence used to reshape Paper A and Paper B into near-submission manuscript packages. Closed PDFs obtained through PSU institutional access may be used locally, but they are not redistributed in Git or OSF.

{md_table(
        ["Source", "Access state", "Use in current package"],
        [
            ["Computers & Education Guide for Authors", "Web verified", "Paper A target-journal package: abstract, keywords, highlights, editable tables, double-anonymized file separation"],
            ["Research Synthesis Methods author instructions", "Web verified", "Paper B target-journal package: 250-word abstract, RSM highlights, repository/data-availability expectations, GenAI evaluation transparency"],
            ["Scherer, Siddiq, and Tondeur (2019), Computers & Education", "Bibliographic record verified; closed PDF needed for exact table/figure replication if not locally available", "Primary Paper A exemplar for TAM/MASEM positioning in C&E"],
            ["Or (2024), OTESSA Journal", "PDF extracted locally, 26 pages", "Open OSMASEM/TAM education exemplar for methods, PRISMA/table/figure spine"],
            ["Schmidt et al. (2024), ALTARS/arXiv", "PDF extracted locally, 30/32 pages", "Paper B exemplar for manual evaluation, domain-specific data extraction accuracy, and caution against autonomous replacement"],
            ["Schroeder et al. (2025), arXiv AIDE/HIL", "PDF extracted locally, 12 pages", "Paper B exemplar for explicit versus derived variable families, human-in-the-loop workflow, and source panel/provenance framing"],
            ["Legate and Nimon F1000 protocol", "Local PDF extracted, 28 pages", "Paper B social-science methods gap and living-review/reporting workflow context"],
            ["RSM GenAI submission guidance", "Web/PMC/RSM metadata verified", "Paper B reporting benchmark for validation, prompt/model versioning, locked outputs, and transparent AI declaration"],
            ["Annals AI-assisted data-extraction SWAR", "Web metadata verified; publisher PDF blocked", "Comparator for workflow validation, time/error outcome reporting, and human verification framing"],
        ],
    )}

## Source Links

- Computers & Education guide: https://www.sciencedirect.com/journal/computers-and-education/publish/guide-for-authors
- Research Synthesis Methods instructions: https://www.cambridge.org/core/journals/research-synthesis-methods/information/author-instructions/preparing-your-materials
- Scherer et al. DOI: https://doi.org/10.1016/j.compedu.2018.09.009
- Or (2024) DOI: https://doi.org/10.18357/otessaj.2024.4.3.66
- Schmidt et al. arXiv: https://arxiv.org/abs/2405.14445
- Schroeder et al. arXiv: https://arxiv.org/abs/2501.11840
- RSM automated meta-analysis review: https://doi.org/10.1017/rsm.2025.10065
- RSM GenAI submission guidance: https://doi.org/10.1017/rsm.2025.10058
- Annals prospective cohort DOI: https://doi.org/10.7326/ANNALS-25-00739
- OSF repository for Paper B: https://osf.io/mkrgd/overview
"""

    decomposition = f"""# Exemplar Decomposition Matrix

Date: 2026-06-12

## Transferable Patterns

{md_table(
        ["Exemplar", "Method spine", "Table/Figure pattern", "How to use it here", "Boundary"],
        [
            ["Scherer et al. C&E TAM-MASEM", "Education technology adoption framed around TAM, MASEM, theory competition", "PRISMA, study-characteristics table, pooled-matrix/path tables, conceptual/path figures", "Use as Paper A target-journal model for C&E positioning and MASEM result architecture", "Need PSU PDF for exact visual replication before final submission"],
            ["Or 2024 OSMASEM", "One-step MASEM of TAM in education using R/metaSEM references", "PRISMA/selection flow, source-study table, model fit table, path/effect table, model diagrams", "Use as Paper A open-PDF reference for OSMASEM reporting and figure spine", "Open exemplar is useful but not the target C&E house style"],
            ["Schmidt et al. 2024", "Prompt development plus manual evaluation by study domain and extraction item", "Prompt table, extraction-variable table, accuracy-by-domain/item tables, cautionary narrative", "Use for Paper B task-family denominators and manual scoring explanation", "Do not import one overall accuracy denominator"],
            ["Schroeder et al. 2025", "Pilot-to-full human-in-the-loop extraction across explicit and derived variables", "Workflow/application figures, exact/accurate match tables, model result table", "Use for Paper B explicit versus derived task-family reporting and HIL figure", "Use as workflow comparator, not a source for replacement claims"],
            ["F1000 social-science protocol", "Living review of semi-automated extraction approaches in social sciences", "Workflow and reporting figures, protocol tables, agreement/reliability plan", "Use for Paper B social-science contribution and broader evidence-synthesis relevance", "Protocol status means use for rationale, not result benchmarking"],
            ["RSM GenAI guidance", "Reporting expectations for GenAI evaluations in SRMA", "Validation/methodology/transparency checklist logic", "Use to structure Paper B methods, AI declaration, model/prompt reproducibility, repository statement", "RSM expects clear model-version and use-case boundary"],
            ["Annals SWAR", "Prospective workflow validation with human verification", "Concordance, accuracy, time-on-task, error severity tables", "Use for Paper B discussion of workflow evaluation outcomes", "Publisher PDF needed for exact table/figure deconstruction"],
        ],
    )}

## Design Consequences

Paper A should look like a C&E empirical synthesis article: concise education-technology contribution, theory-forward introduction, transparent PRISMA and coding workflow, then pooled matrices, path models, moderators, and sensitivity diagnostics.

Paper B should look like an RSM methods article: it should foreground validation design, source-anchored reference construction, denominator-family scoring, locked model outputs, reproducible scripts, and claim boundaries. The central claim is workflow augmentation and review triage, not LLM replacement.
"""

    paper_a_map = f"""# Paper A Target-Journal Structure Map

Target journal: Computers & Education

Decision state: target journal, construct scope, figure spine, and table spine are approved. The actual TSSEM/OSMASEM/sensitivity estimate insertion is approved but cannot be treated as complete from the current input because the model-ready primary file lacks source-supported numeric sample size for most rows.

## Current Analysis Gate

{md_table(
        ["Check", "Current value", "Implication"],
        [
            ["Primary model-ready rows", a["rows"], "Available in the OneDrive Paper1 working folder"],
            ["Usable 10-construct rows after r checks", a["usable_rows"], "No r absolute value >= 1 blocker in the tiered primary file"],
            ["Rows missing numeric N", a["missing_n_rows"], "N-weighted TSSEM/MASEM cannot be claimed until sample size is reconciled or explicit exclusion rule is applied"],
            ["Studies represented", a["studies"], "Correlation evidence is spread across incomplete matrices"],
            ["Covered construct pairs", f'{a["covered_pairs"]}/{a["total_pairs"]}', "Coverage is broad but not complete"],
            ["Complete 10-construct matrices", a["complete_10construct_studies"], "Complete-case 10-construct TSSEM is not feasible as the primary route"],
            ["Studies with 15 or more construct pairs", a["studies_ge_15_pairs"], "FIML/TSSEM route may be possible after N reconciliation"],
        ],
    )}

Least-covered pairs:

{md_table(["Construct pair", "Rows"], least_pair_rows)}

## C&E Submission Components

- Main manuscript: double-anonymized Word file with title page separated at submission.
- Abstract: no more than 250 words.
- Keywords: 1 to 7.
- Highlights: 3 to 5 bullets, no more than 85 characters each, submitted separately.
- Tables: editable text, cited in order, with captions and notes.
- Figures: conceptual model, PRISMA, coverage heatmap, and path model as editable/high-resolution files.

## Proposed Highlights

- Synthesizes AI adoption evidence in higher education with MASEM.
- Integrates TAM/UTAUT predictors with trust and AI anxiety.
- Separates direct-r inputs from converted sensitivity evidence.
- Tests structural and moderator paths across ten constructs.
- Provides reproducible extraction and QC artifacts.

## Section Spine

1. Introduction: higher-education AI adoption problem, why structural synthesis is needed, why AI trust/anxiety expand TAM/UTAUT.
2. Theory and hypotheses: TAM/UTAUT, attitude mediation, self-efficacy, trust, anxiety, use behavior.
3. Method: search, screening, source adjudication, coding, 10-construct harmonization, TSSEM/OSMASEM, sensitivity.
4. Results: PRISMA, study characteristics, input coverage, Stage 1 pooled matrix, Stage 2 model paths, moderators, sensitivity.
5. Discussion: theory, institutional implications, limitations, reproducibility, future AI-adoption measurement.

## Table Spine

{md_table(
        ["Table", "Title", "Status"],
        [
            ["Table 1", "Construct harmonization and operational definitions", "Ready to draft"],
            ["Table 2", "PRISMA and study-characteristics profile", "Needs final inclusion lock"],
            ["Table 3", "Analysis-ready input sets and source-type rules", "Ready with current caveats"],
            ["Table 4", "Construct-pair coverage and missingness", "Ready as pre-model table"],
            ["Table 5", "Stage 1 pooled correlation matrix", "Needs TSSEM run after N gate"],
            ["Table 6", "Stage 2 structural paths and indirect effects", "Needs TSSEM run"],
            ["Table 7", "Moderator and sensitivity results", "Needs OSMASEM/sensitivity run"],
        ],
    )}

## Figure Spine

{md_table(
        ["Figure", "Purpose", "Status"],
        [
            ["Figure 1", "PRISMA flow from 22,166 records to final included studies", "Needs final inclusion lock"],
            ["Figure 2", "Ten-construct conceptual model", "Ready to draw"],
            ["Figure 3", "Construct-pair coverage heatmap", "Ready from current coverage"],
            ["Figure 4", "Final Stage 2 path model", "Needs TSSEM run"],
            ["Figure 5", "Moderator/sensitivity comparison", "Needs OSMASEM/sensitivity run"],
        ],
    )}

## Manuscript Boundary

Do not insert final C&E results claims until numeric N is reconciled or an approved N-eligible subset rule is documented for Paper A. The currently defensible manuscript state is introduction, theory spine, methods, input/QC results, and analysis gate.
"""

    paper_b_map = f"""# Paper B Target-Journal Structure Map

Target journal: Research Synthesis Methods

Decision state: Paper B is a source-anchored, human-adjudicated LLM augmentation/validation study with downstream MASEM/TSSEM diagnostic checks. It is not an LLM replacement paper and not a vendor-ranking benchmark.

## Current Full-Corpus Gate

{md_table(
        ["Check", "Current value", "Implication"],
        [
            ["Full-corpus Step 5 shell exists", b["shell_exists"], "Anchors the completed 2,043-row M1-R run"],
            ["Task rows in shell", b["shell_rows"], "Full-corpus denominator before exception handling"],
            ["Full-corpus M1-R execution", "Completed 2026-06-12", "Nine source-packet-required shards locked and scored"],
            ["Locked/scored rows", "2,043", "0 duplicates, 0 model CLI failures, 15 exception-layer rows"],
            ["SEM reporting lane", "Core-6 diagnostic only", "No all-construct/all-row SEM claim until final specification"],
            ["Post-freeze reference", "213 studies frozen on 2026-06-09", "Governing reference layer"],
            ["OSF archive", "https://osf.io/mkrgd/overview", "Share-safe public repository exists"],
        ],
    )}

Denominator-family shell counts:

{md_table(["Family", "Rows"], shell_family_rows)}

## RSM Submission Components

- Abstract: no more than 250 words and readable to a multidisciplinary audience.
- Keywords: 4 to 6, plus RSM-specific keywords at submission.
- Required highlights in manuscript: What is already known; What is new; Potential impact for Research Synthesis Methods readers.
- Data availability statement with repository DOI/URL or explanation.
- AI-methods transparency: model names/versions, dates, access path, prompts, validation, and output-locking process.

## Proposed RSM Highlights

### What is already known

Systematic review data extraction is labor-intensive, and LLMs can sometimes support extraction workflows. Existing studies show promising but task-dependent accuracy and require human verification.

### What is new

This study validates a locked LLM workflow against a source-anchored adjudicated human reference standard for MASEM-ready extraction. It reports task-family denominators, source conditions, human-disagreement traces, and downstream substitution diagnostics rather than one pooled accuracy score.

### Potential impact for Research Synthesis Methods readers

The design shows how to evaluate LLM extraction as auditable workflow augmentation for complex evidence synthesis, especially when downstream meta-analytic models depend on numeric source accuracy and sample-size eligibility.

## Table Spine

{md_table(
        ["Table", "Title", "Status"],
        [
            ["Table 1", "Data states and claim roles", "Ready"],
            ["Table 2", "Reference construction and source-adjudication workflow", "Ready"],
            ["Table 3", "Task-family denominators and scoring rules", "Ready"],
            ["Table 4", "Locked model outputs and coverage", "Ready for legacy plus completed full-corpus M1-R"],
            ["Table 5", "RQ1 extraction validity by denominator family", "Ready with full-corpus M1-R evidence and exception caveat"],
            ["Table 6", "RQ2 error taxonomy by source condition", "Ready with legacy package"],
            ["Table 7", "RQ3 review-priority triage", "Ready with legacy package"],
            ["Table 8", "Post-freeze full-corpus M1-R outcomes", "Ready as denominator-family result, not pooled accuracy"],
            ["Table 9", "Downstream substitution and TSSEM diagnostics", "Ready as core-6 diagnostic only"],
        ],
    )}

## Figure Spine

{md_table(
        ["Figure", "Purpose", "Status"],
        [
            ["Figure 1", "Five-step source-anchored validation workflow", "Ready"],
            ["Figure 2", "Task-family scoring architecture", "Ready"],
            ["Figure 3", "Locked-output and model-provenance flow", "Ready"],
            ["Figure 4", "Accuracy/abstention profile by denominator family", "Ready with caveats"],
            ["Figure 5", "Source-risk and human-disagreement triage heatmap", "Ready"],
            ["Figure 6", "Downstream substitution diagnostic", "Ready as bounded subset only"],
        ],
    )}

## Manuscript Boundary

Permitted: Paper B can claim a reproducible, source-anchored workflow for evaluating LLM-assisted extraction by task family, report the completed full-corpus `M1-R` outcomes by denominator family and exception-aware gate status, and report bounded diagnostic evidence that the expert-reviewed primary LLM-assisted input did not change the current primary pooled-correlation subset.

Not permitted: one pooled full-corpus accuracy denominator, all-construct/all-row SEM substitution stability, autonomous LLM replacement, silent pooling of direct-r/source-flagged/converted rows, or vendor ranking.
"""

    team_brief = """# Team Writing Brief: Literature Review and Discussion

Date: 2026-06-12

Lead owner: Hosung. Lead retains Introduction, Methods, Results, analysis execution, claim boundary, and final integration. Team members should draft only Literature Review and Discussion materials.

## Paper A Literature Review Assignment

Target length: 1,500 to 2,000 words.

Deliverables:

- Define the 10 constructs: PE, EE, SI, FC, ATT, SE, TRU, ANX, BI, UB.
- Synthesize TAM, UTAUT, TPB/TRA, and AI-specific trust/anxiety literature.
- Explain why attitude remains theoretically important in higher-education AI adoption.
- Position the paper against Scherer et al. and broader education-technology MASEM work.
- Draft hypotheses or research questions without claiming final path estimates.

Do not write:

- Do not report Stage 1 or Stage 2 estimates.
- Do not imply every construct pair has strong coverage.
- Do not treat converted beta/path values as direct correlations.

## Paper A Discussion Assignment

Target length after results insertion: 1,500 to 2,000 words.

Deliverables:

- Interpret final TSSEM/OSMASEM results after lead inserts estimates.
- Discuss theory implications for TAM/UTAUT in AI-specific higher education.
- Discuss practical implications for institutions adopting generative or adaptive AI tools.
- Discuss limitations: missing sample size, incomplete matrices, source-type sensitivity, construct heterogeneity, publication bias.
- Keep implications conditional until final estimates are inserted.

## Paper B Literature Review Assignment

Target length: 1,500 to 2,000 words.

Deliverables:

- Review AI/LLM-assisted data extraction in systematic reviews and meta-analyses.
- Emphasize source-anchored reference standards, locked outputs, prompt/model versioning, and human-in-the-loop validation.
- Explain why MASEM extraction is harder than simple metadata extraction.
- Use task-family and downstream-consequence language rather than one overall denominator.
- Position the study in RSM/evidence-synthesis methodology, not only education.

Do not write:

- Do not say gold standard.
- Do not say the LLM replaces human coders.
- Do not rank Codex, Claude, and Gemini as vendors in the main claim.
- Do not treat 8,783 task units as one accuracy denominator.
- Do not call the 90-row M1-R shard a full-corpus result.

## Paper B Discussion Assignment

Target length after final gate decision: 1,500 to 2,000 words.

Deliverables:

- Interpret augmentation value, review triage, and source-risk visibility.
- Explain where the workflow failed: high-consequence numeric extraction and abstention.
- Discuss why task-family denominators and source conditions matter for research synthesis.
- Discuss OSF/public repository contribution and closed-PDF/private workbook boundaries.
- State next-method implications for evidence synthesis teams adopting LLM workflows.

## Shared Reference Set for Team Authors

- Computers & Education guide: https://www.sciencedirect.com/journal/computers-and-education/publish/guide-for-authors
- RSM instructions: https://www.cambridge.org/core/journals/research-synthesis-methods/information/author-instructions/preparing-your-materials
- Scherer et al. DOI: https://doi.org/10.1016/j.compedu.2018.09.009
- Or (2024) DOI: https://doi.org/10.18357/otessaj.2024.4.3.66
- Schmidt et al. arXiv: https://arxiv.org/abs/2405.14445
- Schroeder et al. arXiv: https://arxiv.org/abs/2501.11840
- RSM GenAI guidance DOI: https://doi.org/10.1017/rsm.2025.10058
- Paper B OSF: https://osf.io/mkrgd/overview
"""

    gap_register = f"""# Submission Readiness Gap Register

Date: 2026-06-12

## P0 Gates

{md_table(
        ["Paper", "Gate", "Current state", "Required before submission claim"],
        [
            ["Paper A", "N-weighted TSSEM/OSMASEM", f'{a["missing_n_rows"]} of {a["usable_rows"]} usable primary rows lack numeric N', "Reconcile sample sizes from source-supported records or approve an N-eligible subset rule specific to Paper A"],
            ["Paper A", "Final inclusion count", "Proposal count and working final count are not harmonized in the manuscript", "Lock final PRISMA counts and final MASEM-eligible study count"],
            ["Paper A", "Final Stage 1/Stage 2 estimates", "Not generated in repo-local submission package", "Run 10-construct TSSEM, sensitivity, and moderator analyses after N gate"],
            ["Paper A", "Closed target exemplar PDF", "Scherer C&E PDF not available in share-safe workspace", "Use PSU access PDF locally for final exact table/figure comparison if needed"],
            ["Paper B", "Full-corpus M1-R expansion", "Completed 2,043-row source-packet-required run with exception-aware scoring", "Report only by denominator family and exception-aware gate status"],
            ["Paper B", "Full-corpus accuracy claim", "Full-corpus denominator-family scoring exists", "Do not collapse into one pooled accuracy denominator"],
            ["Paper B", "SEM reporting lane", "Core-6 diagnostic only selected for current manuscript state", "Complete final TSSEM/MASEM specification before any all-construct/all-row claim"],
            ["Paper B", "Closed comparator PDFs", "Annals/RSM exact PDF layouts not all locally available", "Use PSU access PDFs for exact final table/figure deconstruction if target formatting requires it"],
        ],
    )}

## P1 Build Tasks

- Generate Paper A PRISMA diagram after final inclusion lock.
- Generate Paper A coverage heatmap and conceptual model figure.
- Generate Paper B workflow, scoring architecture, model-provenance, and triage figures.
- Convert table spine into final journal-ready Word tables.
- Run citation manager cleanup and journal reference style conversion.
- Update OSF data-availability text after any new share-safe outputs are added.
"""

    paper_a_gate = f"""# Paper A TSSEM/OSMASEM Run Gate

Date: 2026-06-12

Decision: The lead approved actual TSSEM/OSMASEM/sensitivity estimation for Paper A. The current model-ready primary file is not yet sufficient for a defensible N-weighted metaSEM run because most usable rows do not carry numeric sample size.

## Readiness Snapshot

{md_table(
        ["Metric", "Value"],
        [
            ["Primary rows", a["rows"]],
            ["Usable rows after 10-construct and r checks", a["usable_rows"]],
            ["Rows missing numeric N", a["missing_n_rows"]],
            ["Studies represented", a["studies"]],
            ["Construct-pair coverage", f'{a["covered_pairs"]}/{a["total_pairs"]}'],
            ["Complete 10-construct studies", a["complete_10construct_studies"]],
            ["Studies with 15 or more pairs", a["studies_ge_15_pairs"]],
        ],
    )}

## Stop Condition

Do not run or report a primary N-weighted TSSEM/OSMASEM estimate from this file until one of the following is true:

1. Source-supported numeric N is filled for the rows entering the SEM input; or
2. A Paper A-specific N-eligible subset rule is approved and documented; or
3. A methodological decision is made to run an explicitly labeled non-primary unweighted/pseudo-N diagnostic, with no final substantive path claims.

## Recommended Next Action

Apply a Paper A sample-size reconciliation pass analogous to the Paper B deterministic N reconciliation layer, using source-supported study/sample N from the frozen consensus/reference materials. Then rerun this gate and execute TSSEM/OSMASEM from the reconciled input.
"""

    paper_b_gate = f"""# Full-Corpus M1-R Expansion Completion

Date: 2026-06-12

Decision: The user approved and completed full-corpus expansion centered on the post-freeze 213-study full-corpus gate. The 2,043-row source-packet-required `M1-R` run is locked, registered, and exception-aware scored.

## Current Completion State

{md_table(
        ["Item", "Value", "Status"],
        [
            ["Full-corpus shell", str(PAPER_B_SHELL.relative_to(REPO)), f'{b["shell_rows"]} rows'],
            ["Primary direct/source-r rows", b["family_counts"].get("primary_direct_r_or_source_reported_correlation", 0), "Ready in shell"],
            ["Primary latent/construct correlation rows", b["family_counts"].get("primary_latent_or_construct_correlation_with_source_type_flag", 0), "Ready in shell"],
            ["Secondary beta/path converted rows", b["family_counts"].get("secondary_beta_or_path_converted_effect_size", 0), "Ready in shell"],
            ["Full-corpus M1-R locked rows", "2,043", "Completed across nine shards"],
            ["Exception-aware scorer", "scripts/llm_scoring_20260606/score_full_corpus_m1_r_with_exception_layer.py", "Available"],
            ["Exception-layer rows", "15", "Interpret by gate status, not as generic accuracy"],
            ["SEM reporting lane", "Core-6 diagnostic only", "No all-construct/all-row claim without final specification"],
        ],
    )}

## Required Interpretation Boundary

1. Report the full-corpus `M1-R` result by denominator family.
2. Keep source-reference contract caveats outside the generic full-accuracy numerator.
3. Treat converted beta/path rows as an explicit sensitivity stratum unless a source-type-approved model rebuild is specified.
4. Do not use this run for model-vendor ranking or autonomous replacement claims.

## Claim Boundary

Paper B may report the completed full-corpus `M1-R` denominator-family outcomes and the bounded core-6 TSSEM diagnostic. It may not report one pooled full-corpus accuracy denominator or all-row SEM substitution stability.
"""

    paper_a_target = f"""# AI Adoption in Higher Education: A Meta-Analytic Structural Equation Modeling Study

Target journal: Computers & Education

Draft date: 2026-06-12

## Submission Package State

This target-journal draft updates the existing APA-style shell for Computers & Education. It includes the introduction, methods, input/QC results, table/figure spine, and team insertion points. It deliberately does not invent Stage 1/Stage 2 estimates. The lead-approved TSSEM/OSMASEM run remains gated by Paper A numeric sample-size reconciliation.

## Highlights

- Synthesizes AI adoption evidence in higher education with MASEM.
- Integrates TAM/UTAUT predictors with trust and AI anxiety.
- Separates direct-r inputs from converted sensitivity evidence.
- Tests structural and moderator paths across ten constructs.
- Provides reproducible extraction and QC artifacts.

## Abstract

Artificial intelligence tools are increasingly embedded in higher education, but the empirical adoption literature remains fragmented across constructs, samples, tools, and reporting formats. This study synthesizes higher-education AI adoption evidence using meta-analytic structural equation modeling. The planned model integrates performance expectancy, effort expectancy, social influence, facilitating conditions, attitude, self-efficacy, trust in AI, AI anxiety, behavioral intention, and use behavior. The current analysis-ready package preserves source-reported direct correlations separately from expanded direct-r-form and converted sensitivity evidence. At the 2026-06-05 checkpoint, the model-ready primary input contains 804 rows after tiered source decisions. The current target-journal package documents the final structure, tables, figures, and analysis gate for a Computers & Education submission. Final path estimates, indirect effects, model fit, and moderator results should be inserted only after sample-size reconciliation and the approved TSSEM/OSMASEM run are complete.

Keywords: artificial intelligence; technology acceptance; higher education; MASEM; UTAUT; trust; anxiety

## Introduction

Artificial intelligence tools are now embedded in higher education through large language models, intelligent tutoring systems, automated assessment systems, writing assistants, recommendation tools, and analytics platforms. Their spread has produced a rapidly expanding empirical literature on adoption, acceptance, and use. Yet this literature remains difficult to interpret cumulatively because studies draw from overlapping but nonidentical acceptance frameworks, measure different subsets of constructs, and report evidence in formats that do not directly support a single structural synthesis.

Meta-analytic structural equation modeling is well suited to this problem because it can synthesize study-level correlation matrices and test a theory-guided network of relationships. In AI adoption research, this is especially important. Performance expectancy, effort expectancy, social influence, facilitating conditions, attitude, self-efficacy, behavioral intention, and use behavior provide continuity with TAM, TPB, and UTAUT. Trust in AI and AI anxiety capture psychological features of AI systems that are not fully reducible to general usefulness or ease-of-use beliefs.

The present study develops a MASEM of AI adoption in higher education that integrates traditional technology acceptance constructs with AI-specific psychological constructs. The working model treats attitude as a theoretically meaningful mediator rather than assuming that the parsimonious UTAUT exclusion of attitude applies unchanged to AI adoption contexts. It also tests whether trust and anxiety contribute to behavioral intention beyond standard acceptance predictors.

## Literature Review

[Reserved for team contribution. Use the Team Writing Brief in `docs/07_manuscript_exemplars/20260612/TEAM_WRITING_BRIEF_LIT_REVIEW_DISCUSSION_20260612.md`.]

## Method

### Design and Reporting

Paper A is the parent meta-analysis for the AI adoption evidence-synthesis project. It uses systematic-review procedures to identify eligible studies and applies TSSEM/OSMASEM to synthesize construct-level relationships. Reporting should align with PRISMA 2020 and Computers & Education submission requirements.

### Search, Screening, and Eligibility

The documented search workflow yielded 22,166 records. After deduplication, 16,189 records remained for screening. The final manuscript must harmonize the proposal-stage included-study count with the final locked full-text MASEM-eligible count before submission.

### Constructs and Model Architecture

The primary model uses 10 constructs: performance expectancy, effort expectancy, social influence, facilitating conditions, attitude, self-efficacy, trust in AI, AI anxiety, behavioral intention, and use behavior. The planned architecture places performance expectancy and effort expectancy upstream of attitude, attitude upstream of behavioral intention, behavioral intention upstream of use behavior, and trust/anxiety as AI-specific antecedents of behavioral intention.

### Analysis Plan

The primary analysis will use two-stage MASEM after numeric sample-size reconciliation. Stage 1 will pool study-level correlation matrices with random-effects TSSEM when feasible. Stage 2 will fit the prespecified structural model to the pooled matrix. OSMASEM or equivalent meta-regression will test moderators when moderator data are sufficiently complete. Sensitivity analyses will separate primary direct-r evidence from expanded direct-r-form and converted beta/path/source-statistic evidence.

## Results

### Analysis-Ready Input

{md_table(
        ["Input or gate", "Current value", "Submission interpretation"],
        [
            ["Primary model-ready rows", a["rows"], "Main candidate input after tiered freeze"],
            ["Usable rows after r checks", a["usable_rows"], "Rows available for matrix construction once N is resolved"],
            ["Rows missing numeric N", a["missing_n_rows"], "Primary N-weighted SEM blocker"],
            ["Studies represented", a["studies"], "Current study-level matrix universe"],
            ["Construct-pair coverage", f'{a["covered_pairs"]}/{a["total_pairs"]}', "Broad but incomplete coverage"],
        ],
    )}

### Primary MASEM Results

[Lead insertion point. Insert Stage 1 pooled correlation matrix, heterogeneity estimates, Stage 2 path coefficients, indirect effects, model fit, and sensitivity results after the Paper A sample-size gate and TSSEM/OSMASEM run are complete.]

## Discussion

[Reserved for team contribution after lead inserts final results.]
"""

    paper_b_target = """# Can a Prespecified LLM Workflow Augment MASEM-Ready Evidence Extraction?

Target journal: Research Synthesis Methods

Draft date: 2026-06-12

## Submission Package State

This target-journal draft updates the existing APA-style shell for Research Synthesis Methods. It keeps the claim boundary approved by the lead: Paper B is an LLM augmentation and validation study against a source-anchored adjudicated human reference standard, with bounded downstream diagnostics. It is not an autonomous replacement paper and not a vendor-ranking benchmark.

## Abstract

Data extraction for meta-analytic structural equation modeling requires more than article summarization: reviewers must recover numeric source evidence, map constructs, preserve provenance, distinguish direct correlations from converted statistics, and maintain sample-size eligibility for downstream models. We evaluate a prespecified locked-output LLM workflow against a source-anchored adjudicated human reference standard in an AI adoption evidence-synthesis project. Task units are analyzed by denominator family rather than as one pooled accuracy score. The workflow records model provenance, source conditions, abstentions, human-disagreement traces, and review-priority signals. Current post-freeze evidence includes a 213-study reference standard, clean model-explicit legacy outputs for Codex GPT-5.5, Claude Sonnet, and Gemini 3 Flash, a completed 2,043-row source-rendered full-corpus M1-R run, and a bounded six-construct TSSEM substitution diagnostic. The results support a workflow-augmentation claim: locked LLM outputs and cross-model disagreement can structure review triage, but high-consequence numeric extraction should remain under human review and exception-aware scoring. Full-corpus results are reported by denominator family and gate status; all-row SEM stability and autonomous replacement claims remain outside the current evidence.

Keywords: evidence synthesis; data extraction; large language models; MASEM; validation; human-in-the-loop

## Highlights

### What is already known

LLMs can support parts of systematic review workflows, but extraction performance is task-dependent and usually needs human verification.

### What is new

This study evaluates LLM-assisted extraction for MASEM-ready evidence using source-anchored adjudication, locked outputs, task-family scoring, and downstream substitution diagnostics.

### Potential impact for Research Synthesis Methods readers

The workflow shows how evidence-synthesis teams can evaluate LLM extraction without collapsing heterogeneous tasks into one accuracy denominator or overstating replacement claims.

## Introduction

Evidence synthesis for MASEM requires extraction decisions that are more demanding than simple article summarization. A usable evidence record must distinguish source-reported direct correlations from converted statistics, map constructs consistently, preserve source provenance, identify source-absence cases, handle human-coder disagreement, and maintain enough sample-size and matrix information to support downstream SEM weighting and model fitting.

Large language models may help with this work, but their value depends on the unit of evaluation. Treating thousands of heterogeneous task units as one accuracy denominator would obscure the difference between low-consequence metadata, high-consequence direct-r extraction, source-risk triage, and downstream substitution risk.

## Literature Review

[Reserved for team contribution. Use the Team Writing Brief in `docs/07_manuscript_exemplars/20260612/TEAM_WRITING_BRIEF_LIT_REVIEW_DISCUSSION_20260612.md`.]

## Method

### Corpus and Reference Standard

Paper B uses the validation and extraction subset derived from the AI adoption in higher education MASEM project. The workflow separates raw independent human coder workbooks, pre-adjudication human-human disagreement queues, source-document adjudication decisions, a frozen source-anchored adjudicated human reference layer, locked LLM outputs, and downstream diagnostic analysis files. The post-freeze full-corpus reference contains 213 studies and preserves caveats rather than silently rewriting raw coder workbooks.

### Task Families and Scoring Rules

Task units are not interpreted as one accuracy denominator. Direct-r extraction rows, converted or source-statistic numeric rows, metadata rows, human-review decision rows, source-absence rows, duplicate-source exclusions, blank/absence consensus rows, and trace rows are scored or interpreted separately. Abstentions on scorable rows count as incorrect and are reported as workflow behavior.

### Model Scope and Locked Outputs

Codex GPT-5.5 is the primary prespecified workflow. Claude Sonnet and Gemini 3 Flash are retained as supplementary sensitivity and triage evidence only. Earlier Claude default-unspecified rows are retained only as audit provenance after the Sonnet backfill.

### Analysis Plan

RQ1 evaluates extraction validity by denominator family and task stratum. RQ2 classifies errors by source condition and downstream consequence. RQ3 evaluates whether model behavior, cross-model disagreement, source-risk flags, and human-disagreement traces prioritize expert review. Downstream substitution analyses are reported only as bounded diagnostics.

## Results

### Data States and Claim Roles

| Data state | Current evidence | Claim role |
| --- | --- | --- |
| Frozen full-corpus reference | 213 studies frozen on 2026-06-09 | Current governing reference layer |
| Legacy task-unit package | 8,783 task units | Pre-full-corpus reproducibility and denominator-family evidence |
| Clean model-explicit outputs | 7,859 rows per model | RQ1-RQ3 task-family scoring and sensitivity |
| Bounded source-rendered M1-R shard | 90 rows | Staged diagnostic only |
| Full-corpus source-rendered M1-R | 2,043 rows | Denominator-family and exception-aware Step 5 evidence |
| Core-6 TSSEM diagnostic | 15 complete-case studies | Subset substitution-stability diagnostic |

### Post-Freeze M1-R Full-Corpus Results

The completed full-corpus M1-R run covers 2,043 source-rendered task rows across nine shards, with 0 duplicate task IDs and 0 model CLI failures. Denominator-family outcomes are reported separately: 931 primary latent/source-flagged correlation rows (715 scored, 672 correct, 216 abstentions), 697 primary direct/source-r rows (572 scored, 517 correct, 125 abstentions), and 415 secondary beta/path converted-effect rows (338 scored, 153 correct, 77 abstentions). Fifteen exception-layer rows remain gated by source-reference contract, structural-path evidence, or manual adjudication status.

### Downstream Diagnostic

The bounded core-6 TSSEM diagnostic used PE, EE, SI, FC, BI, and UB in 15 complete-case studies. The current manuscript lane is core-6 diagnostic only. It supports subset diagnostic stability and does not support all-construct or all-row substitution claims.

## Discussion

[Reserved for team contribution after the lead confirms whether the manuscript should remain on the core-6 diagnostic lane or later add a source-type-approved broader TSSEM/MASEM specification.]

## Data Availability

The share-safe Paper B public repository is available at https://osf.io/mkrgd/overview. It excludes raw PDFs, raw human coder workbooks, and private OneDrive-only working materials.
"""

    return {
        OUT / "EXEMPLAR_PDF_MANIFEST_20260612.md": source_manifest,
        OUT / "EXEMPLAR_DECOMPOSITION_MATRIX_20260612.md": decomposition,
        OUT / "PAPER_A_COMPUTERS_AND_EDUCATION_STRUCTURE_MAP_20260612.md": paper_a_map,
        OUT / "PAPER_B_RESEARCH_SYNTHESIS_METHODS_STRUCTURE_MAP_20260612.md": paper_b_map,
        OUT / "TEAM_WRITING_BRIEF_LIT_REVIEW_DISCUSSION_20260612.md": team_brief,
        OUT / "SUBMISSION_READINESS_GAP_REGISTER_20260612.md": gap_register,
        OUT / "PAPER_A_TSSEM_OSMASEM_RUN_GATE_20260612.md": paper_a_gate,
        PAPER_B_RESULTS / "FULL_CORPUS_M1_R_EXPANSION_GATE_20260612.md": paper_b_gate,
        PAPER_A_OUT / "PAPER_A_COMPUTERS_AND_EDUCATION_TARGET_DRAFT_20260612.md": paper_a_target,
        PAPER_B_OUT / "PAPER_B_RESEARCH_SYNTHESIS_METHODS_TARGET_DRAFT_20260612.md": paper_b_target,
        OUT / "PAPER_A_B_TARGET_SUBMISSION_PACKET_20260612.md": "\n\n".join(
            [
                "# Paper A/B Target Submission Packet",
                source_manifest,
                decomposition,
                paper_a_map,
                paper_b_map,
                team_brief,
                gap_register,
            ]
        ),
    }


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for edge, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        element = margins.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[int]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.append(grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Pt(widths[idx] / 20)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def format_table_cell(cell, header: bool = False) -> None:
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(9.25)
            if header:
                run.bold = True


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ["List Bullet", "List Number"]:
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = "Calibri"
            style.font.size = Pt(11)
            style.paragraph_format.space_after = Pt(8)
            style.paragraph_format.line_spacing = 1.167
            style.paragraph_format.left_indent = Inches(0.5)
            style.paragraph_format.first_line_indent = Inches(-0.25)


def clean_inline(text: str) -> str:
    return re.sub(r"`([^`]+)`", r"\1", text.replace("**", ""))


def parse_md_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    table_lines = []
    idx = start
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        table_lines.append(lines[idx].strip())
        idx += 1
    rows = []
    for line in table_lines:
        cells = [clean_inline(cell.strip()) for cell in line.strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells):
            continue
        rows.append(cells)
    return rows, idx


def column_widths(headers: list[str], rows: list[list[str]]) -> list[int]:
    ncols = len(headers)
    weights = []
    for col in range(ncols):
        max_len = max([len(headers[col])] + [len(row[col]) if col < len(row) else 0 for row in rows])
        weights.append(max(1, min(max_len, 45)))
    total = sum(weights)
    raw = [max(900, int(9360 * weight / total)) for weight in weights]
    diff = 9360 - sum(raw)
    raw[-1] += diff
    return raw


def add_md_to_doc(doc: Document, md: str) -> None:
    lines = md.splitlines()
    i = 0
    first_title = True
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith("|"):
            table_rows, i = parse_md_table(lines, i)
            if len(table_rows) >= 2:
                headers = table_rows[0]
                body = table_rows[1:]
                table = doc.add_table(rows=1, cols=len(headers))
                table.alignment = WD_TABLE_ALIGNMENT.LEFT
                table.style = "Table Grid"
                for idx, header in enumerate(headers):
                    cell = table.rows[0].cells[idx]
                    cell.text = header
                    shade_cell(cell, "F2F4F7")
                    format_table_cell(cell, header=True)
                repeat_header(table.rows[0])
                for row in body:
                    cells = table.add_row().cells
                    for idx, value in enumerate(row[: len(headers)]):
                        cells[idx].text = value
                        format_table_cell(cells[idx])
                set_table_width(table, column_widths(headers, body))
                doc.add_paragraph()
            continue
        if stripped.startswith("### "):
            doc.add_paragraph(clean_inline(stripped[4:]), style="Heading 3")
        elif stripped.startswith("## "):
            doc.add_paragraph(clean_inline(stripped[3:]), style="Heading 2")
        elif stripped.startswith("# "):
            text = clean_inline(stripped[2:])
            if first_title:
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.space_after = Pt(10)
                run = para.add_run(text)
                run.font.name = "Calibri"
                run.font.size = Pt(20)
                run.font.color.rgb = RGBColor.from_string("0B2545")
                run.bold = True
                first_title = False
            else:
                doc.add_paragraph(text, style="Heading 1")
        elif stripped.startswith("- "):
            doc.add_paragraph(clean_inline(stripped[2:]), style="List Bullet")
        elif re.match(r"^\d+\.\s+", stripped):
            doc.add_paragraph(clean_inline(re.sub(r"^\d+\.\s+", "", stripped)), style="List Number")
        elif stripped.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            para = doc.add_paragraph()
            run = para.add_run("\n".join(code_lines))
            run.font.name = "Courier New"
            run.font.size = Pt(8.5)
        else:
            doc.add_paragraph(clean_inline(stripped))
        i += 1


def build_docx(md: str, out_path: Path) -> None:
    doc = Document()
    style_doc(doc)
    add_md_to_doc(doc, md)
    doc.save(out_path)


def main() -> None:
    ensure_dirs()
    a = paper_a_readiness()
    b = paper_b_readiness()
    contents = build_contents(a, b)
    for path, text in contents.items():
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(text.rstrip() + "\n", encoding="utf-8")

    build_docx(
        contents[PAPER_A_OUT / "PAPER_A_COMPUTERS_AND_EDUCATION_TARGET_DRAFT_20260612.md"],
        PAPER_A_OUT / "PAPER_A_COMPUTERS_AND_EDUCATION_TARGET_DRAFT_20260612.docx",
    )
    build_docx(
        contents[PAPER_B_OUT / "PAPER_B_RESEARCH_SYNTHESIS_METHODS_TARGET_DRAFT_20260612.md"],
        PAPER_B_OUT / "PAPER_B_RESEARCH_SYNTHESIS_METHODS_TARGET_DRAFT_20260612.docx",
    )
    build_docx(
        contents[OUT / "PAPER_A_B_TARGET_SUBMISSION_PACKET_20260612.md"],
        OUT / "PAPER_A_B_TARGET_SUBMISSION_PACKET_20260612.docx",
    )

    print("Generated target-journal package:")
    for path in sorted(contents):
        print(f"- {path.relative_to(REPO)}")
    print(f"- {PAPER_A_OUT.relative_to(REPO)}/PAPER_A_COMPUTERS_AND_EDUCATION_TARGET_DRAFT_20260612.docx")
    print(f"- {PAPER_B_OUT.relative_to(REPO)}/PAPER_B_RESEARCH_SYNTHESIS_METHODS_TARGET_DRAFT_20260612.docx")
    print(f"- {OUT.relative_to(REPO)}/PAPER_A_B_TARGET_SUBMISSION_PACKET_20260612.docx")


if __name__ == "__main__":
    main()
