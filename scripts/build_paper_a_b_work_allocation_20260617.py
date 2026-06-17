#!/usr/bin/env python3
"""Build Korean Paper A/B work allocation artifacts."""

from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "08_operations"
DOCX_PATH = OUT_DIR / "PAPER_A_B_WORK_ALLOCATION_AND_TRACKING_PLAN_KO_20260617.docx"
MD_PATH = OUT_DIR / "PAPER_A_B_WORK_ALLOCATION_AND_TRACKING_PLAN_KO_20260617.md"
CSV_PATH = OUT_DIR / "PAPER_A_B_TASK_TRACKER_SEED_20260617.csv"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(92, 99, 112)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"


META = [
    ("문서 목적", "Paper A/B 연구자 작업 배분, 추적, 배포 경계, OneDrive 작업 전환 기준 정리"),
    ("기준 시점", "2026-06-17 KST"),
    ("근거 스냅샷", "git HEAD 24dd8dc / Paper A tag paper-a-longtable-panel-submission-draft-20260616"),
    ("주요 작업 위치", "OneDrive root 00_INDEX/2026-06-17_Paper_A_B_work_allocation 및 05_manuscripts"),
    ("관리 원칙", "Git 추적 파일은 repository mirror에 동기화하고, 팀 작업 문서는 OneDrive root canonical tree에 두며, raw/private 자료는 공개 배포물과 분리"),
]

SUMMARY = [
    ("Paper A", "제출 직전 보강 단계",
     "LongTable panel draft가 최신 기준이다. full10은 이론적 target/evidence map이고 core7/trust6만 경험적 model-family 경로로 쓴다.",
     "final full-text eligibility box source-lock, duplicate DOI 해석, Word pagination/figure placement, core7/trust6 영향진단"),
    ("Paper B", "post-freeze Step 5 증거 정리 단계",
     "213-study source-anchored adjudicated human reference는 동결되었고 2,043-row M1-R full-corpus run은 존재한다.",
     "최종 all-construct/all-row SEM 대체 안정성 주장은 matrix sparsity, source-type, model specification 경계가 풀린 뒤에만 가능"),
    ("공통 운영", "OneDrive 작업 전환",
     "OneDrive root canonical tree를 기준으로 추적 문서, 원고 패키지, OSF 산출물을 같은 release register에서 관리한다.",
     "공개 가능 자료, 내부 원천 자료, runtime/private 상태를 분리해서 연구자별 병렬 작업 충돌을 줄인다."),
]

PRINCIPLES = [
    ("주장 경계 우선", "원고 문장보다 claim boundary를 먼저 잠근다.",
     "Paper A full10을 단일 추정 SEM처럼 쓰지 않는다. Paper B Step 5를 최종 SEM substitution 완료로 쓰지 않는다."),
    ("상태 분리", "raw, pairwise disagreement, adjudicated reference, LLM output, manuscript derivation을 별도 상태로 유지한다.",
     "raw coder workbook을 consensus처럼 고쳐 쓰지 않는다. LongTable/.omx runtime state를 연구자 배포물로 취급하지 않는다."),
    ("근거 파일 연결", "모든 작업 항목은 증거 파일과 통과 기준을 가진다.",
     "구두 결정만으로 PRISMA count, reference freeze, OSF 공개자료를 업데이트하지 않는다."),
    ("OneDrive는 작업 공간, OSF는 공개 패키지", "OneDrive root canonical tree는 공동 작업 공간이고 repository mirror는 검토용 복제본이며 OSF는 share-safe release surface다.",
     "local PDFs, private source packets, raw LLM transcripts를 OSF/Git 공개 패키지에 섞지 않는다."),
    ("역할은 산출물 기준", "사람별 소유권은 산출물, gate, 검증 지표로 정의한다.",
     "역할명이 있어도 산출물 없는 검토를 완료로 처리하지 않는다."),
]

ROLES = [
    ("R1 / PI / 총괄",
     "최종 claim boundary, PRISMA/full-text source-lock, duplicate DOI 판단, Paper B adjudication 최종 승인",
     "Paper A PRISMA final lock note; Paper B manuscript claim boundary memo; 주간 release register",
     "source-lock 완료 여부, DOI 중복 판정, 원고 문장과 evidence file 일치"),
    ("R2 / 독립 검토자",
     "Paper A include/conflict/uncertain 검증 표본 또는 지정 범위 독립 검토; Paper B Pair D 또는 cross-check",
     "IRR 표본 판정표; disagreement review notes; source-check 후보 목록",
     "Cohen's kappa 또는 agreement summary, unresolved row count"),
    ("R3 / 독립 검토자·중재 보조",
     "R2와 독립 코딩/IRR, Paper B Pair D, cross-pair adjudication 보조",
     "R2-R3 독립판정 비교표; escalation list; construct remap 후보",
     "불일치 유형별 triage 완료율, source-type mismatch 해소율"),
    ("R4 / 데이터·품질 검토",
     "Paper B Pair C, workbook/package QA, manifest/checksum, source-type boundary audit",
     "package manifest; checksum/run manifest review; data availability exclusion list",
     "0 duplicate task IDs, 0 model CLI failure, manifest/file count 일치"),
    ("Methods Critic",
     "MASEM/TSSEM/OSMASEM 방법론, complete-case k, CI 기반 추론, 영향진단 검토",
     "methods risk memo; sensitivity/influence diagnostic recommendation",
     "core7 k=4, trust6 k=7, incomplete CI 문장에 대한 승인/수정 여부"),
    ("Measurement Auditor",
     "construct mapping, source type, HTMT/Fornell-Larcker/path coefficient 처리 규칙 점검",
     "construct operational definition table; source-type audit notes",
     "ATT/TRU/ANX/SE/FC/PE/EE/SI/BI/UB mapping exception count"),
    ("Voice Keeper / Theory Writer",
     "Paper A Introduction/Theory/Discussion 서사, 'theory-preserving estimability diagnosis' 톤 유지",
     "문헌통합 draft; Korean writing guide 반영표; overclaim 제거표",
     "full10/core7/trust6 구분이 모든 주요 섹션에 일관되게 반영됨"),
    ("Reviewer / Venue Strategist",
     "target journal fit, APA 7, submission checklist, figure/table placement, reviewer objection 대응",
     "journal readiness checklist; cover-letter issue list; Word pagination QA",
     "submission blocker 0개 또는 명시된 residual risk 목록"),
    ("Ethics / Release Manager",
     "OSF/OneDrive/Git 공개 경계, data availability, private material exclusion",
     "Paper A/B release register; EXCLUDED_PRIVATE_MATERIALS review; OSF update checklist",
     "OSF zip에 private raw/PDF/source packet 미포함"),
]

PAPER_A_TASKS = [
    ("A1", "PRISMA/source-lock", "R1 + R2/R3",
     "225 included rows - 1 duplicate DOI = 224 unique reports/studies working lock을 최종 full-text eligibility boxes와 대조",
     "paper_a/PRISMA_COUNTS_LOCK_20260615.md update 또는 confirmation note",
     "duplicate DOI가 same report, metadata error, distinct report 중 무엇인지 결정"),
    ("A2", "Screening/eligibility audit", "R2 + R3",
     "include/conflict/uncertain 검토 표본과 최종판정 기록을 재점검하고 IRR/불일치 요약 생성",
     "IRR summary + unresolved eligibility queue",
     "reviewer가 PRISMA human review 절차를 추적 가능하다고 판단"),
    ("A3", "Model-family methods QA", "Methods Critic + Measurement Auditor",
     "full10/core7/trust6, incomplete CI, ANX/SE underidentification, PE-vs-EE role comparison 문장 점검",
     "methods and measurement risk memo",
     "full10을 추정 SEM으로 쓰는 문장이 0개"),
    ("A4", "Discussion/literature integration", "Voice Keeper + Theory Writer",
     "최근 AI-in-education 문헌, construct genealogy, operational definition table, theory-preserving diagnosis 서사 통합",
     "Discussion/literature revision pack",
     "claim tone이 'definitive model confirmation'이 아니라 'estimability diagnosis'로 유지"),
    ("A5", "Word/APA submission package", "Reviewer + Venue Strategist",
     "PAPER_A_LONGTABLE_PANEL_SUBMISSION_DRAFT_20260616.docx를 기준으로 pagination, figure placement, APA references 점검",
     "journal readiness checklist + final DOCX candidate",
     "caption/table/figure order, reference italics, PRISMA label 점검 완료"),
    ("A6", "OSF/OneDrive release sync", "Ethics / Release Manager",
     "Paper A OSF bwzgc 패키지와 OneDrive root canonical tree의 최신 산출물 차이 확인",
     "release register row + share-safe manifest",
     "public zip에는 raw PDFs/private state 미포함"),
]

PAPER_B_TASKS = [
    ("B1", "Reference-standard protocol guard", "R1 + R4",
     "raw freeze -> disagreement -> source adjudication -> reference freeze -> LLM comparison 순서가 문서/원고에서 유지되는지 확인",
     "protocol compliance memo",
     "raw workbook overwrite나 gold standard 표현 없음"),
    ("B2", "Disagreement and adjudication trace", "R2 + R3 + R4",
     "pairwise disagreement, decision logs, source-type mismatch, construct remap 항목을 원고 표/부록 후보로 정리",
     "disagreement taxonomy table draft",
     "meaningful difference triage와 source-backed decision이 연결됨"),
    ("B3", "Step 5 full-corpus evidence packaging", "R4 + Methods Critic",
     "2,043-row M1-R full-corpus run, exception-aware scorer, denominator-family 결과를 manuscript derivation-ready 형태로 정리",
     "Paper B Step 5 evidence packet",
     "0 duplicate task IDs/0 model CLI failures/exception layer 반영 확인"),
    ("B4", "SEM claim boundary", "Methods Critic + Measurement Auditor",
     "core-6 diagnostic, broader core7/core8 sparse probes, all-construct gate를 분리해 Results 문장 점검",
     "SEM boundary table + Results paragraph",
     "all-construct/all-row SEM substitution stability claim을 하지 않음"),
    ("B5", "RQ3 cross-model triage narrative", "Reviewer + Venue Strategist",
     "cross-model disagreement을 vendor ranking이 아니라 triage/error visibility 신호로 정리",
     "RQ3 figure/table draft",
     "model comparison 문장이 ranking/marketing으로 읽히지 않음"),
    ("B6", "OSF/data availability maintenance", "Ethics / Release Manager",
     "Paper B OSF mkrgd, public_data_repository_20260611, EXCLUDED_PRIVATE_MATERIALS, data availability text 정합성 점검",
     "Paper B release register row",
     "private source packets/raw transcripts/PDFs excluded"),
]

CADENCE = [
    ("월요일", "상태 잠금", "R1이 Paper A/B board를 status, blocker, evidence_path 기준으로 정리한다."),
    ("화요일-수요일", "독립 작업", "R2/R3/R4와 specialist role이 각자 산출물을 만들고 tracker에 evidence path를 적는다."),
    ("목요일", "methods/measurement gate", "Methods Critic과 Measurement Auditor가 claim boundary와 source-type 문제를 승인/수정 요청한다."),
    ("금요일", "release/readiness gate", "Reviewer, Venue Strategist, Ethics/Release Manager가 OneDrive/Git/OSF 상태와 submission blocker를 확인한다."),
    ("상시", "변경 규칙", "Paper B workflow status가 바뀌면 WORKFLOW_STATUS_LOG.md를 같은 commit에서 업데이트한다."),
]

RELEASE_LANES = [
    ("Git", "protocols, scripts, share-safe summaries, manuscripts, public repository packages", "tracked files only; raw/private/runtime 제외"),
    ("OneDrive root canonical tree", "팀 편집용 문서, 원고 패키지, source packages, 작업 추적 폴더", "00_INDEX, 02_source_packages, 05_manuscripts를 기준으로 관리"),
    ("OneDrive repository mirror", "Git 추적 파일 검토용 복제본", "90_repository_mirror/journal_AI-adoption_meta; 실행 repo가 아니라 review mirror로만 사용"),
    ("OSF Paper A", "Paper A public_data_repository_20260615_osf_ready.zip", "bwzgc component; PRISMA/source-lock 확정 후 update 여부 결정"),
    ("OSF Paper B", "Paper B public_data_repository_20260611_osf_upload.zip", "mkrgd component; full-corpus evidence는 경계 문장과 함께만 반영"),
]

TRACKER_ROWS = [
    ("Paper A", "A1", "PRISMA/source-lock", "R1 + R2/R3", "High", "Open",
     "final full-text eligibility boxes confirmed", "paper_a/PRISMA_COUNTS_LOCK_20260615.md",
     "OneDrive/Git/OSF", "duplicate DOI interpretation required"),
    ("Paper A", "A2", "Screening/eligibility audit", "R2 + R3", "High", "Open",
     "IRR/unresolved queue written", "paper_a/DISCUSSION_LOG_KR.md",
     "OneDrive/Git", "use reviewer codes, do not overwrite source logs"),
    ("Paper A", "A3", "Model-family methods QA", "Methods Critic + Measurement Auditor", "High", "Open",
     "overclaim check passes", "CURRENT.md",
     "Git/manuscript", "full10 is target/evidence map"),
    ("Paper A", "A4", "Discussion/literature integration", "Voice Keeper + Theory Writer", "Medium", "Open",
     "Discussion revision pack ready", "paper_a/manuscript/target_journal/apa7_model_family_full_manuscript_scaffold_20260615/PAPER_A_THEORY_DISCUSSION_WRITING_GUIDE_KR_20260616.docx",
     "OneDrive/Git", "preserve researcher voice"),
    ("Paper A", "A5", "Word/APA submission package", "Reviewer + Venue Strategist", "High", "Open",
     "submission blocker checklist cleared", "paper_a/manuscript/target_journal/apa7_model_family_full_manuscript_scaffold_20260615/PAPER_A_LONGTABLE_PANEL_SUBMISSION_DRAFT_20260616.docx",
     "OneDrive/Git", "manual pagination review required"),
    ("Paper A", "A6", "OSF/OneDrive release sync", "Ethics / Release Manager", "Medium", "Open",
     "release register updated", "paper_a/public_data_repository_20260615_osf_ready.zip",
     "OneDrive/OSF", "no private raw materials in public zip"),
    ("Paper B", "B1", "Reference-standard protocol guard", "R1 + R4", "High", "Open",
     "workflow order verified", "data/04_extraction/WORKFLOW_STATUS_LOG.md",
     "Git/manuscript", "source-anchored adjudicated human reference standard"),
    ("Paper B", "B2", "Disagreement/adjudication trace", "R2 + R3 + R4", "Medium", "Open",
     "taxonomy table ready", "data/04_extraction/02_pre_adjudication_disagreement/RATER_COMPARISON_PLAYBOOK.md",
     "Git/manuscript", "meaningful differences only"),
    ("Paper B", "B3", "Step 5 evidence packaging", "R4 + Methods Critic", "High", "Open",
     "full-corpus packet ready", "data/04_extraction/05_llm_masem_substitution/results/PAPER_B_STEP5_FULL_CORPUS_M1R_STATUS_AND_NEXT_WORK_20260612.md",
     "Git/OSF", "denominator-family and exception-aware scoring"),
    ("Paper B", "B4", "SEM claim boundary", "Methods Critic + Measurement Auditor", "High", "Open",
     "SEM boundary table approved", "data/04_extraction/05_llm_masem_substitution/results/r_tssem_substitution_20260611/PAPER2_TSSEM_SUBSTITUTION_DIAGNOSTIC_20260611.md",
     "Git/manuscript", "core-6 diagnostic only unless broader model approved"),
    ("Paper B", "B5", "RQ3 triage narrative", "Reviewer + Venue Strategist", "Medium", "Open",
     "triage narrative ready", "data/04_extraction/05_llm_masem_substitution/results/PAPER2_RQ3_TRIAGE_CROSS_MODEL_SENSITIVITY_20260611.md",
     "Git/manuscript", "no vendor ranking"),
    ("Paper B", "B6", "OSF/data availability maintenance", "Ethics / Release Manager", "Medium", "Open",
     "public/private boundary checked", "paper_b/public_data_repository_20260611/EXCLUDED_PRIVATE_MATERIALS.md",
     "OneDrive/OSF", "raw PDFs/workbooks excluded from public release"),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_in):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths_in):
                cell.width = Inches(widths_in[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def format_paragraph(p, before=0, after=6, line=1.25, keep=False):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.keep_together = keep


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    format_paragraph(p)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    format_paragraph(p, after=4)
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    p = cell.paragraphs[0]
    format_paragraph(p, after=4, line=1.25)
    r = p.add_run(title)
    set_run_font(r, bold=True, color=INK)
    p2 = cell.add_paragraph()
    format_paragraph(p2, after=0, line=1.25)
    r2 = p2.add_run(body)
    set_run_font(r2)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_simple_table(doc, headers, rows, widths, header_fill=LIGHT_BLUE, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_shading(hdr[i], header_fill)
        p = hdr[i].paragraphs[0]
        format_paragraph(p, after=0, line=1.15)
        r = p.add_run(header)
        set_run_font(r, size=font_size, bold=True, color=INK)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_shading(cells[i], WHITE)
            p = cells[i].paragraphs[0]
            format_paragraph(p, after=0, line=1.15)
            r = p.add_run(str(value))
            set_run_font(r, size=font_size)
    return table


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    for list_name in ("List Bullet", "List Number"):
        style = styles[list_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    format_paragraph(header, after=0, line=1.0)
    r = header.add_run("AI Adoption Meta Analysis | Paper A/B Work Allocation")
    set_run_font(r, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_paragraph(footer, after=0, line=1.0)
    fr = footer.add_run("Internal coordination draft | 2026-06-17")
    set_run_font(fr, size=9, color=MUTED)


def build_docx():
    doc = Document()
    configure_document(doc)

    p = doc.add_paragraph()
    format_paragraph(p, before=4, after=3, line=1.15, keep=True)
    r = p.add_run("Paper A/B 연구 작업 배분 및 추적 운영안")
    set_run_font(r, size=23, color=RGBColor(0, 0, 0), bold=True)
    p2 = doc.add_paragraph()
    format_paragraph(p2, before=0, after=14, line=1.15, keep=True)
    r2 = p2.add_run("연구자별 산출물 책임, Paper A/B claim boundary, OneDrive 작업 전환, OSF/Git 배포 관리를 한 문서에서 추적하기 위한 운영 브리프")
    set_run_font(r2, size=12.5, color=MUTED)

    for label, value in META:
        p = doc.add_paragraph()
        format_paragraph(p, after=2, line=1.1)
        lr = p.add_run(f"{label}: ")
        set_run_font(lr, bold=True, color=INK)
        vr = p.add_run(value)
        set_run_font(vr)

    add_callout(
        doc,
        "핵심 판단",
        "Paper A는 제출 전 source-lock과 overclaim 방지가 병목이고, Paper B는 post-freeze Step 5 근거를 manuscript-ready로 정리하되 최종 all-construct SEM substitution claim은 아직 gate로 관리해야 한다. OneDrive는 공동 작업 최신본, Git은 추적 가능한 share-safe 이력, OSF는 공개 배포면으로 분리한다.",
    )

    add_heading(doc, "1. 현재 상태와 즉시 관리 초점", 1)
    add_simple_table(
        doc,
        ["범위", "현재 상태", "사용 가능한 주장", "다음 gate"],
        SUMMARY,
        [1.0, 1.15, 2.35, 2.0],
        font_size=8.7,
    )

    doc.add_page_break()
    add_heading(doc, "2. 운영 원칙", 1)
    add_simple_table(
        doc,
        ["원칙", "적용 방식", "금지/주의"],
        PRINCIPLES,
        [1.25, 2.45, 2.8],
        font_size=8.8,
    )

    add_heading(doc, "3. 역할별 작업 배분", 1)
    add_body(
        doc,
        "아래 역할명은 기존 연구 로그의 R1-R4 및 LongTable panel 역할을 기준으로 정리한 운영 단위다. 실제 성명은 OneDrive tracker의 owner 필드에 추가하면 된다.",
    )
    add_simple_table(
        doc,
        ["역할/담당", "핵심 책임", "즉시 산출물", "관리 지표"],
        ROLES,
        [1.35, 2.05, 1.75, 1.35],
        font_size=7.9,
    )

    add_heading(doc, "4. Paper A 작업 보드", 1)
    add_body(
        doc,
        "Paper A의 최신 원고 기준은 2026-06-16 LongTable panel submission draft다. 작업의 중심은 PRISMA/full-text source-lock, model-family claim boundary, Discussion/literature integration, Word/APA submission readiness다.",
    )
    add_simple_table(
        doc,
        ["ID", "작업 묶음", "Owner", "실행 내용", "산출물", "완료 gate"],
        PAPER_A_TASKS,
        [0.35, 0.85, 1.05, 1.65, 1.35, 1.25],
        font_size=7.4,
    )

    doc.add_page_break()
    add_heading(doc, "5. Paper B 작업 보드", 1)
    add_body(
        doc,
        "Paper B는 source-anchored adjudicated human reference standard 이후의 Step 5 증거를 원고화하는 단계다. 단, evidence packaging과 final SEM claim을 분리해서 관리한다.",
    )
    add_simple_table(
        doc,
        ["ID", "작업 묶음", "Owner", "실행 내용", "산출물", "완료 gate"],
        PAPER_B_TASKS,
        [0.35, 0.85, 1.05, 1.65, 1.35, 1.25],
        font_size=7.4,
    )

    add_heading(doc, "6. 공통 추적 리듬", 1)
    for day, label, detail in CADENCE:
        p = doc.add_paragraph()
        format_paragraph(p, after=4, line=1.25)
        r = p.add_run(f"{day} - {label}: ")
        set_run_font(r, bold=True, color=INK)
        r2 = p.add_run(detail)
        set_run_font(r2)

    add_heading(doc, "7. OneDrive, Git, OSF 배포 경계", 1)
    add_simple_table(
        doc,
        ["면", "넣을 것", "운영 규칙"],
        RELEASE_LANES,
        [1.15, 2.35, 3.0],
        font_size=8.3,
    )
    add_callout(
        doc,
        "OneDrive 전환 규칙",
        "공동 작업 문서는 OneDrive root canonical tree에서 관리하고, repository mirror는 Git 추적 파일 검토용으로만 사용한다. raw/private 파일은 source package/workbook 계열에 둘 수 있지만 Git mirror/OSF public package에는 자동 포함하지 않는다. 공개 배포 전에는 release register, manifest, EXCLUDED_PRIVATE_MATERIALS를 같이 확인한다.",
    )

    add_heading(doc, "8. 실행 체크리스트", 1)
    checks = [
        "OneDrive repository mirror가 Git 추적 파일 기준으로 최신인지 확인하고, 팀 작업 문서는 OneDrive root canonical tree에 있는지 확인한다.",
        "Paper A/B tracker CSV에서 owner, status, next_gate, evidence_path를 매주 갱신한다.",
        "Paper A는 PRISMA/source-lock이 끝나기 전 final PRISMA 2020 flow로 부르지 않는다.",
        "Paper B는 all-construct/all-row SEM stability claim을 최종 gate 전까지 쓰지 않는다.",
        "OSF 업데이트 시 public zip, data availability statement, private material exclusion list를 함께 확인한다.",
        "원고 문장 수정 후에는 claim boundary 담당자가 마지막으로 과장 표현을 제거한다.",
    ]
    for item in checks:
        add_bullet(doc, item)

    doc.save(DOCX_PATH)


def md_table(headers, rows):
    lines = []
    lines.append("| " + " | ".join(headers) + " |")
    lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
    for row in rows:
        lines.append("| " + " | ".join(str(cell).replace("\n", "<br>") for cell in row) + " |")
    return "\n".join(lines)


def build_md():
    lines = [
        "# Paper A/B 연구 작업 배분 및 추적 운영안",
        "",
        "연구자별 산출물 책임, Paper A/B claim boundary, OneDrive 작업 전환, OSF/Git 배포 관리를 한 문서에서 추적하기 위한 운영 브리프입니다.",
        "",
    ]
    for label, value in META:
        lines.append(f"- **{label}:** {value}")
    lines.extend([
        "",
        "## 핵심 판단",
        "",
        "Paper A는 제출 전 source-lock과 overclaim 방지가 병목이고, Paper B는 post-freeze Step 5 근거를 manuscript-ready로 정리하되 최종 all-construct SEM substitution claim은 아직 gate로 관리해야 합니다. OneDrive는 공동 작업 최신본, Git은 추적 가능한 share-safe 이력, OSF는 공개 배포면으로 분리합니다.",
        "",
        "## 1. 현재 상태와 즉시 관리 초점",
        "",
        md_table(["범위", "현재 상태", "사용 가능한 주장", "다음 gate"], SUMMARY),
        "",
        "## 2. 운영 원칙",
        "",
        md_table(["원칙", "적용 방식", "금지/주의"], PRINCIPLES),
        "",
        "## 3. 역할별 작업 배분",
        "",
        "아래 역할명은 기존 연구 로그의 R1-R4 및 LongTable panel 역할을 기준으로 정리한 운영 단위입니다. 실제 성명은 OneDrive tracker의 owner 필드에 추가하면 됩니다.",
        "",
        md_table(["역할/담당", "핵심 책임", "즉시 산출물", "관리 지표"], ROLES),
        "",
        "## 4. Paper A 작업 보드",
        "",
        "Paper A의 최신 원고 기준은 2026-06-16 LongTable panel submission draft입니다. 작업의 중심은 PRISMA/full-text source-lock, model-family claim boundary, Discussion/literature integration, Word/APA submission readiness입니다.",
        "",
        md_table(["ID", "작업 묶음", "Owner", "실행 내용", "산출물", "완료 gate"], PAPER_A_TASKS),
        "",
        "## 5. Paper B 작업 보드",
        "",
        "Paper B는 source-anchored adjudicated human reference standard 이후의 Step 5 증거를 원고화하는 단계입니다. evidence packaging과 final SEM claim은 분리해서 관리합니다.",
        "",
        md_table(["ID", "작업 묶음", "Owner", "실행 내용", "산출물", "완료 gate"], PAPER_B_TASKS),
        "",
        "## 6. 공통 추적 리듬",
        "",
    ])
    for day, label, detail in CADENCE:
        lines.append(f"1. **{day} - {label}:** {detail}")
    lines.extend([
        "",
        "## 7. OneDrive, Git, OSF 배포 경계",
        "",
        md_table(["면", "넣을 것", "운영 규칙"], RELEASE_LANES),
        "",
        "## 8. 실행 체크리스트",
        "",
    ])
    for item in [
        "OneDrive repository mirror가 Git 추적 파일 기준으로 최신인지 확인하고, 팀 작업 문서는 OneDrive root canonical tree에 있는지 확인한다.",
        "Paper A/B tracker CSV에서 owner, status, next_gate, evidence_path를 매주 갱신한다.",
        "Paper A는 PRISMA/source-lock이 끝나기 전 final PRISMA 2020 flow로 부르지 않는다.",
        "Paper B는 all-construct/all-row SEM stability claim을 최종 gate 전까지 쓰지 않는다.",
        "OSF 업데이트 시 public zip, data availability statement, private material exclusion list를 함께 확인한다.",
        "원고 문장 수정 후에는 claim boundary 담당자가 마지막으로 과장 표현을 제거한다.",
    ]:
        lines.append(f"- {item}")
    lines.append("")
    MD_PATH.write_text("\n".join(lines), encoding="utf-8")


def build_csv():
    with CSV_PATH.open("w", encoding="utf-8", newline="") as f:
        writer = csv.writer(f, lineterminator="\n")
        writer.writerow([
            "paper",
            "task_id",
            "workstream",
            "owner",
            "priority",
            "status",
            "next_gate",
            "evidence_path",
            "release_surface",
            "notes",
        ])
        writer.writerows(TRACKER_ROWS)


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    build_md()
    build_csv()
    build_docx()
    print(DOCX_PATH)
    print(MD_PATH)
    print(CSV_PATH)


if __name__ == "__main__":
    main()
