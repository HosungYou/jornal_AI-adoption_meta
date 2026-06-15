#!/usr/bin/env python3
from __future__ import annotations

import csv
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path('/Users/newhosung/Academic/2026/AI Adoption Meta Analysis')
DATE = '20260615'
PKG = ROOT / 'paper_a/manuscript/target_journal/apa7_model_family_full_manuscript_scaffold_20260615'
DATA_PKG = ROOT / 'data/04_extraction/05_llm_masem_substitution/results/paper_a_apa7_model_family_manuscript_package_20260615'
ONEDRIVE = Path('/Users/newhosung/Library/CloudStorage/OneDrive-SharedLibraries-ThePennsylvaniaStateUniversity/AI Adoption Meta Analysis - Documents/05_manuscripts/Paper_A/2026-06-15_apa7_model_family_full_manuscript_scaffold')
MD = PKG / f'PAPER_A_APA7_MODEL_FAMILY_FULL_MANUSCRIPT_SCAFFOLD_{DATE}.md'
FIT = PKG / 'tables/paper_a_model_family_fit_with_n_20260615.csv'
PATHS = PKG / 'tables/paper_a_model_family_structural_paths_ci_inference_20260615.csv'
FIGURES = [
    (PKG / 'figures/figure_1_full10_theoretical_evidence_map_heatmap_ci_20260615.png', 'Figure 1. Full10 theoretical evidence map.'),
    (PKG / 'figures/figure_2_core7_att_mediation_masem_path_ci_20260615.png', 'Figure 2. Core7 attitude-mediation MASEM path diagram.'),
    (PKG / 'figures/figure_3_trust6_mechanism_masem_path_ci_20260615.png', 'Figure 3. Trust6 mechanism MASEM path diagram.'),
]
DOCX = PKG / f'PAPER_A_APA7_MODEL_FAMILY_FULL_MANUSCRIPT_SCAFFOLD_{DATE}.docx'
DATA_DOCX = DATA_PKG / DOCX.name
ONEDRIVE_DOCX = ONEDRIVE / DOCX.name
CURRENT = ROOT / 'CURRENT.md'
README = PKG / f'README_PAPER_A_APA7_MODEL_FAMILY_MANUSCRIPT_PACKAGE_{DATE}.md'


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement('w:fldChar')
    fld_char1.set(qn('w:fldCharType'), 'begin')
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = 'PAGE'
    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def set_cell_text(cell, text, bold=False):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.size = Pt(9)


def read_csv(path):
    with path.open(newline='', encoding='utf-8') as f:
        return list(csv.DictReader(f))


def configure_document(doc: Document):
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    add_page_number(sec.header.paragraphs[0])

    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Times New Roman'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 2
    normal.paragraph_format.space_after = Pt(0)

    for name in ['Title', 'Heading 1', 'Heading 2', 'Heading 3']:
        style = styles[name]
        style.font.name = 'Times New Roman'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        style.font.size = Pt(12)
        style.font.bold = True
        style.paragraph_format.line_spacing = 2
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)


def add_inline_markdown(paragraph, text):
    text = text.replace('  ', ' ')
    pattern = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`)')
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        token = match.group(0)
        if token.startswith('**'):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith('`'):
            run = paragraph.add_run(token[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_paragraph(doc, text, style=None, align=None, indent=False):
    p = doc.add_paragraph(style=style)
    if align:
        p.alignment = align
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.5)
    p.paragraph_format.line_spacing = 2
    p.paragraph_format.space_after = Pt(0)
    add_inline_markdown(p, text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = 2
    p.paragraph_format.space_after = Pt(0)
    add_inline_markdown(p, text)


def clean_markdown_line(line):
    return line.strip().replace('  ', ' ')


def add_markdown_body(doc: Document, md_text: str):
    lines = md_text.splitlines()
    buffer = []
    in_ref = False

    def flush():
        nonlocal buffer
        if not buffer:
            return
        text = ' '.join(buffer).strip()
        buffer = []
        if not text:
            return
        if text.startswith('**Keywords:**'):
            add_paragraph(doc, text, indent=False)
        elif in_ref:
            p = add_paragraph(doc, text, indent=False)
            p.paragraph_format.first_line_indent = Inches(-0.5)
            p.paragraph_format.left_indent = Inches(0.5)
        else:
            add_paragraph(doc, text, indent=True)

    for raw in lines:
        line = raw.rstrip()
        stripped = line.strip()
        if not stripped:
            flush()
            continue
        if stripped.startswith('# '):
            flush()
            title = stripped[2:].strip()
            p = add_paragraph(doc, title, style='Title', align=WD_ALIGN_PARAGRAPH.CENTER)
            p.paragraph_format.line_spacing = 2
            continue
        if stripped.startswith('## '):
            flush()
            heading = stripped[3:].strip()
            in_ref = heading.lower() == 'references'
            add_paragraph(doc, heading, style='Heading 1', align=WD_ALIGN_PARAGRAPH.CENTER)
            continue
        if stripped.startswith('### '):
            flush()
            add_paragraph(doc, stripped[4:].strip(), style='Heading 2', align=WD_ALIGN_PARAGRAPH.LEFT)
            continue
        if stripped.startswith('- '):
            flush()
            add_bullet(doc, stripped[2:].strip())
            continue
        if stripped.startswith('|'):
            continue
        if stripped.startswith('---'):
            continue
        if stripped.endswith('  '):
            stripped = stripped[:-2]
        buffer.append(clean_markdown_line(stripped))
    flush()


def add_csv_table(doc, title, rows, columns, max_rows=None):
    add_paragraph(doc, title, style='Heading 1', align=WD_ALIGN_PARAGRAPH.CENTER)
    show_rows = rows if max_rows is None else rows[:max_rows]
    table = doc.add_table(rows=1, cols=len(columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, col in enumerate(columns):
        set_cell_text(table.rows[0].cells[i], col, bold=True)
    for row in show_rows:
        cells = table.add_row().cells
        for i, col in enumerate(columns):
            set_cell_text(cells[i], row.get(col, ''))
    if max_rows is not None and len(rows) > max_rows:
        add_paragraph(doc, f'Note. Table truncated in Word draft to first {max_rows} rows; full CSV is included in the manuscript package.', indent=False)


def add_figures(doc):
    add_paragraph(doc, 'Figures', style='Heading 1', align=WD_ALIGN_PARAGRAPH.CENTER)
    for idx, (path, caption) in enumerate(FIGURES):
        if idx:
            doc.add_page_break()
        p = add_paragraph(doc, caption, align=WD_ALIGN_PARAGRAPH.CENTER)
        p.runs[0].bold = True
        if path.exists():
            pic_p = doc.add_paragraph()
            pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = pic_p.add_run()
            run.add_picture(str(path), width=Inches(6.5))
        note = 'Note. Solid paths indicate likelihood-based 95% confidence intervals excluding zero; dashed paths indicate intervals including zero; dotted paths indicate incomplete interval bounds and are interpreted descriptively.'
        add_paragraph(doc, note, indent=False)


def build_docx():
    doc = Document()
    configure_document(doc)
    md_text = MD.read_text(encoding='utf-8')
    add_markdown_body(doc, md_text)
    doc.add_page_break()
    fit_rows = read_csv(FIT)
    add_csv_table(doc, 'Table 1\nModel-Family Fit and Data Structure', fit_rows, list(fit_rows[0].keys()))
    doc.add_page_break()
    path_rows = read_csv(PATHS)
    path_cols = ['model_family', 'path', 'estimate', 'ci_lower_95', 'ci_upper_95', 'inference']
    add_csv_table(doc, 'Table 2\nStructural Path Estimates and Confidence-Interval Inference', path_rows, path_cols)
    doc.add_page_break()
    add_figures(doc)
    DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX)
    DATA_PKG.mkdir(parents=True, exist_ok=True)
    shutil.copy2(DOCX, DATA_DOCX)
    ONEDRIVE.mkdir(parents=True, exist_ok=True)
    shutil.copy2(DOCX, ONEDRIVE_DOCX)


def update_docs():
    addition = f"\n- Word manuscript generated: `PAPER_A_APA7_MODEL_FAMILY_FULL_MANUSCRIPT_SCAFFOLD_{DATE}.docx`.\n"
    if README.exists():
        text = README.read_text(encoding='utf-8')
        if DOCX.name not in text:
            text = text.replace('## Main files\n', '## Main files\n\n- `' + DOCX.name + '`\n')
            text = text.rstrip() + addition
            README.write_text(text + '\n', encoding='utf-8')
    current_note = f"""\n## 2026-06-15 Paper A Word manuscript\n\n- Generated Word manuscript: `paper_a/manuscript/target_journal/apa7_model_family_full_manuscript_scaffold_20260615/{DOCX.name}`.\n- Mirrored Word file to OneDrive manuscript package.\n- Word draft includes APA-ish manuscript styling, manuscript text, fit/path tables, and three model-family figures.\n"""
    text = CURRENT.read_text(encoding='utf-8') if CURRENT.exists() else '# Current Project State\n'
    if '2026-06-15 Paper A Word manuscript' not in text:
        CURRENT.write_text(text.rstrip() + '\n' + current_note, encoding='utf-8')


def main():
    build_docx()
    update_docs()
    print(f'Wrote Word manuscript: {DOCX}')
    print(f'Copied data package Word manuscript: {DATA_DOCX}')
    print(f'Copied OneDrive Word manuscript: {ONEDRIVE_DOCX}')


if __name__ == '__main__':
    main()
