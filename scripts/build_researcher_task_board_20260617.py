#!/usr/bin/env python3
"""Build Track Changes-centered Paper A/B researcher workflow artifacts."""

from __future__ import annotations

import shutil
import os
from datetime import date
from pathlib import Path
from tempfile import TemporaryDirectory
from xml.etree import ElementTree as ET
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.formatting.rule import FormulaRule
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.worksheet.datavalidation import DataValidation


DEFAULT_ROOT = Path(
    "/Users/newhosung/Library/CloudStorage/OneDrive-SharedLibraries-"
    "ThePennsylvaniaStateUniversity/AI Adoption Meta Analysis - Documents"
)
ROOT = Path(os.environ.get("AI_ADOPTION_ROOT", str(DEFAULT_ROOT))).expanduser()
STORAGE_LABEL = os.environ.get("AI_ADOPTION_STORAGE_LABEL", "OneDrive")
WORK_DIR = ROOT / "00_INDEX" / "2026-06-17_Paper_A_B_work_allocation"
SHARED = WORK_DIR / "00_shared"
WORKING_MANUSCRIPT_DIR = WORK_DIR / "01_working_manuscript"
TRACKING_DIR = WORK_DIR / "02_tracking"
REFERENCES_DIR = WORK_DIR / "References" / "Paper_A"
ARCHIVE = WORK_DIR / "99_archive"
NOTES_ARCHIVE = ARCHIVE / "notes_workflow_before_track_changes_20260618"

SOURCE_PAPER_A_MANUSCRIPT = ROOT / (
    "05_manuscripts/Paper_A/2026-06-15_apa7_model_family_full_manuscript_scaffold/"
    "PAPER_A_LONGTABLE_PANEL_SUBMISSION_DRAFT_20260616.docx"
)
PAPER_A_ORIGINAL_DOCX = WORKING_MANUSCRIPT_DIR / "Paper_A_원본_DO_NOT_EDIT_20260617.docx"
PAPER_A_TRACK_DOCX = WORKING_MANUSCRIPT_DIR / "Paper_A_작업원고_TRACK_CHANGES_20260617.docx"
TRACK_README_MD = WORKING_MANUSCRIPT_DIR / "README_Track_Changes_작업방법_20260617.md"

BOARD_XLSX = SHARED / "연구자_작업보드_20260617.xlsx"
GUIDE_DOCX = SHARED / "연구자_작업안내_20260617.docx"
GUIDE_MD = SHARED / "연구자_작업안내_20260617.md"
README_MD = SHARED / "README_먼저_읽어주세요.md"
PROCESS_DOCX = SHARED / "Paper_A_포함판단_과정_설명서_20260617.docx"
PROCESS_MD = SHARED / "Paper_A_포함판단_과정_설명서_20260617.md"
EMAIL_DOCX = SHARED / "연구자_안내이메일_초안_20260617.docx"
EMAIL_MD = SHARED / "연구자_안내이메일_초안_20260617.md"
ROOT_README_MD = ROOT / "README_LOCAL_WORKFLOW_20260618.md"

DECISION_LOG_XLSX = TRACKING_DIR / "Paper_A_결정로그_20260617.xlsx"
REFERENCE_MATRIX_XLSX = TRACKING_DIR / "Paper_A_Reference_Matrix_20260617.xlsx"
APA_JARS_XLSX = TRACKING_DIR / "Paper_A_APA_JARS_Checklist_20260617.xlsx"

LOCAL_REPO_ROOT = Path(__file__).resolve().parents[1]
LOCAL_PAPER_A_REFERENCE_PDFS = LOCAL_REPO_ROOT / "references" / "paper_a_apa7_evidence_review_20260615" / "pdfs"

STATUS_VALUES = ["시작 전", "진행 중", "검토 요청", "완료", "막힘"]
PRIORITY_VALUES = ["높음", "보통", "낮음"]
REVIEW_VALUES = ["직접 수정", "댓글만", "직접 수정+댓글", "보조 추적표", "PI 결정 필요"]


def root_rel(path: Path) -> str:
    return path.relative_to(ROOT).as_posix()


REFERENCE_FILES = {
    "REF-A-COUNT": {
        "paper": "Paper A",
        "label": "포함 수/PRISMA 기준 파일",
        "path": ROOT / "90_repository_mirror/journal_AI-adoption_meta",
        "must": "\n".join(
            [
                "루트: 90_repository_mirror/journal_AI-adoption_meta",
                "paper_a/PRISMA_COUNTS_LOCK_20260615.md",
                "paper_a/PRISMA_COUNTS_REVIEW_NEEDED_20260615.md",
                "data/02_screening/screening_summary.json",
                "data/02_screening/human_screening_results_consolidated.csv",
                "paper_a/public_data_repository_20260615/6_PRISMA_Count_Lock/PRISMA_COUNTS_LOCK_20260615.md",
            ]
        ),
        "note": "숫자 차이를 새로 추측하려는 것이 아니라 Methods/Results/PRISMA 표현을 같은 기준으로 잠그기 위한 파일입니다.",
    },
    "REF-A-PROCESS": {
        "paper": "Paper A",
        "label": "포함 판단 과정 설명서",
        "path": PROCESS_DOCX,
        "must": "\n".join(
            [
                "00_shared/Paper_A_포함판단_과정_설명서_20260617.docx",
                "00_shared/Paper_A_포함판단_과정_설명서_20260617.md",
                "03_source_adjudication/Paper_A/2026-06-14_human_process_candidate_audit/REVIEW_THIS_PAPER_A_HUMAN_PROCESS_AUDIT_20260614.docx",
                "03_source_adjudication/Paper_A/2026-06-14_human_style_source_adjudication/PAPER_A_HUMAN_STYLE_SOURCE_ADJUDICATION_GUIDE_20260614.docx",
            ]
        ),
        "note": "A2에서 인간 검토, AI 보조 선별, 최종 포함 기준을 Methods 문장으로 정리할 때 봅니다.",
    },
    "REF-A-METHODS": {
        "paper": "Paper A",
        "label": "Methods/Results 원고와 결과 근거",
        "path": PAPER_A_TRACK_DOCX,
        "must": "\n".join(
            [
                "01_working_manuscript/Paper_A_작업원고_TRACK_CHANGES_20260617.docx",
                "01_working_manuscript/Paper_A_원본_DO_NOT_EDIT_20260617.docx",
                "05_manuscripts/Paper_A/2026-06-15_apa7_model_family_full_manuscript_scaffold/tables/",
                "05_manuscripts/Paper_A/2026-06-15_apa7_model_family_full_manuscript_scaffold/supplemental_diagnostics/",
                "03_source_adjudication/Paper_A/2026-06-15_researcher_approved_s048_analysis/PAPER_A_SOURCE_CORRECTED_COMPLETE_CASE_TSSEM_20260615.md",
            ]
        ),
        "note": "full10은 이론적 목표 구조/근거 지도이고, core7/trust6는 경험적으로 추정 가능한 모델군 경로라는 구분을 확인합니다.",
    },
    "REF-A-INTRO-REFERENCES": {
        "paper": "Paper A",
        "label": "선행연구/References 보강 자료",
        "path": REFERENCES_DIR,
        "must": "\n".join(
            [
                "00_INDEX/2026-06-17_Paper_A_B_work_allocation/References/Paper_A/README.md",
                "00_INDEX/2026-06-17_Paper_A_B_work_allocation/References/Paper_A/pdfs/",
                "02_tracking/Paper_A_Reference_Matrix_20260617.xlsx",
                "90_repository_mirror/journal_AI-adoption_meta/references/paper_a_apa7_evidence_review_20260615/",
                "90_repository_mirror/journal_AI-adoption_meta/references/paper_a_model_family_masem_20260614/",
            ]
        ),
        "note": "R2가 중심이 되어 PDF를 읽고 원고 어느 문장을 보강하는지 Reference Matrix에 남깁니다.",
    },
    "REF-A-DISCUSSION": {
        "paper": "Paper A",
        "label": "Discussion/한계/시사점 보강 자료",
        "path": ROOT
        / "05_manuscripts/Paper_A/2026-06-15_apa7_model_family_full_manuscript_scaffold/"
        / "PAPER_A_THEORY_DISCUSSION_WRITING_GUIDE_KR_20260616.docx",
        "must": "\n".join(
            [
                "01_working_manuscript/Paper_A_작업원고_TRACK_CHANGES_20260617.docx",
                "05_manuscripts/Paper_A/2026-06-15_apa7_model_family_full_manuscript_scaffold/PAPER_A_THEORY_DISCUSSION_WRITING_GUIDE_KR_20260616.docx",
                "05_manuscripts/Paper_A/2026-06-15_apa7_model_family_full_manuscript_scaffold/supplemental_diagnostics/",
                "02_tracking/Paper_A_Reference_Matrix_20260617.xlsx",
            ]
        ),
        "note": "Results 범위를 넘지 않도록 Discussion, limitations, implications를 보강합니다.",
    },
    "REF-A-APA": {
        "paper": "Paper A",
        "label": "APA 7/JARS/표그림/References 점검 자료",
        "path": APA_JARS_XLSX,
        "must": "\n".join(
            [
                "01_working_manuscript/Paper_A_작업원고_TRACK_CHANGES_20260617.docx",
                "02_tracking/Paper_A_APA_JARS_Checklist_20260617.xlsx",
                "05_manuscripts/Paper_A/2026-06-15_apa7_model_family_full_manuscript_scaffold/tables/",
                "05_manuscripts/Paper_A/2026-06-15_apa7_model_family_full_manuscript_scaffold/figures/",
                "90_repository_mirror/journal_AI-adoption_meta/references/paper_a_apa7_evidence_review_20260615/paper_a_crossref_apa7_reference_list_20260615.md",
            ]
        ),
        "note": "R4가 APA 7 professional style, JARS, 표/그림 첫 언급, References 누락, Conclusion 일관성을 봅니다.",
    },
    "REF-B-MAIN": {
        "paper": "Paper B",
        "label": "Paper B 최신 원고 후보",
        "path": ROOT
        / "05_manuscripts/Paper_B/2026-06-12_target_journal/"
        / "PAPER_B_RESEARCH_SYNTHESIS_METHODS_TARGET_DRAFT_20260612.docx",
        "must": "05_manuscripts/Paper_B/2026-06-12_target_journal/PAPER_B_RESEARCH_SYNTHESIS_METHODS_TARGET_DRAFT_20260612.docx",
        "note": "Paper B 모델 비교 결과는 제한된 근거 정리로 읽어야 합니다.",
    },
    "REF-B-REFERENCE": {
        "paper": "Paper B",
        "label": "Paper B source-anchored reference 후보",
        "path": ROOT
        / "03_source_adjudication/Paper_B/reference_standard_candidates/"
        / "Paper2_Human_Final_Consensus_20260605_v2/Paper2_Human_Final_Consensus_Reference_Document_20260605_v2.md",
        "must": "03_source_adjudication/Paper_B/reference_standard_candidates/Paper2_Human_Final_Consensus_20260605_v2/Paper2_Human_Final_Consensus_Reference_Document_20260605_v2.md",
        "note": "'절대적 정답'이 아니라 출처를 확인해 만든 최종 기준표로 설명합니다.",
    },
    "REF-B-WORKBOOKS": {
        "paper": "Paper B",
        "label": "R1-R4 최신 워크북 묶음",
        "path": ROOT / "01_workbooks/latest_collections/20260605_R1_R4",
        "must": "01_workbooks/latest_collections/20260605_R1_R4/",
        "note": "원자료 셀은 덮어쓰지 않고 결정은 별도 로그나 작업보드에 남깁니다.",
    },
    "REF-B-STEP5": {
        "paper": "Paper B",
        "label": "Paper B 모델 비교 분석 결과",
        "path": ROOT / "04_analysis_outputs/Paper_B/analysis_input_20260530",
        "must": "04_analysis_outputs/Paper_B/analysis_input_20260530/",
        "note": "전체 구성개념/전체 행에서 구조모형 대체가 안정적이라는 최종 주장으로 바로 쓰면 안 됩니다.",
    },
}


TASKS = [
    {
        "id": "A1",
        "paper": "Paper A",
        "owner": "R1",
        "support": "R2",
        "priority": "높음",
        "section": "Methods / Results / PRISMA",
        "title": "포함 연구 수와 PRISMA 표현 잠금",
        "do": "225개 포함 행, 중복 DOI 1건 병합, 224개 고유 포함 보고서/연구라는 단위 구분이 Methods/Results/PRISMA에서 일관되도록 확인합니다.",
        "edit_mode": "직접 수정+댓글",
        "track": "작업원고 해당 문장에 [A1-R1] 댓글을 붙이고, 최종 숫자 표현은 결정로그에 남깁니다.",
        "central_output": "01_working_manuscript/Paper_A_작업원고_TRACK_CHANGES_20260617.docx; 02_tracking/Paper_A_결정로그_20260617.xlsx",
        "reference": "REF-A-COUNT",
        "done": "원고 문장, PRISMA 표현, 결정로그가 같은 기준을 쓰면 완료입니다.",
    },
    {
        "id": "A2",
        "paper": "Paper A",
        "owner": "R1",
        "support": "R2",
        "priority": "높음",
        "section": "Methods의 screening/eligibility 절차",
        "title": "인간 판단 과정과 AI 보조 선별 과정 설명",
        "do": "새 판단을 만들지 말고, 이미 진행된 인간 검토와 AI 보조 선별이 독자에게 분리되어 보이도록 Methods 문장을 정리합니다.",
        "edit_mode": "직접 수정+댓글",
        "track": "작업원고 Methods 문단에 [A2-R1] 댓글을 붙이고 포함판단 과정 설명서와 대조합니다.",
        "central_output": "01_working_manuscript/Paper_A_작업원고_TRACK_CHANGES_20260617.docx",
        "reference": "REF-A-PROCESS",
        "done": "인간 검토, AI 보조 역할, 최종 포함 기준이 분리되어 설명되면 완료입니다.",
    },
    {
        "id": "A3",
        "paper": "Paper A",
        "owner": "R1",
        "support": "R3",
        "priority": "높음",
        "section": "Methods / Results / Tables / Figures",
        "title": "방법론 적합성 및 결과 해석 리뷰",
        "do": "모델군 MASEM 접근, full10/core7/trust6 구분, complete-case 근거, 표/그림 해석이 방법론적으로 방어 가능한지 봅니다.",
        "edit_mode": "댓글만",
        "track": "방법론상 PI 판단이 필요한 문장에는 [A3-R1] 댓글을 붙이고, 확실한 표현 수정만 Track Changes로 반영합니다.",
        "central_output": "01_working_manuscript/Paper_A_작업원고_TRACK_CHANGES_20260617.docx; 02_tracking/Paper_A_결정로그_20260617.xlsx",
        "reference": "REF-A-METHODS",
        "done": "방법론상 PI 확인 대상과 즉시 수정 가능한 문장이 구분되면 완료입니다.",
    },
    {
        "id": "A4",
        "paper": "Paper A",
        "owner": "R2",
        "support": "R4",
        "priority": "높음",
        "section": "Introduction / Theoretical Background / References",
        "title": "Introduction과 선행연구 보강",
        "do": "AI adoption, UTAUT/TAM, trust/reliance, attitude, anxiety/self-efficacy, MASEM 선행연구를 읽고 원고 문장을 보강합니다.",
        "edit_mode": "직접 수정+댓글",
        "track": "작업원고 Introduction에 [A4-R2] 댓글 또는 직접 수정으로 남기고, 인용 근거는 Reference Matrix에 정리합니다.",
        "central_output": "01_working_manuscript/Paper_A_작업원고_TRACK_CHANGES_20260617.docx; 02_tracking/Paper_A_Reference_Matrix_20260617.xlsx",
        "reference": "REF-A-INTRO-REFERENCES",
        "done": "원고 보강 문장과 해당 근거 문헌이 Reference Matrix에 연결되면 완료입니다.",
    },
    {
        "id": "A5",
        "paper": "Paper A",
        "owner": "R3",
        "support": "R2",
        "priority": "보통",
        "section": "Discussion / Limitations / Implications",
        "title": "Discussion, 한계, 시사점 보강",
        "do": "Results에서 실제로 말할 수 있는 범위를 넘지 않도록 하면서 이론적 기여, 교육/HRD 시사점, 방법론적 한계를 보강합니다.",
        "edit_mode": "직접 수정+댓글",
        "track": "작업원고 Discussion 이후 문단에 [A5-R3] 댓글을 붙이고, 과장 위험이 있는 문장은 댓글로 표시합니다.",
        "central_output": "01_working_manuscript/Paper_A_작업원고_TRACK_CHANGES_20260617.docx",
        "reference": "REF-A-DISCUSSION",
        "done": "Discussion/limitations/implications 수정이 Results 근거와 연결되면 완료입니다.",
    },
    {
        "id": "A6",
        "paper": "Paper A",
        "owner": "R4",
        "support": "R1",
        "priority": "높음",
        "section": "APA 7 / JARS / Tables / Figures / References / Conclusion",
        "title": "APA 7, JARS, 표/그림, References, Conclusion 점검",
        "do": "APA 7 professional style, JARS 보고 항목, 표/그림 첫 언급, reference list 누락, Conclusion과 Results의 일관성을 확인합니다.",
        "edit_mode": "보조 추적표",
        "track": "작업원고에 [A6-R4] 댓글을 붙이고 반복 점검 항목은 APA/JARS Checklist에 정리합니다.",
        "central_output": "01_working_manuscript/Paper_A_작업원고_TRACK_CHANGES_20260617.docx; 02_tracking/Paper_A_APA_JARS_Checklist_20260617.xlsx",
        "reference": "REF-A-APA",
        "done": "Checklist와 작업원고 댓글/수정이 연결되면 완료입니다.",
    },
    {
        "id": "B1",
        "paper": "Paper B",
        "owner": "R1",
        "support": "R4",
        "priority": "높음",
        "section": "Paper B Methods/Results",
        "title": "Paper B 진행 순서 확인",
        "do": "원 코딩 보존, 코더 차이 정리, 원문 확인 결정, 최종 기준표 확정, 모델 비교 순서가 뒤섞이지 않았는지 확인합니다.",
        "edit_mode": "별도 산출물",
        "track": "Paper B는 이번 Track Changes 전환의 중심 대상이 아니므로 작업보드와 별도 산출물로 추적합니다.",
        "central_output": "R1/B1_PaperB_진행순서_확인메모_20260617.md",
        "reference": "REF-B-REFERENCE",
        "done": "잘못된 순서 표현과 수정 문장이 정리되면 완료입니다.",
    },
    {
        "id": "B2",
        "paper": "Paper B",
        "owner": "R2",
        "support": "R3",
        "priority": "높음",
        "section": "Paper B coder comparison",
        "title": "코더 간 차이 정리",
        "do": "포함/제외, 표본 선택, 구성개념 분류, 근거 유형 차이를 연구자가 읽을 수 있게 정리합니다.",
        "edit_mode": "별도 산출물",
        "track": "Paper B 작업보드와 별도 정리표로 추적합니다.",
        "central_output": "R2/B2_코더차이_정리표_20260617.xlsx",
        "reference": "REF-B-WORKBOOKS",
        "done": "차이 유형, 연구 ID, 원문 확인 필요 여부가 보이면 완료입니다.",
    },
    {
        "id": "B3",
        "paper": "Paper B",
        "owner": "R3",
        "support": "R2, R4",
        "priority": "높음",
        "section": "Paper B source adjudication",
        "title": "원문 확인 결정 누락 항목 찾기",
        "do": "코더 간 차이가 원문 확인 결정으로 연결되었는지 확인하고 누락 항목을 PI 확인 목록으로 올립니다.",
        "edit_mode": "별도 산출물",
        "track": "Paper B 작업보드와 별도 목록으로 추적합니다.",
        "central_output": "R3/B3_원문확인결정_누락목록_20260617.xlsx",
        "reference": "REF-B-REFERENCE",
        "done": "원문 확인 결정이 있음/없음/확인 필요로 나뉘면 완료입니다.",
    },
    {
        "id": "B4",
        "paper": "Paper B",
        "owner": "R4",
        "support": "R1",
        "priority": "높음",
        "section": "Paper B model comparison",
        "title": "모델 비교 결과 원고용 근거 묶음",
        "do": "2,043행 M1-R 결과와 예외 항목을 본문 표/부록 후보로 정리합니다.",
        "edit_mode": "별도 산출물",
        "track": "Paper B 작업보드와 별도 정리표로 추적합니다.",
        "central_output": "R4/B4_모델비교_원고근거묶음_20260617.xlsx",
        "reference": "REF-B-STEP5",
        "done": "원고에 넣을 표 후보와 주의 문장이 함께 있으면 완료입니다.",
    },
    {
        "id": "B5",
        "paper": "Paper B",
        "owner": "R1",
        "support": "R4",
        "priority": "높음",
        "section": "Paper B claim boundary",
        "title": "구조모형 주장 범위 점검",
        "do": "core-6, core7/core8 확장, 전체 구성개념/전체 행 확인 단계를 분리해서 원고 문장을 점검합니다.",
        "edit_mode": "별도 산출물",
        "track": "Paper B 작업보드와 별도 점검표로 추적합니다.",
        "central_output": "R1/B5_구조모형_주장범위_점검표_20260617.docx",
        "reference": "REF-B-MAIN",
        "done": "사용 가능 문장과 금지 문장이 구분되면 완료입니다.",
    },
    {
        "id": "B6",
        "paper": "Paper B",
        "owner": "R4",
        "support": "R1",
        "priority": "보통",
        "section": "Paper B tables/appendices",
        "title": "Paper B 표/부록 후보와 결과 문장 점검",
        "do": "결과표, 부록 후보, 모델 비교 결과 설명이 연구자가 읽기 쉽게 정리되어 있는지 확인합니다.",
        "edit_mode": "별도 산출물",
        "track": "Paper B 작업보드와 별도 점검표로 추적합니다.",
        "central_output": "R4/B6_PaperB_표부록_결과문장_점검표_20260617.xlsx",
        "reference": "REF-B-MAIN",
        "done": "본문 문장 후보, 부록 후보, 수정 필요 표현이 구분되면 완료입니다.",
    },
]


REFERENCES = [
    {
        "short": "PRISMA 2020",
        "use": "Paper A 체계적 문헌고찰/메타분석 보고와 흐름도 확인",
        "citation": "Page MJ et al. (2021). The PRISMA 2020 statement. BMJ, 372, n71.",
        "url": "https://www.bmj.com/content/372/bmj.n71",
    },
    {
        "short": "APA JARS",
        "use": "원고 보고 항목과 양적 연구/메타분석 보고 확인",
        "citation": "APA Style. Journal Article Reporting Standards.",
        "url": "https://apastyle.apa.org/jars",
    },
]


def ensure_dirs() -> None:
    for folder in [
        SHARED,
        WORKING_MANUSCRIPT_DIR,
        TRACKING_DIR,
        REFERENCES_DIR / "pdfs",
        REFERENCES_DIR / "to_request",
        ARCHIVE,
        NOTES_ARCHIVE,
    ]:
        folder.mkdir(parents=True, exist_ok=True)


def archive_notes_workflow_files() -> list[str]:
    moved: list[str] = []
    for role in ["R1", "R2", "R3", "R4"]:
        role_dir = WORK_DIR / role
        if not role_dir.exists():
            continue
        for src in sorted(role_dir.glob("*")):
            if src.is_dir():
                continue
            dst_dir = NOTES_ARCHIVE / role
            dst_dir.mkdir(parents=True, exist_ok=True)
            dst = dst_dir / src.name
            if dst.exists():
                dst = dst_dir / f"{src.stem}_archived{src.suffix}"
            shutil.move(str(src), str(dst))
            moved.append(root_rel(src))
    for legacy in [SHARED / "완료메모_템플릿_20260617.md"]:
        if legacy.exists():
            dst = NOTES_ARCHIVE / legacy.name
            if dst.exists():
                dst = NOTES_ARCHIVE / f"{legacy.stem}_archived{legacy.suffix}"
            shutil.move(str(legacy), str(dst))
            moved.append(root_rel(legacy))
    return moved


def set_track_revisions(docx_path: Path) -> None:
    with TemporaryDirectory() as tmp:
        tmp_dir = Path(tmp)
        with ZipFile(docx_path) as zf:
            zf.extractall(tmp_dir)
        settings = tmp_dir / "word" / "settings.xml"
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        ET.register_namespace("w", ns["w"])
        tree = ET.parse(settings)
        root = tree.getroot()
        if root.find("w:trackRevisions", ns) is None:
            root.append(ET.Element(f"{{{ns['w']}}}trackRevisions"))
        tree.write(settings, encoding="UTF-8", xml_declaration=True)
        with ZipFile(docx_path, "w", ZIP_DEFLATED) as zf:
            for item in sorted(tmp_dir.rglob("*")):
                if item.is_file():
                    zf.write(item, item.relative_to(tmp_dir).as_posix())


def build_working_manuscripts() -> None:
    if not SOURCE_PAPER_A_MANUSCRIPT.exists():
        raise FileNotFoundError(SOURCE_PAPER_A_MANUSCRIPT)
    shutil.copy2(SOURCE_PAPER_A_MANUSCRIPT, PAPER_A_ORIGINAL_DOCX)
    shutil.copy2(SOURCE_PAPER_A_MANUSCRIPT, PAPER_A_TRACK_DOCX)
    set_track_revisions(PAPER_A_TRACK_DOCX)
    TRACK_README_MD.write_text(
        "\n".join(
            [
                "# Paper A Track Changes 작업 방법",
                "",
                "이 폴더의 정본 수정 파일은 `Paper_A_작업원고_TRACK_CHANGES_20260617.docx`입니다.",
                "",
                "## 파일 구분",
                "",
                "- `Paper_A_원본_DO_NOT_EDIT_20260617.docx`: 비교와 복구를 위한 원본입니다. 수정하지 말아 주세요.",
                "- `Paper_A_작업원고_TRACK_CHANGES_20260617.docx`: R1-R4가 함께 수정하는 작업 원고입니다. Word의 Track Changes를 켠 상태로 수정해 주세요.",
                "",
                "## 댓글 표기",
                "",
                "- R1: `[A1-R1]`, `[A2-R1]`, `[A3-R1]`",
                "- R2: `[A4-R2]`",
                "- R3: `[A5-R3]`",
                "- R4: `[A6-R4]`",
                "",
                "## 원칙",
                "",
                "- 원고 문장 수정은 가능하면 이 파일 안에서 직접 해 주세요.",
                "- PI 판단이 필요한 부분은 댓글로 남겨 주세요.",
                "- 포함 수, 방법론 판단, APA/JARS, 참고문헌 근거처럼 표로 봐야 하는 내용만 `02_tracking`의 보조 파일에 남겨 주세요.",
                "- 별도 Notes 파일은 기본 산출물이 아닙니다.",
            ]
        )
        + "\n",
        encoding="utf-8",
    )


def set_korean_font(run, font_name: str = "Malgun Gothic") -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Number"]:
        style = doc.styles[style_name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal = doc.styles["Normal"]
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18
    h1 = doc.styles["Heading 1"]
    h1.font.size = Pt(16)
    h1.font.color.rgb = RGBColor(46, 116, 181)
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(8)
    h2 = doc.styles["Heading 2"]
    h2.font.size = Pt(13)
    h2.font.color.rgb = RGBColor(46, 116, 181)
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(6)


def add_title(doc: Document, title: str, subtitle: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    run.font.size = Pt(21)
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 78, 121)
    set_korean_font(run)
    if subtitle:
        s = doc.add_paragraph()
        sr = s.add_run(subtitle)
        set_korean_font(sr)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_korean_font(r)


def add_numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(text)
    set_korean_font(r)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        p = table.rows[0].cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        r.bold = True
        set_korean_font(r)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            r = p.add_run(value)
            set_korean_font(r)


def build_guide_docx() -> None:
    doc = Document()
    style_doc(doc)
    add_title(
        doc,
        "Paper A/B 연구자 작업안내",
        "Paper A는 하나의 Track Changes 원고를 중심으로 수정하고, 작업보드는 상태와 PI 확인을 추적하기 위한 문서입니다.",
    )

    doc.add_heading("1. 이번 운영 방식의 핵심", level=1)
    add_bullet(doc, "원고 문장 수정은 `01_working_manuscript/Paper_A_작업원고_TRACK_CHANGES_20260617.docx` 한 파일에서 진행합니다.")
    add_bullet(doc, "연구자별 Notes 파일은 기본 산출물이 아닙니다. 문장 수정은 Word Track Changes와 댓글로 남겨 주세요.")
    add_bullet(doc, "작업보드는 누가 어느 섹션을 봤는지, 상태가 무엇인지, PI 확인이 필요한지 추적하는 용도입니다.")
    add_bullet(doc, "결정로그, Reference Matrix, APA/JARS Checklist는 원고 안에 담기 어려운 근거 표만 정리합니다.")

    doc.add_heading("2. 먼저 여실 파일", level=1)
    add_table(
        doc,
        ["구분", "파일", "역할"],
        [
            ["작업 원고", "01_working_manuscript/Paper_A_작업원고_TRACK_CHANGES_20260617.docx", "R1-R4가 실제로 수정하는 중심 파일"],
            ["원본", "01_working_manuscript/Paper_A_원본_DO_NOT_EDIT_20260617.docx", "비교/복구용 원본. 수정하지 않음"],
            ["작업보드", "00_shared/연구자_작업보드_20260617.xlsx", "상태, PI 확인, 원고 내 위치 추적"],
            ["포함 과정", "00_shared/Paper_A_포함판단_과정_설명서_20260617.docx", "A2 Methods 설명 기준"],
            ["결정로그", "02_tracking/Paper_A_결정로그_20260617.xlsx", "A1/A3처럼 PI 결정이 필요한 기준 기록"],
            ["Reference Matrix", "02_tracking/Paper_A_Reference_Matrix_20260617.xlsx", "A4 선행연구와 원고 보강 위치 연결"],
            ["APA/JARS", "02_tracking/Paper_A_APA_JARS_Checklist_20260617.xlsx", "A6 반복 점검 항목 관리"],
        ],
    )

    doc.add_heading("3. 작업 순서", level=1)
    for item in [
        "작업보드에서 본인 담당 작업과 원고 섹션을 확인합니다.",
        "작업원고를 열고 Word의 Track Changes가 켜져 있는지 확인합니다.",
        "원고 문장을 직접 고치거나 댓글을 남깁니다. 댓글 앞에는 `[A4-R2]`처럼 작업 ID와 역할을 붙입니다.",
        "포함 수, 선행연구 근거, APA/JARS처럼 표로 남길 필요가 있는 내용만 `02_tracking` 파일에 기록합니다.",
        "작업보드 상태를 `진행 중`, `검토 요청`, `완료`, `막힘` 중 하나로 업데이트합니다.",
        "PI 판단이 필요한 부분은 작업보드 `PI 확인 필요`와 Word 댓글 양쪽에 남깁니다.",
    ]:
        add_numbered(doc, item)

    doc.add_heading("4. 역할별 Paper A 배분", level=1)
    add_table(
        doc,
        ["역할", "담당", "원고 안에서 할 일"],
        [
            ["R1", "A1-A3", "Methods/Results, 포함 수, PRISMA, 인간 판단/AI 보조 선별 설명, 방법론 적합성"],
            ["R2", "A4", "Introduction, 이론적 배경, 선행연구와 References 보강"],
            ["R3", "A5", "Discussion, Limitations, Implications 보강과 과장 위험 표시"],
            ["R4", "A6", "APA 7, JARS, 표/그림, References, Conclusion 점검"],
            ["PI", "검토 요청/완료 행", "댓글과 결정로그를 보고 승인, 수정 요청, 추가 확인 결정"],
        ],
    )

    doc.add_heading("5. Paper A 작업별 완료 기준", level=1)
    for task in [t for t in TASKS if t["paper"] == "Paper A"]:
        doc.add_heading(f"{task['id']}. {task['title']}", level=2)
        add_bullet(doc, f"담당: {task['owner']} / 함께 볼 사람: {task['support']}")
        add_bullet(doc, f"원고 섹션: {task['section']}")
        add_bullet(doc, f"작업 방식: {task['edit_mode']}")
        add_bullet(doc, f"원고 내 추적: {task['track']}")
        add_bullet(doc, f"완료 기준: {task['done']}")

    doc.add_heading("6. Paper B", level=1)
    add_bullet(doc, "Paper B는 이번 구조 전환의 중심 대상이 아니므로 기존처럼 작업보드와 별도 산출물로 추적합니다.")
    add_bullet(doc, "Paper B에서도 원자료와 기준표는 덮어쓰지 않고, 수정 제안과 판단 근거만 별도 산출물에 남깁니다.")

    doc.add_heading("7. 주의할 표현", level=1)
    add_bullet(doc, "Paper A의 PRISMA/포함 수는 아직 원고 표현을 잠그는 단계입니다. 제출 준비 완료처럼 쓰지 않습니다.")
    add_bullet(doc, "AI 보조 선별은 최종 포함 결정을 AI가 했다는 뜻으로 쓰지 않습니다.")
    add_bullet(doc, "full10은 이론적 목표 구조와 근거 지도이고, core7/trust6는 현재 경험적으로 추정 가능한 모델군 경로입니다.")

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run(f"생성일 {date.today().isoformat()} | AI Adoption Meta Analysis")
    fr.font.size = Pt(8)
    set_korean_font(fr)
    doc.save(GUIDE_DOCX)


def build_process_doc() -> None:
    md = [
        "# Paper A 포함 판단 과정 설명서",
        "",
        "이 문서는 A2 작업을 위한 연구자용 기준 문서입니다. 새 포함/제외 판단을 만들라는 뜻이 아니라, 이미 진행된 인간 검토와 AI 보조 선별 절차를 Methods에 정확히 설명하기 위한 문서입니다.",
        "",
        "## 핵심 구분",
        "",
        "- 인간 검토: 최종 포함/제외 판단과 원문 확인의 책임이 있는 단계입니다.",
        "- AI 보조 선별: 대량 문헌 후보를 줄이고 검토 우선순위를 돕는 단계입니다. 최종 포함 결정을 AI가 했다고 쓰면 안 됩니다.",
        "- 최종 포함 기준: 현재 작업 기준은 225개 포함 행, 중복 DOI 1건 병합, 224개 고유 포함 보고서/연구입니다.",
        "",
        "## 반드시 확인할 파일",
        "",
        "- 루트: `90_repository_mirror/journal_AI-adoption_meta`",
        "- `data/02_screening/screening_summary.json`",
        "- `data/02_screening/human_screening_results_consolidated.csv`",
        "- `paper_a/PRISMA_COUNTS_LOCK_20260615.md`",
        "- `paper_a/PRISMA_COUNTS_REVIEW_NEEDED_20260615.md`",
        "- `03_source_adjudication/Paper_A/2026-06-14_human_process_candidate_audit/REVIEW_THIS_PAPER_A_HUMAN_PROCESS_AUDIT_20260614.docx`",
        "- `03_source_adjudication/Paper_A/2026-06-14_human_style_source_adjudication/PAPER_A_HUMAN_STYLE_SOURCE_ADJUDICATION_GUIDE_20260614.docx`",
        "",
        "## Methods에서 확인할 질문",
        "",
        "1. 검색, 중복 제거, AI 보조 선별, 인간 검토, 원문 확인, 최종 포함이 순서대로 설명되어 있습니까?",
        "2. AI 보조 선별이 최종 포함 판단처럼 읽히지 않습니까?",
        "3. 657개 인간 검토, 225개 포함 행, 224개 고유 포함 보고서/연구의 관계가 독자에게 분명합니까?",
        "4. PRISMA 수치와 Methods 문장이 서로 다른 기준을 쓰지 않습니까?",
        "5. 불확실한 항목은 원고에 단정적으로 쓰지 않고 PI 확인 대상으로 남겨 두었습니까?",
        "",
        "## 피해야 할 표현",
        "",
        "- AI가 최종 포함 연구를 결정했다.",
        "- 225개 연구가 최종 분석에 포함되었다.",
        "- PRISMA 수치가 이미 최종 확정되었다.",
    ]
    PROCESS_MD.write_text("\n".join(md) + "\n", encoding="utf-8")

    doc = Document()
    style_doc(doc)
    add_title(doc, "Paper A 포함 판단 과정 설명서", "A2 작업자가 Methods 문장을 검수할 때 쓰는 기준 문서입니다.")
    doc.add_heading("1. 핵심 구분", level=1)
    add_bullet(doc, "인간 검토: 최종 포함/제외 판단과 원문 확인의 책임이 있는 단계입니다.")
    add_bullet(doc, "AI 보조 선별: 대량 문헌 후보를 줄이고 검토 우선순위를 돕는 단계입니다. 최종 포함 결정을 AI가 했다고 쓰면 안 됩니다.")
    add_bullet(doc, "최종 포함 기준: 현재 작업 기준은 225개 포함 행, 중복 DOI 1건 병합, 224개 고유 포함 보고서/연구입니다.")
    doc.add_heading("2. 반드시 확인할 파일", level=1)
    for line in REFERENCE_FILES["REF-A-PROCESS"]["must"].splitlines():
        add_bullet(doc, line)
    doc.add_heading("3. Methods에서 확인할 질문", level=1)
    for line in md[22:27]:
        add_numbered(doc, line[3:])
    doc.add_heading("4. 피해야 할 표현과 권장 방향", level=1)
    add_table(
        doc,
        ["피해야 할 표현", "권장 방향"],
        [
            ["AI가 최종 포함 연구를 결정했다.", "AI 보조 선별은 인간 검토를 돕는 후보 축소 및 우선순위화 절차로 설명합니다."],
            ["225개 연구가 최종 분석에 포함되었다.", "225개 포함 행과 224개 고유 포함 보고서/연구는 서로 다른 단위라고 씁니다."],
            ["PRISMA 수치가 이미 최종 확정되었다.", "Methods/Results 투고 전 최종 확인 대상으로 남겨 둡니다."],
        ],
    )
    doc.save(PROCESS_DOCX)


def build_markdown_readmes() -> None:
    GUIDE_MD.write_text(
        "\n".join(
            [
                "# Paper A/B 연구자 작업안내",
                "",
                "Paper A는 하나의 Word 작업원고에서 Track Changes로 수정합니다.",
                "",
                "## 먼저 열 파일",
                "",
                "- 작업 원고: `../01_working_manuscript/Paper_A_작업원고_TRACK_CHANGES_20260617.docx`",
                "- 원본: `../01_working_manuscript/Paper_A_원본_DO_NOT_EDIT_20260617.docx`",
                "- 작업보드: `연구자_작업보드_20260617.xlsx`",
                "- 안내 문서: `연구자_작업안내_20260617.docx`",
                "- 포함 판단 과정: `Paper_A_포함판단_과정_설명서_20260617.docx`",
                "",
                "## 원칙",
                "",
                "- 원고 수정은 작업원고 안에서 Track Changes로 진행합니다.",
                "- 별도 Notes 파일은 기본 산출물이 아닙니다.",
                "- 작업보드는 상태와 PI 확인을 남기는 용도입니다.",
                "- 보조 추적표는 `../02_tracking/`에 있습니다.",
            ]
        )
        + "\n",
        encoding="utf-8",
    )
    README_MD.write_text(
        "\n".join(
            [
                "# 먼저 읽어주세요",
                "",
                "이 폴더는 Paper A/B 연구자 작업보드의 최신 공유 위치입니다.",
                "",
                "1. Paper A 문장 수정은 `../01_working_manuscript/Paper_A_작업원고_TRACK_CHANGES_20260617.docx`에서 해 주세요.",
                "2. `../01_working_manuscript/Paper_A_원본_DO_NOT_EDIT_20260617.docx`는 수정하지 않는 비교용 원본입니다.",
                "3. 담당 작업과 완료 상태는 `연구자_작업보드_20260617.xlsx`에서 확인해 주세요.",
                "4. 포함 판단 과정은 `Paper_A_포함판단_과정_설명서_20260617.docx`에서 확인해 주세요.",
                "5. 결정로그, Reference Matrix, APA/JARS Checklist는 `../02_tracking/`에 있습니다.",
                "6. 기존 연구자별 Notes 방식 파일은 `../99_archive/notes_workflow_before_track_changes_20260618/`에 보관했습니다.",
            ]
        )
        + "\n",
        encoding="utf-8",
    )


def build_root_readme() -> None:
    ROOT_README_MD.write_text(
        "\n".join(
            [
                "# AI Adoption Meta Analysis 로컬 작업 시작 안내",
                "",
                f"이 폴더는 {STORAGE_LABEL} 로컬 동기화 폴더에서 Paper A/B 연구자 작업을 진행하기 위한 작업 루트입니다.",
                "",
                "## 바로 열 파일",
                "",
                "- 작업 폴더: `00_INDEX/2026-06-17_Paper_A_B_work_allocation/`",
                "- Paper A 작업원고: `00_INDEX/2026-06-17_Paper_A_B_work_allocation/01_working_manuscript/Paper_A_작업원고_TRACK_CHANGES_20260617.docx`",
                "- Paper A 원본: `00_INDEX/2026-06-17_Paper_A_B_work_allocation/01_working_manuscript/Paper_A_원본_DO_NOT_EDIT_20260617.docx`",
                "- 연구자 작업보드: `00_INDEX/2026-06-17_Paper_A_B_work_allocation/00_shared/연구자_작업보드_20260617.xlsx`",
                "- 연구자 작업안내: `00_INDEX/2026-06-17_Paper_A_B_work_allocation/00_shared/연구자_작업안내_20260617.docx`",
                "- 이메일 초안: `00_INDEX/2026-06-17_Paper_A_B_work_allocation/00_shared/연구자_안내이메일_초안_20260617.docx`",
                "",
                "## 작업 방식",
                "",
                "1. Paper A 문장 수정은 작업원고 하나에서 Track Changes로 진행합니다.",
                "2. 원본 파일은 비교와 복구용이므로 수정하지 않습니다.",
                "3. 담당 작업, 상태, PI 확인 필요 여부는 작업보드에서 업데이트합니다.",
                "4. 포함 수, 선행연구, APA/JARS처럼 표로 봐야 하는 근거만 `02_tracking` 파일에 남깁니다.",
                "5. 로컬 파일 하이퍼링크는 Office 권한 팝업을 만들 수 있어 넣지 않았습니다. 공유가 필요하면 Google Drive 웹에서 링크를 복사해 작업보드의 웹 링크 열에 붙여 주세요.",
                "",
                "## 주의",
                "",
                "- Google Drive 동기화가 완전히 끝나기 전에는 다른 사람이 같은 파일을 열 수 없을 수 있습니다.",
                "- 로컬 Excel 파일만으로는 행 완료 시 자동 이메일 알림이 가지 않습니다. 알림이 필요하면 Google Sheets/Apps Script 또는 Drive 댓글/공유 알림으로 따로 구성해야 합니다.",
            ]
        )
        + "\n",
        encoding="utf-8",
    )


def build_references_folder() -> None:
    pdf_dir = REFERENCES_DIR / "pdfs"
    request_dir = REFERENCES_DIR / "to_request"
    pdf_dir.mkdir(parents=True, exist_ok=True)
    request_dir.mkdir(parents=True, exist_ok=True)
    if LOCAL_PAPER_A_REFERENCE_PDFS.exists():
        for src in sorted(LOCAL_PAPER_A_REFERENCE_PDFS.glob("*.pdf")):
            dst = pdf_dir / src.name
            if not dst.exists():
                shutil.copy2(src, dst)
    (REFERENCES_DIR / "README.md").write_text(
        "\n".join(
            [
                "# Paper A References",
                "",
                "이 폴더는 Paper A의 Introduction, 이론적 배경, Discussion, References를 보강하기 위한 선행연구 PDF 큐입니다.",
                "",
                "## 사용 방식",
                "",
                "- PDF 파일은 `pdfs/`에 둡니다.",
                "- PDF별 메모를 별도 Notes로 흩뜨리지 말고, `02_tracking/Paper_A_Reference_Matrix_20260617.xlsx`에 정리해 주세요.",
                "- 아직 확보하지 못한 문헌은 `to_request/`에 목록으로 남겨 주세요.",
            ]
        )
        + "\n",
        encoding="utf-8",
    )


def style_sheet(ws, title: str) -> None:
    ws.sheet_view.showGridLines = False
    ws.freeze_panes = "A6"
    ws["A1"] = title
    ws["A1"].font = Font(name="Malgun Gothic", size=16, bold=True, color="FFFFFF")
    ws["A1"].fill = PatternFill("solid", fgColor="1F4E79")
    ws["A1"].alignment = Alignment(vertical="center")
    ws.row_dimensions[1].height = 28
    for row in ws.iter_rows():
        for cell in row:
            cell.font = Font(name="Malgun Gothic", size=10, color="1F1F1F")
            cell.alignment = Alignment(vertical="top", wrap_text=True)


def set_widths(ws, widths: list[int]) -> None:
    for idx, width in enumerate(widths, 1):
        ws.column_dimensions[chr(64 + idx)].width = width


def style_header(ws, row: int, cols: int, color: str = "5B9BD5") -> None:
    for col in range(1, cols + 1):
        cell = ws.cell(row, col)
        cell.font = Font(name="Malgun Gothic", bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor=color)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)


def apply_common_sheet_formatting(wb: Workbook) -> None:
    thin = Side(style="thin", color="D9E2F3")
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
                cell.alignment = Alignment(vertical="top", wrap_text=True)
        ws.sheet_properties.pageSetUpPr.fitToPage = True
        ws.page_setup.fitToWidth = 1
        ws.page_setup.fitToHeight = 0


def add_validations(ws, first: int, last: int) -> None:
    status = DataValidation(type="list", formula1=f'"{",".join(STATUS_VALUES)}"', allow_blank=True)
    priority = DataValidation(type="list", formula1=f'"{",".join(PRIORITY_VALUES)}"', allow_blank=True)
    review = DataValidation(type="list", formula1=f'"{",".join(REVIEW_VALUES)}"', allow_blank=True)
    ws.add_data_validation(status)
    ws.add_data_validation(priority)
    ws.add_data_validation(review)
    status.add(f"K{first}:K{last}")
    priority.add(f"E{first}:E{last}")
    review.add(f"I{first}:I{last}")


def write_task_sheet(ws, tasks: list[dict], title: str) -> None:
    style_sheet(ws, title)
    ws["A2"] = "사용 방법"
    ws["B2"] = "Paper A 문장 수정은 작업원고에서 Track Changes로 하고, 이 보드는 상태와 PI 확인을 추적합니다."
    ws["A3"] = "작업 원고"
    ws["B3"] = "01_working_manuscript/Paper_A_작업원고_TRACK_CHANGES_20260617.docx"
    ws["A4"] = "보조 추적표"
    ws["B4"] = "02_tracking/Paper_A_결정로그_20260617.xlsx / Paper_A_Reference_Matrix_20260617.xlsx / Paper_A_APA_JARS_Checklist_20260617.xlsx"
    headers = [
        "작업 ID",
        "논문",
        "담당자",
        "함께 볼 사람",
        "우선순위",
        "원고 섹션",
        "작업 이름",
        "지금 하실 일",
        "수정 방식",
        "원고 내 추적 방법",
        "상태",
        "완료 표시일",
        "PI 확인 필요",
        "PI 확인일",
        "중앙 작업 파일/보조 파일",
        "참고 위치",
        "막힌 점/메모",
        "웹 링크",
    ]
    start = 6
    for col, header in enumerate(headers, 1):
        ws.cell(start, col, header)
    style_header(ws, start, len(headers))
    for r_idx, task in enumerate(tasks, start + 1):
        ref = REFERENCE_FILES[task["reference"]]
        values = [
            task["id"],
            task["paper"],
            task["owner"],
            task["support"],
            task["priority"],
            task["section"],
            task["title"],
            task["do"],
            task["edit_mode"],
            task["track"],
            "시작 전",
            "",
            f'=IF(OR(K{r_idx}="완료",K{r_idx}="검토 요청"),"예","")',
            "",
            task["central_output"],
            ref["label"],
            "",
            "",
        ]
        for c_idx, value in enumerate(values, 1):
            ws.cell(r_idx, c_idx, value)
        ws.row_dimensions[r_idx].height = 72
    if tasks:
        ws.auto_filter.ref = f"A{start}:R{start + len(tasks)}"
        add_validations(ws, start + 1, start + len(tasks))
        done_fill = PatternFill("solid", fgColor="E2F0D9")
        review_fill = PatternFill("solid", fgColor="FFF2CC")
        blocked_fill = PatternFill("solid", fgColor="F4CCCC")
        ws.conditional_formatting.add(f"A{start + 1}:R{start + len(tasks)}", FormulaRule(formula=[f'$K{start + 1}="완료"'], fill=done_fill))
        ws.conditional_formatting.add(f"A{start + 1}:R{start + len(tasks)}", FormulaRule(formula=[f'$K{start + 1}="검토 요청"'], fill=review_fill))
        ws.conditional_formatting.add(f"A{start + 1}:R{start + len(tasks)}", FormulaRule(formula=[f'$K{start + 1}="막힘"'], fill=blocked_fill))
    set_widths(ws, [10, 10, 10, 14, 10, 28, 28, 48, 18, 52, 14, 14, 14, 14, 58, 30, 42, 34])


def write_reference_sheet(ws) -> None:
    style_sheet(ws, "파일 위치")
    headers = ["ID", "논문", "파일/폴더", f"{STORAGE_LABEL} 루트 기준 경로", "반드시 볼 파일", "주의할 점"]
    start = 4
    for col, header in enumerate(headers, 1):
        ws.cell(start, col, header)
    style_header(ws, start, len(headers), "70AD47")
    for row_idx, (ref_id, ref) in enumerate(REFERENCE_FILES.items(), start + 1):
        values = [ref_id, ref["paper"], ref["label"], root_rel(ref["path"]) if ref.get("path") else "", ref["must"], ref["note"]]
        for col_idx, value in enumerate(values, 1):
            ws.cell(row_idx, col_idx, value)
        ws.row_dimensions[row_idx].height = 110
    ws.auto_filter.ref = f"A{start}:F{start + len(REFERENCE_FILES)}"
    set_widths(ws, [18, 12, 36, 70, 74, 56])


def write_use_sheet(ws) -> None:
    style_sheet(ws, "사용방법")
    rows = [
        ("1", "작업보드에서 본인 담당 작업과 원고 섹션을 확인합니다."),
        ("2", "Paper A는 `01_working_manuscript/Paper_A_작업원고_TRACK_CHANGES_20260617.docx`에서 직접 수정합니다."),
        ("3", "댓글 앞에는 `[A4-R2]`처럼 작업 ID와 역할을 붙입니다."),
        ("4", "보조 추적표가 필요한 작업만 `02_tracking` 파일에 기록합니다."),
        ("5", "작업보드에서 상태와 PI 확인 필요 여부를 업데이트합니다."),
        ("6", f"로컬 파일 하이퍼링크는 Grant Access 팝업을 만들 수 있어 넣지 않았습니다. 웹 공유 링크가 필요하면 {STORAGE_LABEL}에서 링크 복사 후 웹 링크 열에 붙여 주세요."),
    ]
    ws["A3"] = "순서"
    ws["B3"] = "하실 일"
    style_header(ws, 3, 2)
    for idx, row in enumerate(rows, 4):
        ws.cell(idx, 1, row[0])
        ws.cell(idx, 2, row[1])
        ws.row_dimensions[idx].height = 36
    set_widths(ws, [12, 120])


def write_citations_sheet(ws) -> None:
    style_sheet(ws, "참고 기준")
    headers = ["기준", "이 보드에서 쓰는 이유", "인용/출처", "URL"]
    start = 4
    for col, header in enumerate(headers, 1):
        ws.cell(start, col, header)
    style_header(ws, start, len(headers), "70AD47")
    for row_idx, ref in enumerate(REFERENCES, start + 1):
        ws.cell(row_idx, 1, ref["short"])
        ws.cell(row_idx, 2, ref["use"])
        ws.cell(row_idx, 3, ref["citation"])
        ws.cell(row_idx, 4, ref["url"])
        ws.cell(row_idx, 4).hyperlink = ref["url"]
        ws.cell(row_idx, 4).style = "Hyperlink"
        ws.row_dimensions[row_idx].height = 48
    set_widths(ws, [18, 56, 70, 52])


def build_board_workbook() -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "전체보기"
    write_task_sheet(ws, TASKS, "전체 작업보드")
    ws["D2"] = "전체 작업 수"
    ws["E2"] = '=COUNTA(A7:A200)'
    ws["D3"] = "완료"
    ws["E3"] = '=COUNTIF(K7:K200,"완료")'
    ws["D4"] = "막힘"
    ws["E4"] = '=COUNTIF(K7:K200,"막힘")'
    for name, paper in [("Paper A", "Paper A"), ("Paper B", "Paper B")]:
        sheet = wb.create_sheet(name)
        write_task_sheet(sheet, [t for t in TASKS if t["paper"] == paper], f"{name} 작업")
    for role in ["R1", "R2", "R3", "R4"]:
        sheet = wb.create_sheet(role)
        write_task_sheet(sheet, [t for t in TASKS if t["owner"] == role], f"{role} 담당 작업")
    write_reference_sheet(wb.create_sheet("파일위치"))
    write_use_sheet(wb.create_sheet("사용방법"))
    write_citations_sheet(wb.create_sheet("참고 기준"))
    apply_common_sheet_formatting(wb)
    wb.calculation.fullCalcOnLoad = True
    wb.calculation.forceFullCalc = True
    wb.properties.title = "Paper A/B Track Changes 연구자 작업보드"
    wb.properties.subject = "AI Adoption Meta Analysis researcher task board"
    wb.properties.creator = "Codex"
    wb.save(BOARD_XLSX)


def build_decision_log() -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "결정로그"
    style_sheet(ws, "Paper A 결정로그")
    headers = ["작업 ID", "결정 주제", "원고 위치", "현재 기준/후보", "근거 파일", "PI 결정", "결정일", "원고 반영 위치", "비고"]
    for col, header in enumerate(headers, 1):
        ws.cell(4, col, header)
    style_header(ws, 4, len(headers), "4472C4")
    rows = [
        ["A1", "포함 수/PRISMA 표현", "Methods/Results/PRISMA", "225개 포함 행, 중복 DOI 1건 병합, 224개 고유 포함 보고서/연구", "screening_summary.json; human_screening_results_consolidated.csv; PRISMA lock docs", "", "", "", ""],
        ["A2", "인간 검토와 AI 보조 선별 설명", "Methods", "AI는 후보 축소/우선순위화 보조, 최종 포함 판단은 인간 검토/원문 확인", "Paper_A_포함판단_과정_설명서_20260617.docx", "", "", "", ""],
        ["A3", "full10/core7/trust6 표현", "Methods/Results", "full10은 이론적 근거 지도, core7/trust6는 추정 가능한 모델군 경로", "S048 analysis; supplemental diagnostics", "", "", "", ""],
    ]
    for r, row in enumerate(rows, 5):
        for c, value in enumerate(row, 1):
            ws.cell(r, c, value)
        ws.row_dimensions[r].height = 54
    ws.auto_filter.ref = f"A4:I{4 + len(rows)}"
    set_widths(ws, [12, 28, 28, 48, 50, 30, 14, 34, 34])
    apply_common_sheet_formatting(wb)
    wb.save(DECISION_LOG_XLSX)


def build_reference_matrix() -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Reference Matrix"
    style_sheet(ws, "Paper A Reference Matrix")
    headers = ["PDF/문헌", "핵심 주장", "원고 보강 위치", "넣을 문장 후보", "APA reference", "인용 우선순위", "확인 상태", "비고"]
    for col, header in enumerate(headers, 1):
        ws.cell(4, col, header)
    style_header(ws, 4, len(headers), "70AD47")
    examples = [
        ["Cheung_2015_metaSEM_Frontiers.pdf", "MASEM 방법론 근거", "Methods", "", "", "반드시 인용", "시작 전", ""],
        ["Dwivedi_2019_reexamining_UTAUT.pdf", "UTAUT 계열 선행연구", "Introduction", "", "", "선택 인용", "시작 전", ""],
        ["Labadze_2023_AI_chatbots_education_review.pdf", "교육 맥락 AI 도입/활용 논의", "Discussion", "", "", "선택 인용", "시작 전", ""],
    ]
    for r, row in enumerate(examples, 5):
        for c, value in enumerate(row, 1):
            ws.cell(r, c, value)
        ws.row_dimensions[r].height = 60
    ws.auto_filter.ref = f"A4:H{4 + len(examples)}"
    set_widths(ws, [38, 42, 24, 52, 52, 18, 16, 32])
    status = DataValidation(type="list", formula1='"시작 전,확인 중,원고 반영,보류,PI 확인"', allow_blank=True)
    ws.add_data_validation(status)
    status.add("G5:G200")
    apply_common_sheet_formatting(wb)
    wb.save(REFERENCE_MATRIX_XLSX)


def build_apa_jars_checklist() -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "APA JARS Checklist"
    style_sheet(ws, "Paper A APA/JARS Checklist")
    headers = ["영역", "점검 항목", "원고 위치", "현재 상태", "수정 필요", "수정 제안", "PI 확인 필요", "완료 여부"]
    for col, header in enumerate(headers, 1):
        ws.cell(4, col, header)
    style_header(ws, 4, len(headers), "C55A11")
    rows = [
        ["APA 7", "제목/저자/초록/키워드 형식", "Title page/Abstract", "", "", "", "", "시작 전"],
        ["JARS", "연구 설계와 표본/검색/선별 절차 보고 충분성", "Methods", "", "", "", "", "시작 전"],
        ["Tables", "표 제목, 주석, 본문 첫 언급 위치", "Tables/Results", "", "", "", "", "시작 전"],
        ["Figures", "그림 캡션, 번호, 본문 첫 언급 위치", "Figures/Results", "", "", "", "", "시작 전"],
        ["References", "본문 인용과 reference list 일치", "References", "", "", "", "", "시작 전"],
        ["Conclusion", "Conclusion이 Results 범위를 넘지 않는지", "Conclusion", "", "", "", "", "시작 전"],
    ]
    for r, row in enumerate(rows, 5):
        for c, value in enumerate(row, 1):
            ws.cell(r, c, value)
        ws.row_dimensions[r].height = 50
    ws.auto_filter.ref = f"A4:H{4 + len(rows)}"
    set_widths(ws, [18, 44, 26, 28, 14, 48, 16, 16])
    for col in ["E", "G"]:
        dv = DataValidation(type="list", formula1='"예,아니오,판단 필요"', allow_blank=True)
        ws.add_data_validation(dv)
        dv.add(f"{col}5:{col}200")
    status = DataValidation(type="list", formula1='"시작 전,진행 중,완료,PI 확인"', allow_blank=True)
    ws.add_data_validation(status)
    status.add("H5:H200")
    apply_common_sheet_formatting(wb)
    wb.save(APA_JARS_XLSX)


def build_email_docs() -> None:
    subject = "[AI Adoption Meta] Paper A Track Changes 작업원고 및 R1-R4 작업보드 공유"
    body = [
        "안녕하세요, 선생님들.",
        "",
        "Paper A/B 작업 배분 방식을 다시 정리했습니다. Paper A는 더 이상 각자 별도 Notes 파일을 중심으로 수정하지 않고, 하나의 Word 작업원고에서 Track Changes와 댓글을 사용해 수정하는 방식으로 진행하겠습니다.",
        "",
        "가장 먼저 열어 보실 파일은 아래 세 가지입니다.",
        "",
        "1. Paper A 작업원고",
        "   - 위치: 01_working_manuscript/Paper_A_작업원고_TRACK_CHANGES_20260617.docx",
        "   - 용도: R1-R4가 실제로 문장을 수정하고 댓글을 남기는 중심 파일입니다.",
        "   - 부탁: Word의 Track Changes가 켜져 있는지 확인하시고, 댓글 앞에는 [A1-R1], [A4-R2]처럼 작업 ID와 역할을 붙여 주세요.",
        "",
        "2. Paper A 원본",
        "   - 위치: 01_working_manuscript/Paper_A_원본_DO_NOT_EDIT_20260617.docx",
        "   - 용도: 비교와 복구를 위한 원본입니다. 이 파일은 수정하지 말아 주세요.",
        "",
        "3. 연구자 작업보드",
        "   - 위치: 00_shared/연구자_작업보드_20260617.xlsx",
        "   - 용도: 본인 담당 작업, 원고 섹션, 수정 방식, 완료 상태, PI 확인 필요 여부를 확인하고 업데이트하는 파일입니다.",
        "",
        "Paper A 작업 배분은 아래와 같습니다.",
        "",
        "- R1: A1-A3. Methods/Results, 포함 수와 PRISMA 표현, 인간 판단/AI 보조 선별 설명, 방법론 적합성 검토를 맡아 주세요.",
        "- R2: A4. Introduction, 이론적 배경, 선행연구와 References 보강을 맡아 주세요. 인용 근거는 02_tracking/Paper_A_Reference_Matrix_20260617.xlsx에 정리해 주세요.",
        "- R3: A5. Discussion, 한계, 시사점 문단을 보강해 주세요. Results 범위를 넘는 표현은 댓글로 표시해 주세요.",
        "- R4: A6. APA 7, JARS, 표/그림, References, Conclusion을 점검해 주세요. 반복 점검 항목은 02_tracking/Paper_A_APA_JARS_Checklist_20260617.xlsx에 남겨 주세요.",
        "",
        "Paper A의 포함 판단 과정과 Methods 문장 기준은 아래 문서를 참고해 주세요.",
        "",
        "- 00_shared/Paper_A_포함판단_과정_설명서_20260617.docx",
        "",
        "특히 포함 수는 새로 추측하거나 다시 계산하자는 의미가 아니라, 현재 기준인 225개 포함 행, 중복 DOI 1건 병합, 224개 고유 포함 보고서/연구라는 단위를 Methods/Results/PRISMA에서 일관되게 쓰기 위한 확인 작업입니다.",
        "",
        "작업을 진행하실 때 부탁드릴 점은 세 가지입니다.",
        "",
        "1. 원고 문장 수정은 가능하면 작업원고 안에서 직접 해 주세요.",
        "2. 별도 Notes 파일은 기본 산출물로 만들지 말아 주세요. 필요한 경우에도 작업보드나 보조 추적표와 연결되게 남겨 주세요.",
        "3. 작업을 마치셨거나 PI 판단이 필요하면 작업보드에서 상태를 '완료' 또는 '검토 요청'으로 바꾸고, 원고 내 위치나 댓글 ID를 메모에 적어 주세요.",
        "",
        "Paper B는 이번 전환의 중심 대상이 아니므로 기존처럼 작업보드와 별도 산출물 방식으로 추적하겠습니다.",
        "",
        "감사합니다.",
    ]
    EMAIL_MD.write_text("제목: " + subject + "\n\n" + "\n".join(body) + "\n", encoding="utf-8")

    doc = Document()
    style_doc(doc)
    add_title(doc, "연구자 안내 이메일 초안", subject)
    for para in body:
        if not para:
            doc.add_paragraph()
            continue
        p = doc.add_paragraph()
        r = p.add_run(para)
        if para.startswith(("1.", "2.", "3.", "- R", "- 00_", "- 위치:", "- 용도:", "- 부탁:")):
            r.bold = para.startswith(("1.", "2.", "3.", "- R"))
        set_korean_font(r)
    doc.save(EMAIL_DOCX)


def build_all_workbooks() -> None:
    build_board_workbook()
    build_decision_log()
    build_reference_matrix()
    build_apa_jars_checklist()


def inspect_hyperlinks(paths: list[Path]) -> dict[str, tuple[int, int]]:
    result: dict[str, tuple[int, int]] = {}
    for path in paths:
        count = 0
        non_web = 0
        with ZipFile(path) as zf:
            for name in zf.namelist():
                if not name.endswith(".rels"):
                    continue
                root = ET.fromstring(zf.read(name))
                for rel in root:
                    if rel.attrib.get("Type", "").endswith("/hyperlink"):
                        count += 1
                        target = rel.attrib.get("Target", "")
                        if not target.startswith(("http://", "https://", "mailto:")):
                            non_web += 1
        result[path.name] = (count, non_web)
    return result


def main() -> None:
    ensure_dirs()
    moved = archive_notes_workflow_files()
    build_working_manuscripts()
    build_references_folder()
    build_process_doc()
    build_guide_docx()
    build_markdown_readmes()
    build_root_readme()
    build_all_workbooks()
    build_email_docs()

    print(f"work_dir={WORK_DIR}")
    print(f"storage_label={STORAGE_LABEL}")
    print(f"working_manuscript={PAPER_A_TRACK_DOCX}")
    print(f"board={BOARD_XLSX}")
    print(f"guide={GUIDE_DOCX}")
    print(f"email={EMAIL_DOCX}")
    print(f"root_readme={ROOT_README_MD}")
    print(f"archived_notes={len(moved)}")
    for name, (count, non_web) in inspect_hyperlinks([BOARD_XLSX, GUIDE_DOCX, PROCESS_DOCX, EMAIL_DOCX]).items():
        print(f"hyperlinks {name}: total={count} non_web={non_web}")


if __name__ == "__main__":
    main()
