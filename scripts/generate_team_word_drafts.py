#!/usr/bin/env python3
"""Generate team-facing Word drafts for Paper 1 and Paper 2.

The documents are deliberately framed as working drafts. They report current
corpus/extraction status and planned result logic without claiming final MASEM
or LLM-validation results before the source-anchored reference is frozen.
"""

from __future__ import annotations

from collections import OrderedDict
from datetime import date
from pathlib import Path
import csv

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DATE = date(2026, 5, 30).isoformat()

SUMMARY_CSV = ROOT / "data/04_extraction/02_pre_adjudication_disagreement/combined/derived/combined_pairwise_disagreement_summary_20260525.csv"
STUDY_QUEUE_CSV = ROOT / "data/04_extraction/02_pre_adjudication_disagreement/combined/derived/combined_study_review_queue_20260525.csv"
CORR_QUEUE_CSV = ROOT / "data/04_extraction/02_pre_adjudication_disagreement/combined/derived/combined_correlation_review_queue_20260525.csv"
CODER_VALUES_CSV = ROOT / "data/04_extraction/02_pre_adjudication_disagreement/combined/derived/combined_coder_values_long_20260525.csv"
CONFIRMED_INCLUDES_CSV = ROOT / "data/02_screening/confirmed_includes.csv"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 89, 89)
HEADER_FILL = "F2F4F7"


def read_csv_dicts(path: Path) -> list[dict[str, str]]:
    with path.open(newline="", encoding="utf-8-sig") as f:
        return list(csv.DictReader(f))


def as_int(value: str | int | float | None) -> int:
    if value in (None, ""):
        return 0
    return int(float(value))


def extraction_stats() -> dict[str, object]:
    summary = read_csv_dicts(SUMMARY_CSV)
    study_queue = read_csv_dicts(STUDY_QUEUE_CSV)
    corr_queue = read_csv_dicts(CORR_QUEUE_CSV)
    values = read_csv_dicts(CODER_VALUES_CSV)
    confirmed = read_csv_dicts(CONFIRMED_INCLUDES_CSV)

    totals_by_family: OrderedDict[str, int] = OrderedDict()
    totals_by_mismatch: OrderedDict[str, int] = OrderedDict()
    totals_by_pair: OrderedDict[tuple[str, str], int] = OrderedDict()
    for row in summary:
        n = as_int(row["n"])
        totals_by_family[row["field_family"]] = totals_by_family.get(row["field_family"], 0) + n
        totals_by_mismatch[row["mismatch_type"]] = totals_by_mismatch.get(row["mismatch_type"], 0) + n
        key = (row["phase_block"], row["pair"])
        totals_by_pair[key] = totals_by_pair.get(key, 0) + n

    studies_by_pair: OrderedDict[tuple[str, str], int] = OrderedDict()
    for row in study_queue:
        key = (row["phase_block"], row["pair"])
        studies_by_pair[key] = studies_by_pair.get(key, 0) + 1

    queue_sums = {
        "difference_rows": sum(as_int(r.get("n_difference_rows")) for r in study_queue),
        "metadata_diff": sum(as_int(r.get("n_metadata_diff")) for r in study_queue),
        "one_coder_only": sum(as_int(r.get("n_one_coder_only")) for r in study_queue),
        "numeric_or_source_diff": sum(as_int(r.get("n_numeric_or_source_diff")) for r in study_queue),
    }

    def top_metadata(field_key: str, limit: int = 5) -> list[tuple[str, int]]:
        counts: OrderedDict[str, int] = OrderedDict()
        for row in values:
            if row.get("field_family") != "metadata" or row.get("field_key") != field_key:
                continue
            value = (row.get("value") or "").strip()
            if not value:
                continue
            counts[value] = counts.get(value, 0) + 1
        return sorted(counts.items(), key=lambda item: (-item[1], item[0]))[:limit]

    return {
        "confirmed_includes": len(confirmed),
        "summary_rows": summary,
        "total_disagreement_rows": sum(as_int(r["n"]) for r in summary),
        "totals_by_family": totals_by_family,
        "totals_by_mismatch": totals_by_mismatch,
        "totals_by_pair": totals_by_pair,
        "studies_by_pair": studies_by_pair,
        "study_queue_rows": len(study_queue),
        "correlation_queue_rows": len(corr_queue),
        "queue_sums": queue_sums,
        "top_country": top_metadata("country", 8),
        "top_sample_type": top_metadata("sample_type", 5),
        "top_education_level": top_metadata("education_level", 5),
        "top_ai_tool_name": top_metadata("ai_tool_name", 6),
        "top_ai_type": top_metadata("ai_type", 5),
        "top_study_design": top_metadata("study_design", 5),
    }


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[float]) -> None:
    table.autofit = False
    table.allow_autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), "9360")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    for col_idx, width in enumerate(widths):
        for cell in table.columns[col_idx].cells:
            cell.width = Inches(width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ["List Bullet", "List Number"]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def add_title_block(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(21)
    run.font.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(subtitle)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(11)
    run.font.color.rgb = MUTED


def add_lead(doc: Document, label: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{label}: ")
    r.bold = True
    r.font.color.rgb = DARK_BLUE
    p.add_run(text)


def add_source_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def add_table(doc: Document, headers: list[str], rows: list[list[object]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    set_row_cant_split(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_shading(cell, HEADER_FILL)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(str(header))
        r.bold = True
        r.font.color.rgb = RGBColor(0, 0, 0)
    for row in rows:
        new_row = table.add_row()
        set_row_cant_split(new_row)
        cells = new_row.cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.add_run(str(value))


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def set_running_header(doc: Document, text: str) -> None:
    section = doc.sections[0]
    paragraph = section.header.paragraphs[0]
    paragraph.text = text
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in paragraph.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.text = "Working draft for internal team discussion"
    for run in footer_p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED


def pair_rows(stats: dict[str, object]) -> list[list[object]]:
    studies_by_pair = stats["studies_by_pair"]
    totals_by_pair = stats["totals_by_pair"]
    rows = []
    for key in [("phase1", "Pair A"), ("phase1", "Pair B"), ("phase2", "Pair C"), ("phase2", "Pair D")]:
        rows.append([
            key[0],
            key[1],
            studies_by_pair.get(key, 0),
            totals_by_pair.get(key, 0),
        ])
    return rows


def metadata_snapshot_rows(stats: dict[str, object]) -> list[list[object]]:
    mapping = [
        ("국가", "top_country"),
        ("표본 유형", "top_sample_type"),
        ("교육 수준", "top_education_level"),
        ("AI 도구", "top_ai_tool_name"),
        ("AI 유형", "top_ai_type"),
        ("연구 설계", "top_study_design"),
    ]
    rows = []
    for label, key in mapping:
        items = stats[key]
        rows.append([label, "; ".join(f"{name} ({count})" for name, count in items)])
    return rows


def make_paper1(stats: dict[str, object]) -> Path:
    out_dir = ROOT / "paper_a/manuscript"
    out_dir.mkdir(parents=True, exist_ok=True)
    out = out_dir / f"Paper_1_Team_Draft_AI_Adoption_MASEM_{OUT_DATE}.docx"

    doc = Document()
    style_doc(doc)
    set_running_header(doc, "Paper 1 Team Draft")
    add_title_block(
        doc,
        "Paper 1 초안: AI Adoption in Higher Education MASEM",
        f"팀 공유용 working draft | Hosung You | {OUT_DATE}",
    )
    add_lead(
        doc,
        "핵심 메시지",
        "Paper 1은 교육 맥락의 AI adoption을 설명하는 실질 MASEM 논문이다. "
        "현재 초안은 이론 모델, 수집된 corpus의 윤곽, 분석 전략, 그리고 예상되는 결과 그림을 공유하기 위한 팀용 문서이며, "
        "source-anchored adjudicated human reference standard가 아직 frozen되지 않았기 때문에 최종 MASEM 결과를 주장하지 않는다.",
    )

    doc.add_heading("1. 논문 방향", level=1)
    doc.add_paragraph(
        "이 논문은 TAM/UTAUT, TRA/TPB, 그리고 AI-specific psychological constructs를 결합해 "
        "AI adoption in higher education의 구조적 관계를 추정한다. 분석의 중심은 개별 연구의 단순 요약이 아니라, "
        "연구 간 상관행렬을 통합해 Performance Expectancy, Effort Expectancy, Attitude, Behavioral Intention, "
        "Use Behavior, Trust in AI, AI Anxiety 등이 어떤 경로로 adoption을 설명하는지 검증하는 것이다."
    )
    add_bullets(
        doc,
        [
            "주요 독자: educational technology, information systems, higher education AI adoption 연구자.",
            "주요 방법: two-stage MASEM(TSSEM)을 기본 분석으로 사용하고, publication year, culture, education level, AI tool type을 OSMASEM moderator로 검토한다.",
            "핵심 기여: AI-specific constructs인 Trust in AI와 AI Anxiety를 기존 TAM/UTAUT 구조 안에 통합한다.",
        ],
    )

    doc.add_heading("2. 현재까지 모인 데이터", level=1)
    add_table(
        doc,
        ["항목", "현재 값", "해석"],
        [
            ["초기 검색", "22,166 records", "Web of Science, Scopus, PsycINFO, IEEE Xplore 중심 검색 기록"],
            ["Deduplication 이후", "16,189 records", "screening master 기준"],
            ["현재 confirmed include", stats["confirmed_includes"], "2026-03-09 screening summary의 confirmed_includes.csv 기준"],
            ["Paper A proposal working corpus", "224 studies", "PROPOSAL_BRIEF.md에 기록된 2015-2025 empirical study working count"],
            ["Paper B validation corpus", "213 studies", "Paper A extraction quality를 검증하는 Phase 1+2 corpus; 10 calibration studies는 별도"],
            ["현재 workflow status", "Step 3 source-document adjudication active", "Step 4 reference freeze 전이므로 final MASEM claim은 보류"],
        ],
        [2.0, 1.7, 2.8],
    )
    add_source_note(
        doc,
        "근거 파일: paper_a/PROPOSAL_BRIEF.md; data/02_screening/screening_summary.json; "
        "data/02_screening/confirmed_includes.csv; data/04_extraction/README.md; data/04_extraction/WORKFLOW_STATUS_LOG.md."
    )

    doc.add_heading("3. Construct model", level=1)
    add_table(
        doc,
        ["Construct", "Abbr.", "현재 working k", "모델 내 역할"],
        [
            ["Performance Expectancy", "PE", 186, "AI 사용이 성과를 높인다는 belief; ATT/BI의 핵심 선행요인"],
            ["Effort Expectancy", "EE", 162, "AI 사용의 용이성; PE와 ATT를 설명"],
            ["Social Influence", "SI", 114, "동료/교수/기관의 사회적 압력 또는 규범"],
            ["Facilitating Conditions", "FC", 105, "자원, 기술지원, 접근성; UB와 연결"],
            ["Attitude", "ATT", 81, "AI 사용에 대한 평가적 태도; PE/EE와 BI 사이의 mediator"],
            ["Self-Efficacy", "SE", 44, "AI 사용 능력에 대한 자기효능감"],
            ["AI Anxiety", "ANX", 40, "AI에 대한 불안, 위협, apprehension; ATT/BI에 부적 영향 예상"],
            ["Trust in AI", "TRU", 36, "AI 판단 또는 산출물에 대한 신뢰; BI에 정적 영향 예상"],
            ["Behavioral Intention", "BI", 185, "주요 adoption intention outcome"],
            ["Use Behavior", "UB", 90, "실제 또는 자기보고 사용 행동"],
        ],
        [1.85, 0.65, 0.9, 3.1],
    )
    add_source_note(doc, "working k 값은 paper_a/PROPOSAL_BRIEF.md의 construct table을 사용했다.")

    doc.add_heading("4. 대략적인 결과 그림", level=1)
    add_lead(
        doc,
        "주의",
        "아래 내용은 최종 분석 결과가 아니라, 현재 이론 모델과 data structure가 그릴 가능성이 높은 결과 narrative이다. "
        "pooled correlation matrix와 Stage 2 SEM 결과가 나온 뒤 숫자와 문장은 교체되어야 한다.",
    )
    add_bullets(
        doc,
        [
            "PE와 EE는 여전히 adoption intention의 핵심 설명축이 될 가능성이 높다. 특히 EE는 PE를 통해 간접적으로 작동할 수 있다.",
            "ATT는 UTAUT에서 생략되었지만 AI adoption corpus에서는 충분히 자주 측정되어, PE/EE -> ATT -> BI mediation을 검증할 수 있다.",
            "Trust in AI는 generative AI/LLM 맥락에서 instrumental usefulness만으로 설명되지 않는 relational belief로 다룬다.",
            "AI Anxiety는 낮은 self-efficacy와 동일한 개념으로 처리하지 않고, threat appraisal 또는 loss-of-agency 변수로 분리한다.",
            "데이터는 2024-2025, generative AI, student sample, cross-sectional survey가 강하게 많으므로, moderator 및 sensitivity 분석에서 이 편중을 드러내야 한다.",
        ],
    )

    doc.add_heading("5. 현재 corpus profile snapshot", level=1)
    add_table(
        doc,
        ["Field", "현재 raw coding values에서 자주 보이는 값"],
        metadata_snapshot_rows(stats),
        [1.4, 5.1],
    )
    add_source_note(
        doc,
        "이 snapshot은 combined_coder_values_long_20260525.csv의 raw coder values를 단순 집계한 것이며, "
        "중복 coder 값과 pre-adjudication 상태를 포함한다. 최종 study-level 분포는 reference freeze 후 재산출해야 한다."
    )

    doc.add_heading("6. 팀이 다음에 해야 할 일", level=1)
    add_numbered(
        doc,
        [
            "Paper B Step 3 source-document adjudication을 완료해 Paper A의 분석 입력값으로 쓸 source-anchored adjudicated human reference를 freeze한다.",
            "상관행렬 입력에서 zero-order r, Fornell-Larcker off-diagonal latent correlation, beta/path coefficient, excluded row를 명확히 구분한다.",
            "pooled correlation matrix를 구성하고 positive-definite, sample-size, source-type sensitivity를 점검한다.",
            "Stage 1 TSSEM, Stage 2 SEM, OSMASEM moderator, beta/path sensitivity 분석 순서로 결과를 만든다.",
            "최종 결과가 나오기 전 팀 공유 문장에는 'planned', 'working', 'expected pattern'을 쓰고 confirmed result처럼 표현하지 않는다.",
        ],
    )

    doc.add_heading("7. 회의용 질문", level=1)
    add_bullets(
        doc,
        [
            "ATT를 main mediator로 유지할 것인가, 아니면 UTAUT-style direct model을 주요 비교 모델로 둘 것인가?",
            "Trust와 Anxiety를 Tier 2 AI-specific constructs로 main model에 넣을지, sensitivity/extended model로 분리할지 결정해야 한다.",
            "beta/path coefficient 기반 값은 main matrix에 포함할지, 별도 sensitivity로 제한할지 Paper A/B가 같은 기준을 써야 한다.",
        ],
    )

    doc.save(out)
    return out


def make_paper2(stats: dict[str, object]) -> Path:
    out_dir = ROOT / "paper_b/manuscript"
    out_dir.mkdir(parents=True, exist_ok=True)
    out = out_dir / f"Paper_2_Team_Draft_LLM_MASEM_Extraction_{OUT_DATE}.docx"

    doc = Document()
    style_doc(doc)
    set_running_header(doc, "Paper 2 Team Draft")
    add_title_block(
        doc,
        "Paper 2 초안: LLM-assisted MASEM-ready Extraction",
        f"팀 공유용 working draft | Research Synthesis Methods 방향 | {OUT_DATE}",
    )
    add_lead(
        doc,
        "핵심 메시지",
        "Paper 2는 LLM이 인간을 대체하는지 평가하는 논문이 아니라, MASEM-ready extraction에서 LLM이 어떤 task family에서는 prefill/triage를 돕고, "
        "어떤 task family에서는 expert adjudication이 필요한지를 검증하는 methods paper이다.",
    )

    doc.add_heading("1. 논문 방향", level=1)
    doc.add_paragraph(
        "현재 framing은 task-contingent LLM augmentation이다. 즉, LLM의 전체 정확도 하나를 보고 자동 대체 가능성을 주장하지 않는다. "
        "대신 raw human-human disagreement, source-document adjudication, source-anchored adjudicated human reference standard, "
        "LLM comparison, downstream MASEM substitution stability를 순서대로 분리한다."
    )
    add_bullets(
        doc,
        [
            "RQ0: 인간 코더들이 adjudication 전에 어디서 왜 불일치하는가?",
            "RQ1: frozen human reference와 비교했을 때 LLM workflow가 task family별로 얼마나 정확한가?",
            "RQ2: 어떤 source format, construct ambiguity, sample definition이 오류를 설명하는가?",
            "RQ3: human disagreement와 LLM flags가 expert review triage에 도움이 되는가?",
            "RQ4: human-supervised LLM-assisted input이 MASEM pooled correlations/path conclusions를 보존하는가?",
        ],
    )

    doc.add_heading("2. 현재까지 모인 데이터", level=1)
    add_table(
        doc,
        ["데이터 상태", "현재 값", "의미"],
        [
            ["Validation corpus", "213 studies", "Phase 1+2 Paper B validation corpus"],
            ["Phase 1", "100 studies", "Pair A(R1+R2) 50; Pair B(R3+R4) 50"],
            ["Phase 2", "113 studies", "Pair C(R1+R4) 57; Pair D(R2+R3) 56"],
            ["Calibration block", "10 studies", "training/calibration으로 별도 취급"],
            ["Current adjudication stage", "Step 3 active", "source-document adjudication 진행 중; Step 4 reference freeze 전"],
            ["Correlation review queue", stats["correlation_queue_rows"], "source-type/numeric/coder-only issue가 있는 study-level review rows"],
        ],
        [2.0, 1.3, 3.2],
    )
    add_source_note(
        doc,
        "근거 파일: data/04_extraction/README.md; WORKFLOW_STATUS_LOG.md; "
        "combined_study_review_queue_20260525.csv; combined_correlation_review_queue_20260525.csv."
    )

    doc.add_heading("3. Pre-adjudication disagreement snapshot", level=1)
    add_table(
        doc,
        ["구분", "Count", "해석"],
        [
            ["전체 disagreement rows", stats["total_disagreement_rows"], "source adjudication 전에 이미 상당한 review load가 존재"],
            ["Metadata disagreement", stats["totals_by_family"].get("metadata", 0), "대부분은 표준화 가능하지만 sample/country/tool type은 moderator에 영향"],
            ["Correlation disagreement", stats["totals_by_family"].get("correlation", 0), "MASEM matrix에 직접 연결되는 high-consequence layer"],
            ["Numeric/source differences", stats["totals_by_mismatch"].get("numeric_or_source_diff", 0), "r/beta/Fornell-Larcker/sample/source table 확인 필요"],
            ["One-coder-only rows", stats["totals_by_mismatch"].get("one_coder_only", 0), "한 코더만 evidence를 찾은 case; missingness와 source recovery issue"],
        ],
        [2.3, 1.0, 3.2],
    )
    add_table(
        doc,
        ["Phase", "Pair", "Studies", "Disagreement rows"],
        pair_rows(stats),
        [1.2, 1.3, 1.2, 2.8],
    )
    add_source_note(doc, "위 수치는 combined_pairwise_disagreement_summary_20260525.csv에서 산출했다.")

    doc.add_heading("4. 새 adjudication 기록 원칙", level=1)
    add_lead(
        doc,
        "중요한 분리",
        "consensus_source_value에는 PDF/source가 실제 보고한 값과 evidence type을 기록하고, "
        "analysis_note에는 그 값이 zero-order r, latent/Fornell-Larcker off-diagonal correlation, beta/path coefficient, HTMT-only, excluded row 중 무엇인지 표시한다.",
    )
    add_bullets(
        doc,
        [
            "zero-order r와 latent correlation은 downstream sensitivity에서 분리할 수 있도록 source type을 남긴다.",
            "standardized path coefficient 또는 beta-converted value는 correlation처럼 직접 섞지 않고, path/beta source로 표시한다.",
            "Fornell-Larcker diagonal, HTMT-only value, wrong sample/construct, untraceable source는 target MASEM correlation row로 자동 채택하지 않는다.",
            "source-check 결과가 exclude_row 또는 exclude_study이면 숫자 자체보다 exclusion rationale과 source location을 우선 기록한다.",
        ],
    )

    doc.add_heading("5. 대략적인 결과 그림", level=1)
    add_lead(
        doc,
        "주의",
        "아래는 현재 pre-adjudication evidence에 기반한 방향성이다. LLM accuracy, triage yield, substitution stability는 frozen human reference와 locked LLM outputs가 생긴 뒤에만 결과로 보고할 수 있다.",
    )
    add_bullets(
        doc,
        [
            "RQ0 결과는 'MASEM-ready extraction은 trained human coders에게도 어려운 task family가 있다'는 것을 보여줄 가능성이 크다.",
            "metadata는 표준화 부담이 크지만 대체로 low consequence로 분류될 수 있다. 다만 country, education level, AI type은 moderator로 연결되므로 완전 자동화하면 안 된다.",
            "correlation/path recovery는 source-type confusion, one-coder-only evidence, Fornell-Larcker/HTMT 구분 때문에 expert adjudication required category가 많이 나올 가능성이 높다.",
            "LLM의 주된 기여는 final answer 자동 대체가 아니라 source span 찾기, structured prefill, human review queue prioritization일 가능성이 높다.",
            "최종 Paper 2의 강점은 element-level accuracy보다 downstream inference stability를 같이 본다는 점이다.",
        ],
    )

    doc.add_heading("6. 팀 공유용 결과 테이블 설계", level=1)
    add_table(
        doc,
        ["Table", "현재 채울 수 있는 내용", "freeze 이후 채울 내용"],
        [
            ["Table 1: Dataset states", "raw human data, pairwise diff data, adjudication status", "frozen reference file/date"],
            ["Table 2: Task taxonomy", "bibliographic, sample, construct, numeric, matrix, moderator task families", "task-family decision category calibration"],
            ["Table 3: RQ0 disagreement", "2,949 disagreement rows and pair/field breakdown", "adjudicated resolution categories"],
            ["Table 4: LLM validity", "shell only", "task-family LLM agreement/error metrics"],
            ["Table 5: Substitution stability", "shell only", "pooled r/path/fit conclusion deltas"],
        ],
        [1.6, 2.55, 2.35],
    )

    doc.add_heading("7. Stop rules for team communication", level=1)
    add_numbered(
        doc,
        [
            "Step 4 reference freeze 전에는 LLM accuracy, model ranking, or substitution stability를 결과처럼 말하지 않는다.",
            "path coefficient를 zero-order correlation과 같은 statistic으로 직접 비교하거나 섞지 않는다.",
            "HTMT-only와 Fornell-Larcker diagonal은 MASEM correlation으로 쓰지 않는다.",
            "raw coder disagreement는 지우지 않고 RQ0 evidence로 보존한다.",
            "최종 Word/논문 draft에는 source-anchored adjudicated human reference standard라는 용어를 사용하고 gold standard라는 표현은 피한다.",
        ],
    )

    doc.add_heading("8. 다음 작업", level=1)
    add_bullets(
        doc,
        [
            "combined_correlation_review_queue_20260525.csv에서 source-type mismatch와 one-coder-only cases를 먼저 adjudicate한다.",
            "S014, S021, S056, S092, S121, S202 같은 include-candidate/review-source studies의 construct/sample/source-type 결정을 기록한다.",
            "S195/S206 duplicate-source/exclusion decision은 final reference-freeze log에 반영한다.",
            "reference freeze 후 LLM outputs를 열고 RQ1-RQ4 분석을 시작한다.",
            "Paper 2 Methods 문단, Table 1-3, Figure 1 design flow는 지금부터 manuscript-ready로 작성 가능하다.",
        ],
    )

    doc.save(out)
    return out


def main() -> None:
    stats = extraction_stats()
    paper1 = make_paper1(stats)
    paper2 = make_paper2(stats)
    print(paper1)
    print(paper2)


if __name__ == "__main__":
    main()
