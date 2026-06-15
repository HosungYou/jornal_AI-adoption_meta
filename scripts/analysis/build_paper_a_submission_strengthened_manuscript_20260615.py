#!/usr/bin/env python3
from __future__ import annotations

import csv
import json
import re
import shutil
import textwrap
import time
import urllib.parse
import urllib.request
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path('/Users/newhosung/Academic/2026/AI Adoption Meta Analysis')
DATE = '20260615'
PKG = ROOT / 'paper_a/manuscript/target_journal/apa7_model_family_full_manuscript_scaffold_20260615'
DATA_PKG = ROOT / 'data/04_extraction/05_llm_masem_substitution/results/paper_a_apa7_model_family_manuscript_package_20260615'
SUPP = ROOT / 'data/04_extraction/05_llm_masem_substitution/results/paper_a_model_family_supplemental_diagnostics_20260615'
PUBLIC = ROOT / 'paper_a/public_data_repository_20260615'
ONEDRIVE = Path('/Users/newhosung/Library/CloudStorage/OneDrive-SharedLibraries-ThePennsylvaniaStateUniversity/AI Adoption Meta Analysis - Documents/05_manuscripts/Paper_A/2026-06-15_apa7_model_family_full_manuscript_scaffold')
REF_BANK = PKG / f'PAPER_A_EXPANDED_REFERENCE_BANK_{DATE}.csv'
REF_REVIEW = ROOT / 'references/paper_a_apa7_evidence_review_20260615'
FIG_DIR = PKG / 'figures'
OUT_MD = PKG / f'PAPER_A_APA7_SUBMISSION_STRENGTHENED_INLINE_MANUSCRIPT_{DATE}.md'
OUT_DOCX = PKG / f'PAPER_A_APA7_SUBMISSION_STRENGTHENED_INLINE_MANUSCRIPT_{DATE}.docx'
DATA_MD = DATA_PKG / OUT_MD.name
DATA_DOCX = DATA_PKG / OUT_DOCX.name
ONEDRIVE_MD = ONEDRIVE / OUT_MD.name
ONEDRIVE_DOCX = ONEDRIVE / OUT_DOCX.name
FIG0_PNG = PKG / 'figures/figure_0_paper_a_research_procedure_black_font_20260615.png'
FIG0_SVG = PKG / 'figures/figure_0_paper_a_research_procedure_black_font_20260615.svg'
CURRENT = ROOT / 'CURRENT.md'
README = PKG / f'README_PAPER_A_APA7_MODEL_FAMILY_MANUSCRIPT_PACKAGE_{DATE}.md'
BLACK = RGBColor(0, 0, 0)

CONSTRUCT_TABLE = [
    ['PE', 'Performance expectancy / perceived usefulness', 'TAM, TAM2, UTAUT', 'Instrumental outcome belief: AI improves learning, teaching, productivity, or task performance.', 'full10; core7; trust6'],
    ['EE', 'Effort expectancy / perceived ease of use', 'TAM, computer self-efficacy, UTAUT', 'Operational-friction belief: AI is manageable, learnable, and low burden.', 'full10; core7; trust6'],
    ['SI', 'Social influence', 'UTAUT', 'Normative/institutional pressure and endorsement mechanism.', 'full10; core7; trust6'],
    ['FC', 'Facilitating conditions', 'UTAUT', 'Resource and infrastructure mechanism enabling evaluation and enacted use.', 'full10; core7'],
    ['ATT', 'Attitude', 'TRA/TPB, TAM', 'Evaluative mediator translating beliefs into intention.', 'full10; core7'],
    ['TRU', 'Trust', 'Trust in automation, trust in IS, AI reliance', 'AI-specific reliance mechanism under opacity, autonomy, uncertainty, and vulnerability.', 'full10; trust6'],
    ['ANX', 'Anxiety', 'Technology readiness, affective threat', 'Threat/unease mechanism retained but underidentified for primary complete-case MASEM.', 'full10; future mechanism'],
    ['SE', 'Self-efficacy', 'Social cognitive theory, computer self-efficacy', 'Capability mechanism; feasible mainly in smaller supplemental sets.', 'full10; future mechanism'],
    ['BI', 'Behavioral intention', 'TRA/TPB, TAM, UTAUT', 'Proximal motivational outcome.', 'full10; core7; trust6'],
    ['UB', 'Use behavior', 'TAM, UTAUT', 'Behavioral adoption outcome.', 'full10; core7; trust6'],
]

PROCEDURE_STEPS = [
    ('1. Theory reconstruction', 'TRA/TPB, TAM, TAM2, UTAUT, trust in automation, self-efficacy, and anxiety traditions were synthesized into an initial AI adoption construct universe.'),
    ('2. Inductive construct mapping', 'Study-level variables were harmonized into construct families observed in the extracted AI adoption literature.'),
    ('3. Source-anchored adjudication', 'Human-coded correlations and researcher-approved source-traced additions formed the analysis input; raw PDFs and private workbooks remain non-public.'),
    ('4. Full10 evidence-map diagnosis', 'The 10-construct target was evaluated for 45/45 pairwise coverage and complete-case matrix feasibility.'),
    ('5. Empirical model-family MASEM', 'Core7 and trust6 were estimated as complete-case model-family descendants using two-stage MASEM.'),
    ('6. Supplemental diagnostics', 'Reduced models, PE-vs-EE role comparison, omitted-pair diagnosis, and ANX/SE feasibility scans tested robustness and boundaries.'),
    ('7. Open science package', 'Share-safe aggregate results, figures, scripts, and manifests were prepared for OSF; raw PDFs and private source files are excluded.'),
]


def doi_from_url(value: str) -> str:
    value = (value or '').strip()
    if not value:
        return ''
    if 'doi.org/' in value:
        return value.split('doi.org/', 1)[1].strip()
    if value.lower().startswith('10.'):
        return value
    return ''


def initials(given: str) -> str:
    if not given:
        return ''
    parts = re.split(r'[\s-]+', given.replace('.', ''))
    out = []
    for part in parts:
        if part:
            out.append(part[0].upper() + '.')
    return ' '.join(out)


def format_authors(authors) -> str:
    if not authors:
        return ''
    names = []
    for a in authors[:20]:
        family = a.get('family', '')
        given = a.get('given', '')
        if family:
            names.append(f"{family}, {initials(given)}".strip())
    if len(authors) > 20:
        names = names[:19] + ['... ' + names[-1]]
    if len(names) == 1:
        return names[0]
    if len(names) == 2:
        return names[0] + ', & ' + names[1]
    return ', '.join(names[:-1]) + ', & ' + names[-1]


def get_year(msg) -> str:
    for key in ['published-print', 'published-online', 'issued', 'created']:
        try:
            parts = msg[key]['date-parts'][0]
            if parts:
                return str(parts[0])
        except Exception:
            pass
    return 'n.d.'


def sentence(text: str) -> str:
    text = re.sub(r'\s+', ' ', (text or '').strip())
    return text


def fetch_crossref(doi: str):
    url = 'https://api.crossref.org/works/' + urllib.parse.quote(doi, safe='')
    req = urllib.request.Request(url, headers={'User-Agent': 'PaperAReferenceFormatter/1.0 (mailto:newhosung@psu.edu)'})
    try:
        with urllib.request.urlopen(req, timeout=20) as resp:
            data = json.loads(resp.read().decode('utf-8'))
        return data.get('message', {})
    except Exception:
        return None


def apa_from_crossref(msg, doi: str, fallback: str):
    if not msg:
        return {'apa_text': fallback, 'authors': '', 'year': '', 'title': '', 'container': '', 'volume': '', 'issue': '', 'pages': '', 'doi': doi, 'source': 'fallback'}
    authors = format_authors(msg.get('author'))
    year = get_year(msg)
    title = sentence((msg.get('title') or [''])[0])
    container = sentence((msg.get('container-title') or [''])[0])
    volume = sentence(msg.get('volume', ''))
    issue = sentence(msg.get('issue', ''))
    pages = sentence(msg.get('page', '')) or sentence(msg.get('article-number', ''))
    parts = []
    if authors:
        parts.append(f'{authors} ({year}).')
    elif year:
        parts.append(f'({year}).')
    if title:
        parts.append(f'{title}.')
    journal = container
    vol_issue = ''
    if volume and issue:
        vol_issue = f'{volume}({issue})'
    elif volume:
        vol_issue = volume
    source_part = ''
    if journal and vol_issue and pages:
        source_part = f'{journal}, {vol_issue}, {pages}.'
    elif journal and vol_issue:
        source_part = f'{journal}, {vol_issue}.'
    elif journal and pages:
        source_part = f'{journal}, {pages}.'
    elif journal:
        source_part = f'{journal}.'
    if source_part:
        parts.append(source_part)
    if doi:
        parts.append('https://doi.org/' + doi)
    return {'apa_text': ' '.join(parts), 'authors': authors, 'year': year, 'title': title, 'container': container, 'volume': volume, 'issue': issue, 'pages': pages, 'doi': doi, 'source': 'crossref'}


def build_reference_metadata():
    out_csv = REF_REVIEW / 'paper_a_crossref_apa7_reference_metadata_20260615.csv'
    rows = []
    with REF_BANK.open(newline='', encoding='utf-8') as f:
        bank_rows = list(csv.DictReader(f))
    for row in bank_rows:
        doi = doi_from_url(row.get('doi_or_url', ''))
        fallback = row.get('citation', '').strip() + ((' ' + row.get('doi_or_url', '').strip()) if row.get('doi_or_url', '').strip() else '')
        msg = fetch_crossref(doi) if doi else None
        meta = apa_from_crossref(msg, doi, fallback)
        meta.update({'category': row.get('category', ''), 'fallback_citation': row.get('citation', ''), 'doi_or_url': row.get('doi_or_url', '')})
        rows.append(meta)
        if doi:
            time.sleep(0.05)
    out_csv.parent.mkdir(parents=True, exist_ok=True)
    with out_csv.open('w', newline='', encoding='utf-8') as f:
        fieldnames = ['category','apa_text','authors','year','title','container','volume','issue','pages','doi','doi_or_url','source','fallback_citation']
        w = csv.DictWriter(f, fieldnames=fieldnames)
        w.writeheader(); w.writerows(rows)
    out_md = REF_REVIEW / 'paper_a_crossref_apa7_reference_list_20260615.md'
    lines = ['# Paper A DOI-based APA 7 reference list', '', 'Generated: 2026-06-15', '', 'Entries marked `crossref` were rebuilt from Crossref DOI metadata; fallback entries require manual verification before final submission.', '']
    for r in sorted(rows, key=lambda x: x['apa_text']):
        lines.append(f"- [{r['source']}] {r['apa_text']}")
    out_md.write_text('\n'.join(lines) + '\n', encoding='utf-8')
    return rows, out_csv, out_md


def fmt(x, digits=3):
    try:
        if pd.isna(x):
            return 'NA'
        return f'{float(x):.{digits}f}'.replace('-0.000', '0.000')
    except Exception:
        return str(x)


def fmt_p(x):
    try:
        if pd.isna(x):
            return 'NA'
        v = float(x)
        if v < .001:
            return '< .001'
        return f'{v:.3f}'.lstrip('0')
    except Exception:
        return str(x)


def ci(lo, hi):
    return f'[{fmt(lo)}, {fmt(hi)}]'


def md_table(headers, rows):
    lines = ['| ' + ' | '.join(headers) + ' |', '| ' + ' | '.join(['---'] * len(headers)) + ' |']
    for row in rows:
        lines.append('| ' + ' | '.join(str(v).replace('|', '\\|') for v in row) + ' |')
    return '\n'.join(lines)


def build_tables():
    primary_fit = pd.read_csv(PKG / 'tables/paper_a_model_family_fit_with_n_20260615.csv')
    primary_rows = []
    for _, r in primary_fit.iterrows():
        primary_rows.append([r['model_family'], int(r['complete_case_k']), int(r['effective_sample_size']), fmt(r['chisq']), fmt(r['df'],0), fmt_p(r['p']), fmt(r['CFI']), fmt(r['TLI']), fmt(r['RMSEA']), fmt(r['SRMR'])])
    model_comp = pd.read_csv(SUPP / 'paper_a_supplemental_model_comparison_20260615.csv')
    keep = ['core7_full','core6_no_ATT_direct_beliefs','core7_pure_ATT_mediation_no_direct_belief_BI','trust6_full','trust5_no_TRU_direct_acceptance','trust6_trust_mediator_no_direct_belief_BI','se4_capability_effort_intention']
    supp_rows = []
    for _, r in model_comp[model_comp['model_id'].isin(keep)].iterrows():
        supp_rows.append([r['model_id'], int(r['positive_definite_complete_case_studies']), r['stage2_status'], fmt(r['chisq']), fmt(r['df'],0), fmt_p(r['p']), fmt(r['CFI']), fmt(r['TLI']), fmt(r['RMSEA']), fmt(r['SRMR']), fmt(r['AIC'])])
    pe_ee = pd.read_csv(SUPP / 'paper_a_pe_vs_ee_role_comparison_20260615.csv')
    pe_rows = []
    for _, r in pe_ee[pe_ee['source'] == 'primary_path'].iterrows():
        pe_rows.append([r['family'], r['predictor'], r['target'], fmt(r['estimate']), ci(r['ci_low'], r['ci_high']), r['inference_class']])
    anx = pd.read_csv(SUPP / 'paper_a_anx_se_targeted_model_attempts_20260615.csv')
    anx_rows = []
    for _, r in anx.iterrows():
        anx_rows.append([r['model_id'], r['constructs'], int(r['positive_definite_complete_case_studies']), r['stage2_status'], fmt(r['CFI']), fmt(r['RMSEA'])])
    return primary_rows, supp_rows, pe_rows, anx_rows


def make_figure0():
    import subprocess
    import tempfile

    titles = [t for t, _ in PROCEDURE_STEPS]
    descs = [d for _, d in PROCEDURE_STEPS]
    r_titles = 'c(' + ', '.join(json.dumps(t) for t in titles) + ')'
    r_descs = 'c(' + ', '.join(json.dumps(d) for d in descs) + ')'
    r_code = f"""
titles <- {r_titles}
descs <- {r_descs}
draw_fig <- function() {{
  par(mar=c(0.3,0.3,0.6,0.3), family='serif')
  plot.new()
  plot.window(xlim=c(0,1), ylim=c(0,1))
  text(0.5, 0.965, 'Paper A research procedure', cex=1.15, font=2, col='black')
  x <- 0.08; w <- 0.84; h <- 0.095; gap <- 0.027; y_top <- 0.88
  for (i in seq_along(titles)) {{
    y <- y_top - (i - 1) * (h + gap)
    rect(x, y, x+w, y+h, border='black', lwd=1.2, col='white')
    text(x+0.02, y+h-0.028, titles[i], adj=c(0,1), cex=0.82, font=2, col='black')
    wrapped <- strwrap(descs[i], width=110)
    text(x+0.02, y+h-0.058, paste(wrapped, collapse='\\n'), adj=c(0,1), cex=0.66, col='black')
    if (i < length(titles)) arrows(0.5, y-0.002, 0.5, y-gap+0.006, length=0.06, lwd=1, col='black')
  }}
}}
png({json.dumps(str(FIG0_PNG))}, width=10.5, height=7.2, units='in', res=300)
draw_fig()
dev.off()
svg({json.dumps(str(FIG0_SVG))}, width=10.5, height=7.2)
draw_fig()
dev.off()
"""
    with tempfile.NamedTemporaryFile('w', suffix='.R', delete=False) as f:
        f.write(r_code)
        script = f.name
    subprocess.run(['Rscript', script], check=True)


def build_markdown(refs):
    primary_rows, supp_rows, pe_rows, anx_rows = build_tables()
    primary_fit = pd.read_csv(PKG / 'tables/paper_a_model_family_fit_with_n_20260615.csv')
    core = primary_fit[primary_fit['route'] == 'paper_a_core7_att_mediation'].iloc[0]
    trust = primary_fit[primary_fit['route'] == 'paper_a_trust6_mechanism'].iloc[0]
    omitted = pd.read_csv(SUPP / 'paper_a_full10_omitted_pair_diagnostic_20260615.csv')
    future_pairs = int((omitted['diagnostic_priority'] == 'future_mechanism_feasibility_or_pairwise_only').sum())
    ref_list = '\n\n'.join([r['apa_text'] for r in sorted(refs, key=lambda x: x['apa_text'])])
    text = f"""\
# From Theoretical Coverage to Estimable Model Families: A Meta-Analytic Structural Equation Modeling Study of AI Adoption

**Author Note**  
Author affiliations, acknowledgments, funding, conflicts of interest, data availability, and repository policy will be completed after the team confirms the final target journal and OSF component.

## Abstract

AI adoption research in education draws on technology acceptance, unified acceptance, trust, self-efficacy, and anxiety traditions, but individual studies rarely report the complete correlation matrices needed to test the whole theoretical system. This study reconstructs AI adoption as a full 10-construct theoretical target and evaluates which parts of that target are empirically estimable using model-family meta-analytic structural equation modeling (MASEM). The full10 target was generated by combining model-history reconstruction, inductive construct mapping from the extracted literature, and matrix-feasibility diagnosis. It achieved complete pairwise coverage across all 45 construct pairs but had no complete same-study 10-construct matrices. Therefore, empirical MASEM was conducted through estimable model-family descendants: a core7 attitude-mediation model and a trust6 AI-reliance model. The core7 model fit well, chi-square(5) = {fmt(core['chisq'])}, p = {fmt_p(core['p'])}, CFI = {fmt(core['CFI'])}, RMSEA = {fmt(core['RMSEA'])}. The trust6 model also fit well, chi-square(4) = {fmt(trust['chisq'])}, p = {fmt_p(trust['p'])}, CFI = {fmt(trust['CFI'])}, RMSEA = {fmt(trust['RMSEA'])}. Supplemental diagnostics showed that performance expectancy and effort expectancy operate as distinct usefulness and effort mechanisms, that trust is best interpreted as a direct AI reliance mechanism rather than a fully mediating mechanism, and that anxiety/self-efficacy remain theoretically important but underidentified for primary complete-case MASEM. The study contributes a transparent framework for preserving broad AI adoption theory while reporting only empirically defensible structural models.

**Keywords:** artificial intelligence adoption, technology acceptance, trust in AI systems, attitude mediation, MASEM, TSSEM

## Introduction

AI-supported systems increasingly mediate writing, feedback, tutoring, learning analytics, assessment support, and administrative decision making in education. This growth has produced many empirical studies of AI acceptance, but the literature remains structurally uneven. Primary studies typically estimate selected parts of technology acceptance or AI trust models rather than the full set of relations implied by the accumulated theory. The result is a field with broad pairwise evidence but limited complete-matrix evidence for full-system structural synthesis.

The present study treats that limitation as a scientific finding. Rather than forcing a theoretically attractive but empirically underidentified omnibus model, Paper A asks how AI adoption theory develops, which constructs form the full theoretical target, which model-family descendants are estimable, and which mechanisms remain future work because the matrix structure is insufficient. This approach aligns with MASEM guidance that structural claims must respect the available correlation matrix, heterogeneity, and model identification conditions.

### Theoretical Evolution of Adoption Models

The adoption model lineage begins with belief-attitude-intention-behavior logic in reasoned action and planned behavior. TAM translated that logic into technology use by distinguishing perceived usefulness and perceived ease of use as central belief mechanisms. TAM2 and UTAUT extended the theory by emphasizing performance expectancy, effort expectancy, social influence, facilitating conditions, intention, and use. AI adoption adds further mechanisms because AI systems are opaque, probabilistic, autonomous, and often consequential. Users must judge not only whether AI is useful and easy to use, but also whether it can be trusted, whether they are capable of using it, and whether it evokes anxiety or threat.

### Constructing the Full10 Theoretical Target

The full10 target was produced through three linked steps. First, theory reconstruction identified constructs repeatedly used in TRA/TPB, TAM, TAM2, UTAUT, trust in automation, self-efficacy, and anxiety traditions. Second, inductive construct mapping harmonized study-level variables into construct families actually observed in the AI adoption literature. Third, meta-analytic matrix diagnosis evaluated whether those construct families had pairwise coverage and complete-case feasibility. Table 1 summarizes the resulting construct genealogy.

### Table 1

*Construct Genealogy and Model-Family Role in Paper A*

{md_table(['Construct','Label','Origin','AI-Adoption Function','Role'], CONSTRUCT_TABLE)}

The full10 target includes performance expectancy/perceived usefulness, effort expectancy/perceived ease of use, social influence, facilitating conditions, attitude, trust, anxiety, self-efficacy, behavioral intention, and use behavior. It achieved complete pairwise coverage across all 45 construct pairs, but no study reported a complete same-study 10-construct matrix. Therefore, full10 is a theoretical target and evidence map, not an estimated complete-case SEM.

### Figure 0

*Paper A Research Procedure*

![Figure 0. Paper A research procedure.]({FIG0_PNG.relative_to(ROOT)})

### Core7 and Trust6 as Model-Family Descendants

Core7 is the estimable acceptance backbone. It preserves PE, EE, social influence, facilitating conditions, attitude, behavioral intention, and use behavior. Trust6 is the AI-specific reliance descendant. It preserves PE, EE, social influence, trust, behavioral intention, and use behavior. These are not arbitrary reduced models; they are complete-case model-family descendants of full10 that preserve theory while satisfying MASEM requirements.

### Performance Expectancy Versus Effort Expectancy

PE and EE are compared as distinct mechanisms rather than as a simple PE-EE association. PE captures expected value and performance improvement. EE captures operational friction and ease of use. Their relative roles can differ across attitude, intention, and trust. This distinction is central because an AI system may be powerful but difficult to use, or easy to use but not valuable enough for adoption.

### Anxiety and Self-Efficacy as Future Mechanisms

Anxiety and self-efficacy remain in the theory because they represent threat and capability mechanisms. However, the current data structure does not support their inclusion in the primary complete-case MASEM. This is a matrix-identification result, not a theoretical rejection. The full10 omitted-pair diagnostic identified {future_pairs} future-mechanism pairs involving anxiety or self-efficacy.

## Method

### Design Overview

The study used a model-family TSSEM/MASEM design. Figure 0 shows the full research procedure. The workflow intentionally separated theory construction, source-anchored data preparation, matrix-feasibility diagnosis, primary model-family MASEM, supplemental diagnostics, and open-science packaging.

### Search, Screening, and Eligibility

The broader project identified and screened AI adoption studies in education and adjacent AI-use contexts. Eligible studies reported quantitative associations among constructs relevant to AI adoption, acceptance, intention, or use. The final PRISMA counts should be inserted from the locked screening workbook before submission. The current manuscript focuses on the analysis-ready source-anchored correlation dataset.

### Source-Anchored Construct Coding

The analysis used the source-anchored adjudicated human reference standard with researcher-approved S048 additions. Study-level variables were mapped to PE, EE, SI, FC, ATT, TRU, ANX, SE, BI, and UB. Correlations were retained when they were traceable to primary sources or approved adjudication records and had usable sample-size information. Raw PDFs, private source documents, and raw coder workbooks are not redistributed.

### Matrix-Feasibility Diagnosis

For each candidate model, the analysis evaluated required construct pairs, observed pairwise coverage, complete-case study availability, and positive-definiteness of study-level matrices. The full10 target reached all 45 pairwise relations but zero complete same-study matrices. Core7 and trust6 had enough positive-definite complete-case matrices for empirical MASEM.

### Primary TSSEM/MASEM

Stage 1 synthesized complete-case correlation matrices using random-effects TSSEM. Stage 2 estimated prespecified structural paths for core7 and trust6. Fit was evaluated using chi-square, CFI, TLI, RMSEA, SRMR, AIC, and BIC. Path-level support was classified by likelihood-based 95% confidence intervals because finite standard errors and z-based p values were not available for all path estimates.

### Supplemental Diagnostics

Supplemental diagnostics tested whether the model-family decisions were scientifically defensible. Reduced models removed PE, EE, SI, FC, ATT, or TRU; alternative models changed direct versus mediated path assumptions; PE and EE were compared as usefulness and effort mechanisms; omitted full10 pairs were catalogued; and ANX/SE-inclusive construct sets were scanned for complete-case feasibility. These diagnostics are not definitive nested chi-square tests because construct removal can alter k, degrees of freedom, and matrix structure.

### Open Science and Reproducibility

A share-safe OSF-ready package has been prepared locally. It includes aggregate results, original figures, scripts, the manuscript draft, and manifests. It excludes raw PDFs, private source documents, and raw coder workbooks. The intended OSF route is a Paper A component under the existing AI adoption OSF project (`https://osf.io/mkrgd/overview`), pending researcher upload authorization and component creation.

## Results

### Full10 Evidence Map

The full10 target achieved complete pairwise coverage across all 45 construct pairs. This supports its role as the theoretical evidence map. However, zero studies contained a complete 10-construct matrix, so the full10 target was not estimable as a single complete-case SEM.

### Primary Model-Family Fit

Table 2 reports the primary empirical MASEM results. Both model families showed excellent approximate fit. The results support the model-family route: Paper A can report estimable structural evidence without overstating full10 as an omnibus SEM.

### Table 2

*Primary Model-Family MASEM Fit*

{md_table(['Model','k','N_eff','chi-square','df','p','CFI','TLI','RMSEA','SRMR'], primary_rows)}

### Figure 1

*Full10 Theoretical Evidence Map*

![Figure 1. Full10 theoretical evidence map.]({(FIG_DIR / 'figure_1_full10_theoretical_evidence_map_heatmap_ci_20260615.png').relative_to(ROOT)})

### Primary Structural Paths

In core7, supported paths included FC -> ATT, SI -> BI, ATT -> BI, FC -> UB, and BI -> UB. PE -> ATT and SI -> ATT had intervals that included zero, and EE -> ATT, PE -> BI, and EE -> BI had incomplete likelihood-based intervals. In trust6, supported paths included EE -> BI, TRU -> BI, and BI -> UB. EE -> TRU, SI -> TRU, and SI -> BI had intervals that included zero, whereas PE -> TRU and PE -> BI were positive but interval-incomplete.

### Figure 2

*Core7 Attitude-Mediation MASEM Path Diagram*

![Figure 2. Core7 attitude-mediation MASEM path diagram.]({(FIG_DIR / 'figure_2_core7_att_mediation_masem_path_ci_20260615.png').relative_to(ROOT)})

### Figure 3

*Trust6 Mechanism MASEM Path Diagram*

![Figure 3. Trust6 mechanism MASEM path diagram.]({(FIG_DIR / 'figure_3_trust6_mechanism_masem_path_ci_20260615.png').relative_to(ROOT)})

### Supplemental Model-Family Diagnostics

Table 3 reports selected supplemental models. Removing FC weakened core7 fit, supporting retention of facilitating conditions. A direct-belief model without ATT fit well but used a much larger complete-case set, so it is diagnostic rather than a replacement for the attitude-mediation model. A pure attitude-mediation model without direct PE/EE/SI -> BI paths fit worse, supporting retention of direct belief-intention paths. In trust6, the trust-only mediation model fit poorly, so trust should be framed as a direct AI reliance mechanism operating alongside usefulness, effort, and social influence, rather than as the sole mediator.

### Table 3

*Supplemental Reduced and Alternative Model-Family Diagnostics*

{md_table(['Model','k','Status','chi-square','df','p','CFI','TLI','RMSEA','SRMR','AIC'], supp_rows)}

### PE Versus EE Role Comparison

Table 4 supports the conceptual distinction between PE and EE. EE -> BI was supported in trust6, whereas several PE paths were positive but interval-incomplete. The manuscript should not claim universal PE or EE dominance. The defensible claim is that usefulness and effort are distinct adoption mechanisms whose empirical roles depend on target outcome and model-family context.

### Table 4

*Performance Expectancy and Effort Expectancy as Distinct Mechanisms*

{md_table(['Model','Predictor','Target','Estimate','95% CI','Inference'], pe_rows)}

### Anxiety and Self-Efficacy Feasibility

Table 5 shows that self-efficacy has more near-term empirical potential than anxiety in the current complete-case structure. The SE -> EE -> BI/UB targeted model was minimally estimable but fragile. Anxiety-inclusive models involving ATT, TRU, BI, and UB had zero positive-definite complete-case matrices and should not be forced into primary MASEM.

### Table 5

*Targeted Anxiety and Self-Efficacy Feasibility Attempts*

{md_table(['Model','Constructs','k','Status','CFI','RMSEA'], anx_rows)}

## Discussion

### Theoretical Contribution

Paper A contributes a theory-preserving approach to AI adoption meta-analysis. It reconstructs AI adoption as a full10 target while refusing to overclaim an omnibus SEM that the matrix structure cannot support. This distinction improves scientific transparency because it shows what the field has measured, what can be structurally estimated, and which mechanisms require more complete future reporting.

### Trust in AI Systems

Trust is a central AI-specific mechanism because AI systems require reliance under opacity, autonomy, uncertainty, and vulnerability. The trust6 result supports TRU -> BI, and the supplemental diagnostics show that trust-only mediation is not adequate. Thus, trust should be described as a direct reliance predictor that operates alongside usefulness and effort rather than as a simple extension variable or the only mediator.

### Performance Expectancy and Effort Expectancy

The PE-vs-EE comparison strengthens the manuscript because it avoids treating acceptance beliefs as interchangeable. PE and EE represent different practical questions: whether AI is worth using and whether AI is manageable to use. Their empirical roles differed across model families, which suggests that AI adoption theory should preserve this distinction rather than collapse both into a general positive-belief factor.

### Anxiety and Self-Efficacy

Anxiety and self-efficacy are future mechanisms, not failed mechanisms. Pairwise evidence exists, but complete-case structural evidence is currently insufficient for primary MASEM. This finding can guide future primary studies: researchers should report full correlation matrices that include anxiety, self-efficacy, trust, attitude, intention, and use behavior if the field wants to test these mechanisms structurally.

### Methodological Contribution

The manuscript also contributes a replicable procedure for sparse-matrix MASEM. The workflow first preserves broad theory, then diagnoses pair coverage and complete-case feasibility, then estimates defensible model-family members, and finally reports supplemental diagnostics. This avoids two common errors: reducing the theory merely to what is easy to estimate, and forcing a theory-rich model into data that cannot support it.

### Limitations and Future Work

The main limitation is matrix sparsity. Complete-case MASEM necessarily privileges studies reporting all required construct pairs within a model family. Reduced model comparisons should be interpreted diagnostically because k and matrix structure can change when constructs are removed. Future work should update the source-anchored dataset, verify final PRISMA counts, strengthen ANX/SE complete-case evidence, and preregister the final model-family hierarchy before journal submission if the team wants stronger confirmatory language.

## Data Availability

A share-safe Paper A public repository package has been prepared locally at `paper_a/public_data_repository_20260615/` and as `paper_a/public_data_repository_20260615_osf_ready.zip`. The intended route is to upload this package to a Paper A component under the existing OSF project `https://osf.io/mkrgd/overview`. Raw PDFs, private source documents, raw coder workbooks, and runtime files are excluded from the public package.

## References

{ref_list}
"""
    OUT_MD.write_text(text, encoding='utf-8')
    shutil.copy2(OUT_MD, DATA_MD)
    shutil.copy2(OUT_MD, ONEDRIVE_MD)


def set_run(run, bold=False, italic=False, size=12):
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.color.rgb = BLACK


def configure_doc(doc):
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    p = sec.header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run()
    fld1 = OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'PAGE'
    fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'end')
    r._r.append(fld1); r._r.append(instr); r._r.append(fld2)
    for style_name in ['Normal','Title','Heading 1','Heading 2','Heading 3','List Bullet']:
        st = doc.styles[style_name]
        st.font.name = 'Times New Roman'
        st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        st.font.size = Pt(12)
        st.font.color.rgb = BLACK
        st.paragraph_format.line_spacing = 2
        st.paragraph_format.space_after = Pt(0)
    for style_name in ['Title','Heading 1','Heading 2','Heading 3']:
        doc.styles[style_name].font.bold = True


def add_p(doc, text='', style=None, align=None, indent=False, bold=False, italic=False):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.line_spacing = 2
    p.paragraph_format.space_after = Pt(0)
    if indent:
        p.paragraph_format.first_line_indent = Inches(.5)
    r = p.add_run(text)
    set_run(r, bold=bold, italic=italic)
    return p


def set_cell(cell, text, bold=False):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 1
    r = p.add_run(str(text))
    set_run(r, bold=bold, size=9)


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = table._tbl.tblPr
    borders = tblPr.first_child_found_in('w:tblBorders')
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tblPr.append(borders)
    for edge in ['top','bottom','insideH']:
        e = borders.find(qn('w:'+edge)) or OxmlElement('w:'+edge)
        if e.getparent() is None:
            borders.append(e)
        e.set(qn('w:val'), 'single'); e.set(qn('w:color'), '000000'); e.set(qn('w:sz'), '6')
    for edge in ['left','right','insideV']:
        e = borders.find(qn('w:'+edge)) or OxmlElement('w:'+edge)
        if e.getparent() is None:
            borders.append(e)
        e.set(qn('w:val'), 'nil')


def add_table(doc, num, title, headers, rows, note):
    add_p(doc, f'Table {num}', bold=True)
    add_p(doc, title, italic=True)
    t = doc.add_table(rows=1, cols=len(headers))
    style_table(t)
    for i,h in enumerate(headers):
        set_cell(t.rows[0].cells[i], h, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i,v in enumerate(row):
            set_cell(cells[i], v)
    add_p(doc, 'Note. ' + note)


def add_fig(doc, num, title, path, note, width=6.3):
    add_p(doc, f'Figure {num}', bold=True)
    add_p(doc, title, italic=True)
    if path.exists():
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(width))
    add_p(doc, 'Note. ' + note)


def add_reference(doc, r):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2
    p.paragraph_format.first_line_indent = Inches(-0.5)
    p.paragraph_format.left_indent = Inches(0.5)
    text = r['apa_text']
    # Simple structured APA italicization: title not italicized for articles; container and volume italicized.
    container = r.get('container') or ''
    volume = r.get('volume') or ''
    pos = 0
    spans = []
    for token in [container, volume]:
        if token:
            idx = text.find(token)
            if idx >= 0:
                spans.append((idx, idx+len(token)))
    spans = sorted(spans)
    for start, end in spans:
        if start < pos:
            continue
        if start > pos:
            run = p.add_run(text[pos:start]); set_run(run)
        run = p.add_run(text[start:end]); set_run(run, italic=True)
        pos = end
    if pos < len(text):
        run = p.add_run(text[pos:]); set_run(run)


def build_docx(refs):
    primary_rows, supp_rows, pe_rows, anx_rows = build_tables()
    doc = Document(); configure_doc(doc)
    add_p(doc, 'From Theoretical Coverage to Estimable Model Families: A Meta-Analytic Structural Equation Modeling Study of AI Adoption', style='Title', align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, 'Author Note', bold=True)
    add_p(doc, 'Author affiliations, acknowledgments, funding, conflicts of interest, data availability, and repository policy will be completed after the team confirms the final target journal and OSF component.')
    sections = [
        ('Abstract', 'AI adoption research in education draws on technology acceptance, unified acceptance, trust, self-efficacy, and anxiety traditions, but individual studies rarely report the complete correlation matrices needed to test the whole theoretical system. This study reconstructs AI adoption as a full 10-construct theoretical target and evaluates which parts of that target are empirically estimable using model-family MASEM. The full10 target was generated by combining model-history reconstruction, inductive construct mapping, and matrix-feasibility diagnosis. Empirical MASEM was conducted through core7 and trust6 model-family descendants. Supplemental diagnostics examined reduced models, PE-versus-EE roles, omitted full10 pairs, and ANX/SE feasibility.'),
        ('Introduction', 'AI-supported systems increasingly mediate writing, feedback, tutoring, learning analytics, assessment support, and administrative decision making in education. This study treats sparse matrix structure as a scientific finding rather than a nuisance, asking which parts of AI adoption theory are theoretically covered and which parts are empirically estimable.'),
        ('Theoretical Evolution of Adoption Models', 'The adoption model lineage begins with belief-attitude-intention-behavior logic, then develops through TAM, TAM2, UTAUT, trust in automation, self-efficacy, and anxiety traditions. AI adoption extends this lineage because AI systems are opaque, probabilistic, autonomous, and consequential.'),
        ('Constructing the Full10 Theoretical Target', 'The full10 target was produced through theory reconstruction, inductive construct mapping, and meta-analytic matrix diagnosis. Table 1 summarizes the construct genealogy.'),
    ]
    for heading, body in sections:
        add_p(doc, heading, style='Heading 1' if heading in ['Abstract','Introduction'] else 'Heading 2', align=WD_ALIGN_PARAGRAPH.CENTER if heading in ['Abstract','Introduction'] else None)
        add_p(doc, body, indent=(heading not in ['Abstract']))
    add_table(doc, 1, 'Construct Genealogy and Model-Family Role in Paper A', ['Construct','Label','Origin','AI-Adoption Function','Role'], CONSTRUCT_TABLE, 'PE = performance expectancy; EE = effort expectancy; ATT = attitude; TRU = trust; ANX = anxiety; SE = self-efficacy; BI = behavioral intention; UB = use behavior.')
    add_fig(doc, 0, 'Paper A Research Procedure', FIG0_PNG, 'The procedure separates theory reconstruction, source-anchored extraction, matrix diagnosis, model-family MASEM, supplemental diagnostics, and open science packaging.', width=6.1)
    add_p(doc, 'Method', style='Heading 1', align=WD_ALIGN_PARAGRAPH.CENTER)
    for heading, body in [
        ('Design Overview', 'The study used a model-family TSSEM/MASEM design. The workflow intentionally separated theory construction, source-anchored data preparation, matrix-feasibility diagnosis, primary model-family MASEM, supplemental diagnostics, and open-science packaging.'),
        ('Source-Anchored Construct Coding', 'The analysis used the source-anchored adjudicated human reference standard with researcher-approved S048 additions. Raw PDFs, private source documents, and raw coder workbooks are not redistributed.'),
        ('Matrix-Feasibility Diagnosis', 'For each candidate model, the analysis evaluated required construct pairs, observed pairwise coverage, complete-case study availability, and positive-definiteness of study-level matrices.'),
        ('Primary TSSEM/MASEM', 'Stage 1 synthesized complete-case correlation matrices using random-effects TSSEM. Stage 2 estimated prespecified structural paths for core7 and trust6. Path-level support was classified by likelihood-based 95% confidence intervals.'),
        ('Supplemental Diagnostics', 'Supplemental diagnostics tested reduced models, PE-vs-EE role comparison, omitted full10 pairs, and ANX/SE-inclusive feasibility scans. They are sensitivity diagnostics, not definitive nested model-selection tests.'),
    ]:
        add_p(doc, heading, style='Heading 2')
        add_p(doc, body, indent=True)
    add_p(doc, 'Results', style='Heading 1', align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, 'Full10 Evidence Map', style='Heading 2')
    add_p(doc, 'The full10 target achieved complete pairwise coverage across all 45 construct pairs. However, zero studies contained a complete 10-construct matrix, so full10 was not estimable as a single complete-case SEM.', indent=True)
    add_table(doc, 2, 'Primary Model-Family MASEM Fit', ['Model','k','N_eff','χ²','df','p','CFI','TLI','RMSEA','SRMR'], primary_rows, 'Full10 is reported as a theoretical evidence map rather than an estimable complete-case SEM.')
    add_fig(doc, 1, 'Full10 Theoretical Evidence Map', FIG_DIR / 'figure_1_full10_theoretical_evidence_map_heatmap_ci_20260615.png', 'Cells show pairwise random-effects pooled correlations and k. This is an evidence map, not a full10 SEM estimate.')
    add_p(doc, 'Primary Structural Paths', style='Heading 2')
    add_p(doc, 'Core7 supported FC -> ATT, SI -> BI, ATT -> BI, FC -> UB, and BI -> UB. Trust6 supported EE -> BI, TRU -> BI, and BI -> UB. Several PE paths were positive but interval-incomplete.', indent=True)
    add_fig(doc, 2, 'Core7 Attitude-Mediation MASEM Path Diagram', FIG_DIR / 'figure_2_core7_att_mediation_masem_path_ci_20260615.png', 'Solid paths indicate likelihood-based 95% CIs excluding zero; dashed paths include zero; dotted paths have incomplete intervals.')
    add_fig(doc, 3, 'Trust6 Mechanism MASEM Path Diagram', FIG_DIR / 'figure_3_trust6_mechanism_masem_path_ci_20260615.png', 'Solid paths indicate likelihood-based 95% CIs excluding zero; dashed paths include zero; dotted paths have incomplete intervals.')
    add_p(doc, 'Supplemental Model-Family Diagnostics', style='Heading 2')
    add_p(doc, 'Supplemental diagnostics strengthened the model-family argument. Trust-only mediation fit poorly, supporting trust as a direct AI reliance mechanism rather than the sole mediator. Pure attitude mediation fit worse than the baseline core7 route, supporting direct belief-intention paths.', indent=True)
    add_table(doc, 3, 'Supplemental Reduced and Alternative Model-Family Diagnostics', ['Model','k','Status','χ²','df','p','CFI','TLI','RMSEA','SRMR','AIC'], supp_rows, 'Reduced models are diagnostic and not definitive nested tests because construct removal can change complete-case k and matrix structure.')
    add_p(doc, 'PE Versus EE Role Comparison', style='Heading 2')
    add_p(doc, 'PE and EE are distinct usefulness and effort mechanisms. EE -> BI was supported in trust6, whereas several PE paths were positive but interval-incomplete; therefore the manuscript should not claim universal PE or EE dominance.', indent=True)
    add_table(doc, 4, 'Performance Expectancy and Effort Expectancy as Distinct Mechanisms', ['Model','Predictor','Target','Estimate','95% CI','Inference'], pe_rows, 'The comparison concerns PE and EE roles across targets, not the PE-EE correlation.')
    add_p(doc, 'Anxiety and Self-Efficacy Feasibility', style='Heading 2')
    add_p(doc, 'Self-efficacy has more near-term empirical potential than anxiety in the current complete-case structure, but ANX/SE models remain exploratory or underidentified.', indent=True)
    add_table(doc, 5, 'Targeted Anxiety and Self-Efficacy Feasibility Attempts', ['Model','Constructs','k','Status','CFI','RMSEA'], anx_rows, 'ANX/SE models are feasibility diagnostics and should not replace the primary model-family MASEM.')
    add_p(doc, 'Discussion', style='Heading 1', align=WD_ALIGN_PARAGRAPH.CENTER)
    for heading, body in [
        ('Theoretical Contribution', 'Paper A contributes a theory-preserving approach to AI adoption meta-analysis: it reconstructs full10 while refusing to overclaim an omnibus SEM that the matrix structure cannot support.'),
        ('Trust in AI Systems', 'Trust is a central AI-specific reliance mechanism under opacity, autonomy, uncertainty, and vulnerability. It should not be framed as the only mediator of PE/EE/SI effects.'),
        ('Performance Expectancy and Effort Expectancy', 'PE and EE represent different practical questions: whether AI is worth using and whether AI is manageable to use. Their roles differ across model families.'),
        ('Anxiety and Self-Efficacy', 'Anxiety and self-efficacy are future mechanisms, not failed mechanisms. Pairwise evidence exists, but complete-case structural evidence is not yet sufficient for primary MASEM.'),
        ('Methodological Contribution', 'The workflow preserves broad theory, diagnoses pair coverage and complete-case feasibility, estimates defensible model-family members, and reports supplemental diagnostics.'),
        ('Limitations and Future Work', 'The main limitation is matrix sparsity. Future work should verify final PRISMA counts, strengthen ANX/SE complete-case evidence, and preregister the final model-family hierarchy before submission if stronger confirmatory language is desired.'),
        ('Data Availability', 'A share-safe Paper A public repository package has been prepared locally and is intended for upload to a Paper A component under https://osf.io/mkrgd/overview. Raw PDFs, private source documents, raw coder workbooks, and runtime files are excluded.'),
    ]:
        add_p(doc, heading, style='Heading 2')
        add_p(doc, body, indent=True)
    add_p(doc, 'References', style='Heading 1', align=WD_ALIGN_PARAGRAPH.CENTER)
    for ref in sorted(refs, key=lambda x: x['apa_text']):
        add_reference(doc, ref)
    doc.save(OUT_DOCX)
    shutil.copy2(OUT_DOCX, DATA_DOCX)
    shutil.copy2(OUT_DOCX, ONEDRIVE_DOCX)


def update_osf_package(ref_csv, ref_md):
    (PUBLIC / '1_Manuscript').mkdir(parents=True, exist_ok=True)
    (PUBLIC / '5_Reference_Review_Manifest').mkdir(parents=True, exist_ok=True)
    shutil.copy2(OUT_MD, PUBLIC / '1_Manuscript' / OUT_MD.name)
    shutil.copy2(OUT_DOCX, PUBLIC / '1_Manuscript' / OUT_DOCX.name)
    shutil.copy2(ref_csv, PUBLIC / '5_Reference_Review_Manifest' / ref_csv.name)
    shutil.copy2(ref_md, PUBLIC / '5_Reference_Review_Manifest' / ref_md.name)
    readme = PUBLIC / 'README.md'
    txt = readme.read_text(encoding='utf-8') if readme.exists() else '# Paper A public repository package\n'
    marker = '\n## 2026-06-15 submission-strengthened update\n\n- Added submission-strengthened manuscript with Figure 0 procedure diagram and inline APA 7 tables/figures.\n- Added Crossref DOI-based APA 7 reference metadata/list for final reference verification.\n- Intended OSF target: Paper A component under https://osf.io/mkrgd/overview, pending researcher authorization.\n'
    if 'submission-strengthened update' not in txt:
        readme.write_text(txt.rstrip() + marker + '\n', encoding='utf-8')
    data_avail = PUBLIC / 'DATA_AVAILABILITY_STATEMENT.md'
    dav = data_avail.read_text(encoding='utf-8') if data_avail.exists() else '# Data availability statement draft\n'
    add = '\n## Recommended OSF route\n\nThe recommended route is to upload this share-safe package to a new `Paper A` component under the existing OSF project: https://osf.io/mkrgd/overview. The component should exclude raw PDFs, private source documents, raw coder workbooks, and runtime files.\n'
    if 'Recommended OSF route' not in dav:
        data_avail.write_text(dav.rstrip() + add + '\n', encoding='utf-8')
    # Manifest after additions
    import hashlib
    rows = []
    for p in sorted(PUBLIC.rglob('*')):
        if p.is_file() and p.name != 'MANIFEST_SHA256.csv':
            rows.append({'path': str(p.relative_to(PUBLIC)), 'bytes': p.stat().st_size, 'sha256': hashlib.sha256(p.read_bytes()).hexdigest()})
    with (PUBLIC / 'MANIFEST_SHA256.csv').open('w', newline='', encoding='utf-8') as f:
        w = csv.DictWriter(f, fieldnames=['path','bytes','sha256'])
        w.writeheader(); w.writerows(rows)
    zip_path = ROOT / 'paper_a/public_data_repository_20260615_osf_ready.zip'
    if zip_path.exists():
        zip_path.unlink()
    shutil.make_archive(str(zip_path).replace('.zip',''), 'zip', ROOT, 'paper_a/public_data_repository_20260615')


def update_docs(ref_csv, ref_md):
    note = f"""\n## 2026-06-15 Paper A submission-strengthened manuscript and reference rebuild\n\n- Generated submission-strengthened manuscript with Figure 0 research procedure, stronger Methods sequence, expanded Discussion, body-embedded APA 7 tables/figures, black font, and DOI-based reference rebuild.\n- Markdown: `paper_a/manuscript/target_journal/apa7_model_family_full_manuscript_scaffold_20260615/{OUT_MD.name}`.\n- Word: `paper_a/manuscript/target_journal/apa7_model_family_full_manuscript_scaffold_20260615/{OUT_DOCX.name}`.\n- Figure 0: `paper_a/manuscript/target_journal/apa7_model_family_full_manuscript_scaffold_20260615/figures/{FIG0_PNG.name}`.\n- Crossref reference metadata: `{ref_csv.relative_to(ROOT)}` and `{ref_md.relative_to(ROOT)}`.\n- OSF-ready package refreshed: `paper_a/public_data_repository_20260615_osf_ready.zip`; recommended upload target remains a Paper A component under `https://osf.io/mkrgd/overview`.\n"""
    txt = CURRENT.read_text(encoding='utf-8') if CURRENT.exists() else '# CURRENT\n'
    if '2026-06-15 Paper A submission-strengthened manuscript and reference rebuild' not in txt:
        CURRENT.write_text(txt.rstrip() + '\n' + note, encoding='utf-8')
    rtxt = README.read_text(encoding='utf-8') if README.exists() else '# README\n'
    add = f"\n- `{OUT_MD.name}`\n- `{OUT_DOCX.name}`\n- `figures/{FIG0_PNG.name}`\n"
    if OUT_DOCX.name not in rtxt:
        README.write_text(rtxt.rstrip() + '\n\n## Submission-strengthened manuscript\n' + add, encoding='utf-8')
    decision = ROOT / 'docs/06_decisions/2026-06-15_Paper_A_Submission_Strengthening_and_OSF_Route.md'
    decision.write_text(f"""# Paper A submission strengthening and OSF route\n\nDate: 2026-06-15\n\n## Decisions accepted\n\n- Keep Paper A as one model-family MASEM manuscript.\n- Treat PE and EE as distinct usefulness and effort mechanisms, not as a PE-EE correlation question.\n- Use external reference PDFs as reporting/method exemplars, not as copied table/figure content.\n- Strengthen the manuscript in this order: Methods procedure, Figure 0 research procedure, DOI-based APA references, Discussion.\n- Recommended OSF target: a new Paper A component under the existing OSF project `https://osf.io/mkrgd/overview`.\n\n## Outputs\n\n- Manuscript Markdown: `{OUT_MD.relative_to(ROOT)}`\n- Manuscript Word: `{OUT_DOCX.relative_to(ROOT)}`\n- Figure 0: `{FIG0_PNG.relative_to(ROOT)}`\n- DOI metadata: `{ref_csv.relative_to(ROOT)}`\n- DOI reference list: `{ref_md.relative_to(ROOT)}`\n- OSF-ready zip: `paper_a/public_data_repository_20260615_osf_ready.zip`\n\n## OSF boundary\n\nThe package is share-safe and excludes raw PDFs, private source documents, raw coder workbooks, and runtime state. Actual OSF upload remains credential-gated and should be done after the researcher creates or authorizes the Paper A component.\n""", encoding='utf-8')


def main():
    make_figure0()
    refs, ref_csv, ref_md = build_reference_metadata()
    build_markdown(refs)
    build_docx(refs)
    update_osf_package(ref_csv, ref_md)
    update_docs(ref_csv, ref_md)
    print(f'Wrote Figure 0: {FIG0_PNG}')
    print(f'Wrote manuscript Markdown: {OUT_MD}')
    print(f'Wrote manuscript Word: {OUT_DOCX}')
    print(f'Wrote Crossref metadata: {ref_csv}')
    print(f'Refreshed OSF-ready package: {ROOT / "paper_a/public_data_repository_20260615_osf_ready.zip"}')

if __name__ == '__main__':
    main()
