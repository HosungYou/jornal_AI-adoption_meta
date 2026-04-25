#!/usr/bin/env python3
"""Render Markdown to a conservative DOCX layout.

This renderer intentionally avoids Word tables because several wide protocol
tables render as vertical text in the artifact-tool/Word pipeline. Pipe tables
are converted to compact row blocks so the shared DOCX remains readable across
machines.
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


PIPE_SEP_RE = re.compile(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$")
IMAGE_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def clean_inline(text: str) -> str:
    text = LINK_RE.sub(r"\1", text)
    text = text.replace("**", "")
    text = text.replace("__", "")
    text = text.replace("`", "")
    return text.strip()


def split_pipe_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [clean_inline(cell) for cell in line.split("|")]


def is_pipe_table_start(lines: list[str], idx: int) -> bool:
    return (
        idx + 1 < len(lines)
        and lines[idx].lstrip().startswith("|")
        and "|" in lines[idx]
        and bool(PIPE_SEP_RE.match(lines[idx + 1]))
    )


def add_table_as_blocks(doc: Document, table_lines: list[str]) -> None:
    headers = split_pipe_row(table_lines[0])
    rows = [split_pipe_row(line) for line in table_lines[2:] if line.strip()]

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("Table: " + " | ".join(headers))
    run.bold = True
    run.font.size = Pt(9)

    for row in rows:
        pairs = []
        for idx, value in enumerate(row):
            header = headers[idx] if idx < len(headers) else f"Field {idx + 1}"
            if value:
                pairs.append(f"{header}: {value}")
        p = doc.add_paragraph("; ".join(pairs), style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(1)
        for run in p.runs:
            run.font.size = Pt(8.5)


def add_code_line(doc: Document, line: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(line)
    run.font.name = "Consolas"
    run.font.size = Pt(8)


def add_paragraph(doc: Document, text: str) -> None:
    text = clean_inline(text)
    if not text:
        return
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size = Pt(10)


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Calibri"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 3"].font.size = Pt(11)
    return doc


def render_markdown(md_path: Path, out_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = setup_document()
    in_code = False
    first_heading = True
    i = 0

    while i < len(lines):
        line = lines[i].rstrip()

        if line.strip().startswith("```"):
            in_code = not in_code
            i += 1
            continue

        if in_code:
            add_code_line(doc, line)
            i += 1
            continue

        if is_pipe_table_start(lines, i):
            table_lines = [lines[i], lines[i + 1]]
            i += 2
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_table_as_blocks(doc, table_lines)
            continue

        image_match = IMAGE_RE.search(line)
        if image_match:
            alt, image_ref = image_match.groups()
            image_path = (md_path.parent / image_ref).resolve()
            if image_path.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(str(image_path), width=Inches(5.8))
                if alt:
                    cap = doc.add_paragraph(clean_inline(alt))
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in cap.runs:
                        run.italic = True
                        run.font.size = Pt(9)
            else:
                add_paragraph(doc, f"[Missing image: {image_ref}]")
            i += 1
            continue

        if line.startswith("#"):
            level = min(len(line) - len(line.lstrip("#")), 3)
            heading = clean_inline(line[level:].strip())
            if level == 1 and not first_heading:
                doc.add_page_break()
            doc.add_heading(heading, level=level)
            first_heading = False
            i += 1
            continue

        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        if stripped.startswith(">"):
            p = doc.add_paragraph(clean_inline(stripped.lstrip("> ")))
            p.paragraph_format.left_indent = Inches(0.25)
            for run in p.runs:
                run.italic = True
                run.font.size = Pt(9.5)
            i += 1
            continue

        bullet_match = re.match(r"^([-*]|\d+\.)\s+(.*)$", stripped)
        if bullet_match:
            p = doc.add_paragraph(clean_inline(bullet_match.group(2)), style="List Bullet")
            p.paragraph_format.space_after = Pt(1)
            for run in p.runs:
                run.font.size = Pt(9.5)
            i += 1
            continue

        if stripped == "---":
            p = doc.add_paragraph("_" * 72)
            for run in p.runs:
                run.font.size = Pt(8)
            i += 1
            continue

        add_paragraph(doc, stripped)
        i += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("markdown", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    render_markdown(args.markdown, args.output)


if __name__ == "__main__":
    main()
