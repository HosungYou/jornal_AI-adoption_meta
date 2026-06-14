#!/usr/bin/env python3
"""Build a human-coding-language source-adjudication packet for Paper A AI candidates.

This script does not update analytic matrices. It rewrites AI/source-trace
candidate rows into human coding review language and creates a DOCX, Markdown,
and CSV decision table for human adjudication.
"""
from __future__ import annotations

from collections import defaultdict
from pathlib import Path
import csv

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / 'data/04_extraction/05_llm_masem_substitution/results/paper_a_human_style_source_adjudication_packet_20260614'
TRACE = ROOT / 'data/04_extraction/05_llm_masem_substitution/results/paper_a_ai_candidate_source_trace_20260614/paper_a_ai_candidate_full10_densification_trace_20260614.csv'
CSV_OUT = OUT / 'paper_a_priority_source_adjudication_decision_table_20260614.csv'
MD_OUT = OUT / 'PAPER_A_HUMAN_STYLE_SOURCE_ADJUDICATION_GUIDE_20260614.md'
DOCX_OUT = OUT / 'PAPER_A_HUMAN_STYLE_SOURCE_ADJUDICATION_GUIDE_20260614.docx'

PRIORITY_STUDIES = {'S057', 'S138', 'S176'}


def classify(row: dict[str, str]) -> tuple[str, str, str, str]:
    sid = row['study_id'].strip()
    pair = row['missing_pair'].strip()
    c1 = row['construct_1'].strip()
    c2 = row['construct_2'].strip()
    constructs = {c1, c2}

    if sid == 'S057':
        return (
            'MEDIUM_REVIEW_PRIORITY',
            'review_before_add_no_numeric_value_yet',
            '최신 human-coded matrix에는 이 missing pair가 없고 AI trace는 source-review 후보로 올렸다. 그러나 candidate_value가 비어 있으므로 numeric cell은 아직 추출/확정되지 않았다. PDF/source 표에서 두 construct가 실제 target construct로 같은 correlation matrix 안에 있는지 확인해야 한다.',
            '아직 “인간이 놓친 값”으로 단정하지 않는다. source table의 실제 construct label과 numeric cell을 사람이 확인해야 한다.',
        )

    if sid == 'S138' and 'ANX' in constructs:
        return (
            'LOW_PROBABLE_CONSTRUCT_MISMATCH',
            'do_not_add_now_unless_anxiety_construct_confirmed',
            'AI term hit는 risk/fear 계열 단어에서 발생했을 가능성이 높다. Paper A의 ANX로 코딩하려면 source가 anxiety 또는 동등한 anxiety construct를 명시하고 상관값을 제공해야 한다.',
            'human coder가 누락했다기보다 target construct 부재 또는 construct remap 미확정일 가능성이 높다.',
        )

    if sid == 'S138' and 'FC' in constructs:
        return (
            'LOW_TO_MEDIUM_REMAP_REVIEW',
            'do_not_add_now_unless_facilitating_conditions_confirmed',
            'AI term hit는 resources/support 계열 단어에서 발생했을 가능성이 있다. FC로 추가하려면 source가 facilitating conditions 또는 명백한 동등 construct를 쓰고 동일 source matrix에 numeric cell이 있어야 한다.',
            '현재로서는 human omission보다 construct mapping 미확정 후보로 처리한다.',
        )

    if sid == 'S176' and 'ANX' in constructs:
        return (
            'LOW_PROBABLE_NOT_TARGET_CONSTRUCT',
            'do_not_add_probable_construct_absent',
            'PDF text에서 확인된 Table 4 construct set은 HM, UB, BI, EE, FC, HA, PE, PI, SI, TR이며 ANX가 보이지 않는다.',
            'ANX target construct가 source matrix에 없어서 human coder가 넣지 않은 것으로 보는 것이 현재 가장 타당하다.',
        )

    if sid == 'S176' and 'SE' in constructs:
        return (
            'LOW_PROBABLE_NOT_TARGET_CONSTRUCT',
            'do_not_add_probable_construct_absent_or_mismatch',
            'PDF text에서 확인된 Table 4 construct set은 HM, UB, BI, EE, FC, HA, PE, PI, SI, TR이며 SE가 보이지 않는다. PI/HA/TR 등은 Paper A의 SE로 자동 remap하면 안 된다.',
            'SE target construct가 source matrix에 없거나 명시적 동등 construct가 아니어서 human coder가 넣지 않은 것으로 보는 것이 현재 가장 타당하다.',
        )

    return (
        'LOW_REVIEW_ONLY',
        'do_not_add_now',
        f'{pair}는 source table과 construct label 확인 전에는 보충값으로 추가하지 않는다.',
        '누락인지 제외인지 판단 보류.',
    )


def load_rows() -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    with TRACE.open(newline='', encoding='utf-8-sig') as f:
        reader = csv.DictReader(f)
        for source_row in reader:
            sid = source_row['study_id'].strip()
            if sid not in PRIORITY_STUDIES:
                continue
            if source_row['ai_trace_status'] != 'possible_densification_source_review_candidate':
                continue
            confidence, decision, reason, why_missing = classify(source_row)
            rows.append({
                'study_id': sid,
                'missing_pair': source_row['missing_pair'].strip(),
                'construct_1': source_row['construct_1'].strip(),
                'construct_2': source_row['construct_2'].strip(),
                'ai_trace_status': source_row['ai_trace_status'],
                'candidate_status': source_row['candidate_status'],
                'candidate_value': source_row['candidate_value'],
                'confidence_tier': confidence,
                'provisional_decision': decision,
                'human_coding_reason': reason,
                'why_not_in_human_matrix_language': why_missing,
                'source_trace_terms_c1': source_row['construct_1_terms_found'],
                'source_trace_terms_c2': source_row['construct_2_terms_found'],
                'source_trace_table_terms': source_row['table_terms_found'],
                'paper_b_boundary': source_row['paper_b_boundary'],
            })
    return sorted(rows, key=lambda r: (r['study_id'], r['missing_pair']))


def write_csv(rows: list[dict[str, str]]) -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    with CSV_OUT.open('w', newline='', encoding='utf-8-sig') as f:
        writer = csv.DictWriter(f, fieldnames=list(rows[0].keys()))
        writer.writeheader()
        writer.writerows(rows)


def summary_blocks() -> dict[str, dict[str, str]]:
    return {
        'S057': {
            'label': 'Manual source-table review priority, not add-now',
            'evidence': '최신 human-coded direct input에는 이 17개 missing-pair 후보가 없고, trace 파일의 candidate_value도 비어 있다. 따라서 AI가 숫자를 추가한 것이 아니라 사람이 source table에서 확인해야 하는 상태다.',
            'decision': 'source matrix에서 두 construct가 실제 target construct이고 같은 상관행렬의 numeric cell이 보이면 source_confirmed_add 후보가 될 수 있다. 그 전에는 add 금지.',
        },
        'S138': {
            'label': 'Mostly construct-remap review, not add-now',
            'evidence': '최신 human-coded input은 Table 5 기반 28개 상관쌍을 이미 포함하지만 ANX/FC 관련 후보는 없다. source text의 risk/fear/resources/support term hit만으로 ANX/FC target construct를 확정할 수 없다.',
            'decision': 'ANX는 perceived risk/fear를 anxiety로 오코딩할 위험이 있어 낮은 confidence. FC는 resources/support term hit만으로는 부족하므로 facilitating conditions label 확인 전 add 금지.',
        },
        'S176': {
            'label': 'Probable do-not-add due to target-construct absence',
            'evidence': 'PDF text로 확인된 Table 4 construct set은 HM, UB, BI, EE, FC, HA, PE, PI, SI, TR이다. ANX와 SE는 보이지 않는다.',
            'decision': 'ANX/SE 후보 17개는 현재 human omission이 아니라 target construct absent 또는 construct mismatch로 처리한다. PI/HA/TR을 SE/ANX로 자동 remap하지 않는다.',
        },
    }


def write_md(rows: list[dict[str, str]]) -> None:
    blocks = summary_blocks()
    md: list[str] = []
    md.append('# Paper A Source-Adjudication Review Packet: AI-Candidate Rows in Human Coding Terms')
    md.append('')
    md.append('Date: 2026-06-14')
    md.append('')
    md.append('## Plain answer')
    md.append('')
    md.append('These rows should not be described as values that human coders definitely failed to find. The current label is: values not present in the latest human-coded Paper A matrix, but flagged by the AI/source-trace workflow as possible source-adjudication candidates.')
    md.append('')
    md.append('Important: candidate_value is blank for these rows, so the AI did not add numeric values. It created a review queue for human source adjudication.')
    md.append('')
    md.append('## Human-coding terminology')
    md.append('')
    terms = [
        ('raw coder value', 'a value entered by the original human coding workflow'),
        ('source-check queue', 'an item that must be checked against PDF/source before it can affect the matrix'),
        ('source_confirmed_add', 'a value absent from the current human matrix but visible in an acceptable source matrix'),
        ('source_corrected_add', 'a value requiring correction of construct label, row/column orientation, or source location before adding'),
        ('exclude_no_target_construct', 'the candidate term appears in text, but the source does not contain the Paper A target construct'),
        ('exclude_no_usable_r', 'the paper discusses the construct but does not provide a usable correlation/latent correlation cell'),
        ('exclude_source_type_mismatch', 'the value is beta/path coefficient, HTMT, loading, reliability, or another non-target statistic'),
        ('defer_unclear', 'source or OCR/table structure is not clear enough for a coding decision'),
    ]
    for term, meaning in terms:
        md.append(f'- {term}: {meaning}.')
    md.append('')
    md.append('## Current prioritized decision')
    md.append('')
    md.append('No high-confidence add-now value is identified in this priority subset. The correct next step is human source adjudication, not direct matrix insertion.')
    md.append('')
    for sid in ['S057', 'S138', 'S176']:
        md.append(f'### {sid}: {blocks[sid]["label"]}')
        md.append('')
        md.append(f'Evidence summary: {blocks[sid]["evidence"]}')
        md.append('')
        md.append(f'Coding decision: {blocks[sid]["decision"]}')
        md.append('')
    md.append('## Priority candidate table')
    md.append('')
    md.append('| study_id | missing_pair | confidence_tier | provisional_decision | human coding reason |')
    md.append('|---|---|---|---|---|')
    for row in rows:
        md.append(f"| {row['study_id']} | {row['missing_pair']} | {row['confidence_tier']} | {row['provisional_decision']} | {row['human_coding_reason']} |")
    md.append('')
    MD_OUT.write_text('\n'.join(md), encoding='utf-8')


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement('w:tblHeader')
    tbl_header.set(qn('w:val'), 'true')
    tr_pr.append(tbl_header)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 8) -> None:
    cell.text = ''
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Aptos'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def write_docx(rows: list[dict[str, str]]) -> None:
    blocks = summary_blocks()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = doc.styles
    styles['Normal'].font.name = 'Aptos'
    styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    styles['Normal'].font.size = Pt(10)
    for style_name in ['Title', 'Heading 1', 'Heading 2', 'Heading 3']:
        styles[style_name].font.name = 'Aptos Display'
        styles[style_name]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Paper A Source-Adjudication Review Packet')
    run.bold = True
    run.font.size = Pt(20)
    run.font.name = 'Aptos Display'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('AI-Candidate Rows Rewritten in Human Coding Terms').italic = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Date: 2026-06-14 | Scope: priority source-adjudication subset S057, S138, S176')

    doc.add_heading('1. Plain answer', level=1)
    doc.add_paragraph('현재 표기된 후보값은 “인간 코더가 확실히 찾지 못한 값”으로 단정하면 안 됩니다. 정확한 현재 지위는 “최신 human-coded Paper A matrix에는 없지만 AI/source-trace workflow가 source-adjudication 후보로 올린 값”입니다.')
    doc.add_paragraph('Trace CSV의 candidate_value가 비어 있으므로 AI가 numeric value를 추가한 것이 아닙니다. 값을 실제 matrix에 추가하려면 사람이 PDF/source package에서 같은 source matrix 안에 두 target construct와 numeric correlation cell이 존재함을 확인해야 합니다.')

    callout = doc.add_table(rows=1, cols=1)
    callout.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = callout.cell(0, 0)
    set_cell_shading(cell, 'FFF2CC')
    set_cell_text(cell, 'Decision headline: 이 priority subset에서는 HIGH confidence add-now 값이 아직 없습니다. S057은 manual review priority이고, S138/S176은 대부분 construct mismatch 또는 target-construct absence 가능성이 큽니다.', bold=True, size=10)

    doc.add_heading('2. Human-coding terminology to use', level=1)
    terms = [
        ('raw coder value', 'original human coding workflow가 입력한 값'),
        ('source-check queue', 'PDF/source 확인 전까지 matrix에 영향을 주지 않는 검토 대기 항목'),
        ('source_confirmed_add', '현재 human matrix에는 없지만 acceptable source matrix에서 두 construct와 numeric cell이 확인된 값'),
        ('source_corrected_add', 'construct label, row/column orientation, source location 등을 보정한 후 추가 가능한 값'),
        ('exclude_no_target_construct', '용어는 보이지만 Paper A target construct가 source matrix에 없는 경우'),
        ('exclude_no_usable_r', 'construct 논의는 있으나 usable correlation/latent correlation cell이 없는 경우'),
        ('exclude_source_type_mismatch', 'beta/path coefficient, HTMT, loading, reliability 등 비대상 statistic인 경우'),
        ('defer_unclear', 'OCR/table layout/source identity가 불명확하여 판단 보류하는 경우'),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    set_repeat_table_header(table.rows[0])
    for i, header in enumerate(['Human coding term', 'Meaning in this packet']):
        set_cell_shading(table.rows[0].cells[i], '1F4E79')
        set_cell_text(table.rows[0].cells[i], header, bold=True, size=9)
    for term, meaning in terms:
        cells = table.add_row().cells
        set_cell_text(cells[0], term, bold=True, size=8)
        set_cell_text(cells[1], meaning, size=8)

    doc.add_heading('3. Confirmation procedure', level=1)
    for step in [
        'Open the PDF/source package for the study ID and locate the exact reported matrix/table.',
        'Verify that both constructs are Paper A target constructs or explicitly approved remaps. Do not rely on generic word hits such as use, support, risk, or trust in narrative text.',
        'Confirm that the value is a usable correlation/latent-correlation cell, not a path coefficient, loading, HTMT, reliability, descriptive statistic, or theory-only statement.',
        'Confirm that the sample/study identity matches the coded study and that no duplicate row already exists in the current human-coded matrix.',
        'Assign one status: source_confirmed_add, source_corrected_add, exclude_no_target_construct, exclude_no_usable_r, exclude_source_type_mismatch, or defer_unclear.',
    ]:
        doc.add_paragraph(step, style='List Number')

    doc.add_heading('4. Confidence rubric', level=1)
    rubric = [
        ('HIGH_ADD_CANDIDATE', 'Exact target constructs or approved remaps plus a visible numeric cell in an acceptable source matrix. These are the rows one can ask: “왜 안 넣었어요?” in an audit sense.'),
        ('MEDIUM_REVIEW_PRIORITY', 'The source may contain the construct pair, but abbreviation, source type, table alignment, or remap must be checked before adding.'),
        ('LOW_REVIEW_ONLY', 'The hit is likely conceptual/theoretical text, references, generic wording, or a loose remap. Do not add without stronger PDF/table evidence.'),
        ('PROBABLE_DO_NOT_ADD', 'Target construct appears absent from the source matrix or the remap is unsafe.'),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_repeat_table_header(table.rows[0])
    for i, header in enumerate(['Tier', 'Operational meaning']):
        set_cell_shading(table.rows[0].cells[i], '1F4E79')
        set_cell_text(table.rows[0].cells[i], header, bold=True, size=9)
    for tier, meaning in rubric:
        cells = table.add_row().cells
        set_cell_text(cells[0], tier, bold=True, size=8)
        set_cell_text(cells[1], meaning, size=8)

    doc.add_heading('5. Priority studies: current adjudication notes', level=1)
    for sid in ['S057', 'S138', 'S176']:
        doc.add_heading(f'{sid}: {blocks[sid]["label"]}', level=2)
        doc.add_paragraph(f'Evidence summary: {blocks[sid]["evidence"]}')
        doc.add_paragraph(f'Coding decision: {blocks[sid]["decision"]}')

    doc.add_heading('6. Reviewer language for “why was this not added?”', level=1)
    for phrase in [
        'Potential omission: the source matrix appears to contain both target constructs and a numeric cell, but the current human-coded matrix does not include the pair. Human confirmation required before adding.',
        'Not added because the target construct is absent from the coded source matrix.',
        'Not added because the source provides conceptual discussion only, not a usable matrix value.',
        'Not added because the candidate depends on an unapproved construct remap.',
        'Not added because the visible statistic is not an acceptable correlation/latent-correlation value.',
    ]:
        doc.add_paragraph(phrase, style='List Bullet')

    doc.add_page_break()
    doc.add_heading('7. Priority candidate decision table', level=1)
    doc.add_paragraph(
        'The Word packet keeps the adjudication table readable by grouping repeated missing-pair patterns. '
        'The full row-level reason text is preserved in the companion CSV decision table.'
    )
    group_rows = [
        (
            'S057',
            '17 pairs involving TRU and/or UB',
            'MEDIUM_REVIEW_PRIORITY',
            'review_before_add_no_numeric_value_yet',
            'Candidate_value is blank. Treat as a source-check queue: add only if the PDF/source matrix visibly contains both target constructs and a usable numeric correlation cell.',
        ),
        (
            'S138',
            '9 ANX-related pairs',
            'LOW_PROBABLE_CONSTRUCT_MISMATCH',
            'do_not_add_now_unless_anxiety_construct_confirmed',
            'Risk/fear term hits are not enough to code ANX. Add only if anxiety or an approved anxiety-equivalent construct is explicitly present in the source matrix.',
        ),
        (
            'S138',
            '8 FC-related pairs',
            'LOW_TO_MEDIUM_REMAP_REVIEW',
            'do_not_add_now_unless_facilitating_conditions_confirmed',
            'Resources/support term hits are not enough to code FC. Add only if facilitating conditions or an approved equivalent is explicitly present in the source matrix.',
        ),
        (
            'S176',
            '9 ANX-related pairs',
            'LOW_PROBABLE_NOT_TARGET_CONSTRUCT',
            'do_not_add_probable_construct_absent',
            'Current PDF text shows Table 4 constructs as HM, UB, BI, EE, FC, HA, PE, PI, SI, TR. ANX is not visible in that matrix.',
        ),
        (
            'S176',
            '8 SE-related pairs',
            'LOW_PROBABLE_NOT_TARGET_CONSTRUCT',
            'do_not_add_probable_construct_absent_or_mismatch',
            'Current PDF text shows Table 4 constructs as HM, UB, BI, EE, FC, HA, PE, PI, SI, TR. SE is not visible; PI/HA/TR must not be auto-remapped to SE.',
        ),
    ]
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_repeat_table_header(table.rows[0])
    for i, header in enumerate(['Study', 'Candidate group', 'Confidence', 'Provisional decision', 'Human coding reason']):
        set_cell_shading(table.rows[0].cells[i], '1F4E79')
        set_cell_text(table.rows[0].cells[i], header, bold=True, size=8)
    for row in group_rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], row[0], bold=True, size=8)
        set_cell_text(cells[1], row[1], size=8)
        set_cell_text(cells[2], row[2], size=8)
        set_cell_text(cells[3], row[3], size=8)
        set_cell_text(cells[4], row[4], size=8)

    doc.add_heading('Appendix A. Pair-level source-check queue', level=1)
    doc.add_paragraph(
        'These are not matrix-ready values. Each row remains candidate-only until a human reviewer records source value, evidence type, and source location.'
    )
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_repeat_table_header(table.rows[0])
    for i, header in enumerate(['Study', 'Missing pair', 'Confidence', 'Provisional decision']):
        set_cell_shading(table.rows[0].cells[i], '1F4E79')
        set_cell_text(table.rows[0].cells[i], header, bold=True, size=8)
    confidence_short = {
        'MEDIUM_REVIEW_PRIORITY': 'MEDIUM review',
        'LOW_PROBABLE_CONSTRUCT_MISMATCH': 'LOW mismatch',
        'LOW_TO_MEDIUM_REMAP_REVIEW': 'LOW/MED remap',
        'LOW_PROBABLE_NOT_TARGET_CONSTRUCT': 'LOW not target',
    }
    decision_short = {
        'review_before_add_no_numeric_value_yet': 'review before add',
        'do_not_add_now_unless_anxiety_construct_confirmed': 'do not add unless ANX confirmed',
        'do_not_add_now_unless_facilitating_conditions_confirmed': 'do not add unless FC confirmed',
        'do_not_add_probable_construct_absent': 'probable do not add',
        'do_not_add_probable_construct_absent_or_mismatch': 'probable do not add/remap',
    }
    for row in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], row['study_id'], bold=True, size=7)
        set_cell_text(cells[1], row['missing_pair'], size=7)
        set_cell_text(cells[2], confidence_short.get(row['confidence_tier'], row['confidence_tier']), size=7)
        set_cell_text(cells[3], decision_short.get(row['provisional_decision'], row['provisional_decision']), size=7)

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.text = 'Paper A source-adjudication packet | candidate rows only | not a matrix update'

    doc.save(DOCX_OUT)


def main() -> None:
    rows = load_rows()
    if not rows:
        raise SystemExit('No priority candidate rows loaded; check trace CSV schema or filters.')
    write_csv(rows)
    write_md(rows)
    write_docx(rows)
    print(DOCX_OUT)
    print(MD_OUT)
    print(CSV_OUT)
    print(f'rows={len(rows)}')


if __name__ == '__main__':
    main()
