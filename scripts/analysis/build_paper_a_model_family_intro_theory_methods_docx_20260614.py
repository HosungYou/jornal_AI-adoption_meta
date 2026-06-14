#!/usr/bin/env python3
"""Render Paper A model-family Intro/Theory/Methods draft to DOCX."""

from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

REPO = Path(__file__).resolve().parents[2]
ONEDRIVE = Path(
    "/Users/newhosung/Library/CloudStorage/OneDrive-SharedLibraries-ThePennsylvaniaStateUniversity/"
    "AI Adoption Meta Analysis - Documents"
)
SRC_MD = REPO / "paper_a/manuscript/model_family_masem_20260614/PAPER_A_MODEL_FAMILY_INTRO_THEORY_METHODS_DRAFT_20260614.md"
OUT_DOCX = REPO / "paper_a/manuscript/model_family_masem_20260614/PAPER_A_MODEL_FAMILY_INTRO_THEORY_METHODS_DRAFT_20260614.docx"
ONEDRIVE_DIR = ONEDRIVE / "Meta/AI Adoption/05_manuscript/Paper_A/2026-06-14_model_family_masem_intro_theory_methods"

ACCENT = RGBColor(46, 116, 181)
DARK_ACCENT = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 89, 89)
BODY = RGBColor(32, 32, 32)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        elem = borders.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "4")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), "D9E2EC")


def set_table_width(table, width_dxa: int = 9360) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def setup_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = BODY
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, ACCENT, 16, 8),
        ("Heading 2", 13, ACCENT, 12, 6),
        ("Heading 3", 12, DARK_ACCENT, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.10


def add_run_text(paragraph, text: str, bold: bool = False, italic: bool = False, color=None) -> None:
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            run.font.size = Pt(10)
            run.font.color.rgb = DARK_ACCENT
        else:
            run = paragraph.add_run(part)
            run.bold = bold
            run.italic = italic
            if color:
                run.font.color.rgb = color


def add_metadata_table(doc: Document, date_line: str, working_title: str) -> None:
    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_table_width(table)
    set_table_borders(table)
    rows = [
        ("Document type", "Paper A manuscript draft: Introduction, theoretical background, and methods"),
        ("Date", date_line.replace("Date:", "").strip()),
        ("Working title", working_title.replace("Working title:", "").strip()),
    ]
    for i, (label, value) in enumerate(rows):
        cells = table.rows[i].cells
        cells[0].text = label
        cells[1].text = value
        set_cell_width(cells[0], 2100)
        set_cell_width(cells[1], 7260)
        set_cell_shading(cells[0], "F2F4F7")
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)
                    if cell is cells[0]:
                        run.bold = True
                        run.font.color.rgb = DARK_ACCENT
    doc.add_paragraph()


def parse_markdown_lines(text: str) -> tuple[str, str, str, list[str]]:
    lines = text.splitlines()
    title = "Paper A Draft"
    date_line = "Date: 2026-06-14"
    working_title = "Working title: Understanding AI Adoption in Education"
    body_start = 0
    for i, line in enumerate(lines):
        if line.startswith("# "):
            title = line[2:].strip()
            body_start = i + 1
            break
    for line in lines[body_start:body_start + 8]:
        if line.startswith("Date:"):
            date_line = line.strip()
        if line.startswith("Working title:"):
            working_title = line.strip()
    return title, date_line, working_title, lines[body_start:]


def build_docx() -> None:
    text = SRC_MD.read_text(encoding="utf-8")
    title, date_line, working_title, lines = parse_markdown_lines(text)
    doc = Document()
    setup_styles(doc)

    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(6)
    run = title_p.add_run(title)
    run.font.name = "Calibri Light"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri Light")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(11, 37, 69)
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    add_run_text(subtitle, "Model-family MASEM route: full10 theory target, core7/trust6 empirical family", color=MUTED)

    add_metadata_table(doc, date_line, working_title)

    skip_front = True
    for raw in lines:
        line = raw.rstrip()
        if skip_front and (not line or line.startswith("Date:") or line.startswith("Working title:")):
            continue
        skip_front = False
        if not line:
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif re.match(r"^\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            add_run_text(p, re.sub(r"^\d+\.\s+", "", line))
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_run_text(p, line[2:].strip())
        else:
            p = doc.add_paragraph()
            add_run_text(p, line)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("Paper A model-family MASEM draft | 2026-06-14")
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    ONEDRIVE_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copy2(OUT_DOCX, ONEDRIVE_DIR / OUT_DOCX.name)
    print(f"docx={OUT_DOCX}")
    print(f"onedrive={ONEDRIVE_DIR / OUT_DOCX.name}")


if __name__ == "__main__":
    build_docx()
