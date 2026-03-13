#!/usr/bin/env python3
"""
Generate a professionally formatted Word document from the MASEM Coding Manual v2.
Designed for high visual clarity for human coders (R1-R4).
"""

import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from latex2omml import add_inline_equation

# ── Constants ──────────────────────────────────────────
FONT_BODY = "Calibri"
FONT_HEADING = "Calibri"
FONT_MONO = "Consolas"

COLOR_PRIMARY = RGBColor(0x1B, 0x3A, 0x5C)      # Dark navy
COLOR_SECONDARY = RGBColor(0x2E, 0x74, 0xB5)     # Blue accent
COLOR_ACCENT = RGBColor(0xD4, 0x5B, 0x07)        # Orange accent
COLOR_LIGHT_BG = "D6E4F0"                         # Light blue bg
COLOR_TABLE_HEADER = "1B3A5C"                      # Navy header
COLOR_TABLE_ALT = "F2F7FB"                         # Alternating row
COLOR_WHITE = "FFFFFF"
COLOR_BORDER = "B4C7DC"                            # Table border
COLOR_TIP_BG = "FFF3E0"                            # Tip box bg
COLOR_CODE_BG = "F5F5F5"                           # Code block bg

OUTPUT_PATH = "/Volumes/External SSD/Projects/Meta-Analysis/jornal_AI-adoption_meta/data/04_extraction/AI_Adoption_MASEM_Coding_Manual_v2.docx"


def create_document():
    doc = Document()

    # ── Page setup ──
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.header_distance = Cm(1.0)
        section.footer_distance = Cm(1.0)

    # ── Define custom styles ──
    setup_styles(doc)
    return doc


def setup_styles(doc):
    styles = doc.styles

    # Body text
    style = styles["Normal"]
    font = style.font
    font.name = FONT_BODY
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.15

    # Heading 1
    h1 = styles["Heading 1"]
    h1.font.name = FONT_HEADING
    h1.font.size = Pt(22)
    h1.font.bold = True
    h1.font.color.rgb = COLOR_PRIMARY
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.keep_with_next = True

    # Heading 2
    h2 = styles["Heading 2"]
    h2.font.name = FONT_HEADING
    h2.font.size = Pt(16)
    h2.font.bold = True
    h2.font.color.rgb = COLOR_SECONDARY
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(8)
    h2.paragraph_format.keep_with_next = True

    # Heading 3
    h3 = styles["Heading 3"]
    h3.font.name = FONT_HEADING
    h3.font.size = Pt(13)
    h3.font.bold = True
    h3.font.color.rgb = COLOR_PRIMARY
    h3.paragraph_format.space_before = Pt(14)
    h3.paragraph_format.space_after = Pt(6)
    h3.paragraph_format.keep_with_next = True


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" w:space="0" '
            f'w:color="{val.get("color", COLOR_BORDER)}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def add_styled_table(doc, headers, rows, col_widths=None, first_col_bold=False):
    """Create a visually polished table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Style header row
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = hdr_cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        run.font.name = FONT_HEADING
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, COLOR_TABLE_HEADER)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row_data):
            cell = row_cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            # Parse bold markers
            add_formatted_runs(p, str(val), base_size=Pt(10))
            if c_idx == 0 and first_col_bold:
                for run in p.runs:
                    run.bold = True
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            # Alternating row shading
            if r_idx % 2 == 1:
                set_cell_shading(cell, COLOR_TABLE_ALT)

    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Inches(width)

    # Add spacing after table
    doc.add_paragraph()
    return table


def add_formatted_runs(paragraph, text, base_size=Pt(11), base_color=None):
    """Parse markdown-style bold (**text**) and code (`text`) in text."""
    parts = re.split(r'(\*\*.*?\*\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.size = base_size
            run.font.name = FONT_BODY
            if base_color:
                run.font.color.rgb = base_color
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = FONT_MONO
            run.font.size = Pt(base_size.pt - 1) if hasattr(base_size, 'pt') else Pt(9)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        else:
            run = paragraph.add_run(part)
            run.font.size = base_size
            run.font.name = FONT_BODY
            if base_color:
                run.font.color.rgb = base_color


def add_body_text(doc, text, bold=False):
    p = doc.add_paragraph()
    add_formatted_runs(p, text)
    if bold:
        for run in p.runs:
            run.bold = True
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    add_formatted_runs(p, text)
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    return p


def add_numbered(doc, text, level=0):
    p = doc.add_paragraph(style="List Number")
    p.clear()
    add_formatted_runs(p, text)
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    return p


def add_tip_box(doc, text, label="TIP"):
    """Add a highlighted callout box."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, COLOR_TIP_BG)
    p = cell.paragraphs[0]
    run = p.add_run(f"  {label}: ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_ACCENT
    run.font.name = FONT_BODY
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = FONT_BODY
    doc.add_paragraph()


def add_important_box(doc, text):
    """Add an important/critical callout box (blue)."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, COLOR_LIGHT_BG)
    p = cell.paragraphs[0]
    run = p.add_run("  IMPORTANT: ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_PRIMARY
    run.font.name = FONT_BODY
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = FONT_BODY
    doc.add_paragraph()


def add_code_block(doc, lines):
    """Add a code/workflow block with monospace font and gray background."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, CODE_COLOR_BG if 'CODE_COLOR_BG' in dir() else "F5F5F5")
    for i, line in enumerate(lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        run = p.add_run(line)
        run.font.name = FONT_MONO
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.0
    doc.add_paragraph()


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="{COLOR_BORDER}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_page_number(doc):
    """Add page numbers to footer."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Add "Page X" with field code
        run = p.add_run("Page ")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run.font.name = FONT_BODY
        fld_xml = (
            f'<w:fldSimple {nsdecls("w")} w:instr=" PAGE "><w:r><w:t>1</w:t></w:r></w:fldSimple>'
        )
        p._p.append(parse_xml(fld_xml))


def add_header_bar(doc, text):
    """Add a running header."""
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(text)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run.font.name = FONT_BODY
        run.font.italic = True
        # Bottom border for header
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="{COLOR_BORDER}"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)


# ══════════════════════════════════════════════════════
# DOCUMENT CONTENT
# ══════════════════════════════════════════════════════

def build_document():
    doc = create_document()
    add_header_bar(doc, "AI Adoption in Education \u2014 MASEM Coding Manual v2.1")
    add_page_number(doc)

    # ── COVER PAGE ──
    for _ in range(4):
        doc.add_paragraph()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("AI Adoption in Education")
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = COLOR_PRIMARY
    run.font.name = FONT_HEADING

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("MASEM Coding Manual")
    run.font.size = Pt(26)
    run.font.color.rgb = COLOR_SECONDARY
    run.font.name = FONT_HEADING

    # Decorative line
    add_horizontal_rule(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Version 2.1  |  March 2026")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.name = FONT_BODY

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Author: Hosung You")
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.font.name = FONT_BODY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Pennsylvania State University, College of Education")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    run.font.name = FONT_BODY

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Meta-Analytic Structural Equation Modeling (MASEM)")
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("12 Target Constructs  \u00b7  4 Competing SEM Models  \u00b7  ~250 Studies")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_page_break()

    # ── VERSION HISTORY ──
    doc.add_heading("Version History", level=2)
    add_styled_table(doc,
        ["Version", "Date", "Changes"],
        [
            ["1.0", "2026-02-16", "Initial release"],
            ["2.0", "2026-03-09", "Major revision: year range 2022\u20132026; education-only scope; independent coding workflow; AI metadata pre-coding; 3 CLI models; full-text exclusion codes E-FT1\u2013E-FT6"],
            ["2.1", "2026-03-10", "Paper A+B integrated design; 2-pair ICR (R1+R2, R3+R4); Phase 1 100-study dual + Phase 2 150-study single; calibration 10 studies; cross-pair adjudication"],
        ],
        col_widths=[0.7, 1.0, 4.8],
        first_col_bold=True,
    )

    # ── TABLE OF CONTENTS ──
    doc.add_heading("Table of Contents", level=2)
    toc_items = [
        "1. Introduction and Purpose",
        "2. Eligibility Criteria",
        "3. Coding Workflow Overview",
        "4. Coder Training Protocol",
        "5. Study-Level Coding Instructions",
        "6. Correlation Matrix Coding Instructions",
        "7. Construct Harmonization Instructions",
        "8. Moderator Variable Coding",
        "9. Reliability Data Extraction",
        "10. AI-Assisted Coding Protocol",
        "11. Inter-Coder Reliability Protocol",
        "12. Discrepancy Resolution Protocol",
        "13. Study Exclusion Protocol",
        "14. Quality Assurance Checklist",
        "15. References",
        "Appendix A: Decision Trees",
        "Appendix B: FAQ",
        "Appendix C: Excel Template Sheet Descriptions",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_SECONDARY
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════
    # CHAPTER 1: Introduction and Purpose
    # ═══════════════════════════════════════════════════
    doc.add_heading("1. Introduction and Purpose", level=1)

    doc.add_heading("1.1 MASEM Study Overview", level=3)
    add_body_text(doc, "This meta-analytic structural equation modeling (MASEM) study synthesizes empirical research on AI technology adoption in educational contexts (K-12, higher education, and vocational) to:")
    add_numbered(doc, "Test the validity of traditional technology acceptance models (TAM/UTAUT) in the educational AI context")
    add_numbered(doc, "Evaluate whether AI-specific constructs (Trust, Anxiety, Transparency, Autonomy) provide incremental explanatory power")
    add_numbered(doc, "Identify moderators of AI adoption relationships (AI type, education level, culture, temporal period)")
    add_numbered(doc, "Compare competing theoretical models using two-stage SEM (TSSEM)")
    add_body_text(doc, "The study extracts correlation matrices from primary studies to build pooled correlation matrices, then fits structural models. This requires precise extraction of Pearson r values (or standardized \u03b2 for conversion) and careful harmonization of construct labels across studies.")

    doc.add_heading("1.2 How to Use This Manual", level=3)
    add_numbered(doc, "Read the entire manual before beginning coding")
    add_numbered(doc, "Complete the training protocol (Chapter 4) before independent coding")
    add_numbered(doc, "Keep the Excel codebook (`AI_Adoption_MASEM_Coding_v2.xlsx`) open alongside this manual")
    add_numbered(doc, "When uncertain, consult the decision trees (Appendix A)")
    add_numbered(doc, "Document all ambiguous decisions in the DISCREPANCY_LOG sheet")

    doc.add_heading("1.3 Canonical Source of Truth", level=3)
    add_important_box(doc, "This document is the canonical coding protocol for the project. The Excel codebook is the data-entry implementation. AI screening decisions are advisory only \u2014 they do not finalize exclusions. All raw AI outputs, human coding records, and adjudication logs must be preserved.")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════
    # CHAPTER 2: Eligibility Criteria
    # ═══════════════════════════════════════════════════
    doc.add_heading("2. Eligibility Criteria", level=1)

    doc.add_heading("2.1 PICOS Framework", level=3)
    add_styled_table(doc,
        ["Element", "Criterion", "Notes"],
        [
            ["Population", "Students, teachers, instructors, faculty, or administrators in educational settings", "K-12, higher education, vocational training"],
            ["Intervention/Exposure", "AI technology adoption/acceptance", "Generative AI, predictive AI, ITS, chatbots, AI-powered LMS, etc."],
            ["Comparison", "Not required (correlational studies)", "Most studies are single-group surveys"],
            ["Outcomes", "Correlation matrix or standardized path coefficients (\u03b2)", "Must include \u2265 2 of 12 target constructs"],
            ["Study Design", "Quantitative empirical studies", "Cross-sectional surveys, longitudinal, experimental"],
        ],
        col_widths=[1.3, 2.8, 2.4],
        first_col_bold=True,
    )

    doc.add_heading("2.2 Inclusion Criteria", level=3)
    inclusion_items = [
        "Empirical quantitative study with primary data",
        "AI technology adoption/acceptance as focal phenomenon",
        "**Educational setting/population** (K-12, higher education, or vocational)",
        "Correlation matrix OR standardized path coefficients (\u03b2) reported",
        "At least 2 of 12 target constructs measured",
        "**Published 2022\u20132026**",
        "English language",
        "Sample size n \u2265 50",
        "Peer-reviewed journal article or full conference paper",
    ]
    for i, item in enumerate(inclusion_items, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. ")
        run.bold = True
        run.font.color.rgb = COLOR_SECONDARY
        add_formatted_runs(p, item)

    doc.add_heading("2.3 Exclusion Criteria", level=3)
    add_styled_table(doc,
        ["Code", "Definition"],
        [
            ["E1", "Not empirical/quantitative (qualitative, conceptual, review)"],
            ["E2", "AI not focal (general IT, e-commerce without AI)"],
            ["E3", "Not educational context (healthcare, manufacturing, finance, etc.)"],
            ["E4", "No adoption/acceptance outcome measured"],
            ["E5", "No effect size data (no correlation or \u03b2)"],
            ["E6", "Not English"],
            ["E7", "Outside 2022\u20132026"],
            ["E8", "n < 50"],
            ["E9", "Not peer-reviewed"],
            ["E10", "Duplicate sample"],
            ["E11", "Qualitative/review only"],
            ["E12", "Other"],
        ],
        col_widths=[0.8, 5.7],
        first_col_bold=True,
    )

    doc.add_heading("2.4 Full-Text Exclusion Criteria (Phase 4)", level=3)
    add_styled_table(doc,
        ["Code", "Definition"],
        [
            ["E-FT1", "Reports < 2 construct-pair statistics (r or \u03b2)"],
            ["E-FT2", "Constructs do not map to the 12-construct model"],
            ["E-FT3", "Not educational context (upon full-text review)"],
            ["E-FT4", "Duplicate sample (same data in multiple publications)"],
            ["E-FT5", "Conference abstract only (no full paper available)"],
            ["E-FT6", "Full-text inaccessible"],
        ],
        col_widths=[0.8, 5.7],
        first_col_bold=True,
    )

    doc.add_heading("2.5 Borderline Decision Rules", level=3)
    borderline = [
        "**Chatbots/virtual assistants:** INCLUDE if AI-powered and adoption is measured in educational setting",
        "**Recommendation systems:** INCLUDE if AI/ML-based and study frames as AI adoption in education",
        "**Mixed-methods studies:** INCLUDE if quantitative portion meets all criteria",
        "**Partial correlation reporting:** INCLUDE (use available-case principle)",
        "**Conference papers later published as journal articles:** Use journal version only",
        "**K-12 vs. Higher Ed ambiguity:** Code the education_level moderator; include both",
    ]
    for item in borderline:
        add_bullet(doc, item)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════
    # CHAPTER 3: Coding Workflow Overview
    # ═══════════════════════════════════════════════════
    doc.add_heading("3. Coding Workflow Overview", level=1)

    doc.add_heading("3.1 Workflow Diagram", level=3)
    workflow_lines = [
        "Phase 0: Calibration (10 studies)",
        "   \u2502  All 4 coders (R1-R4) code same 10 studies",
        "   \u2502  \u25ba Calculate inter-pair consistency",
        "   \u2502  \u25ba Resolve disagreements, refine rules",
        "   \u2193",
        "Phase 1: Dual Coding \u2014 100 studies",
        "   \u2502  = Paper B Gold Standard + Paper A ICR sample",
        "   \u2502  Pair A (R1 + R2): 50 studies independently (blinded)",
        "   \u2502  Pair B (R3 + R4): 50 studies independently (blinded)",
        "   \u2502  \u25ba Cross-pair adjudication for discrepancies",
        "   \u2502  \u25ba ICR targets: \u03ba \u2265 .85, ICC \u2265 .90, MAE \u2264 .03",
        "   \u2193",
        "Phase 2: Single Coding \u2014 ~150 studies",
        "   \u2502  R1: ~38 studies + ~6 spot-checks",
        "   \u2502  R2: ~38 studies + ~6 spot-checks",
        "   \u2502  R3: ~37 studies + ~6 spot-checks",
        "   \u2502  R4: ~37 studies + ~6 spot-checks",
        "   \u2502  \u25ba 15-20% cross-checked by another coder",
        "   \u2193",
        "Phase 3 (parallel): AI Extraction",
        "   \u2502  \u25ba 3-model consensus: Claude + Gemini + Codex",
        "   \u2502  \u25ba Results NOT shown to human coders until Phase 4",
        "   \u2193",
        "Phase 4: ICR & AI-Human Comparison",
        "   \u2193",
        "Phase 5: Discrepancy Resolution",
        "   \u2193",
        "Phase 6: QA Final (6 Gates) \u2192 Final Validated Dataset",
    ]
    add_code_block(doc, workflow_lines)

    doc.add_heading("3.2 Key Workflow Rules", level=3)
    rules = [
        "**Calibration (Phase 0):** All 4 coders code the same 10 studies to establish inter-pair consistency before Phase 1 begins.",
        "**Dual coding (Phase 1):** Two independent pairs (R1+R2, R3+R4) each code 50 studies. This 100-study set serves dual purpose: Paper B gold standard AND Paper A ICR validation.",
        "**Cross-pair adjudication:** Discrepancies within Pair A (R1-R2) are adjudicated by R3 or R4. Discrepancies within Pair B (R3-R4) are adjudicated by R1 or R2.",
        "**Single coding (Phase 2):** Remaining ~150 studies divided equally among R1-R4 (~38 each), with 15-20% spot-checked by a different coder.",
        "**AI extraction (Phase 3):** Runs in parallel with human coding. Results are NOT shown to human coders until Phase 4.",
        "**Human-coded data is the gold standard.** AI output is used for comparison, validation, and Paper B analysis.",
    ]
    for i, rule in enumerate(rules, 1):
        add_numbered(doc, rule)

    doc.add_heading("3.3 Estimated Timeline", level=3)
    add_styled_table(doc,
        ["Phase", "Duration", "Personnel", "Output"],
        [
            ["Phase 0: Calibration", "3 days", "R1, R2, R3, R4 (all)", "Inter-pair consistency report"],
            ["Phase 1: Dual coding (100)", "3 weeks", "Pair A (R1+R2), Pair B (R3+R4)", "Paper B gold standard + ICR"],
            ["Phase 2: Single coding (~150)", "2 weeks", "R1, R2, R3, R4 (equal split)", "Paper A remaining data"],
            ["Phase 3: AI extraction", "1 week", "AI pipeline", "AI consensus dataset"],
            ["Phase 4: ICR calculation", "3 days", "PI", "ICR metrics report"],
            ["Phase 5: Discrepancy resolution", "1 week", "All coders + PI", "Resolved dataset"],
            ["Phase 6: QA finalization", "3 days", "PI", "Final validated dataset"],
        ],
        col_widths=[1.8, 0.9, 1.8, 2.0],
        first_col_bold=True,
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════
    # CHAPTER 4: Coder Training Protocol
    # ═══════════════════════════════════════════════════
    doc.add_heading("4. Coder Training Protocol", level=1)

    add_styled_table(doc,
        ["Phase", "Duration", "Activities", "Success Criterion"],
        [
            ["Orientation", "2 hours", "Read full manual; review Excel codebook; understand MASEM basics", "Can explain MASEM data requirements"],
            ["Practice \u2014 Metadata", "1 day", "Code 3 practice studies independently", ">90% agreement with gold standard"],
            ["Practice \u2014 Correlations", "2 days", "Extract correlations from 3 diverse studies (full matrix, partial, \u03b2-only)", ">95% agreement on r values (within .02)"],
            ["Practice \u2014 Harmonization", "1 day", "Harmonize constructs from 3 studies using mapping tables", ">85% agreement on construct mapping"],
            ["Calibration session", "2 hours", "Compare practice coding; discuss disagreements; refine rules", "All disagreements resolved"],
            ["Certification", "\u2014", "Code 2 new studies independently; compare with gold standard", "All criteria met"],
            ["Calibration (Phase 0)", "3 days", "All 4 coders code same 10 studies; calculate inter-pair \u03ba", "Inter-pair \u03ba \u2265 .80; all disagreements resolved"],
        ],
        col_widths=[1.4, 0.8, 2.5, 1.8],
        first_col_bold=True,
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════
    # CHAPTER 5: Study-Level Coding Instructions
    # ═══════════════════════════════════════════════════
    doc.add_heading("5. Study-Level Coding Instructions", level=1)
    add_body_text(doc, "All study-level variables are coded in the **STUDY_METADATA** sheet. AI pre-codes identification and demographic fields; human coders verify and complete remaining fields.")

    doc.add_heading("5.1 Identification Variables (AI Pre-Coded, Human Verified)", level=3)
    add_styled_table(doc,
        ["Variable", "Type", "Definition", "Example"],
        [
            ["study_id", "string", "Unique ID in S001 format", "S001"],
            ["first_author", "string", "Last name of first author", "Kim"],
            ["year", "integer", "Publication year (2022\u20132026)", "2024"],
            ["title", "string", "Full title of the paper", "AI Adoption in Higher Education..."],
            ["doi", "string", "Digital Object Identifier", "10.1000/xyz123"],
            ["source_type", "categorical", "journal, conference", "journal"],
        ],
        col_widths=[1.3, 1.0, 2.5, 1.7],
        first_col_bold=True,
    )

    doc.add_heading("5.2 Study Design Variables (AI Pre-Coded, Human Verified)", level=3)
    add_styled_table(doc,
        ["Variable", "Type", "Definition", "Example"],
        [
            ["study_design", "categorical", "cross_sectional, longitudinal, experimental", "cross_sectional"],
            ["data_collection", "categorical", "survey, experiment, mixed", "survey"],
            ["theoretical_framework", "categorical", "TAM, UTAUT, UTAUT2, TAM_AI, UTAUT_AI, TPB, SCT, other", "UTAUT2"],
        ],
        col_widths=[1.6, 1.0, 2.5, 1.4],
        first_col_bold=True,
    )

    doc.add_heading("5.3 Sample & Context Variables (AI Pre-Coded, Human Verified)", level=3)
    add_styled_table(doc,
        ["Variable", "Type", "Definition", "Example"],
        [
            ["sample_size", "integer", "Total valid sample size (N). Use final analytic sample.", "384"],
            ["sample_type", "categorical", "students, instructors, mixed, administrators", "students"],
            ["country", "string", "Country where data was collected", "South Korea"],
            ["culture_cluster", "categorical", "individualist (IDV\u226550) or collectivist (IDV<50) per Hofstede", "collectivist"],
            ["education_level", "categorical", "K-12, undergraduate, graduate, vocational, mixed", "undergraduate"],
        ],
        col_widths=[1.4, 1.0, 2.5, 1.6],
        first_col_bold=True,
    )

    doc.add_heading("5.4 AI Technology Variables (AI Pre-Coded, Human Verified)", level=3)
    add_styled_table(doc,
        ["Variable", "Type", "Definition", "Example"],
        [
            ["ai_type", "categorical", "generative, predictive, decision_support, conversational, robotic, general", "generative"],
            ["ai_tool_name", "string", "Specific AI tool if named", "ChatGPT"],
        ],
        col_widths=[1.3, 1.0, 2.8, 1.4],
        first_col_bold=True,
    )

    doc.add_heading("5.5 Quality Assessment Variables (Human Coded)", level=3)
    add_styled_table(doc,
        ["Variable", "Type", "Definition", "Example"],
        [
            ["n_constructs_measured", "integer", "Number of our 12 constructs measured", "6"],
            ["n_correlations_reported", "integer", "Number of pairwise correlations involving our constructs", "15"],
            ["matrix_completeness", "float", "Reported pairs / possible pairs (0\u20131)", "0.75"],
            ["common_method_bias", "categorical", "addressed, not_addressed, partial", "addressed"],
        ],
        col_widths=[1.8, 1.0, 2.5, 1.2],
        first_col_bold=True,
    )

    doc.add_heading("5.6 Source Management Variables", level=3)
    add_styled_table(doc,
        ["Variable", "Type", "Definition", "Example"],
        [
            ["human_coder", "string", "Initials of the human coder", "HY"],
            ["coding_date", "date", "Date coding completed (YYYY-MM-DD)", "2026-04-01"],
            ["ai_precoded", "boolean", "Whether AI pre-coded metadata", "TRUE"],
            ["ai_precoded_verified", "boolean", "Whether human verified AI metadata", "TRUE"],
        ],
        col_widths=[1.6, 0.9, 2.5, 1.5],
        first_col_bold=True,
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════
    # CHAPTER 6: Correlation Matrix Coding
    # ═══════════════════════════════════════════════════
    doc.add_heading("6. Correlation Matrix Coding Instructions", level=1)

    add_important_box(doc, "This is the most critical chapter. MASEM requires pairwise Pearson correlations between constructs. All correlation data is coded in the CORRELATION_MATRIX sheet.")

    doc.add_heading("6.1 Variable Reference", level=3)
    add_styled_table(doc,
        ["Variable", "Type", "Definition", "Example"],
        [
            ["study_id", "string", "Links to STUDY_METADATA", "S001"],
            ["construct_row", "categorical", "Row construct (one of 12)", "PE"],
            ["construct_col", "categorical", "Column construct (one of 12)", "BI"],
            ["r_value", "float", "Pearson correlation (-1 to 1)", "0.52"],
            ["r_source", "categorical", "direct, beta_converted, author_provided", "direct"],
            ["original_beta", "float", "Original \u03b2 if r_source=beta_converted, else NA", "NA"],
            ["r_sample_size", "integer", "Pairwise N if differs from study N, else NA", "NA"],
            ["p_value", "float", "Reported p-value if available", "0.001"],
            ["significance", "categorical", "p<.001, p<.01, p<.05, ns, NR", "p<.001"],
            ["source_location", "string", "Location in paper", "Table 3"],
            ["extraction_notes", "string", "Notes about extraction decisions", "Lower triangle only"],
        ],
        col_widths=[1.3, 1.0, 2.7, 1.5],
        first_col_bold=True,
    )

    doc.add_heading("6.2 Correlation Extraction Protocol", level=3)
    steps = [
        "**Step 1:** Identify the correlation matrix table. Look for tables titled \u201cCorrelations,\u201d \u201cDescriptive Statistics and Correlations,\u201d or \u201cInter-Construct Correlations.\u201d The diagonal should be 1.00 or reliability coefficients.",
        "**Step 2:** Map study constructs to our 12 target constructs using the harmonization rules in Chapter 7. Record the mapping in the CONSTRUCT_MAPPING sheet.",
        "**Step 3:** Extract all pairwise correlations involving our constructs. Record Pearson r values to 2 decimal places. Ignore significance asterisks (record effect size, not p-value). If only lower triangle shown, mirror values.",
        "**Step 4:** Record sample size per correlation. Use overall study N as default. If pairwise N differs, record pairwise N.",
        "**Step 5:** Document source location and any extraction decisions.",
    ]
    for step in steps:
        add_body_text(doc, step)

    doc.add_heading("6.3 Beta-to-r Conversion (Peterson & Brown, 2005)", level=3)
    add_body_text(doc, "When a study reports only standardized path coefficients (\u03b2) without a correlation matrix:")

    # Formula with equation
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    add_inline_equation(p, r"r = \beta + 0.05 \times \lambda")
    run = p.add_run(",  where ")
    run.font.size = Pt(11)
    add_inline_equation(p, r"\lambda = 1")
    run = p.add_run(" if ")
    run.font.size = Pt(11)
    add_inline_equation(p, r"\beta \geq 0")
    run = p.add_run(", ")
    run.font.size = Pt(11)
    add_inline_equation(p, r"\lambda = -1")
    run = p.add_run(" if ")
    run.font.size = Pt(11)
    add_inline_equation(p, r"\beta < 0")

    add_body_text(doc, "Examples:")
    add_bullet(doc, "\u03b2 = 0.40 \u2192 r = 0.40 + 0.05(1) = **0.45**")
    add_bullet(doc, "\u03b2 = \u22120.30 \u2192 r = \u22120.30 + 0.05(\u22121) = **\u22120.35**")

    add_tip_box(doc, "Always set r_source = \"beta_converted\" and record original_beta. Mark the study for sensitivity analysis.")

    doc.add_heading("6.4 Incomplete Matrix Handling", level=3)
    add_body_text(doc, "Apply the **available-case principle**:")
    add_bullet(doc, "Extract all reported correlations; leave unreported cells as NA")
    add_bullet(doc, "Do NOT assume unreported or non-significant correlations are zero")
    add_bullet(doc, "Record `matrix_completeness = reported_pairs / possible_pairs`")
    add_bullet(doc, "Stage 1 TSSEM handles missing cells via FIML")

    doc.add_heading("6.5 Where to Find Correlation Data", level=3)
    add_styled_table(doc,
        ["Priority", "Location", "Notes"],
        [
            ["1", "Main text tables (Results section)", "Most common"],
            ["2", "Appendices", "Often \u201cAppendix: Correlation Matrix\u201d"],
            ["3", "Supplementary materials / OSF", "Online supplements"],
            ["4", "SEM output tables", "May contain implied correlations"],
            ["5", "Author contact", "If \u201cavailable upon request\u201d"],
        ],
        col_widths=[0.8, 2.5, 3.2],
        first_col_bold=True,
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════
    # CHAPTER 7: Construct Harmonization
    # ═══════════════════════════════════════════════════
    doc.add_heading("7. Construct Harmonization Instructions", level=1)
    add_body_text(doc, "Harmonization maps diverse construct labels onto our 12 standard constructs. **Prioritize conceptual alignment over label matching.** All mappings go in the CONSTRUCT_MAPPING sheet.")

    doc.add_heading("7.1 The 12 Target Constructs", level=3)
    add_styled_table(doc,
        ["Abbr", "Construct", "Definition", "Origin"],
        [
            ["PE", "Performance Expectancy", "Belief that AI helps attain performance gains", "UTAUT / TAM (PU)"],
            ["EE", "Effort Expectancy", "Perceived ease of using AI", "UTAUT / TAM (PEOU)"],
            ["SI", "Social Influence", "Important others believe one should use AI", "UTAUT / TRA (SN)"],
            ["FC", "Facilitating Conditions", "Organizational/technical infrastructure for AI use", "UTAUT / TPB (PBC)"],
            ["BI", "Behavioral Intention", "Strength of intention to adopt/use AI", "TAM / UTAUT"],
            ["UB", "Use Behavior", "Actual use of AI technology", "TAM / UTAUT"],
            ["ATT", "Attitude", "Overall evaluative judgment about using AI", "TRA / TAM"],
            ["SE", "Self-Efficacy", "Belief in own capability to use AI", "SCT (Bandura)"],
            ["TRU", "AI Trust", "Willingness to be vulnerable based on positive AI expectations", "Trust theory (Mayer)"],
            ["ANX", "AI Anxiety", "Apprehension or fear about AI", "Computer anxiety lit."],
            ["TRA", "AI Transparency", "Perceived ability to understand AI decisions", "XAI literature"],
            ["AUT", "Perceived AI Autonomy", "Perceived degree of AI independent operation", "Automation literature"],
        ],
        col_widths=[0.6, 1.5, 2.6, 1.5],
        first_col_bold=True,
    )

    doc.add_heading("7.2 TAM/UTAUT Cross-Reference Table", level=3)
    add_styled_table(doc,
        ["Study Construct Label", "Model Origin", "Maps to", "Confidence"],
        [
            ["Perceived Usefulness (PU)", "TAM", "PE", "Exact"],
            ["Perceived Ease of Use (PEOU)", "TAM", "EE", "Exact"],
            ["Attitude Toward Using", "TAM", "ATT", "Exact"],
            ["Behavioral Intention to Use", "TAM/UTAUT", "BI", "Exact"],
            ["Actual System Use", "TAM/UTAUT", "UB", "Exact"],
            ["Performance Expectancy", "UTAUT", "PE", "Exact"],
            ["Effort Expectancy", "UTAUT", "EE", "Exact"],
            ["Social Influence", "UTAUT", "SI", "Exact"],
            ["Facilitating Conditions", "UTAUT", "FC", "Exact"],
            ["Subjective Norm", "TRA/TPB", "SI", "High"],
            ["Perceived Behavioral Control", "TPB", "FC or SE", "High (check items)"],
            ["Relative Advantage", "DOI", "PE", "High"],
            ["Complexity", "DOI", "EE (reverse)", "High"],
            ["Compatibility", "DOI", "FC", "High"],
            ["Computer Self-Efficacy", "SCT", "SE", "High"],
            ["Anxiety (computer/technology)", "Various", "ANX", "High"],
            ["Hedonic Motivation", "UTAUT2", "ATT", "Moderate"],
        ],
        col_widths=[2.2, 1.2, 1.2, 1.6],
        first_col_bold=True,
    )

    doc.add_heading("7.3 AI-Specific Construct Mappings", level=3)
    add_styled_table(doc,
        ["Study Construct Label", "Maps to", "Confidence", "Notes"],
        [
            ["Trust in AI / AI Trust", "TRU", "Exact", "Direct match"],
            ["Algorithmic Trust", "TRU", "Exact", "Trust in algorithm"],
            ["Automation Trust", "TRU", "High", "Trust in automated system"],
            ["Perceived Reliability", "TRU", "Moderate", "Competence dimension of trust"],
            ["AI Anxiety", "ANX", "Exact", "Direct match"],
            ["Technology Anxiety", "ANX", "High", "If AI context"],
            ["Technostress", "ANX", "Moderate", "Only anxiety component"],
            ["Explainability / Interpretability", "TRA", "Exact", "XAI construct"],
            ["Algorithmic Transparency", "TRA", "Exact", "Direct match"],
            ["Black Box Perception", "TRA (reverse)", "High", "Reverse-coded transparency"],
            ["AI Autonomy / Machine Autonomy", "AUT", "Exact", "Direct match"],
            ["Perceived AI Agency", "AUT", "High", "Agency implies autonomy"],
            ["Automation Level", "AUT", "High", "Degree of autonomy"],
        ],
        col_widths=[2.0, 1.0, 1.0, 2.5],
        first_col_bold=True,
    )

    doc.add_heading("7.4 Harmonization Decision Tree", level=3)
    harmonization_steps = [
        "**Exact label match** to our 12? \u2192 Code directly (confidence = exact)",
        "**TAM/UTAUT family** label? \u2192 Use cross-reference table \u00a77.2 (confidence = high)",
        "**AI-specific** label? \u2192 Use AI mapping table \u00a77.3 (confidence = high/exact)",
        "**Definition aligns?** \u2192 Compare study definition to our 12 (confidence = moderate)",
        "**Items align?** \u2192 Review scale items (confidence = moderate)",
        "**Expert review** \u2192 Flag for PI discussion (confidence = low)",
        "**No match** \u2192 Exclude construct from coding",
    ]
    for i, step in enumerate(harmonization_steps, 1):
        add_numbered(doc, step)

    doc.add_heading("7.5 Ambiguous Cases Requiring Item Review", level=3)
    add_styled_table(doc,
        ["Study Construct", "Possible Mappings", "Decision Rule"],
        [
            ["AI Value", "PE or ATT", "\u201cvaluable for tasks\u201d \u2192 PE; \u201cvaluable overall\u201d \u2192 ATT"],
            ["AI Confidence", "SE or TRU", "\u201cconfident in my ability\u201d \u2192 SE; \u201cconfident in AI\u201d \u2192 TRU"],
            ["AI Understanding", "TRA or EE", "\u201cunderstand how AI works\u201d \u2192 TRA; \u201cunderstand how to use\u201d \u2192 EE"],
            ["AI Support", "FC or SI", "\u201corganizational support\u201d \u2192 FC; \u201cpeer encouragement\u201d \u2192 SI"],
            ["Perceived Control", "SE or FC or AUT(rev)", "\u201cI can control\u201d \u2192 SE/FC; \u201cAI controls\u201d \u2192 AUT"],
            ["AI Quality", "PE or TRU", "\u201coutput quality\u201d \u2192 PE; \u201ctrustworthy\u201d \u2192 TRU"],
            ["AI Capability", "PE or AUT", "\u201chelps me\u201d \u2192 PE; \u201coperates independently\u201d \u2192 AUT"],
        ],
        col_widths=[1.3, 1.5, 3.7],
        first_col_bold=True,
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════
    # CHAPTER 8: Moderator Variable Coding
    # ═══════════════════════════════════════════════════
    doc.add_heading("8. Moderator Variable Coding", level=1)

    doc.add_heading("8.1 Variable Reference", level=3)
    add_styled_table(doc,
        ["Variable", "Type", "Definition", "Levels", "Example"],
        [
            ["study_id", "string", "Links to STUDY_METADATA", "\u2014", "S001"],
            ["ai_type", "categorical", "AI technology type", "generative, predictive, decision_support, conversational, robotic, general", "generative"],
            ["education_level", "categorical", "Education context", "K-12, undergraduate, graduate, vocational, mixed", "undergraduate"],
            ["temporal_period", "categorical", "Publication era", "pre_chatgpt (2022), post_chatgpt (2023\u20132026)", "post_chatgpt"],
            ["culture_cluster", "categorical", "Hofstede-based", "individualist (IDV\u226550), collectivist (IDV<50)", "collectivist"],
            ["sample_type", "categorical", "Respondent type", "students, instructors, mixed, administrators", "students"],
            ["user_role", "categorical", "User role studied", "student, instructor, both", "student"],
        ],
        col_widths=[1.2, 0.8, 1.2, 2.0, 1.3],
        first_col_bold=True,
    )

    doc.add_heading("8.2 Moderator Coding Rules", level=3)
    mod_rules = [
        "**AI Type:** Code based on the specific AI technology described. If multiple AI types, code as \u201cgeneral.\u201d",
        "**Education Level:** Code based on the sample\u2019s educational context. If cross-level, code as \u201cmixed.\u201d",
        "**Temporal Period:** Based on publication year. 2022 = pre_chatgpt; 2023\u20132026 = post_chatgpt.",
        "**Culture:** Use Hofstede individualism score. IDV \u2265 50 = individualist; IDV < 50 = collectivist.",
        "**User Role:** Code primary respondent role.",
    ]
    for rule in mod_rules:
        add_bullet(doc, rule)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════
    # CHAPTER 9: Reliability Data Extraction
    # ═══════════════════════════════════════════════════
    doc.add_heading("9. Reliability Data Extraction", level=1)

    doc.add_heading("9.1 Variable Reference", level=3)
    add_styled_table(doc,
        ["Variable", "Type", "Definition", "Example"],
        [
            ["study_id", "string", "Links to STUDY_METADATA", "S001"],
            ["construct", "categorical", "One of 12 target constructs", "PE"],
            ["study_label", "string", "Original construct label in the study", "Perceived Usefulness"],
            ["cronbach_alpha", "float", "Cronbach\u2019s alpha (0\u20131)", "0.89"],
            ["composite_reliability", "float", "Composite reliability / CR (0\u20131)", "0.92"],
            ["ave", "float", "Average Variance Extracted (0\u20131)", "0.65"],
            ["n_items", "integer", "Number of scale items", "5"],
            ["scale_source", "string", "Original scale reference", "Venkatesh et al. (2003)"],
            ["harmonization_confidence", "categorical", "exact, high, moderate, low", "exact"],
        ],
        col_widths=[1.6, 0.9, 2.4, 1.6],
        first_col_bold=True,
    )

    doc.add_heading("9.2 Reliability Extraction Protocol", level=3)
    rel_steps = [
        "Check diagonal of correlation matrix for reliability values (often \u03b1 or CR)",
        "Check measurement model tables (CFA results) for CR and AVE",
        "Check \u201cMeasures\u201d or \u201cInstruments\u201d section for \u03b1 values",
        "If multiple estimates reported, prefer: CR > \u03b1 > test-retest",
        "Record **all** available reliability indicators",
    ]
    for i, step in enumerate(rel_steps, 1):
        add_numbered(doc, step)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════
    # CHAPTER 10: AI-Assisted Coding Protocol
    # ═══════════════════════════════════════════════════
    doc.add_heading("10. AI-Assisted Coding Protocol", level=1)

    doc.add_heading("10.1 Dual-Track Design", level=3)
    add_body_text(doc, "The AI-assisted coding protocol operates on **two separate tracks**:")

    add_body_text(doc, "**Track 1 \u2014 AI Metadata Pre-Coding** (before human coding):")
    add_bullet(doc, "AI extracts non-critical metadata from PDFs")
    add_bullet(doc, "Fields: author, year, title, DOI, sample_size, country, study_design, ai_type, education_level, theoretical_framework")
    add_bullet(doc, "Humans verify and correct during their independent coding")
    add_bullet(doc, "This saves time on clerical tasks without introducing bias")

    add_body_text(doc, "**Track 2 \u2014 AI Independent Extraction** (parallel with human coding):")
    add_bullet(doc, "AI extracts core MASEM data (correlations, construct mappings)")
    add_bullet(doc, "Results are **NOT shown to human coders** until human coding is complete")
    add_bullet(doc, "Used for AI-Human comparison metrics and validation")
    add_bullet(doc, "3-model consensus determines AI extraction quality")

    doc.add_heading("10.2 Models", level=3)
    add_styled_table(doc,
        ["Model", "CLI Tool", "Version (March 2026)", "Role"],
        [
            ["Claude", "Claude CLI", "claude-sonnet-4-6", "Primary AI coder"],
            ["Gemini", "Gemini CLI", "gemini-2.5-flash", "Secondary AI coder"],
            ["Codex", "Codex CLI", "Latest available", "Tertiary AI coder"],
        ],
        col_widths=[1.0, 1.2, 2.0, 2.3],
        first_col_bold=True,
    )

    doc.add_heading("10.3 Execution Rules", level=3)
    exec_rules = [
        "Each AI model processes the same study **independently** (no communication between models)",
        "**Human gold-standard coding is completed independently** \u2014 humans do NOT see AI extraction results for correlations and construct mappings",
        "Humans MAY see AI-precoded metadata (Track 1) as this covers non-critical clerical fields",
        "All AI raw outputs are preserved in the AI_EXTRACTION_PROVENANCE sheet",
        "AI outputs are compared to human coding **after** human coding is complete (Phase D)",
    ]
    for i, rule in enumerate(exec_rules, 1):
        add_numbered(doc, rule)

    doc.add_heading("10.4 Consensus Methods", level=3)
    add_bullet(doc, "**Full agreement** (2+ models agree exactly): Use that value. Confidence = High.")
    add_bullet(doc, "**Close agreement** (all within .05): Use median. Confidence = Moderate. Flag for review.")
    add_bullet(doc, "**Disagreement** (range > .05): Consensus = NULL. Must review.")
    add_bullet(doc, "**Construct mapping:** Majority vote (2/3 agree). If no majority, flag for human review.")

    doc.add_heading("10.5 AI-Human Comparison Metrics", level=3)
    add_styled_table(doc,
        ["Metric", "Applied To", "Minimum", "Target"],
        [
            ["Exact match rate", "Categorical variables", "80%", "90%"],
            ["ICC(2,1)", "Continuous variables (r values)", ".80", ".90"],
            ["Cohen\u2019s \u03ba", "Categorical variables", ".70", ".80"],
            ["Mean absolute error", "r values", "<.05", "<.02"],
        ],
        col_widths=[1.5, 2.0, 1.3, 1.3],
        first_col_bold=True,
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════
    # CHAPTER 11: ICR Protocol
    # ═══════════════════════════════════════════════════
    doc.add_heading("11. Inter-Coder Reliability Protocol", level=1)

    doc.add_heading("11.1 When to Calculate", level=3)
    add_body_text(doc, "Calculate ICR on the **Phase 1 dual-coded set (100 studies)**. Two independent coder pairs (Pair A: R1+R2; Pair B: R3+R4) each code 50 studies.")
    add_body_text(doc, "**ICR is calculated at three levels:**")
    add_numbered(doc, "**Within-pair:** R1 vs R2 (50 studies), R3 vs R4 (50 studies)")
    add_numbered(doc, "**Inter-pair consistency:** Compare Pair A gold standard values to Pair B gold standard values on the calibration set (10 studies coded by all 4)")
    add_numbered(doc, "**AI-Human:** 3-model AI consensus vs human gold standard (100 studies)")

    add_body_text(doc, "**Stratification of Phase 1 sample (100 studies):**")
    add_bullet(doc, "Publication year (2022, 2023\u20132024, 2025\u20132026)")
    add_bullet(doc, "AI type (generative vs. non-generative)")
    add_bullet(doc, "Education level (K-12, undergraduate, graduate+)")

    doc.add_heading("11.2 Metrics and Thresholds", level=3)
    add_styled_table(doc,
        ["Metric", "Applied To", "Minimum", "Target", "If Below Minimum"],
        [
            ["Cohen\u2019s \u03ba", "Categorical variables", ".70", ".85", "Retrain; re-code"],
            ["ICC(2,1)", "Continuous variables (r)", ".90", ".95", "Review procedure; re-code"],
            ["Exact agreement %", "All coded variables", "85%", "95%", "Identify systematic errors"],
            ["r value MAE", "Correlation values", "<.05", "<.02", "Re-extract from source tables"],
        ],
        col_widths=[1.2, 1.5, 0.9, 0.8, 2.1],
        first_col_bold=True,
    )

    doc.add_heading("11.3 Reporting Format", level=3)
    add_styled_table(doc,
        ["Variable Category", "Metric", "Value", "95% CI", "n items"],
        [
            ["Construct harmonization", "Cohen\u2019s \u03ba", "[calculated]", "[CI]", "[n]"],
            ["Correlation r values", "ICC(2,1)", "[calculated]", "[CI]", "[n]"],
            ["Moderator coding", "Cohen\u2019s \u03ba", "[calculated]", "[CI]", "[n]"],
            ["Quality assessment", "Cohen\u2019s \u03ba", "[calculated]", "[CI]", "[n]"],
            ["Beta conversion", "ICC(2,1)", "[calculated]", "[CI]", "[n]"],
        ],
        col_widths=[1.8, 1.2, 1.2, 1.0, 1.0],
        first_col_bold=True,
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════
    # CHAPTER 12: Discrepancy Resolution
    # ═══════════════════════════════════════════════════
    doc.add_heading("12. Discrepancy Resolution Protocol", level=1)

    doc.add_heading("12.1 Discrepancy Classification", level=3)
    add_styled_table(doc,
        ["Classification", "Definition", "Examples", "Resolution"],
        [
            ["Clerical", "Typo, data entry error", "r = .45 vs .54 (transposition)", "Verify against source; correct"],
            ["Interpretive", "Different but defensible readings", "Mapping \u201cAI Value\u201d to PE vs ATT", "Discuss; check items; decide with PI"],
            ["Substantive", "Fundamental disagreement", "Including vs excluding a study", "Full review; PI makes final decision"],
        ],
        col_widths=[1.2, 1.5, 1.8, 2.0],
        first_col_bold=True,
    )

    doc.add_heading("12.2 Resolution Workflow", level=3)
    res_steps = [
        "**Identify** all discrepancies by comparing coder sheets",
        "**Classify** each discrepancy (clerical, interpretive, substantive)",
        "**Resolve** clerical errors by re-checking source",
        "**Discuss** interpretive differences; reach consensus or escalate",
        "**Escalate** substantive disagreements to PI",
        "**Document** all resolutions in DISCREPANCY_LOG",
    ]
    for i, step in enumerate(res_steps, 1):
        add_numbered(doc, step)

    doc.add_heading("12.3 Resolution Hierarchy", level=3)
    add_numbered(doc, "**Original study text** (highest authority)")
    add_numbered(doc, "**Human consensus** (both human coders agree)")
    add_numbered(doc, "**AI consensus** (2+ models agree AND human review confirms)")
    add_numbered(doc, "**PI adjudication** (if still unclear)")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════
    # CHAPTER 13: Study Exclusion Protocol
    # ═══════════════════════════════════════════════════
    doc.add_heading("13. Study Exclusion Protocol", level=1)

    doc.add_heading("13.1 Variable Reference", level=3)
    add_styled_table(doc,
        ["Variable", "Type", "Definition", "Example"],
        [
            ["study_id", "string", "Original study ID", "S045"],
            ["first_author", "string", "First author last name", "Kim"],
            ["year", "integer", "Publication year", "2023"],
            ["title", "string", "Paper title", "AI Acceptance..."],
            ["exclusion_stage", "categorical", "title_abstract, full_text, data_extraction", "full_text"],
            ["exclusion_code", "categorical", "E1\u2013E12 or E-FT1\u2013E-FT6", "E-FT1"],
            ["detailed_rationale", "string", "Free-text explanation", "Only reports regression \u03b2 without correlation table"],
            ["final_decision", "categorical", "exclude", "exclude"],
        ],
        col_widths=[1.4, 1.0, 2.3, 1.8],
        first_col_bold=True,
    )

    doc.add_heading("13.2 When to Log", level=3)
    add_bullet(doc, "Log every study excluded at full-text screening or later")
    add_bullet(doc, "Title/abstract exclusions are tracked in PRISMA counts (not individually logged here)")
    add_bullet(doc, "If initially included then excluded during data extraction, log with stage = \u201cdata_extraction\u201d")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════
    # CHAPTER 14: Quality Assurance Checklist
    # ═══════════════════════════════════════════════════
    doc.add_heading("14. Quality Assurance Checklist", level=1)

    qa_items = [
        "All study-level variables completed for every included study",
        "All correlation matrix cells extracted (or marked NA with reason)",
        "Construct harmonization mapping documented for every study",
        "Beta-to-r conversions checked with formula (r = \u03b2 + 0.05 \u00d7 \u03bb)",
        "Matrix symmetry verified (r(PE,BI) = r(BI,PE))",
        "Sample sizes recorded and plausible",
        "Reliability data extracted where available",
        "Moderator variables coded consistently",
        "AI provenance data preserved for AI-coded studies",
        "Discrepancy log completed for all disagreements",
        "Inter-coder reliability calculated and meets thresholds",
        "Exclusion log completed with reasons",
        "No duplicate studies in dataset",
        "Construct labels in CORRELATION_MATRIX match CONSTRUCT_MAPPING",
    ]

    # Create a checklist table
    table = doc.add_table(rows=len(qa_items), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, item in enumerate(qa_items):
        row = table.rows[i]
        # Checkbox column
        cell0 = row.cells[0]
        cell0.width = Inches(0.4)
        p0 = cell0.paragraphs[0]
        run0 = p0.add_run("\u2610")  # Ballot box
        run0.font.size = Pt(14)
        run0.font.color.rgb = COLOR_SECONDARY
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Item column
        cell1 = row.cells[1]
        cell1.width = Inches(6.1)
        p1 = cell1.paragraphs[0]
        add_formatted_runs(p1, item)
        if i % 2 == 1:
            set_cell_shading(cell1, COLOR_TABLE_ALT)
            set_cell_shading(cell0, COLOR_TABLE_ALT)

    doc.add_paragraph()
    doc.add_page_break()

    # ═══════════════════════════════════════════════════
    # CHAPTER 15: References
    # ═══════════════════════════════════════════════════
    doc.add_heading("15. References", level=1)

    refs = [
        "Ajzen, I. (1991). The theory of planned behavior. OBHDP, 50(2), 179\u2013211.",
        "Bandura, A. (1986). Social foundations of thought and action. Prentice-Hall.",
        "Cheung, M. W.-L. (2015). Meta-analysis: A structural equation modeling approach. Wiley.",
        "Cheung, M. W.-L., & Chan, W. (2005). Meta-analytic structural equation modeling: A two-stage approach. Psychological Methods, 10(1), 40\u201364.",
        "Davis, F. D. (1989). Perceived usefulness, perceived ease of use, and user acceptance of information technology. MIS Quarterly, 13(3), 319\u2013340.",
        "Hofstede, G. (2001). Culture\u2019s consequences (2nd ed.). Sage.",
        "Jak, S., & Cheung, M. W.-L. (2020). Meta-analytic structural equation modeling with moderating effects. Psychological Methods, 25(4), 430\u2013455.",
        "Mayer, R. C., Davis, J. H., & Schoorman, F. D. (1995). An integrative model of organizational trust. AMR, 20(3), 709\u2013734.",
        "Peterson, R. A., & Brown, S. P. (2005). On the use of beta coefficients in meta-analysis. JAP, 90(1), 175\u2013181.",
        "Venkatesh, V., Morris, M. G., Davis, G. B., & Davis, F. D. (2003). User acceptance of information technology. MIS Quarterly, 27(3), 425\u2013478.",
        "Venkatesh, V., Thong, J. Y. L., & Xu, X. (2012). Consumer acceptance and use of information technology. MIS Quarterly, 36(1), 157\u2013178.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        run = p.add_run(ref)
        run.font.size = Pt(10)
        run.font.name = FONT_BODY

    doc.add_page_break()

    # ═══════════════════════════════════════════════════
    # APPENDIX A: Decision Trees
    # ═══════════════════════════════════════════════════
    doc.add_heading("Appendix A: Decision Trees", level=1)

    doc.add_heading("A.1 Inclusion/Exclusion Decision", level=3)
    inc_steps = [
        "Is it an empirical quantitative study? \u2192 NO \u2192 Exclude (E1)",
        "Does it study AI technology? \u2192 NO \u2192 Exclude (E2)",
        "Is it in an **educational setting**? \u2192 NO \u2192 Exclude (E3)",
        "Does it measure adoption/acceptance? \u2192 NO \u2192 Exclude (E4)",
        "Does it report correlations or betas? \u2192 NO \u2192 Exclude (E5)",
        "Does it measure \u2265 2 of our 12 constructs? \u2192 NO \u2192 Exclude (E-FT1)",
        "Is N \u2265 50? \u2192 NO \u2192 Exclude (E8)",
        "Is it in English? \u2192 NO \u2192 Exclude (E6)",
        "Is it peer-reviewed (2022\u20132026)? \u2192 NO \u2192 Exclude (E7/E9)",
        "Is it a duplicate sample? \u2192 YES \u2192 Exclude (E10)",
        "**INCLUDE**",
    ]
    for i, step in enumerate(inc_steps, 1):
        add_numbered(doc, step)

    doc.add_heading("A.2 AI Technology Classification", level=3)
    add_bullet(doc, "**Generative AI:** ChatGPT, GPT-4, Claude, Gemini, DALL-E, Midjourney, Copilot")
    add_bullet(doc, "**Predictive AI:** ML-based forecasting, diagnostic AI, risk scoring")
    add_bullet(doc, "**Decision Support:** AI-assisted decision making, recommendation systems")
    add_bullet(doc, "**Conversational AI:** Chatbots, virtual assistants (Siri, Alexa), AI tutors")
    add_bullet(doc, "**Robotic AI:** Physical robots, surgical robots, educational robots")
    add_bullet(doc, "**General:** Study discusses \u201cAI\u201d without specifying type")

    doc.add_heading("A.3 Construct Harmonization Decision", level=3)
    harm_steps = [
        "Exact label match to our 12? \u2192 Code directly (exact confidence)",
        "TAM/UTAUT family label? \u2192 Use cross-reference table (high confidence)",
        "AI-specific label? \u2192 Use AI mapping table (high/exact confidence)",
        "Definition aligns? \u2192 Map to closest construct (moderate confidence)",
        "Items align? \u2192 Map based on items (moderate confidence)",
        "Ambiguous? \u2192 Flag for expert review (low confidence)",
        "No match? \u2192 Exclude construct from coding",
    ]
    for i, step in enumerate(harm_steps, 1):
        add_numbered(doc, step)

    doc.add_heading("A.4 Correlation Source Selection", level=3)
    add_numbered(doc, "Direct Pearson r available? \u2192 Use direct r (best quality)")
    add_numbered(doc, "Only standardized \u03b2 available? \u2192 Convert using Peterson & Brown (2005)")
    add_numbered(doc, "Only unstandardized b? \u2192 Cannot convert; mark as missing")
    add_numbered(doc, "Implied correlations from SEM? \u2192 Use if no direct r; flag quality")
    add_numbered(doc, "No numeric data? \u2192 Contact authors or exclude")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════
    # APPENDIX B: FAQ
    # ═══════════════════════════════════════════════════
    doc.add_heading("Appendix B: FAQ", level=1)

    faqs = [
        ("What if a study reports both Pearson r and SEM-implied correlations?",
         "Use Pearson r (direct correlations). SEM-implied correlations are model-dependent."),
        ("What if a study measures a construct but does not include it in the correlation matrix?",
         "Record the construct as measured but leave correlations as NA."),
        ("How do I handle reverse-coded constructs?",
         "Flip the sign. E.g., Complexity (reverse EE) with r = \u22120.35 to BI becomes EE-BI r = 0.35. Document the reversal."),
        ("What if a study reports correlation but not reliability?",
         "Code the correlations. Leave reliability fields as NA."),
        ("What if sample sizes differ across correlation pairs?",
         "Record pairwise N in r_sample_size. Use the minimum N for conservative analysis."),
        ("How do I handle multiple samples in one paper?",
         "Treat each independent sample as a separate entry (e.g., S005a, S005b)."),
        ("What if a study uses a construct not in our 12?",
         "Do not code it. Only code correlations involving our 12 target constructs."),
        ("What about second-order constructs?",
         "If a study measures PE as a second-order factor, use the second-order correlation. If only first-order correlations available, average and flag quality."),
        ("Can AI pre-code metadata fields?",
         "Yes. AI pre-codes non-critical fields (identification, demographics). Humans verify during their independent coding. AI does NOT pre-code correlations or construct mappings \u2014 those must be human-coded independently."),
    ]

    for q, a in faqs:
        p = doc.add_paragraph()
        run = p.add_run("Q: " + q)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_PRIMARY
        p = doc.add_paragraph()
        run = p.add_run("A: " + a)
        run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(10)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════
    # APPENDIX C: Excel Template
    # ═══════════════════════════════════════════════════
    doc.add_heading("Appendix C: Excel Template Sheet Descriptions", level=1)

    add_styled_table(doc,
        ["Sheet", "Purpose"],
        [
            ["CODEBOOK", "Master reference for all variables across all sheets. Contains variable_name, sheet, type, valid_values, coding_rules, example, and notes."],
            ["STUDY_METADATA", "One row per included study. Identification, design, sample, AI technology, quality assessment, and source management variables."],
            ["CORRELATION_MATRIX", "One row per construct pair per study. study_id, construct_row, construct_col, r_value, r_source, and extraction metadata."],
            ["CONSTRUCT_MAPPING", "Documents how each study construct maps to our 12 target constructs. Includes harmonization confidence level."],
            ["MODERATOR_VARIABLES", "Moderator values for each study. Used in OSMASEM and subgroup analyses."],
            ["AI_EXTRACTION_PROVENANCE", "Raw AI model outputs for AI-assisted coding. Preserves full provenance chain for Claude, Gemini, and Codex."],
            ["DISCREPANCY_LOG", "All inter-coder disagreements with classification, resolution method, and final values."],
            ["EXCLUSION_LOG", "All studies excluded at full-text or data extraction stage with reasons and codes."],
        ],
        col_widths=[2.2, 4.3],
        first_col_bold=True,
    )

    # ── Final save ──
    doc.save(OUTPUT_PATH)
    print(f"Document saved to: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
